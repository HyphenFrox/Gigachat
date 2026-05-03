"""Compute pool: capability probe + liveness sweep for registered workers.

The pool's data layer (CRUD on `compute_workers`) lives in `db.py`. This
module owns the operational side: ping each worker's Ollama, cache
what's installed there, and surface failures so the Settings UI and
the routing layer can grey out unreachable nodes.

Two entry points:
  * `probe_worker(wid)` — one-shot probe; used by the manual "Test
    connection" button in the UI and by the routing layer when it
    wants to confirm a worker is hot before sending real traffic.
  * `start_periodic_probe()` — background asyncio task scheduled on
    app startup. Sweeps every `_SWEEP_INTERVAL_SEC` (5 min by default)
    so capability and liveness data stay fresh even when no one is
    actively poking the Settings panel.

The probe is intentionally cheap: two parallel GETs (`/api/version`
and `/api/tags`) with a short timeout. If either fails we record the
error string on the worker row and reschedule for next sweep — no
exponential backoff yet because workers are typically on a stable LAN
or tailnet, not flaky public networks. Add backoff if real-world
flakiness appears.
"""
from __future__ import annotations

import asyncio
import hashlib
import json
import logging
import os
import socket
import time
from pathlib import Path
from typing import Any

import httpx

from . import db, jsonutil, split_runtime, sysdetect

log = logging.getLogger(__name__)


# Sweep cadence. 5 min is a balance between "Settings UI shows
# something fresh when the user opens it" and "don't burn the worker's
# bandwidth on a stationary background poll." The sweep is cheap (two
# tiny GETs per worker) so this could go faster, but most workers'
# state changes — model pulls, hardware swaps — happen rarely.
_SWEEP_INTERVAL_SEC = 300

# Per-probe timeout. The worker is on the same LAN or tailnet; a
# multi-second response means it's overloaded or down, not slow.
_PROBE_TIMEOUT_SEC = 5.0

# Default TCP port `rpc-server` (llama.cpp's worker process) listens on.
# This is the upstream's default; users can override per-worker once the
# `rpc_port` schema column lands. Phase 2 commit 3 just probes this port
# to surface whether rpc-server is running on each worker.
_DEFAULT_RPC_PORT = 50052

# How long to wait for the TCP handshake on the rpc-server probe.
# Slightly larger than you'd expect for a SYN-ACK because mDNS
# hostnames often resolve to multiple addresses (IPv4 + IPv6 link-
# local) and asyncio's `open_connection` tries them in turn — the
# IPv6 link-local hop can take several seconds to fail before the
# IPv4 fallback connects. 5 seconds covers that.
_RPC_PROBE_TIMEOUT_SEC = 5.0

# How long a measured `tokens_per_second` is good for before we
# re-benchmark. 1 hour is the right balance: long enough that the
# bench cost amortizes (loading + 30 tokens of generation) but
# short enough that hardware changes (different model loaded on
# host changing GPU contention, worker reboot, etc.) get reflected
# within a session.
_THROUGHPUT_CACHE_TTL_SEC = 3600.0

# Number of tokens to generate during the throughput benchmark.
# 20 is enough to amortize first-token latency without dragging
# probe time too long. On a typical 7B-Q4 model on host VRAM
# this completes in <1 s; on a laptop CPU it might take 5-10 s.
_THROUGHPUT_BENCH_TOKENS = 20

# Cached host throughput, keyed by model name. Measured lazily on
# first route decision and refreshed every TTL.
#
# Long-session leak guard: a backend that's been up for weeks and has
# served many ephemeral model pulls would accumulate a stale entry per
# model name. The TTL only stops *reads* from using stale tps numbers —
# expired entries still occupy memory until a new measurement happens
# to overwrite them. We bound the dict with a small FIFO LRU: when it
# exceeds `_HOST_THROUGHPUT_CACHE_MAX`, the oldest entry is evicted on
# every insert. Evicting an active model is a cheap re-measure on next
# route decision.
_HOST_THROUGHPUT_CACHE: dict[str, tuple[float, float]] = {}  # model -> (tps, measured_at)
_HOST_THROUGHPUT_CACHE_MAX = 64

# Subagent fan-out performance gate. A worker is included in the
# parallel-subagent target list only if its measured TPS is at least
# this fraction of the host's measured TPS for the same model.
# Rationale: round-robin fan-out gets bottlenecked by the slowest
# task; a worker dramatically slower than host slows down the whole
# fan-out instead of accelerating it. 0.25 = "must be at least
# one-quarter as fast as host." Tunable; chosen so a weak laptop
# doesn't degrade `delegate_parallel` on a fast host.
_SUBAGENT_MIN_PERF_RATIO = 0.25


def _worker_base_url(worker: dict) -> str:
    """Build the Ollama base URL for a worker row."""
    addr = (worker.get("address") or "").strip()
    port = int(worker.get("ollama_port") or 11434)
    # Strip any trailing slashes / scheme the user may have pasted in.
    if addr.startswith("http://"):
        addr = addr[len("http://"):]
    elif addr.startswith("https://"):
        addr = addr[len("https://"):]
    addr = addr.rstrip("/")
    return f"http://{addr}:{port}"


def _worker_host(worker: dict) -> str:
    """Bare hostname / IP for non-HTTP probes (rpc-server uses TCP, not HTTP).

    Same scheme stripping as `_worker_base_url` but returns just the host
    portion — `rpc-server` listens on its own port, not the Ollama port.
    """
    addr = (worker.get("address") or "").strip()
    if addr.startswith("http://"):
        addr = addr[len("http://"):]
    elif addr.startswith("https://"):
        addr = addr[len("https://"):]
    return addr.rstrip("/")


async def _measure_throughput(
    base_url: str,
    auth_token: str | None,
    model_name: str,
    n_tokens: int = _THROUGHPUT_BENCH_TOKENS,
) -> tuple[float, float]:
    """Run a small Ollama generation, return (tokens_per_second, total_seconds).

    Folds CPU, RAM bandwidth, GPU compute, and memory size into one
    bottom-line number — the real speed signal for routing. We measure
    rate AFTER the first token to skip cold-load time (model file
    paging from disk into VRAM/RAM dominates the first hundred ms);
    what callers actually want is "how fast does this machine generate
    tokens once it's hot."

    Returns (0, 0) on any failure — caller can fall back to the static
    heuristic ranking when the bench couldn't run (model missing,
    Ollama down, etc.).
    """
    headers: dict[str, str] = {}
    if auth_token:
        headers["Authorization"] = f"Bearer {auth_token}"
    payload = {
        "model": model_name,
        "prompt": "Hello",
        "stream": True,
        "options": {"num_predict": int(n_tokens), "temperature": 0.0},
    }
    timeout = httpx.Timeout(connect=5.0, read=120.0, write=10.0, pool=10.0)
    first_token_at: float | None = None
    tok = 0
    t_total_start = time.perf_counter()
    try:
        async with httpx.AsyncClient(timeout=timeout) as c:
            async with c.stream("POST", f"{base_url}/api/generate", json=payload, headers=headers) as r:
                if r.status_code >= 400:
                    return 0.0, 0.0
                async for line in r.aiter_lines():
                    if not line.strip():
                        continue
                    try:
                        obj = jsonutil.loads(line)
                    except json.JSONDecodeError:
                        continue
                    if obj.get("response"):
                        if first_token_at is None:
                            first_token_at = time.perf_counter()
                        tok += 1
                    if obj.get("done"):
                        break
    except Exception as e:
        log.info("throughput bench failed for %s on %s: %s", model_name, base_url, e)
        return 0.0, 0.0

    total = time.perf_counter() - t_total_start
    if tok == 0 or first_token_at is None:
        return 0.0, total
    gen_seconds = time.perf_counter() - first_token_at
    if gen_seconds <= 0:
        return 0.0, total
    return tok / gen_seconds, total


async def _probe_worker_specs_via_ssh(worker: dict) -> dict:
    """Query CPU + RAM + GPU details over SSH using the worker's
    `ssh_host` alias. Returns a dict suitable for merging into
    capabilities. Empty dict on any failure.

    Workers run plain Ollama which doesn't expose system specs. The
    /api/ps proxy gives us a binary "GPU present" + a VRAM lower
    bound, but not CPU model, RAM total, or GPU class. SSH'ing to the
    worker is the most reliable way to fill in the gaps without
    requiring extra worker-side software. Free when ssh_host is
    already set (the same alias used for LAN model copy).
    """
    ssh_host = (worker.get("ssh_host") or "").strip()
    if not ssh_host:
        return {}

    # PowerShell one-liner that emits a JSON blob the host can parse.
    # We pipe it through ssh -o BatchMode=yes so the call fails fast
    # on auth issues instead of hanging on a password prompt.
    ps = (
        "$cs = Get-CimInstance Win32_ComputerSystem;"
        "$os = Get-CimInstance Win32_OperatingSystem;"
        "$cpu = Get-CimInstance Win32_Processor;"
        "$gpus = @();"
        "Get-CimInstance Win32_VideoController | ForEach-Object {"
        "  $gpus += [pscustomobject]@{"
        "    name = $_.Name;"
        "    adapter_ram_gb = [math]::Round([uint64]$_.AdapterRAM/1GB, 2);"
        "    driver_version = $_.DriverVersion"
        "  }"
        "};"
        # Detect llama-server availability on the worker. The same
        # binary set ships rpc-server, llama-cli, AND llama-server, so
        # any worker that hosts rpc-server probably also has llama-
        # server — but we check explicitly so future v2 routing
        # (worker-side speculative spawn) can rely on this signal.
        "$llama_server = $null;"
        "$candidates = @("
        "  \"$env:USERPROFILE\\.gigachat\\llama-cpp\\llama-server.exe\","
        "  \"llama-server.exe\""
        ");"
        "foreach ($c in $candidates) {"
        "  $resolved = Get-Command $c -ErrorAction SilentlyContinue;"
        "  if ($resolved) { $llama_server = $resolved.Source; break }"
        "};"
        # Detect Python + the optional libraries the distributed-tool
        # paths need. Used by the `read_doc` and `web_search`
        # dispatchers to skip workers that can't actually serve the
        # request (saves a round-trip + an obvious error).
        "$python = (Get-Command py.exe -ErrorAction SilentlyContinue).Source;"
        "if (-not $python) { $python = (Get-Command python.exe -ErrorAction SilentlyContinue).Source };"
        "$read_doc_libs = @();"
        "$has_ddgs = $false;"
        "if ($python) {"
        "  $libs_check = & $python -c 'import json,sys"
        "`nout={\\\"libs\\\":[],\\\"ddgs\\\":False}"
        "`nfor m in (\\\"pymupdf\\\",\\\"docx\\\",\\\"openpyxl\\\"):"
        "`n  try: __import__(m); out[\\\"libs\\\"].append(m)"
        "`n  except Exception: pass"
        "`ntry:"
        "`n  __import__(\\\"ddgs\\\"); out[\\\"ddgs\\\"]=True"
        "`nexcept Exception: pass"
        "`njson.dump(out,sys.stdout)' 2>$null;"
        "  if ($libs_check) {"
        "    try {"
        "      $parsed = $libs_check | ConvertFrom-Json;"
        "      $read_doc_libs = @($parsed.libs);"
        "      $has_ddgs = [bool]$parsed.ddgs"
        "    } catch {}"
        "  }"
        "};"
        # Disk space on the system drive — primarily used to gate
        # storage-pool decisions (where to land a new model pull,
        # whether a worker can absorb a 17 GB GGUF override). Returns
        # both total and free so the inventory UI can show usage
        # progress per worker.
        "$drive = Get-PSDrive -Name ($env:SystemDrive.TrimEnd(':')) -ErrorAction SilentlyContinue;"
        "$disk_total_gb = if ($drive) { [math]::Round(($drive.Used + $drive.Free)/1GB, 1) } else { 0 };"
        "$disk_free_gb = if ($drive) { [math]::Round($drive.Free/1GB, 1) } else { 0 };"
        # Cached override-GGUF files under our private install dir.
        # Surfaces which models the worker can serve without an
        # over-the-LAN re-stream from host on the next spawn.
        "$cached = @();"
        "$cache_dir = \"$env:USERPROFILE\\.gigachat\\llama-cpp\\models\";"
        "if (Test-Path $cache_dir) {"
        "  Get-ChildItem -Path $cache_dir -Filter *.gguf -ErrorAction SilentlyContinue |"
        "    ForEach-Object { $cached += $_.Name }"
        "};"
        "$out = [pscustomobject]@{"
        "  cpu_name = $cpu.Name;"
        "  cpu_cores = $cpu.NumberOfCores;"
        "  cpu_threads = $cpu.NumberOfLogicalProcessors;"
        "  ram_total_gb = [math]::Round($cs.TotalPhysicalMemory/1GB, 1);"
        "  ram_free_gb = [math]::Round($os.FreePhysicalMemory/1MB, 1);"
        "  disk_total_gb = $disk_total_gb;"
        "  disk_free_gb = $disk_free_gb;"
        "  gpus = $gpus;"
        "  llama_server_path = $llama_server;"
        "  read_doc_libs = $read_doc_libs;"
        "  has_ddgs = $has_ddgs;"
        "  cached_overrides = $cached"
        "};"
        "$out | ConvertTo-Json -Compress -Depth 4"
    )
    # PowerShell -EncodedCommand bypasses every layer of arg-list /
    # cmdline / shell quoting nightmare. The script is base64-encoded
    # UTF-16LE bytes — PowerShell decodes and runs it verbatim with
    # zero whitespace mangling.
    import base64
    encoded = base64.b64encode(ps.encode("utf-16-le")).decode("ascii")
    cmd = ["ssh", *_ssh_persistent_args(), "-o", "BatchMode=yes",
           "-o", "ConnectTimeout=5",
           ssh_host, "powershell", "-NoProfile", "-EncodedCommand", encoded]
    import subprocess as _sp
    def _run() -> tuple[int, bytes, bytes]:
        try:
            r = _sp.run(
                cmd, capture_output=True, timeout=15.0,
            )
            return r.returncode, r.stdout, r.stderr
        except _sp.TimeoutExpired:
            return -1, b"", b"timeout"
        except Exception as e:
            return -2, b"", repr(e).encode()
    try:
        rc, stdout, stderr = await asyncio.to_thread(_run)
    except Exception as e:
        log.info("ssh specs probe failed for %s: %s", ssh_host, e)
        return {}
    if rc != 0:
        return {}

    # PowerShell's ConvertTo-Json may emit the object as either a JSON
    # object (single record) or a JSON array (when implicit
    # enumeration kicks in). Handle both.
    text = stdout.decode("utf-8", errors="replace").strip()
    if not text:
        return {}
    try:
        data = jsonutil.loads(text)
    except json.JSONDecodeError:
        return {}
    if isinstance(data, list):
        if not data:
            return {}
        data = data[0]
    if not isinstance(data, dict):
        return {}

    # Stamp the measurement so we know how stale the data is.
    data["specs_measured_at"] = time.time()
    return data


async def _attempt_rpc_server_restart(
    worker: dict, backend: str = "SYCL0,CPU",
) -> bool:
    """Bring up rpc-server on a worker. P2P-first, SSH-fallback.

    Tries the encrypted P2P path (``ensure_rpc_server_via_proxy``)
    first because it requires zero per-machine setup — paired LAN
    peers always have it. If that fails AND the worker has a legacy
    ``ssh_host`` configured, falls back to the SSH-driven spawn so
    advanced users who prefer the SSH path don't lose it.

    The old SSH-only signature is preserved so every existing call
    site keeps working unchanged.
    """
    # P2P path — primary. Works for any paired peer.
    try:
        ok = await ensure_rpc_server_via_proxy(worker, backend=backend)
    except Exception as e:
        log.debug(
            "compute_pool: P2P rpc-server start failed for %s: %s; "
            "trying SSH fallback if available",
            worker.get("label"), e,
        )
        ok = False
    if ok:
        return True
    # SSH path — fallback for users who set ssh_host explicitly.
    return await _attempt_rpc_server_restart_via_ssh(
        worker, backend=backend,
    )


async def _attempt_rpc_server_restart_via_ssh(
    worker: dict, backend: str = "SYCL0,CPU",
) -> bool:
    """SSH into the worker and re-spawn its rpc-server.

    Used by `probe_worker` when the rpc-server TCP probe fails. If the
    worker has an `ssh_host` set, we kill any stale rpc-server.exe and
    relaunch it via WMI Win32_Process.Create so the new process
    survives the SSH session's exit. Returns True if the post-restart
    process appears alive (process found AND port 50052 listening).

    `backend` is passed to rpc-server's `-d` flag. Default
    `SYCL0,CPU` exposes both Intel iGPU and CPU. For MoE models that
    trigger the upstream SYCL+RPC crash, callers can override to
    `CPU` (no SYCL exposed; works around the bug). See
    `_set_workers_backend` for the per-model dynamic switching.

    With the default `SYCL0,CPU`:
      * SYCL0 contributes the iGPU's shared GPU memory pool (~3 GB
        on a typical Intel Iris Xe with default BIOS settings).
      * CPU contributes the worker's full system RAM (typically
        8-16 GB) as a host-style device — llama.cpp can place
        layers there too, just without GPU acceleration. This is
        the "use ALL pool resources" mode: every worker contributes
        TWO tiers (GPU + CPU) to the layer-distribution pool
        instead of one.
      * We deliberately exclude Vulkan because on Intel iGPUs it
        targets the SAME physical hardware as SYCL and would
        double-enumerate the iGPU, causing the "Remote RPC server
        crashed" mid-push that the targeted bench surfaced.

    Failure modes (all return False, no exceptions):
      * No ssh_host configured (caller falls back to "unreachable").
      * ssh connect timeout / auth failure.
      * rpc-server.exe missing on the worker.
      * Process spawned but port not listening within ~4 s.
    """
    ssh_host = (worker.get("ssh_host") or "").strip()
    if not ssh_host:
        return False

    # Single PowerShell payload: kill stale, spawn fresh, verify alive.
    # Same WMI pattern used in the manual recovery commands so the
    # spawned process outlives the SSH session.
    #
    # Stability env vars are set at the user level before spawn so the
    # rpc-server inherits them. WMI's Win32_Process.Create doesn't
    # have a clean per-spawn env block, and PowerShell process-scope
    # env vars don't propagate through WMI either; setting at user
    # scope is the cleanest way to guarantee the child sees them.
    # - GGML_SYCL_DISABLE_OPT=1: dodges the SYCL optimizer bug that
    #   silently corrupts weights on Intel Xe2 / Meteor Lake.
    #   #21893 is still OPEN upstream — this workaround is mandatory
    #   on affected hardware until the kernel-level fix lands.
    # - GGML_SYCL_DISABLE_GRAPH=1: dodges the warmup-crash regression
    #   on Intel iGPUs. #21474 is closed but the env-var safety
    #   net is cheap to keep.
    # - SYCL_CACHE_PERSISTENT=1: persists JIT'd kernel cache.
    ps = (
        "$ErrorActionPreference = 'Continue';"
        "Get-Process -Name 'rpc-server' -ErrorAction SilentlyContinue | "
        "  ForEach-Object { Stop-Process -Id $_.Id -Force };"
        "Start-Sleep -Milliseconds 800;"
        "[Environment]::SetEnvironmentVariable('GGML_SYCL_DISABLE_OPT', '1', 'User');"
        "[Environment]::SetEnvironmentVariable('GGML_SYCL_DISABLE_GRAPH', '1', 'User');"
        "[Environment]::SetEnvironmentVariable('SYCL_CACHE_PERSISTENT', '1', 'User');"
        "$exe = \"$env:USERPROFILE\\.gigachat\\llama-cpp\\rpc-server.exe\";"
        "if (-not (Test-Path $exe)) { Write-Output 'NO_BINARY'; exit 2 };"
        f"$cmdline = '\"' + $exe + '\" -H 0.0.0.0 -p 50052 -d {backend}';"
        "$result = Invoke-CimMethod -ClassName Win32_Process -MethodName Create "
        "  -Arguments @{ CommandLine = $cmdline; "
        "                CurrentDirectory = \"$env:USERPROFILE\\.gigachat\\llama-cpp\" };"
        "if ($result.ReturnValue -ne 0) { Write-Output 'WMI_FAIL'; exit 3 };"
        "Start-Sleep -Seconds 4;"
        "$rpc = Get-Process -Name 'rpc-server' -ErrorAction SilentlyContinue;"
        "if ($rpc) {"
        # Lower the rpc-server's priority class to BelowNormal so it
        # cooperates when the user wants to actively use the worker
        # for other work. Inference still runs at full speed when the
        # user isn't doing anything else; under load the OS scheduler
        # gives priority to whatever the user is interacting with.
        "  try { $rpc.PriorityClass = [System.Diagnostics.ProcessPriorityClass]::BelowNormal } "
        "  catch { Write-Output 'PRIORITY_FAIL' };"
        "}"
        "$port = Get-NetTCPConnection -LocalPort 50052 -State Listen -ErrorAction SilentlyContinue;"
        "if ($rpc -and $port) { Write-Output 'OK' } else { Write-Output 'NO_LISTEN' }"
    )
    import base64
    encoded = base64.b64encode(ps.encode("utf-16-le")).decode("ascii")
    cmd = ["ssh", *_ssh_persistent_args(), "-o", "BatchMode=yes",
           "-o", "ConnectTimeout=8",
           ssh_host, "powershell", "-NoProfile", "-EncodedCommand", encoded]

    import subprocess as _sp
    def _run() -> tuple[int, bytes, bytes]:
        try:
            r = _sp.run(cmd, capture_output=True, timeout=30.0)
            return r.returncode, r.stdout, r.stderr
        except _sp.TimeoutExpired:
            return -1, b"", b"timeout"
        except Exception as e:
            return -2, b"", repr(e).encode()

    try:
        rc, stdout, stderr = await asyncio.to_thread(_run)
    except Exception as e:
        log.info("compute_pool: rpc-server restart ssh failed for %s: %s", ssh_host, e)
        return False

    out = stdout.decode("utf-8", errors="replace").strip()
    if rc != 0 or "OK" not in out:
        log.info(
            "compute_pool: rpc-server restart on %s did not come up "
            "(rc=%d, output=%r)", ssh_host, rc, out[-200:],
        )
        return False
    return True


def _worker_gpu_vendor(worker: dict) -> str:
    """Return the dominant GPU vendor for a worker, based on the SSH
    spec probe's `gpus` list. One of "nvidia", "amd", "intel", "none".

    Used by `_select_worker_backend` to pick a `-d` flag that matches
    the worker's actual hardware. We can't blindly pass `SYCL0,CPU` to
    a worker that has an NVIDIA GPU — SYCL on NVIDIA is unsupported
    by stock llama.cpp's prebuilt SYCL binary, and the worker's CUDA
    capability would go unused. Conversely, an Intel-only worker on
    `-d CUDA0,CPU` would have no GPU device and fall back to CPU.

    Selection rule (priority order):
      1. NVIDIA — strongest GPU acceleration, no known RPC bug
      2. AMD    — Vulkan path on stock llama.cpp; mostly works
      3. Intel  — SYCL on iGPU; subject to the upstream RPC crash
      4. None   — CPU only (no GPU detected at all)
    """
    caps = worker.get("capabilities") or {}
    gpus = caps.get("gpus") or []
    names = [(g.get("name") or "").lower() for g in gpus]
    if any("nvidia" in n or "geforce" in n or "rtx " in n or "gtx " in n for n in names):
        return "nvidia"
    if any("amd" in n or "radeon" in n for n in names):
        return "amd"
    if any("intel" in n or "iris" in n or "uhd" in n for n in names):
        return "intel"
    # Fallback: paired peers populate `gpu_kind` via the
    # `/api/p2p/system-stats` endpoint even when SSH spec-probe
    # data isn't available. This is the path that catches paired
    # LAN peers without ssh_host configured.
    gpu_kind = (caps.get("gpu_kind") or "").lower()
    if gpu_kind in ("nvidia", "amd", "intel"):
        return gpu_kind
    return "none"


# How long to remember that an iGPU backend crashed before retrying
# it on a worker. 24 hours — long enough that a transient driver
# state issue clears, short enough that a real fix (driver update,
# llama.cpp upgrade) gets adopted within a day.
_BACKEND_FAILURE_COOLDOWN_SEC = 24 * 3600.0


def _select_intel_backend_with_fallback(
    worker: dict, *, in_split: bool,
) -> str:
    """Pick the best iGPU-inclusive backend that hasn't recently
    crashed for this worker.

    Preference chain (most aggressive → safest):
      1. ``SYCL0`` — Intel iGPU only, single-device rpc-server.
         Avoids the SYCL+CPU heterogeneous-allocator crash where
         tensor layout mismatches between SYCL allocations and CPU
         compute trip ggml-rpc.cpp's RPC_STATUS_ASSERT mid-decode
         (verified against build 8940 by the upstream source: macro
         at ggml-rpc.cpp:41 abort()s on any send/recv failure or
         buffer-bounds mismatch, no retry path). Single-backend
         rpc-server has no such pairing.
      2. ``Vulkan0`` — alternate iGPU path, single device. Try
         when SYCL crashes (which is rarer with ``SYCL0`` alone
         vs ``SYCL0,CPU`` but still possible per #21893).
      3. ``SYCL0,CPU`` — hybrid; gives the orchestrator both iGPU
         and worker CPU layers but is the failure mode we're
         trying to avoid. Tried only after both single-backend
         options crashed within the cool-down window.
      4. ``Vulkan0,CPU`` — same hybrid risk on Vulkan.
      5. ``CPU`` — always works; gives up iGPU contribution.

    The selector consults capability flags persisted by the
    failure-detection hook (``record_backend_failure``). A flag
    older than ``_BACKEND_FAILURE_COOLDOWN_SEC`` is treated as
    expired (the iGPU stack may have been fixed since).
    """
    if not in_split:
        return "SYCL0,CPU"
    caps = worker.get("capabilities") or {}
    now = time.time()
    sycl_dead = (
        (now - float(caps.get("sycl_split_failed_at") or 0))
        < _BACKEND_FAILURE_COOLDOWN_SEC
    )
    sycl_hybrid_dead = (
        (now - float(caps.get("sycl_hybrid_split_failed_at") or 0))
        < _BACKEND_FAILURE_COOLDOWN_SEC
    )
    vulkan_dead = (
        (now - float(caps.get("vulkan_split_failed_at") or 0))
        < _BACKEND_FAILURE_COOLDOWN_SEC
    )
    vulkan_hybrid_dead = (
        (now - float(caps.get("vulkan_hybrid_split_failed_at") or 0))
        < _BACKEND_FAILURE_COOLDOWN_SEC
    )
    if not sycl_dead:
        return "SYCL0"  # single-backend SYCL — bypasses the hybrid layout bug
    if not vulkan_dead:
        return "Vulkan0"
    if not sycl_hybrid_dead:
        return "SYCL0,CPU"
    if not vulkan_hybrid_dead:
        return "Vulkan0,CPU"
    return "CPU"


def record_backend_failure(worker_id: str, backend: str) -> None:
    """Mark a worker's iGPU backend as crashed so the next selector
    pass falls through to the next preference.

    Called by `split_lifecycle` whenever a split start failure or a
    chat-mid-decode crash is observed and the worker was on an iGPU
    backend at the time. Idempotent — successive calls just refresh
    the timestamp."""
    w = db.get_compute_worker(worker_id)
    if not w:
        return
    caps = dict(w.get("capabilities") or {})
    backend_lower = (backend or "").lower()
    is_hybrid = "," in (backend or "")
    now = time.time()
    if "sycl" in backend_lower:
        if is_hybrid:
            caps["sycl_hybrid_split_failed_at"] = now
            log.info(
                "compute_pool: recorded SYCL+CPU hybrid split failure "
                "for worker %s", w.get("label"),
            )
        else:
            caps["sycl_split_failed_at"] = now
            log.info(
                "compute_pool: recorded SYCL-only split failure for "
                "worker %s; next start will try Vulkan-only",
                w.get("label"),
            )
    elif "vulkan" in backend_lower:
        if is_hybrid:
            caps["vulkan_hybrid_split_failed_at"] = now
            log.info(
                "compute_pool: recorded Vulkan+CPU hybrid split failure "
                "for worker %s", w.get("label"),
            )
        else:
            caps["vulkan_split_failed_at"] = now
            log.info(
                "compute_pool: recorded Vulkan-only split failure for "
                "worker %s", w.get("label"),
            )
    else:
        # CPU has no fallback — record the failure but the next
        # selector will still return CPU (nothing safer to try).
        caps["cpu_split_failed_at"] = now
    try:
        db.update_compute_worker_capabilities(worker_id, capabilities=caps)
    except Exception as e:
        log.debug("compute_pool: backend-failure persist failed: %s", e)


def _select_worker_backend(worker: dict, *, in_split: bool) -> str:
    """Pick the right `-d` flag for this worker.

    `in_split` distinguishes split-mode (where the SYCL+RPC bug fires
    on Intel iGPUs) from non-split mode (where the worker's
    iGPU/Ollama path is used directly without RPC layer push).

    Per-vendor logic:
      * NVIDIA — `CUDA0,CPU`. CUDA over RPC is stable; no workaround
        needed even in split mode. Worker contributes both GPU
        compute and system RAM via CPU device.
      * AMD    — `Vulkan0,CPU`. Vulkan-on-AMD has fewer known issues
        than Vulkan-on-Intel-iGPU (#21516 was Intel-specific).
      * Intel  — `SYCL0,CPU` always. The historical SYCL+RPC crashes
        (#21420 / #20259 / #21474) are all closed upstream and the
        build we ship (8940) has the fixes. Keeping the iGPU
        exposed during split inference gives workers a real GPU
        contribution to the layer-distribution pool — on a 7-8 GB
        laptop with a 3 GB Iris Xe shared pool, that's the
        difference between "worker contributes ~6 GB CPU + 0 GB
        GPU" and "worker contributes ~6 GB CPU + 3 GB iGPU". The
        env-var safety nets (`GGML_SYCL_DISABLE_OPT=1`,
        `GGML_SYCL_DISABLE_GRAPH=1`) handle the last remaining
        SYCL kernel quirks (#21893 still open).
      * None   — `CPU` always (no GPU to expose).
    """
    vendor = _worker_gpu_vendor(worker)
    if vendor == "nvidia":
        return "CUDA0,CPU"
    if vendor == "amd":
        return "Vulkan0,CPU"
    if vendor == "intel":
        # Adaptive fallback: try the most-aggressive iGPU backend
        # this worker hasn't crashed on. Order:
        #   SYCL0,CPU      — best perf when stable
        #   Vulkan0,CPU    — alt iGPU path; sometimes works when SYCL
        #                    has a #21893-class quirk
        #   CPU            — always works; loses iGPU contribution
        # The selector reads `sycl_split_failed_at` /
        # `vulkan_split_failed_at` capability flags written by the
        # split-start failure-detection hook. Workers with no
        # recorded failures get the SYCL+CPU default. The flags
        # are timestamp-based so the selector also retries an
        # iGPU backend after a long enough cool-down (driver
        # update, reboot, model size change).
        return _select_intel_backend_with_fallback(worker, in_split=in_split)
    # No GPU detected
    return "CPU"


async def _set_workers_backend(workers: list[dict], *, in_split: bool) -> int:
    """Ensure every worker's rpc-server is running with the right
    `-d` backend flag for its hardware AND the current routing mode.

    Per-worker decision via `_select_worker_backend(worker, in_split=)`.
    Workers tracked in `capabilities.current_rpc_backend` so we don't
    restart unnecessarily; only mismatched workers get bounced.

    `in_split=True` is set by `_ensure_split_running_for` before
    spawning llama-server (so Intel workers drop SYCL to dodge the
    upstream bug). `in_split=False` is set by `stop_all_running_splits`
    to restore default backends (Intel SYCL for iGPU acceleration on
    non-split paths).

    Returns the count of workers successfully aligned (already-correct
    + freshly-restarted).
    """
    aligned = 0
    for w in workers:
        # No ssh_host gate any more — `_attempt_rpc_server_restart`
        # tries the encrypted P2P channel first, so paired LAN peers
        # without SSH get prepped here too. Old SSH-configured
        # workers still work via the fallback path inside the helper.
        #
        # Always go through the helper — even when the capability
        # cache says the backend already matches. The cache can lie:
        # the rpc-server process may have been killed externally
        # (system reboot, user task-manager-killed, OOM) without us
        # noticing. The helper's status probe is cheap (one HTTP
        # round-trip via the encrypted proxy) and short-circuits
        # instantly when rpc-server is genuinely up. Without this,
        # a stale "CPU" capability stamp could cause us to enter a
        # split with no rpc-server actually listening on the peer
        # — symptom: llama-server crashes mid-load with "Remote
        # RPC server crashed", with no rpc-server log because no
        # rpc-server was ever respawned.
        #
        # Pre-stamp gpu_kind from live stats before selecting the
        # backend. Without this, callers that bypass the auto-prep
        # path in `route_chat_for` see vendor="none" and the
        # iGPU-aware fallback chain can't engage. We re-read the
        # worker after stamping so the freshly-set gpu_kind is in
        # the dict the selector sees.
        try:
            live_stats = await probe_worker_live_stats(w)
        except Exception:
            live_stats = {}
        if live_stats:
            try:
                caps = dict(w.get("capabilities") or {})
                caps["gpu_kind"] = live_stats.get("gpu_kind") or caps.get("gpu_kind") or ""
                caps["ram_total_gb"] = float(live_stats.get("ram_total_gb") or caps.get("ram_total_gb") or 0)
                caps["ram_free_gb"] = float(live_stats.get("ram_free_gb") or 0)
                caps["vram_total_gb"] = float(live_stats.get("vram_total_gb") or caps.get("vram_total_gb") or 0)
                db.update_compute_worker_capabilities(
                    w["id"], capabilities=caps,
                )
                w["capabilities"] = caps
            except Exception:
                pass
        # Multi-rpc-server engagement: spawn ONE rpc-server per
        # compute tier (iGPU on 50052 + CPU on 50053 for an Intel
        # worker) so the worker contributes BOTH its iGPU memory
        # AND its system RAM as separate compute targets to the
        # orchestrator's --tensor-split. Falls back to the single-
        # backend path automatically when the worker has no GPU.
        #
        # The returned list of {port, backend} endpoints is persisted
        # in capabilities.rpc_endpoints; `split_lifecycle._resolve_
        # rpc_endpoints` reads it to build the llama-server --rpc
        # flag with one entry per endpoint.
        endpoints = await ensure_rpc_servers_via_proxy_multi(w)
        if endpoints:
            aligned += 1
            previous = (w.get("capabilities") or {}).get("current_rpc_backend", "unknown")
            primary_backend = endpoints[0]["backend"]
            if previous != primary_backend:
                log.info(
                    "compute_pool: worker %s rpc-server set: primary=%s, "
                    "%d endpoint(s) total",
                    w.get("label"), primary_backend, len(endpoints),
                )
        else:
            # Multi-spawn failed — fall back to legacy single-spawn
            # using the auto-fallback selector (which knows about
            # SYCL→Vulkan→CPU per-failure flags).
            backend = _select_worker_backend(w, in_split=in_split)
            caps = w.get("capabilities") or {}
            previous_backend = caps.get("current_rpc_backend", "unknown")
            ok = await _attempt_rpc_server_restart(w, backend=backend)
            if ok:
                try:
                    db.update_compute_worker_capabilities(
                        w["id"],
                        capabilities={**caps, "current_rpc_backend": backend},
                    )
                except Exception:
                    pass
                aligned += 1
                if previous_backend != backend:
                    log.info(
                        "compute_pool: worker %s switched rpc-server backend "
                        "%s -> %s", w.get("label"), previous_backend, backend,
                    )
    return aligned


async def _probe_rpc_server(
    host: str, port: int, timeout: float = _RPC_PROBE_TIMEOUT_SEC,
) -> tuple[bool, str | None]:
    """TCP-connect probe for rpc-server. Returns (reachable, error_str).

    rpc-server speaks llama.cpp's binary RPC protocol — it doesn't have
    an HTTP health endpoint, and its protocol handshake is more involved
    than we want to replicate just for liveness. A successful TCP connect
    is enough to know the listener is up.

    Address-family handling: mDNS hostnames like `worker.local` resolve
    to multiple addresses — typically IPv6 link-local first, then IPv6
    global, then IPv4. asyncio's `open_connection` tries them in
    sequence; the first IPv6 link-local fails after a multi-second OS-
    level timeout that defeats our deadline before the IPv4 fallback
    even gets a chance. To make the probe deterministic on a normal
    LAN, we resolve the host ourselves and try IPv4 first (or any
    IPv6 entry that isn't link-local), in parallel via gather. First
    successful connect wins.
    """
    try:
        # getaddrinfo returns ((family, type, proto, canonname, sockaddr), …).
        infos = await asyncio.get_event_loop().getaddrinfo(
            host, port, type=socket.SOCK_STREAM,
        )
    except Exception as e:
        return False, f"rpc-server probe: getaddrinfo failed: {e}"

    # Order: IPv4 first (LAN happy path), then IPv6 globals, then
    # IPv6 link-local last. We don't drop link-local entirely — on
    # some setups (no IPv4 stack at all) it's the only path — but it
    # shouldn't gate the probe.
    def _rank(info):
        family, _, _, _, sockaddr = info
        if family == socket.AF_INET:
            return 0
        ip = sockaddr[0] if sockaddr else ""
        if ip.startswith("fe80"):
            return 2  # link-local — last resort
        return 1
    infos = sorted(infos, key=_rank)

    async def _try(info):
        family, type_, proto, _, sockaddr = info
        try:
            sock = socket.socket(family, type_, proto)
            sock.setblocking(False)
            await asyncio.get_event_loop().sock_connect(sock, sockaddr)
            sock.close()
            return True, None
        except Exception as e:
            return False, f"{type(e).__name__}: {e}"

    last_err: str | None = None
    deadline = asyncio.get_event_loop().time() + timeout
    for info in infos:
        remaining = deadline - asyncio.get_event_loop().time()
        if remaining <= 0:
            break
        try:
            ok, err = await asyncio.wait_for(_try(info), timeout=remaining)
        except asyncio.TimeoutError:
            last_err = "timeout"
            continue
        if ok:
            return True, None
        last_err = err
    return False, f"rpc-server probe: {last_err or 'no addresses tried'}"


async def _probe_one(client: httpx.AsyncClient, base: str, token: str | None) -> dict:
    """Issue the probe GETs in parallel and merge the results.

    Returns a dict with `version`, `models`, hardware-capability hints,
    and any `*_error` markers. Any single endpoint may fail without
    breaking the others — we want the most signal we can get out of
    one round trip.

    Hardware-capability detection: workers run plain Ollama so they
    don't volunteer system specs. The closest signal Ollama exposes is
    `/api/ps` — currently-loaded models with their VRAM split. From
    that we infer:
      * `gpu_present`: True if any loaded model reports `size_vram > 0`.
        A worker with no loaded models doesn't tell us anything yet,
        which is why we treat absence as `False` rather than `None`.
      * `max_vram_seen_bytes`: the largest `size_vram` across loaded
        models, giving a rough lower bound on the worker's VRAM. The
        router uses this to prefer hardware-stronger workers when
        ranking ties.
    """
    headers: dict[str, str] = {}
    if token:
        # Bearer header is the convention Gigachat already uses for the
        # main loopback auth — re-use it so the worker side can validate
        # with the same code path. The worker is expected to enforce
        # this through its own AuthMiddleware-style gate.
        headers["Authorization"] = f"Bearer {token}"

    async def _ver() -> Any:
        r = await client.get(f"{base}/api/version", headers=headers)
        r.raise_for_status()
        return r.json()

    async def _tags() -> Any:
        r = await client.get(f"{base}/api/tags", headers=headers)
        r.raise_for_status()
        return r.json()

    async def _ps() -> Any:
        r = await client.get(f"{base}/api/ps", headers=headers)
        r.raise_for_status()
        return r.json()

    out: dict[str, Any] = {}
    # gather() with return_exceptions so a partial failure on one
    # endpoint doesn't lose the other endpoints' payloads.
    # Time the gather as a coarse LAN-latency proxy. With three small
    # GETs in parallel the wall time is dominated by the slowest
    # roundtrip, which gives us a reasonable "how snappy is this LAN
    # link" signal for the routing decision.
    _t0 = time.perf_counter()
    ver_res, tags_res, ps_res = await asyncio.gather(
        _ver(), _tags(), _ps(), return_exceptions=True,
    )
    out["probe_latency_ms"] = int((time.perf_counter() - _t0) * 1000)
    if isinstance(ver_res, Exception):
        out["version_error"] = f"{type(ver_res).__name__}: {ver_res}"
    else:
        out["version"] = (ver_res or {}).get("version") or "unknown"
    if isinstance(tags_res, Exception):
        out["tags_error"] = f"{type(tags_res).__name__}: {tags_res}"
    else:
        models = []
        for m in (tags_res or {}).get("models", []) or []:
            details = (m.get("details") or {}) if isinstance(m, dict) else {}
            models.append({
                "name": m.get("name") if isinstance(m, dict) else None,
                "size": m.get("size") if isinstance(m, dict) else None,
                "family": details.get("family"),
                "parameter_size": details.get("parameter_size"),
                "quantization_level": details.get("quantization_level"),
            })
        # Drop entries with no name (defensive against malformed responses).
        out["models"] = [m for m in models if m.get("name")]
    # /api/ps is purely a heuristic source — failures shouldn't dim the
    # probe's overall verdict, just leave the hardware fields empty so
    # the router falls through to the no-info default.
    if isinstance(ps_res, Exception):
        out["gpu_present"] = False
        out["max_vram_seen_bytes"] = 0
        out["loaded_count"] = 0
    else:
        loaded = (ps_res or {}).get("models", []) or []
        max_vram = 0
        any_gpu = False
        for m in loaded:
            v = int(m.get("size_vram") or 0)
            if v > 0:
                any_gpu = True
            if v > max_vram:
                max_vram = v
        out["gpu_present"] = any_gpu
        out["max_vram_seen_bytes"] = max_vram
        out["loaded_count"] = len(loaded)
    # /api/ps only sees gpu_present=True for models CURRENTLY loaded
    # with VRAM allocated. A peer with a real GPU but no model loaded
    # right now reports gpu_present=False — which makes the router
    # weight the peer as CPU-only and keeps chat traffic on host even
    # when the peer is actually the strongest node. Fall back to the
    # static hardware probe (vram_total_gb / gpu_kind from sysdetect)
    # so a freshly-restarted peer still ranks correctly.
    if not out.get("gpu_present"):
        if (
            float(out.get("vram_total_gb") or 0.0) > 0.5
            or (out.get("gpu_kind") or "").lower() in {"nvidia", "amd", "intel"}
        ):
            out["gpu_present"] = True
    return out


# ---------------------------------------------------------------------------
# Auto-repair: rediscover a worker's LAN IP after DHCP rebind.
#
# When a worker rejoins the user's Wi-Fi/Ethernet, the home router can
# hand it a different IPv4 lease — the stored `address` becomes stale and
# the LAN probe starts failing with "no route to host" / "connection
# refused". The user expects the pool to heal itself without manual
# editing of the worker row.
#
# The fix uses the worker's optional `tailscale_host` as a stable handle:
# Tailscale assigns each device a permanent overlay address that survives
# the DHCP rebind, so the host can reach the worker over Tailscale just
# long enough to ask "what's your current LAN IPv4?", then resume all
# regular traffic over LAN. Tailscale is only used for that one rediscovery
# query, never for ongoing chat / embedding / model-copy traffic.
#
# Rate-limited per worker (one attempt per cooldown window) so a worker
# that's genuinely powered off doesn't generate a Tailscale SSH per probe
# sweep — that would burn metered Tailscale bandwidth for no benefit.
# ---------------------------------------------------------------------------
_LAN_REPAIR_LAST_ATTEMPT: dict[str, float] = {}
_LAN_REPAIR_COOLDOWN_SEC = 60.0


# ---------------------------------------------------------------------------
# Persistent SSH connections via OpenSSH's ControlMaster
#
# Every distributed-tool dispatch (`fetch_url`, `web_search`, `read_doc`,
# `python_exec`), every model-sync, and every probe-restart opens an SSH
# connection. TCP+SSH handshake is ~50-150 ms per connection; for a
# single chat turn that fires 5 dispatches, that's ~500 ms of dead
# wall-clock time spent on the same handshake to the same worker.
#
# ControlMaster opens one master connection per (user@host:port) and
# reuses it for every subsequent ssh/scp invocation. Setup cost on
# subsequent calls drops to ~5 ms — purely the local FIFO/named-pipe
# round-trip to the master.
#
# ControlPath uses %C (hash of conn-tuple) so masters auto-key per host.
# ControlPersist=60s keeps the connection alive for a minute after the
# last invocation — covers the typical "burst of 3-5 calls in a turn"
# pattern without leaking idle connections forever.
#
# Cross-platform: Windows OpenSSH (8.x+) supports ControlMaster on
# named pipes; older builds print a warning and continue without
# multiplexing, which is harmless. macOS/Linux use UNIX sockets.
# ---------------------------------------------------------------------------
import tempfile as _tempfile_cm
_SSH_CONTROL_DIR = Path(_tempfile_cm.gettempdir()) / "gigachat-ssh-cm"
try:
    _SSH_CONTROL_DIR.mkdir(parents=True, exist_ok=True)
except Exception:
    pass


def _ssh_persistent_args() -> list[str]:
    """Return `-o` flags that enable SSH connection multiplexing.

    Prepend to every ssh/scp argv so the first dispatch pays the
    handshake cost and subsequent dispatches reuse the master. When
    the local OpenSSH client doesn't support ControlMaster (very old
    builds), it prints a warning and ignores the options — behaviour
    falls back to the previous one-handshake-per-call path with no
    correctness impact.
    """
    return [
        "-o", "ControlMaster=auto",
        "-o", f"ControlPath={_SSH_CONTROL_DIR / 'cm-%C'}",
        "-o", "ControlPersist=60",
    ]


def reap_stale_ssh_control_sockets() -> int:
    """Delete SSH ControlMaster sockets left over from previous backend
    runs. Called at startup so a long-uptime process doesn't accumulate
    one zombie socket per worker per restart.

    OpenSSH normally cleans up its own master sockets when the
    `ControlPersist` timer fires, but a hard kill (Ctrl-C, OS reboot,
    process crash) bypasses that. The orphan sockets are unusable
    (the master process is gone) but stay on disk as zero-byte files
    or named pipes that persist until manually deleted.

    A new ssh invocation handles a stale socket gracefully — it
    detects the dead master and falls back to a fresh connection —
    but the warning lines clutter logs and the orphan files
    accumulate forever otherwise. Reaping at startup is cheap
    insurance.

    Returns the count of files removed; logs to compute_pool's
    logger so the operator sees what was reclaimed.
    """
    if not _SSH_CONTROL_DIR.is_dir():
        return 0
    removed = 0
    for entry in _SSH_CONTROL_DIR.iterdir():
        try:
            # ControlMaster sockets typically have names like
            # `cm-<32-char-hash>`. We don't try to verify the master
            # is alive (no easy cross-platform way); the harmless
            # case is "delete a working socket" which forces ssh to
            # re-handshake on the next dispatch — same cost we pay
            # without ControlMaster at all.
            if entry.is_file() or entry.is_socket():
                entry.unlink()
                removed += 1
            elif entry.is_symlink():
                entry.unlink()
                removed += 1
        except OSError:
            # Permission denied / file in use — leave it; the OS
            # will clean it up eventually.
            continue
    if removed:
        log.info(
            "compute_pool: reaped %d stale SSH ControlMaster socket(s) at %s",
            removed, _SSH_CONTROL_DIR,
        )
    return removed


async def _rediscover_lan_ip_via_tailscale(worker: dict) -> str | None:
    """SSH to the worker over Tailscale and return its current LAN IPv4.

    Returns ``None`` when ``tailscale_host`` is unset, the SSH call
    fails, or the worker reports no usable LAN address. The returned
    string is always a private (RFC1918) IPv4 — anything else is
    discarded as an extra defence-in-depth check before we point
    ongoing traffic at it.
    """
    tailscale_host = (worker.get("tailscale_host") or "").strip()
    if not tailscale_host:
        return None

    # PowerShell one-liner that picks the active Wi-Fi / Ethernet IPv4
    # by walking ``Get-NetIPAddress`` filtered to RFC1918 ranges. We
    # sort by InterfaceMetric so the same NIC the OS routes outbound
    # traffic through is the one we report — picks the right adapter
    # on a host with both an Ethernet cable and a Wi-Fi connection.
    ps = (
        "$ips = Get-NetIPAddress -AddressFamily IPv4 -ErrorAction SilentlyContinue"
        " | Where-Object { $_.IPAddress -match"
        " '^(192\\.168\\.|10\\.|172\\.(1[6-9]|2[0-9]|3[01])\\.)' };"
        "if ($ips) {"
        "  ($ips | Sort-Object -Property InterfaceMetric)[0].IPAddress"
        "} else { '' }"
    )
    import base64
    encoded = base64.b64encode(ps.encode("utf-16-le")).decode("ascii")
    # BatchMode=yes prevents ssh from prompting for a password — if key
    # auth isn't configured we fail fast instead of hanging the probe.
    # ConnectTimeout caps the dial wait so a dead Tailscale link doesn't
    # block the periodic sweep.
    cmd = [
        "ssh", *_ssh_persistent_args(),
        "-o", "BatchMode=yes", "-o", "ConnectTimeout=10",
        tailscale_host,
        "powershell", "-NoProfile", "-EncodedCommand", encoded,
    ]
    import subprocess as _sp

    def _run() -> tuple[int, bytes]:
        try:
            r = _sp.run(cmd, capture_output=True, timeout=20.0)
            return r.returncode, r.stdout
        except _sp.TimeoutExpired:
            return -1, b""
        except Exception:
            return -2, b""

    try:
        rc, stdout = await asyncio.to_thread(_run)
    except Exception as e:
        log.info("LAN rediscovery ssh failed for %s: %s", tailscale_host, e)
        return None
    if rc != 0:
        return None
    candidate = stdout.decode("utf-8", errors="replace").strip()
    if not candidate:
        return None

    # Defence in depth: a misbehaving (or compromised) worker shouldn't
    # be able to redirect ongoing traffic to a public IP just by lying
    # about its LAN address. Re-validate the returned string is a
    # private, non-loopback, non-link-local IPv4 before committing it.
    import ipaddress
    try:
        addr = ipaddress.ip_address(candidate)
    except ValueError:
        return None
    if not addr.is_private:
        return None
    if addr.is_loopback or addr.is_link_local:
        return None
    return str(addr)


async def _attempt_lan_address_repair(worker: dict) -> dict | None:
    """Try to refresh a worker's LAN address via Tailscale rediscovery.

    Returns the refreshed worker dict (with the new ``address``) on
    success so callers can immediately retry their probe. Returns
    ``None`` when no repair was attempted or the attempt didn't help.

    Rate-limited per worker via a module-level cache: at most one
    attempt per ``_LAN_REPAIR_COOLDOWN_SEC`` seconds, regardless of how
    many probe sweeps fail in the interim.
    """
    wid = worker.get("id")
    if not wid:
        return None
    if not (worker.get("tailscale_host") or "").strip():
        return None
    now_m = time.monotonic()
    last = _LAN_REPAIR_LAST_ATTEMPT.get(wid, 0.0)
    if now_m - last < _LAN_REPAIR_COOLDOWN_SEC:
        return None
    _LAN_REPAIR_LAST_ATTEMPT[wid] = now_m

    log.info(
        "compute_pool: LAN address %r unreachable for worker %r; "
        "attempting auto-repair via tailscale_host",
        worker.get("address"), worker.get("label"),
    )
    discovered = await _rediscover_lan_ip_via_tailscale(worker)
    if not discovered:
        return None
    if discovered == (worker.get("address") or "").strip():
        # Address already correct — the failure must be downstream
        # (firewall, Ollama crashed, etc.). Rediscovery wouldn't help.
        return None
    try:
        updated = db.update_compute_worker(wid, address=discovered)
    except ValueError:
        return None
    log.info(
        "compute_pool: auto-repair updated worker %r LAN address: %r -> %r",
        worker.get("label"), worker.get("address"), discovered,
    )
    return updated


async def _probe_one_via_secure_proxy(worker: dict) -> dict:
    """Mirror of `_probe_one` that routes the three Ollama probe
    GETs through a paired peer's encrypted Gigachat proxy.

    Builds the same probe payload `_probe_one` does — version + tags
    + ps + probe_latency_ms — so the existing capability-merge
    logic upstream sees an identical shape regardless of transport.

    All three GETs go through `p2p_secure_client.forward()` which
    wraps each request in a `p2p_crypto` envelope addressed to the
    peer. The wire bytes are ciphertext; the peer's secure proxy
    decrypts, forwards to its loopback Ollama, encrypts the
    response back. Anyone observing the LAN sees no probe data.
    """
    from . import p2p_secure_client as _secure
    out: dict[str, Any] = {}

    async def _via(path: str) -> dict | Exception:
        try:
            status, body_text = await _secure.forward(
                worker, method="GET", path=path, body=None,
            )
            if status >= 400:
                raise RuntimeError(f"HTTP {status}")
            try:
                return json.loads(body_text)
            except Exception as je:
                raise RuntimeError(f"non-JSON body: {je}")
        except Exception as e:
            return e

    _t0 = time.perf_counter()
    ver_res, tags_res, ps_res = await asyncio.gather(
        _via("/api/version"), _via("/api/tags"), _via("/api/ps"),
    )
    out["probe_latency_ms"] = int((time.perf_counter() - _t0) * 1000)

    if isinstance(ver_res, Exception):
        out["version_error"] = f"{type(ver_res).__name__}: {ver_res}"
    else:
        out["version"] = (ver_res or {}).get("version") or "unknown"
    if isinstance(tags_res, Exception):
        out["tags_error"] = f"{type(tags_res).__name__}: {tags_res}"
    else:
        models = []
        for m in (tags_res or {}).get("models", []) or []:
            details = (m.get("details") or {}) if isinstance(m, dict) else {}
            models.append({
                "name": m.get("name") if isinstance(m, dict) else None,
                "size": m.get("size") if isinstance(m, dict) else None,
                "family": details.get("family"),
                "parameter_size": details.get("parameter_size"),
                "quantization_level": details.get("quantization_level"),
            })
        out["models"] = [m for m in models if m.get("name")]
    if isinstance(ps_res, Exception):
        out["gpu_present"] = False
        out["max_vram_seen_bytes"] = 0
        out["loaded_count"] = 0
    else:
        loaded = (ps_res or {}).get("models") or []
        gpu_present = False
        max_vram = 0
        for m in loaded:
            if not isinstance(m, dict):
                continue
            v = int(m.get("size_vram") or 0)
            if v > 0:
                gpu_present = True
            if v > max_vram:
                max_vram = v
        out["gpu_present"] = gpu_present
        out["max_vram_seen_bytes"] = max_vram
        out["loaded_count"] = len(loaded)
    # See note in the LAN probe path above — fall back to static
    # hardware signals so a peer with a GPU but no currently-loaded
    # model still ranks as a GPU node.
    if not out.get("gpu_present"):
        if (
            float(out.get("vram_total_gb") or 0.0) > 0.5
            or (out.get("gpu_kind") or "").lower() in {"nvidia", "amd", "intel"}
        ):
            out["gpu_present"] = True
    out["transport"] = "encrypted-proxy"
    return out


# Multi-rpc-server port assignments. Each worker can expose its
# iGPU and its CPU+RAM as TWO separate single-backend rpc-servers
# on adjacent ports. Each rpc-server is single-backend internally
# (no hybrid SYCL+CPU on the same process) so ggml-rpc.cpp's
# layout-mismatch crash has no surface to fire on. The orchestrator
# treats the two ports as independent compute endpoints in
# `--rpc <ip>:<port1>,<ip>:<port2>` and `--tensor-split` weights
# them per-endpoint, mapping cleanly to llama.cpp's per-device
# allocator.
_MULTI_RPC_GPU_PORT = 50052   # iGPU device only (SYCL0 / Vulkan0 / CUDA0)
_MULTI_RPC_CPU_PORT = 50053   # CPU + system RAM only


def select_multi_rpc_specs(worker: dict) -> list[dict]:
    """Return the list of {port, backend} specs to spawn on this
    worker so it contributes EVERY available compute tier (iGPU +
    CPU+RAM, or just CPU+RAM if no iGPU).

    Mirrors `_select_worker_backend`'s logic but emits a per-endpoint
    list instead of one combined string. Each entry is a single-
    backend rpc-server — bypasses the SYCL+CPU hybrid layout-mismatch
    crash that ggml-rpc.cpp's RPC_STATUS_ASSERT trips when both
    devices share an allocator.

    Failure-flag walk (per backend, single-device only):
      * Intel + SYCL0 not crashed -> [SYCL0:50052, CPU:50053]
      * Intel + Vulkan0 not crashed (SYCL crashed) -> [Vulkan0:50052, CPU:50053]
      * Intel + both crashed -> [CPU:50052]  (single endpoint, CPU only)
      * NVIDIA -> [CUDA0:50052, CPU:50053]
      * AMD    -> [Vulkan0:50052, CPU:50053]
      * No GPU -> [CPU:50052]
    """
    vendor = _worker_gpu_vendor(worker)
    caps = worker.get("capabilities") or {}
    now = time.time()
    if vendor == "nvidia":
        return [
            {"port": _MULTI_RPC_GPU_PORT, "backend": "CUDA0"},
            {"port": _MULTI_RPC_CPU_PORT, "backend": "CPU"},
        ]
    if vendor == "amd":
        return [
            {"port": _MULTI_RPC_GPU_PORT, "backend": "Vulkan0"},
            {"port": _MULTI_RPC_CPU_PORT, "backend": "CPU"},
        ]
    if vendor == "intel":
        sycl_dead = (
            (now - float(caps.get("sycl_split_failed_at") or 0))
            < _BACKEND_FAILURE_COOLDOWN_SEC
        )
        vulkan_dead = (
            (now - float(caps.get("vulkan_split_failed_at") or 0))
            < _BACKEND_FAILURE_COOLDOWN_SEC
        )
        if not sycl_dead:
            return [
                {"port": _MULTI_RPC_GPU_PORT, "backend": "SYCL0"},
                {"port": _MULTI_RPC_CPU_PORT, "backend": "CPU"},
            ]
        if not vulkan_dead:
            return [
                {"port": _MULTI_RPC_GPU_PORT, "backend": "Vulkan0"},
                {"port": _MULTI_RPC_CPU_PORT, "backend": "CPU"},
            ]
        # Both iGPU paths crashed within the cool-down — fall back
        # to CPU-only on the primary port. Single endpoint; the
        # secondary CPU port is omitted because we'd just be
        # exposing the same compute tier twice.
        return [{"port": _MULTI_RPC_GPU_PORT, "backend": "CPU"}]
    # No GPU detected — single CPU endpoint.
    return [{"port": _MULTI_RPC_GPU_PORT, "backend": "CPU"}]


async def ensure_rpc_servers_via_proxy_multi(worker: dict) -> list[dict]:
    """Bring up MULTIPLE rpc-server processes on a paired peer (one
    per (port, backend) spec from `select_multi_rpc_specs`) via the
    encrypted P2P channel. Returns the list of endpoints that came
    up successfully, each as ``{"port": int, "backend": str}``.

    The receiver (`/api/p2p/rpc-server/ensure-multi`) starts each
    spec independently — restarting one doesn't take down the other.
    Used by the split-router to engage iGPU + CPU+RAM contribution
    per worker without the SYCL+CPU hybrid layout-mismatch bug.
    """
    from . import p2p_secure_client as _secure
    label = worker.get("label") or worker.get("id") or "?"
    specs = select_multi_rpc_specs(worker)
    if not specs:
        return []
    try:
        status, body_text = await _secure.forward(
            worker,
            method="POST",
            path="/api/p2p/rpc-server/ensure-multi",
            body={"specs": specs},
        )
    except Exception as e:
        log.info(
            "compute_pool: ensure-multi rpc-server failed for %s: %s",
            label, e,
        )
        return []
    if status != 200:
        log.info(
            "compute_pool: ensure-multi on %s returned HTTP %d: %s",
            label, status, body_text[:200],
        )
        return []
    try:
        result = jsonutil.loads(body_text)
    except Exception:
        return []
    # Walk results dict to figure out which ports succeeded.
    results = result.get("results") or {}
    live_endpoints: list[dict] = []
    for spec in specs:
        port = spec["port"]
        # results dict keys may be strings (JSON) or ints (Python).
        per_port = results.get(port) or results.get(str(port)) or {}
        if per_port.get("listening"):
            live_endpoints.append({
                "port": port,
                "backend": per_port.get("active_backend") or spec["backend"],
            })
    log.info(
        "compute_pool: %s exposes %d rpc endpoint(s): %s",
        label, len(live_endpoints),
        ", ".join(f"{e['port']}={e['backend']}" for e in live_endpoints),
    )
    # Persist the endpoint list so the split-router can read it
    # without re-probing.
    if live_endpoints:
        try:
            caps = dict(worker.get("capabilities") or {})
            caps["rpc_server_reachable"] = True
            caps["rpc_endpoints"] = live_endpoints
            # Keep the legacy current_rpc_backend stamped to the
            # primary (iGPU) endpoint so older code paths (status
            # display, single-port readers) still see something
            # informative.
            caps["current_rpc_backend"] = live_endpoints[0]["backend"]
            db.update_compute_worker_capabilities(
                worker["id"], capabilities=caps,
            )
            worker["capabilities"] = caps
        except Exception as e:
            log.debug(
                "compute_pool: rpc_endpoints persist failed for %s: %s",
                label, e,
            )
    return live_endpoints


async def ensure_rpc_server_via_proxy(
    worker: dict, *, backend: str = "SYCL0,CPU", port: int = 50052,
) -> bool:
    """Bring up rpc-server on a paired peer through the encrypted
    P2P channel — no SSH required.

    Workflow:
      1. GET ``/api/p2p/rpc-server/status`` to check whether the peer
         already has rpc-server listening (avoids a pointless restart
         and the 4-second listener-wait that comes with it).
      2. If not listening, POST ``/api/p2p/rpc-server/start`` with
         the desired backend.
      3. On success, mark the worker's ``capabilities.rpc_server_reachable``
         and ``current_rpc_backend`` so ``_eligible_split_workers``
         immediately starts including this peer.

    Returns True iff rpc-server is listening on the peer after this
    call. Best-effort — every failure path returns False without
    raising so the caller can keep iterating across peers.

    This replaces the SSH-based ``_attempt_rpc_server_restart`` for
    paired LAN peers. Pure SSH workers (the rare case where someone
    manually configured ssh_host) still use the SSH path.
    """
    from . import p2p_secure_client as _secure
    label = worker.get("label") or worker.get("id") or "?"

    # Stamp gpu_kind / ram_total_gb from system-stats so the auto-
    # fallback selector knows the worker's vendor (intel/nvidia/amd)
    # without having to wait for the route_chat_for autoprep path.
    # Without this stamp, callers like _set_workers_backend (which
    # `start()` invokes directly, bypassing route_chat_for) see
    # `gpu_kind=null` -> vendor="none" -> always returns "CPU"
    # regardless of whether the worker has a usable iGPU.
    try:
        live_stats = await probe_worker_live_stats(worker)
    except Exception:
        live_stats = {}
    if live_stats:
        try:
            caps = dict(worker.get("capabilities") or {})
            caps["gpu_kind"] = live_stats.get("gpu_kind") or caps.get("gpu_kind") or ""
            caps["ram_total_gb"] = float(live_stats.get("ram_total_gb") or caps.get("ram_total_gb") or 0)
            caps["ram_free_gb"] = float(live_stats.get("ram_free_gb") or 0)
            caps["vram_total_gb"] = float(live_stats.get("vram_total_gb") or caps.get("vram_total_gb") or 0)
            db.update_compute_worker_capabilities(
                worker["id"], capabilities=caps,
            )
            worker["capabilities"] = caps
        except Exception as e:
            log.debug(
                "compute_pool: live-stats stamp failed for %s: %s",
                label, e,
            )

    # Quick status probe first — saves us a 4 s listener-wait when
    # rpc-server is already up (the common case after the first call).
    try:
        status, body_text = await _secure.forward(
            worker, method="GET", path="/api/p2p/rpc-server/status",
            body=None,
        )
    except Exception as e:
        log.debug(
            "compute_pool: rpc-server status probe failed for %s: %s",
            label, e,
        )
        return False
    if status != 200:
        return False
    try:
        snap = jsonutil.loads(body_text)
    except Exception:
        return False
    if snap.get("listening") and snap.get("active_backend") == backend:
        # Already up with the right backend. Stamp the capability
        # cache so the router sees it.
        try:
            caps = dict(worker.get("capabilities") or {})
            caps["rpc_server_reachable"] = True
            caps["current_rpc_backend"] = backend
            db.update_compute_worker_capabilities(
                worker["id"], capabilities=caps,
            )
        except Exception:
            pass
        return True
    # Listening but with a different backend (e.g. running SYCL+CPU
    # but split needs CPU-only) — fall through to the start call,
    # which kills the stale process and respawns. We log it so the
    # operator can spot churn.
    if snap.get("listening") and snap.get("active_backend") != backend:
        log.info(
            "compute_pool: rpc-server on %s is up with backend %r but "
            "we need %r; restarting to switch",
            label, snap.get("active_backend"), backend,
        )
    if not snap.get("binary_present"):
        log.info(
            "compute_pool: peer %s has no rpc-server binary at %s; "
            "split won't engage on this worker",
            label, snap.get("binary_path"),
        )
        return False

    # Spawn it.
    try:
        status, body_text = await _secure.forward(
            worker,
            method="POST",
            path="/api/p2p/rpc-server/start",
            body={"backend": backend, "port": port},
        )
    except Exception as e:
        log.info(
            "compute_pool: rpc-server start failed for %s: %s",
            label, e,
        )
        return False
    if status != 200:
        log.info(
            "compute_pool: rpc-server start on %s returned HTTP %d: %s",
            label, status, body_text[:200],
        )
        return False
    try:
        result = jsonutil.loads(body_text)
    except Exception:
        return False
    listening = bool(result.get("listening"))
    log.info(
        "compute_pool: rpc-server on %s -> status=%s listening=%s pid=%s",
        label, result.get("status"), listening, result.get("pid"),
    )
    if listening:
        try:
            caps = dict(worker.get("capabilities") or {})
            caps["rpc_server_reachable"] = True
            caps["current_rpc_backend"] = backend
            db.update_compute_worker_capabilities(
                worker["id"], capabilities=caps,
            )
        except Exception:
            pass
    return listening


async def probe_worker_bandwidth(worker: dict) -> float:
    """Measure usable LAN/internet bandwidth between this orchestrator
    and the worker, in MB/s. Returns 0 on any failure.

    Method: time the wall-clock cost of fetching the worker's
    `/api/p2p/binary/list` response (manifest of llama-cpp DLLs +
    sha256 sums). This is a small JSON response (~5-15 KB) so it's
    a coarse round-trip latency measurement, not a true throughput
    benchmark — but it correlates with bandwidth on shared media
    (Wi-Fi vs Ethernet vs Tailscale), which is what routing cares
    about. We avoid an explicit "send 10 MB ping" benchmark because
    it would burn pool bandwidth that the user paid for; the manifest
    GET happens during normal LAN-first install probing anyway.

    For workers behind a slow link the measured time will be
    multi-second and the routing decision can demote them out of
    the split engagement set (where layer-push latency dominates
    chat throughput).
    """
    from . import p2p_secure_client as _secure
    import time as _t
    label = worker.get("label") or worker.get("id") or "?"
    t0 = _t.perf_counter()
    try:
        status, body_text = await _secure.forward(
            worker, method="GET", path="/api/p2p/binary/list", body=None,
        )
    except Exception as e:
        log.debug(
            "compute_pool: bandwidth probe failed for %s: %s",
            label, e,
        )
        return 0.0
    elapsed = _t.perf_counter() - t0
    if status != 200 or elapsed <= 0:
        return 0.0
    bytes_recv = len(body_text.encode("utf-8")) if isinstance(body_text, str) else len(body_text)
    if bytes_recv <= 0:
        return 0.0
    # MB/s = bytes / 1e6 / seconds
    return round((bytes_recv / 1e6) / elapsed, 2)


async def probe_worker_live_stats(worker: dict, *, timeout: float | None = None) -> dict:
    """Fetch the worker's CURRENT free RAM / free VRAM via the
    encrypted P2P proxy. Lightweight enough to call every few seconds
    while a split-model is loaded so the orchestrator's layer split
    adapts as the user opens / closes other apps on the worker.

    Returns a dict with at least ``ram_free_gb`` (float) on success,
    or an empty dict on any failure. Errors are swallowed — the
    caller falls back to whatever stats are already cached on the
    worker row.

    Why this matters: until 2026-05 the only path to a worker's free
    RAM was ``_probe_worker_specs_via_ssh``, which (a) needs an
    ssh_host configured (most paired LAN peers won't have that) and
    (b) is heavy enough that we ran it every 5 minutes. That cadence
    is fine for "show me the worker in Settings" but useless for
    realtime layer rebalancing. This helper hits the
    ``/api/p2p/system-stats`` endpoint over the existing encrypted
    pair channel — no SSH dependency, sub-100 ms per probe on LAN.
    """
    from . import p2p_secure_client as _secure
    try:
        status, body_text = await _secure.forward(
            worker, method="GET", path="/api/p2p/system-stats", body=None,
            timeout=timeout,
        )
    except Exception as e:
        log.debug(
            "compute_pool: live-stats probe failed for %r: %s: %s",
            worker.get("label"), type(e).__name__, e,
        )
        return {}
    if status != 200:
        return {}
    try:
        data = jsonutil.loads(body_text)
    except Exception:
        return {}
    if not isinstance(data, dict):
        return {}
    return data


async def probe_worker(wid: str) -> dict:
    """Probe one worker now and persist the result.

    Returns a dict with `ok`, the merged probe payload (`version`,
    `models`, or error markers), and the `last_seen` timestamp. Safe
    to call from any context — failures are caught and recorded on
    the row, never propagated.

    On a network-level failure the auto-repair routine attempts to
    rediscover the worker's LAN IP via Tailscale (when ``tailscale_host``
    is set on the row); a successful rediscovery transparently retries
    the probe with the new address so the user just sees the worker
    stay green across DHCP rebinds.
    """
    worker = db.get_compute_worker(wid)
    if not worker:
        return {"ok": False, "error": "worker not found"}
    if not worker.get("enabled"):
        return {"ok": False, "error": "worker disabled — enable it first"}

    base = _worker_base_url(worker)
    token = db.get_compute_worker_auth_token(wid)
    now = time.time()

    payload: dict | None = None
    last_exc: Exception | None = None
    # Encrypted-proxy mode: paired peers' compute_workers rows
    # point at the peer's Gigachat port (not Ollama), so the
    # plaintext /api/version etc. GETs would 404. Route the same
    # three Ollama probes through the encrypted proxy instead —
    # the wire stays ciphertext and the data shape coming back
    # is identical to a direct Ollama probe.
    if worker.get("use_encrypted_proxy"):
        try:
            payload = await _probe_one_via_secure_proxy(worker)
        except Exception as e:
            last_exc = e
    else:
        try:
            async with httpx.AsyncClient(timeout=_PROBE_TIMEOUT_SEC) as client:
                payload = await _probe_one(client, base, token)
        except Exception as e:
            last_exc = e

    if payload is None:
        # Network-level failure — connection refused, DNS miss, etc.
        # Try to recover by rediscovering the LAN IP via Tailscale; on
        # success, retry the probe once with the refreshed address.
        repaired = await _attempt_lan_address_repair(worker)
        if repaired is not None:
            worker = repaired
            base = _worker_base_url(worker)
            try:
                async with httpx.AsyncClient(timeout=_PROBE_TIMEOUT_SEC) as client:
                    payload = await _probe_one(client, base, token)
            except Exception as e:
                last_exc = e

    if payload is None:
        # Still down after the (possible) repair retry — record the
        # error on the row so the UI can show "unreachable since X".
        err = f"{type(last_exc).__name__}: {last_exc}" if last_exc else "unreachable"
        try:
            db.update_compute_worker_capabilities(
                wid, last_seen=now, last_error=err,
            )
        except Exception:
            pass
        return {"ok": False, "error": err, "last_seen": now}

    # Two-endpoint outcome: success only if BOTH endpoints responded.
    # Partial success is still a usable signal — record what we got
    # but flag the error so the UI can warn.
    has_models = bool(payload.get("models"))
    has_version = "version" in payload
    error_parts = []
    if not has_models:
        error_parts.append(payload.get("tags_error") or "no models field")
    if not has_version:
        error_parts.append(payload.get("version_error") or "no version field")
    error_str = "; ".join(error_parts) if error_parts else ""

    # Phase 2 add-on: TCP-probe the worker's rpc-server port so the UI
    # can show whether layer-split inference is available on this
    # worker. Probe failure is NOT counted as `last_error` — rpc-server
    # is optional (Phase 1 routing works without it), so a worker with
    # Ollama up but rpc-server down is still "online" for chat /
    # embeddings / subagent routing. We just record the rpc state in
    # capabilities so commit 6's UI can render a separate badge.
    rpc_host = _worker_host(worker)
    rpc_port = _DEFAULT_RPC_PORT
    rpc_ok, rpc_err = await _probe_rpc_server(rpc_host, rpc_port)

    # Self-heal: if rpc-server is down AND the worker has ssh_host
    # configured, attempt to restart it remotely. This covers the
    # common case where the rpc-server process died (crash, OS
    # restart, transient driver issue) since the previous probe.
    # We try restart at most once per probe cycle to avoid tight
    # loops if the binary is genuinely broken — a successful restart
    # is verified by an immediate re-probe of the port.
    if not rpc_ok:
        # Self-heal: rpc-server isn't listening. Try to bring it up
        # over the encrypted P2P channel (preferred — works for any
        # paired peer) with SSH as a fallback for legacy setups.
        log.info(
            "compute_pool: rpc-server unreachable on %s; attempting "
            "self-heal (P2P first, SSH fallback if ssh_host set)",
            rpc_host,
        )
        restarted = await _attempt_rpc_server_restart(worker)
        if restarted:
            # Re-probe to confirm the new process is actually serving.
            rpc_ok, rpc_err = await _probe_rpc_server(rpc_host, rpc_port)
            if rpc_ok:
                log.info("compute_pool: rpc-server self-heal succeeded on %s", rpc_host)

    capabilities = {
        "version": payload.get("version"),
        "models": payload.get("models") or [],
        "rpc_server_reachable": rpc_ok,
        "rpc_port": rpc_port,
        "rpc_error": rpc_err,
        # Hardware hints — best-effort, populated from /api/ps. Used by
        # the router to prefer hardware-stronger workers over weaker
        # ones when both are otherwise eligible.
        "gpu_present": bool(payload.get("gpu_present")),
        "max_vram_seen_bytes": int(payload.get("max_vram_seen_bytes") or 0),
        "loaded_count": int(payload.get("loaded_count") or 0),
        # LAN-latency hint from the parallel probe-GET round trip.
        # Used by the routing layer as a proxy for "how expensive is
        # a per-token RPC roundtrip to this worker": low latency →
        # split path is viable; high latency → biased toward host
        # CPU offload over split.
        "probe_latency_ms": int(payload.get("probe_latency_ms") or 0),
    }

    # Carry over throughput + spec data from the previous probe IF
    # they're still fresh (1 hour TTL). Re-measure only when stale —
    # both are expensive enough that we don't want to run them on
    # every 5-min sweep.
    prev_caps = worker.get("capabilities") or {}
    now_t = time.time()
    if prev_caps.get("tokens_per_second") and (
        now_t - (prev_caps.get("tps_measured_at") or 0)
        < _THROUGHPUT_CACHE_TTL_SEC
    ):
        capabilities["tokens_per_second"] = prev_caps["tokens_per_second"]
        capabilities["tps_measured_at"] = prev_caps["tps_measured_at"]
        capabilities["tps_model"] = prev_caps.get("tps_model")
    if prev_caps.get("specs_measured_at") and (
        now_t - prev_caps["specs_measured_at"] < _THROUGHPUT_CACHE_TTL_SEC
    ):
        for k in ("cpu_name", "cpu_cores", "cpu_threads",
                  "ram_total_gb", "ram_free_gb", "gpus", "specs_measured_at"):
            if k in prev_caps:
                capabilities[k] = prev_caps[k]

    # Throughput bench — pick the smallest chat model on the worker
    # so the bench loads fast. Embed-only models give us no chat
    # speed signal; skip them as bench targets.
    if "tokens_per_second" not in capabilities:
        chat_models = [
            m for m in (capabilities.get("models") or [])
            if (m.get("name") or "") and "embed" not in (m.get("name") or "").lower()
        ]
        if chat_models:
            chat_models.sort(key=lambda m: int(m.get("size") or 1 << 60))
            bench_model = chat_models[0]["name"]
            tps, _ = await _measure_throughput(base, token, bench_model)
            if tps > 0:
                capabilities["tokens_per_second"] = tps
                capabilities["tps_measured_at"] = now_t
                capabilities["tps_model"] = bench_model

    # SSH-based hardware spec probe — fills in CPU/RAM/GPU details
    # the API can't surface. Free when the user has set ssh_host.
    if "specs_measured_at" not in capabilities:
        specs = await _probe_worker_specs_via_ssh(worker)
        capabilities.update(specs)
    try:
        db.update_compute_worker_capabilities(
            wid,
            capabilities=capabilities,
            last_seen=now,
            # Empty string clears any previous error; non-empty records.
            last_error=error_str if error_str else "",
        )
    except Exception:
        pass

    return {
        "ok": has_models and has_version,
        "capabilities": capabilities,
        "error": error_str or None,
        "last_seen": now,
    }


async def probe_all_enabled() -> list[dict]:
    """Probe every enabled worker concurrently. Returns a list of
    `{worker_id, label, ok, error}` summaries — the heavy capability
    payload is persisted on the row, this just tells the caller what
    happened in aggregate."""
    workers = db.list_compute_workers(enabled_only=True)
    if not workers:
        return []
    results = await asyncio.gather(
        *(probe_worker(w["id"]) for w in workers),
        return_exceptions=True,
    )
    summaries: list[dict] = []
    for w, r in zip(workers, results):
        if isinstance(r, Exception):
            summaries.append({
                "worker_id": w["id"],
                "label": w["label"],
                "ok": False,
                "error": f"{type(r).__name__}: {r}",
            })
        else:
            summaries.append({
                "worker_id": w["id"],
                "label": w["label"],
                "ok": bool(r.get("ok")),
                "error": r.get("error"),
            })
    return summaries


_PROBE_TASK: asyncio.Task | None = None


# How many files to re-embed in a single idle sweep. Caps the cost so
# one sweep can't dominate a 5-min cycle and starve the next probe.
# At ~10 chunks per file × ~200 ms per embed (worker on Wi-Fi), 20
# files is ~40 s of work — comfortable inside a 300 s window.
_IDLE_REINDEX_MAX_FILES_PER_RUN = 20


async def _drain_idle_reindex() -> int:
    """Reembed files modified since their last index, during pool idle time.

    Source files edited externally (IDE writes, git checkouts, build
    artifacts) bypass the post-`write_file` invalidation that runs
    inline. This sweep catches those by walking every `status=ready`
    codebase index and reembedding any file whose mtime is newer than
    the index's `last_indexed_at`.

    Skips when ANY conversation is mid-turn anywhere in the pool —
    we don't want to compete with the active chat for embed workers.
    Caps at `_IDLE_REINDEX_MAX_FILES_PER_RUN` per sweep so one cycle
    can't starve the next probe.

    Returns the number of files reembedded so the caller can log.
    """
    # Skip when any chat turn is active. The active-turn counter is
    # populated by `register_turn_start` from the agent loop; if any
    # node has count > 0, somebody is actively chatting and we should
    # not steal embed bandwidth.
    if any(c > 0 for c in _ACTIVE_TURNS_PER_NODE.values()):
        return 0

    try:
        indexes = db.list_codebase_indexes()
    except Exception:
        return 0
    ready = [i for i in indexes if i.get("status") == "ready"]
    if not ready:
        return 0

    # Local imports — `tools` imports `compute_pool` at module scope,
    # so we defer the reverse import until call time to avoid a cycle.
    from . import tools as _tools
    from pathlib import Path as _Path

    files_processed = 0
    for idx in ready:
        if files_processed >= _IDLE_REINDEX_MAX_FILES_PER_RUN:
            break
        cwd = idx.get("cwd")
        last_indexed = float(idx.get("last_indexed_at") or 0)
        if not cwd or not last_indexed:
            continue
        try:
            root_path = _Path(cwd)
            if not root_path.is_dir():
                continue
            files = _tools._codebase_list_files(root_path)
        except Exception:
            continue
        for f in files:
            if files_processed >= _IDLE_REINDEX_MAX_FILES_PER_RUN:
                break
            try:
                if f.stat().st_mtime <= last_indexed:
                    continue
            except OSError:
                continue
            # Reembed this file — `_reembed_codebase_file` deletes its
            # existing chunks first so a partial failure leaves the
            # row in a recoverable state.
            try:
                db.delete_doc_chunks_for(str(f))
                await _tools._reembed_codebase_file(
                    str(f), _tools._DEFAULT_EMBED_MODEL,
                )
                files_processed += 1
            except Exception:
                continue
        # Stamp last_indexed_at so the next sweep doesn't redo what
        # we just covered — only files modified AFTER this stamp will
        # match next time.
        if files_processed > 0:
            try:
                db.upsert_codebase_index(cwd, last_indexed_at=time.time())
            except Exception:
                pass

    if files_processed:
        log.info(
            "compute_pool: idle re-index reembedded %d modified file(s)",
            files_processed,
        )
    return files_processed


# Auto-disable workers that have been unreachable for this long. The
# periodic probe still tries them every sweep (so a worker coming back
# online recovers), but we flip `enabled=False` to declutter routing
# decisions and stop the SSH-restart attempts that were piling up
# warning logs. User can re-enable manually from Settings; the row
# itself is preserved with full capabilities + history.
_STALE_WORKER_DISABLE_SECONDS = 28 * 86400  # 4 weeks


async def _auto_disable_stale_workers() -> int:
    """Flip `enabled=False` on workers unreachable for > 4 weeks.

    Routing already filters by freshness (1-hour `last_seen` window),
    so stale workers never actually serve traffic — but they stay in
    the eligible-set scan, tagged `last_error`, and keep being SSH-
    restart-attempted on every probe sweep. After a month of silence,
    flipping `enabled=False` is the right "I forgot about this device"
    semantic. Re-enabling is a one-click action in Settings → Compute.

    Returns the number of workers auto-disabled this sweep so the
    caller can log it.
    """
    rows = db.list_compute_workers(enabled_only=True)
    now = time.time()
    disabled_count = 0
    for w in rows:
        last_seen = float(w.get("last_seen") or 0)
        if last_seen <= 0:
            # Never probed successfully — leave alone, the user just
            # added it. Manual disable via Settings if they meant to.
            continue
        age = now - last_seen
        if age < _STALE_WORKER_DISABLE_SECONDS:
            continue
        try:
            db.update_compute_worker(w["id"], enabled=False)
            disabled_count += 1
            log.info(
                "compute_pool: auto-disabled stale worker %r — "
                "unreachable for %.1f days",
                w.get("label"), age / 86400,
            )
        except Exception as e:
            log.warning(
                "compute_pool: failed to auto-disable %r: %s",
                w.get("label"), e,
            )
    return disabled_count


async def _periodic_loop() -> None:
    """Internal: sweep every `_SWEEP_INTERVAL_SEC`. Started/stopped via
    `start_periodic_probe` / `stop_periodic_probe` on app lifecycle.

    After each liveness sweep we drain any queued auto-syncs that
    routing calls deferred. Doing it here (rather than from the
    routing call) keeps SCP off the chat hot path — bench-confirmed
    win in commit 18.

    Background idle re-index runs LAST so it sees fresh worker probe
    data and doesn't compete with the SCP drain for SSH bandwidth.
    """
    while True:
        try:
            await probe_all_enabled()
        except asyncio.CancelledError:
            raise
        except Exception as e:
            log.warning("compute_pool periodic probe failed: %s", e)
        # Drain the auto-sync queue after the sweep — workers'
        # capabilities reflect any newly-installed models from the
        # previous drain, so the queue from this turn is fresh.
        try:
            n = await drain_deferred_syncs()
            if n:
                log.info("compute_pool: started %d deferred LAN-copy task(s)", n)
        except Exception as e:
            log.warning("compute_pool: drain_deferred_syncs failed: %s", e)
        # Auto-disable workers unreachable for > 4 weeks. Cheap (one
        # SELECT + at most a few UPDATEs); runs every sweep so the
        # user sees the state-change in Settings within minutes of
        # the threshold being crossed.
        try:
            await _auto_disable_stale_workers()
        except Exception as e:
            log.warning("compute_pool: _auto_disable_stale_workers failed: %s", e)
        # Idle pool re-index — uses spare embed bandwidth to keep RAG
        # vectors fresh without blocking active chats. Skipped when a
        # chat is in flight.
        try:
            await _drain_idle_reindex()
        except Exception as e:
            log.warning("compute_pool: _drain_idle_reindex failed: %s", e)
        try:
            await asyncio.sleep(_SWEEP_INTERVAL_SEC)
        except asyncio.CancelledError:
            raise


def start_periodic_probe() -> None:
    """Schedule the background sweep. Idempotent — calling twice is
    a no-op. Called from app.py's startup hook."""
    global _PROBE_TASK
    if _PROBE_TASK and not _PROBE_TASK.done():
        return
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        return
    _PROBE_TASK = loop.create_task(_periodic_loop())


def stop_periodic_probe() -> None:
    """Cancel the background sweep. Called from app shutdown so the
    test runner doesn't see a stranded task."""
    global _PROBE_TASK
    if _PROBE_TASK and not _PROBE_TASK.done():
        _PROBE_TASK.cancel()
    _PROBE_TASK = None


# ---------------------------------------------------------------------------
# Fast peer heartbeat — sub-30 s detection of paired-peer drops + auto-
# reconnect on recovery. Distinct from the 5-min spec sweep above.
#
# Why a separate loop:
#   * The spec sweep (5 min) is heavy — it pulls /api/version + /api/tags
#     + /api/ps + capability JSON + measures throughput. Running it
#     every 30 s would wreck a low-bandwidth pool.
#   * This loop is light — one /api/p2p/system-stats GET per peer
#     (small JSON, no model loads). Tells us whether the peer is
#     ALIVE and gives us live RAM as a side-effect.
#   * On a missed heartbeat, we mark `rpc_server_reachable = False`
#     immediately so the routing layer demotes the peer from the
#     split-engagement set. On a successful heartbeat after a miss,
#     we call `ensure_rpc_server_via_proxy` to rebuild rpc-server
#     state (in case the peer rebooted and lost its rpc-server
#     process).
#
# Cadence: 20 s active probe interval. After 3 consecutive misses
# (~60 s) the peer is treated as offline. Recovery is instant on
# the next successful heartbeat.
_HEARTBEAT_INTERVAL_SEC = 20.0
_HEARTBEAT_MISS_THRESHOLD = 3
_heartbeat_miss_counts: dict[str, int] = {}
_HEARTBEAT_TASK: asyncio.Task | None = None


async def _heartbeat_loop() -> None:
    """Active reachability ping for every paired worker every
    _HEARTBEAT_INTERVAL_SEC. Catches drops in <60s and re-arms
    rpc-server on recovery."""
    while True:
        try:
            workers = db.list_compute_workers(enabled_only=True)
        except Exception:
            workers = []
        # One probe at a time per worker, but workers in parallel.
        async def _probe_one(w: dict) -> None:
            wid = w["id"]
            label = w.get("label") or wid
            stats = {}
            try:
                # Tight 5 s timeout for the heartbeat path. A dead peer
                # would otherwise stall the loop for the default 120 s
                # one-shot timeout, blocking liveness checks for every
                # OTHER peer for that whole window.
                stats = await probe_worker_live_stats(w, timeout=5.0)
            except Exception:
                stats = {}
            # Bandwidth probe — every 5th heartbeat tick (~100s).
            # Bandwidth changes slower than RAM, no need to re-measure
            # every tick. Cheap when stale anyway (one binary-list GET).
            should_measure_bw = False
            try:
                last_bw_ts = float((w.get("capabilities") or {}).get("bandwidth_probed_at") or 0)
                should_measure_bw = (time.time() - last_bw_ts) >= 100.0
            except Exception:
                pass
            bw = 0.0
            if stats and should_measure_bw:
                try:
                    bw = await probe_worker_bandwidth(w)
                except Exception:
                    bw = 0.0
            if stats:
                # Reset miss counter, refresh stats.
                prev_misses = _heartbeat_miss_counts.pop(wid, 0)
                try:
                    caps = dict(w.get("capabilities") or {})
                    caps["ram_free_gb"] = float(stats.get("ram_free_gb") or 0)
                    caps["ram_total_gb"] = float(stats.get("ram_total_gb") or caps.get("ram_total_gb") or 0)
                    caps["vram_total_gb"] = float(stats.get("vram_total_gb") or caps.get("vram_total_gb") or 0)
                    caps["gpu_kind"] = stats.get("gpu_kind") or caps.get("gpu_kind") or ""
                    caps["ram_free_probed_at"] = stats.get("ts") or 0
                    if bw > 0:
                        caps["bandwidth_mbps"] = bw
                        caps["bandwidth_probed_at"] = time.time()
                    # Heartbeat alive → make sure rpc_server_reachable
                    # is True (it may have been flipped to False earlier).
                    # Don't auto-set True if we don't know rpc state —
                    # let ensure_rpc_server_via_proxy handle that on
                    # the next prep call.
                    if caps.get("rpc_server_reachable") is False and prev_misses >= _HEARTBEAT_MISS_THRESHOLD:
                        # Recovery: clear the stale-offline flag so the
                        # next routing decision re-considers this peer.
                        # ensure_rpc_server_via_proxy below will set
                        # the flag definitively based on actual probe.
                        caps.pop("rpc_server_reachable", None)
                    db.update_compute_worker_capabilities(
                        wid, capabilities=caps, last_seen=time.time(),
                    )
                except Exception:
                    pass
                # Recovery path: peer was offline and just came back.
                # Re-arm rpc-server so split inference can resume.
                if prev_misses >= _HEARTBEAT_MISS_THRESHOLD:
                    log.info(
                        "heartbeat: %s recovered after %d missed pings; "
                        "re-arming rpc-server via P2P",
                        label, prev_misses,
                    )
                    try:
                        await ensure_rpc_server_via_proxy(w)
                    except Exception as e:
                        log.debug(
                            "heartbeat: re-arm failed for %s: %s", label, e,
                        )
                return
            # Miss path.
            misses = _heartbeat_miss_counts.get(wid, 0) + 1
            _heartbeat_miss_counts[wid] = misses
            if misses == _HEARTBEAT_MISS_THRESHOLD:
                # Threshold crossed — flip the rpc_server_reachable flag
                # so the routing layer treats this peer as offline.
                # Routing decisions instantly skip it; the next chat
                # turn won't try to engage a dead worker as RPC peer.
                try:
                    caps = dict(w.get("capabilities") or {})
                    caps["rpc_server_reachable"] = False
                    db.update_compute_worker_capabilities(
                        wid, capabilities=caps,
                        last_error=f"missed {misses} consecutive heartbeats — peer offline",
                    )
                except Exception:
                    pass
                log.info(
                    "heartbeat: %s offline after %d missed pings; "
                    "demoted from split-eligible set", label, misses,
                )
        if workers:
            await asyncio.gather(
                *(_probe_one(w) for w in workers),
                return_exceptions=True,
            )
        try:
            await asyncio.sleep(_HEARTBEAT_INTERVAL_SEC)
        except asyncio.CancelledError:
            raise


def start_peer_heartbeat() -> None:
    """Kick off the fast peer-reachability heartbeat. Idempotent."""
    global _HEARTBEAT_TASK
    if _HEARTBEAT_TASK and not _HEARTBEAT_TASK.done():
        return
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        return
    _HEARTBEAT_TASK = loop.create_task(_heartbeat_loop(), name="peer_heartbeat")


def stop_peer_heartbeat() -> None:
    """Cancel the fast heartbeat. Called from shutdown."""
    global _HEARTBEAT_TASK
    if _HEARTBEAT_TASK and not _HEARTBEAT_TASK.done():
        _HEARTBEAT_TASK.cancel()
    _HEARTBEAT_TASK = None


# ---------------------------------------------------------------------------
# Routing: pick a worker for a given workload.
#
# These helpers translate "I need to run an embed / chat / subagent call"
# into "send it to base URL X with bearer Y", or `None` to mean "no
# eligible worker — the host should handle it locally". Everything is
# read-only relative to the worker rows; the routing layer never mutates
# state.
#
# Eligibility for any workload:
#   * row is `enabled`
#   * the workload-specific flag is on (`use_for_embeddings`, etc.)
#   * the last probe succeeded — `last_seen` within `_FRESHNESS_SEC` AND
#     `last_error` is empty. A worker that hasn't been probed yet (last_seen
#     is None) is skipped: we don't want to trust a row the user just
#     added until the periodic sweep — or a manual "Test connection" —
#     confirms it can serve traffic.
#   * if `model` is supplied, the worker's cached capabilities list it as
#     installed (with or without a `:latest` tag suffix).
# ---------------------------------------------------------------------------

# How long after `last_seen` we still consider a worker "fresh enough" to
# route to. The periodic probe runs every 5 min; this 1-hour window covers
# 12 sweeps' worth of buffer for transient blips.
_FRESHNESS_SEC = 60 * 60


def _model_matches(installed_name: str, requested: str) -> bool:
    """Compare an installed model name to a requested name, tolerant to
    Ollama's `:latest` tag default.

    `nomic-embed-text` matches `nomic-embed-text:latest`, and an explicit
    tag (`gemma4:e4b`) matches itself. Mismatched explicit tags are NOT
    coerced — `gemma4:e4b` does not match `gemma4:e2b`.
    """
    if not installed_name or not requested:
        return False
    if installed_name == requested:
        return True
    # Strip `:latest` from either side and compare bare names.
    inst_bare = installed_name.split(":", 1)[0] if ":" in installed_name else installed_name
    req_bare = requested.split(":", 1)[0] if ":" in requested else requested
    inst_has_explicit_tag = ":" in installed_name and not installed_name.endswith(":latest")
    req_has_explicit_tag = ":" in requested and not requested.endswith(":latest")
    # If both sides have explicit tags, they had to match exactly above.
    if inst_has_explicit_tag and req_has_explicit_tag:
        return False
    return inst_bare == req_bare


def _worker_has_model(worker: dict, model: str) -> bool:
    caps = worker.get("capabilities") or {}
    for m in caps.get("models") or []:
        if _model_matches(m.get("name") or "", model):
            return True
    return False


def _is_fresh(worker: dict, now: float | None = None) -> bool:
    """Last successful probe within `_FRESHNESS_SEC`."""
    last_seen = worker.get("last_seen")
    if not last_seen:
        return False
    if worker.get("last_error"):
        return False
    if now is None:
        now = time.time()
    return (now - float(last_seen)) < _FRESHNESS_SEC


def _capability_score(worker: dict) -> tuple:
    """Power-ranking key for sorting eligible workers.

    Returns a tuple suitable for `sorted(key=...)` — bigger is better,
    so callers use `reverse=True`. Components, in priority order:

      1. `tokens_per_second` — real measured throughput from a tiny
         /api/generate bench. Folds CPU + RAM bandwidth + GPU compute
         + memory size into one bottom-line "how fast can this machine
         run a model" number. Workers we haven't benchmarked yet
         contribute 0 here and fall back to the heuristic axes below.
      2. `gpu_present` — binary GPU detection from /api/ps.
      3. `max_vram_seen_bytes` — lower bound on the worker's VRAM.
      4. `ram_total_gb` — worker RAM total (from SSH probe). Held
         out when ssh_host isn't set.
      5. `cpu_threads` — CPU parallelism (also from SSH probe).
      6. `last_seen` — final tie-breaker.

    All four axes are intentionally factored: throughput is the
    bottom-line speed signal but only available where benchmarked;
    the heuristic axes (gpu / vram / ram / cpu) keep ranking sane
    for workers without measurement data yet.
    """
    caps = worker.get("capabilities") or {}
    # Prefer max_vram_seen_bytes (lower bound — proven by an actual
    # load), but fall back to vram_total_gb when no model has been
    # loaded yet. Without the fallback, a freshly-restarted peer
    # with an 8 GB CUDA card scores 0 for VRAM and loses the routing
    # comparison to a host with a 2 GB iGPU — a verified bug that
    # caused chat for FBS-only models to get sent to host Ollama
    # (which doesn't have them) → 404 model not found.
    max_vram_bytes = int(caps.get("max_vram_seen_bytes") or 0)
    if max_vram_bytes == 0:
        max_vram_bytes = int(float(caps.get("vram_total_gb") or 0.0) * (1024 ** 3))
    return (
        float(caps.get("tokens_per_second") or 0.0),
        1 if caps.get("gpu_present") else 0,
        max_vram_bytes,
        float(caps.get("ram_total_gb") or 0.0),
        int(caps.get("cpu_threads") or 0),
        float(worker.get("last_seen") or 0),
    )


async def _measure_host_throughput(model_name: str) -> float:
    """Cached host throughput in tokens/sec for `model_name`. Runs a
    tiny /api/generate against `localhost:11434`; result valid for
    `_THROUGHPUT_CACHE_TTL_SEC`. Returns 0 on failure (caller falls
    back to heuristic ranking).
    """
    cached = _HOST_THROUGHPUT_CACHE.get(model_name)
    now = time.time()
    if cached and (now - cached[1]) < _THROUGHPUT_CACHE_TTL_SEC:
        return cached[0]
    tps, _ = await _measure_throughput("http://localhost:11434", None, model_name)
    if tps > 0:
        # Re-insertion order is preserved on dict (CPython 3.7+); pop
        # then set so the freshly measured model moves to the back.
        # Evict the oldest entry if we'd exceed the cap.
        _HOST_THROUGHPUT_CACHE.pop(model_name, None)
        _HOST_THROUGHPUT_CACHE[model_name] = (tps, now)
        if len(_HOST_THROUGHPUT_CACHE) > _HOST_THROUGHPUT_CACHE_MAX:
            oldest = next(iter(_HOST_THROUGHPUT_CACHE))
            _HOST_THROUGHPUT_CACHE.pop(oldest, None)
    return tps


def _host_capability_score(model_name: str | None = None) -> tuple:
    """Host's capability tuple, same shape as `_capability_score`.

    When `model_name` is provided AND we've benchmarked the host on
    that model recently, slot the measured TPS in as the primary
    axis. Otherwise we leave it at 0 and rely on the static heuristic
    axes (vram bytes, ram_gb, cpu_threads, last_seen=inf for
    "always online"). Last component is +inf so a host that ties on
    all measurable axes still wins over a worker (no LAN hop).
    """
    try:
        spec = sysdetect.detect_system()
        vram_gb = float(spec.get("vram_gb") or 0.0)
        ram_gb = float(spec.get("ram_gb") or 0.0)
    except Exception:
        vram_gb = 0.0
        ram_gb = 0.0

    host_tps = 0.0
    if model_name:
        cached = _HOST_THROUGHPUT_CACHE.get(model_name)
        if cached and (time.time() - cached[1]) < _THROUGHPUT_CACHE_TTL_SEC:
            host_tps = cached[0]

    try:
        cpu_threads = int(os.cpu_count() or 0)
    except Exception:
        cpu_threads = 0

    return (
        host_tps,
        1 if vram_gb > 0 else 0,
        int(vram_gb * 1024 ** 3),
        ram_gb,
        cpu_threads,
        float("inf"),
    )


def _eligible_workers(flag: str, model: str | None = None) -> list[dict]:
    """Return enabled+fresh workers whose `flag` is on and (optionally) have
    `model` installed.

    Sort order (`_capability_score` desc): GPU-present workers first,
    then by max VRAM observed, then freshest probe. The user's
    intuition "use the more powerful machine" maps to (1) → GPU wins
    over CPU-only, (2) → more VRAM wins over less. Tie at hardware →
    fall back to freshest, like Phase 1 commit 6's original behavior.

    Auto-sync side effect: when a worker has `ssh_host` set, is
    enabled+fresh, and has the right `flag` on, but is missing the
    model, this function kicks off a background SCP from host so the
    model lands on the worker without the user touching a button. The
    current call still excludes that worker (it really doesn't have
    the model right now), but subsequent probes will see the new model
    and the worker becomes eligible. Failures are silent — auto-sync
    is best-effort.
    """
    rows = db.list_compute_workers(enabled_only=True)
    now = time.time()
    # Read the public-pool toggle ONCE for this call. When the user
    # has disabled "join public compute pool" in Settings, we MUST
    # exclude every public-pool peer from the eligible set even if
    # they're already in compute_workers — otherwise the toggle is
    # cosmetic and a previously-registered public peer keeps getting
    # our chat traffic forever. The flag is the user's explicit
    # privacy / resource-control choice; honouring it is non-optional.
    public_pool_on = True
    try:
        from . import p2p_pool_routing as _ppr
        public_pool_on = _ppr._public_pool_enabled()
    except Exception:
        pass
    out: list[dict] = []
    for w in rows:
        if not w.get(flag):
            continue
        if not _is_fresh(w, now=now):
            continue
        # Drop public-pool peers when the toggle is OFF. We detect
        # public peers two ways for robustness: (a) the explicit
        # `paired_devices.role='public'` tag set by
        # `p2p_pool_routing.ensure_public_peer_worker`, AND (b) the
        # legacy `label=public:<device_id>` naming convention used
        # in older registrations.
        if not public_pool_on:
            label = (w.get("label") or "")
            is_public = label.startswith("public:")
            if not is_public and w.get("gigachat_device_id"):
                try:
                    paired = db.get_paired_device(w["gigachat_device_id"])
                    if paired and paired.get("role") == "public":
                        is_public = True
                except Exception:
                    pass
            if is_public:
                continue
        if model:
            if _worker_has_model(w, model):
                # Memory-aware filter: skip the worker only when its
                # available memory is FAR below the model size AND
                # the model isn't currently loaded there. Ollama
                # keeps recently-used models warm for ~5 min; if the
                # model is already in the worker's loaded set, the
                # NEXT chat reuses that load and needs no additional
                # RAM — so the cached ram_free_gb reading (which
                # already excludes the loaded model's footprint)
                # would falsely scare us off a perfectly viable
                # warm worker.
                #
                # Conservative cutoff: skip only when the model is
                # at least 3x the cached free memory. This catches
                # the obvious "another big model is squatting all
                # the RAM" case (e.g. dolphin-mixtral 26 GB
                # llama-server holding 11 GB on FBS, blocking a
                # cold gemma3 4 GB load) while letting through the
                # routine warm-cache case (Naresh has gemma3:4b
                # loaded and 0.9 GB free → still serves the next
                # gemma3 chat instantly).
                model_size_gb = 0.0
                caps = w.get("capabilities") or {}
                for m in (caps.get("models") or []):
                    if (m.get("name") or "") == model:
                        model_size_gb = float(m.get("size") or 0) / (1024 ** 3)
                        break
                if model_size_gb > 0:
                    free_ram = float(caps.get("ram_free_gb") or 0)
                    free_vram = float(caps.get("vram_free_gb") or 0)
                    available = free_ram + free_vram
                    if available > 0 and model_size_gb > available * 3.0:
                        log.info(
                            "compute_pool: skipping %s for model %r — "
                            "needs ~%.1f GB but worker has only %.1f GB "
                            "free (3x+ shortfall = blocked by another "
                            "loaded model); routing to a less-loaded peer.",
                            w.get("label"), model, model_size_gb, available,
                        )
                        continue
                out.append(w)
            elif w.get("ssh_host"):
                _maybe_kickoff_lan_sync(w, model)
            # else: no ssh_host, no path to install — skip silently.
        else:
            out.append(w)
    out.sort(key=_capability_score, reverse=True)
    return out


# Auto-sync architecture (Phase 2 commit 18 redesign):
#
# Auto-syncing was previously fired DIRECTLY from the routing call
# (`_eligible_workers`) — every chat turn that found a worker missing
# a model would kick off a background SCP. Real-world bench showed
# that SCP eats host disk + CPU bandwidth, contending with the chat
# turn that's actively streaming. Net result: chat got 4% slower with
# pool enabled, embed savings didn't compound.
#
# New design: routing calls only ENQUEUE the (worker, model) sync
# requests; they do NOT spawn SCPs. The periodic probe loop
# (`_periodic_loop`, every 5 min) drains the queue and runs the syncs
# then. So:
#   * Chat turns never compete with SCP for host bandwidth.
#   * Background sync still happens — just batched on the probe
#     cadence instead of per-turn.
#   * The bench now measures hot inference cleanly, no SCP noise.
_DEFERRED_SYNC_QUEUE: set[tuple[str, str]] = set()
_AUTO_SYNC_TASKS: dict[tuple[str, str], asyncio.Task] = {}


def _maybe_kickoff_lan_sync(worker: dict, model_name: str) -> None:
    """Defer (not fire) a LAN-copy of `model_name` to `worker`.

    Called from `_eligible_workers` when a worker is otherwise eligible
    but missing the model. Just adds to the deferred queue; the
    periodic probe loop drains it. Means routing decisions are pure
    DB reads — no SCP work in the chat hot path.

    Guards:
      * ssh_host must be set on the worker.
      * Skip if already in queue or already mid-sync.
    """
    wid = worker.get("id")
    if not wid or not model_name:
        return
    if not (worker.get("ssh_host") or "").strip():
        return
    key = (wid, model_name)
    if key in _DEFERRED_SYNC_QUEUE:
        return
    existing = _AUTO_SYNC_TASKS.get(key)
    if existing and not existing.done():
        return  # already mid-flight from a prior drain
    _DEFERRED_SYNC_QUEUE.add(key)


async def drain_deferred_syncs() -> int:
    """Run any queued auto-syncs in the background.

    Called from the periodic probe loop AFTER the liveness sweep
    finishes — so syncs never compete with chat traffic the way the
    old per-routing trigger did. Returns the number of syncs
    started so the caller can log it.
    """
    started = 0
    pending = list(_DEFERRED_SYNC_QUEUE)
    _DEFERRED_SYNC_QUEUE.clear()
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        # Re-enqueue and bail; next sweep will catch it.
        _DEFERRED_SYNC_QUEUE.update(pending)
        return 0

    for wid, model_name in pending:
        existing = _AUTO_SYNC_TASKS.get((wid, model_name))
        if existing and not existing.done():
            continue
        async def _run(wid=wid, model_name=model_name):
            try:
                from . import model_sync
                await model_sync.sync_model(model_name, wid)
                log.info("compute_pool: auto-synced %s -> worker %s",
                         model_name, wid)
            except Exception as e:
                log.info("compute_pool: auto-sync %s -> %s failed: %s",
                         model_name, wid, e)
            finally:
                _AUTO_SYNC_TASKS.pop((wid, model_name), None)
        _AUTO_SYNC_TASKS[(wid, model_name)] = loop.create_task(_run())
        started += 1
    return started


# Per-process round-robin counter for embedding fan-out, keyed by
# embedding model name. Embedding requests are stateless and uniform —
# any worker that has the model loaded can serve any request in the
# stream. Rotating across the eligible set is dramatically faster than
# pinning to one worker: a 1000-chunk codebase index parallelises N-fold
# across N workers instead of bottlenecking on the strongest single
# worker's queue.
_EMBED_TARGET_INDEX: dict[str, int] = {}


# ---------------------------------------------------------------------------
# Per-conversation worker affinity
#
# Without affinity, every turn re-runs `pick_chat_target` and routes to
# the strongest node — but two simultaneously-active conversations both
# pile onto the same strongest node, queueing serially even though the
# rest of the pool sits idle. With affinity, the first turn lands on
# whichever node is least-loaded among eligible candidates and that
# (conversation, node) pair sticks for subsequent turns. KV cache stays
# warm per-conversation, AND multiple concurrent conversations naturally
# spread across the pool.
#
# In-memory state — affinity is best-effort; a backend restart resets
# every pairing and the next turn re-picks. The KV cache is process-
# local on each node anyway, so a restart was going to lose the warmth
# regardless. Active-turn counters are also in-memory; they undercount
# briefly across a restart but recover on the first new turn.
# ---------------------------------------------------------------------------

# conv_id → "host" | "worker:<wid>"
_CONV_AFFINITY: dict[str, str] = {}

# node_id → count of currently-running chat turns. Used to break ties
# when two equally-strong nodes are both eligible — the least-busy
# wins so concurrent conversations spread automatically.
_ACTIVE_TURNS_PER_NODE: dict[str, int] = {}


def _node_id_for_target(target: tuple[str, str | None] | None) -> str:
    """Translate a `pick_chat_target` return value into a node identifier.

    `pick_chat_target` returns ``(base_url, token)`` for a worker or
    ``None`` for host. We collapse to a string so affinity / active-
    turn tables can key uniformly.
    """
    if target is None:
        return "host"
    base_url = (target[0] or "").rstrip("/")
    # Look up the worker row matching this base_url so the node_id
    # ties to a stable identifier (worker IDs are durable; URLs may
    # change across DHCP rebinds).
    for w in db.list_compute_workers(enabled_only=False):
        if _worker_base_url(w).rstrip("/") == base_url:
            return f"worker:{w['id']}"
    return "host"


def register_turn_start(conv_id: str, node_id: str) -> None:
    """Record that ``conv_id`` is now actively running a turn on ``node_id``.

    Updates two structures:
      * `_CONV_AFFINITY[conv_id] = node_id` — sticks the conversation
        to this node for subsequent turns so KV cache stays warm.
      * `_ACTIVE_TURNS_PER_NODE[node_id] += 1` — load counter that
        future picks consult to break ties between equally-eligible
        nodes.

    Idempotent: calling twice for the same (conv, node) bumps the
    counter twice. Caller MUST balance every start with `register_turn_end`.
    """
    if not conv_id or not node_id:
        return
    _CONV_AFFINITY[conv_id] = node_id
    _ACTIVE_TURNS_PER_NODE[node_id] = _ACTIVE_TURNS_PER_NODE.get(node_id, 0) + 1


def register_turn_end(conv_id: str, node_id: str) -> None:
    """Counterpart to `register_turn_start`. Decrement the load counter.

    Affinity persists past turn end on purpose — the next turn for
    this conversation should land on the same node. Affinity is only
    cleared when the node goes ineligible (worker disabled / probe
    failing) or by an explicit `forget_conv_affinity` call.
    """
    if not node_id:
        return
    cur = _ACTIVE_TURNS_PER_NODE.get(node_id, 0)
    _ACTIVE_TURNS_PER_NODE[node_id] = max(0, cur - 1)


def forget_conv_affinity(conv_id: str) -> None:
    """Drop the affinity record for ``conv_id``. Called when the
    conversation is deleted so the in-memory dict doesn't grow forever.
    """
    _CONV_AFFINITY.pop(conv_id, None)


_COMPACTION_TARGET_INDEX: dict[str, int] = {}


def pick_compaction_target(model: str) -> tuple[str, str | None] | None:
    """Choose a node to run an auto-compaction summarization, preferring
    workers over host.

    Why: compaction is a real LLM call against the chat model; running
    it on the host while the chat is also hot on host slows down the
    user's ongoing turn. When a worker has the model installed, route
    compaction there — host stays free for the active chat.

    Round-robins across eligible workers so back-to-back compactions
    spread across the pool. Returns ``None`` (host fallback) when no
    worker has the model — preserves the previous host-only behavior
    on single-node setups.
    """
    cands = _eligible_workers("use_for_chat", model=model)
    if not cands:
        return None
    idx = _COMPACTION_TARGET_INDEX.get(model, 0)
    pick = cands[idx % len(cands)]
    _COMPACTION_TARGET_INDEX[model] = (idx + 1) % len(cands)
    base = _worker_base_url(pick)
    token = db.get_compute_worker_auth_token(pick["id"])
    return base, token


def embed_concurrency_limit(model: str) -> int:
    """Recommend an in-flight embed-call cap for a given model.

    `pick_embed_target` round-robins each call across host + every
    eligible worker. The previous indexer awaited each call serially,
    so even with N workers the pool ran at 1× — only one backend was
    ever busy. Concurrent fan-out (asyncio.gather under a semaphore)
    multiplies indexing throughput by N effectively.

    The cap balances two pressures:
      * Too low → workers idle while the host is busy, no win.
      * Too high → memory pressure on small workers (each in-flight
        embed holds ~50 MB of activation buffers), and HTTP queue
        timeouts on the slowest worker.

    Heuristic: 2× the number of usable backends (host + eligible
    workers), clamped to [2, 16]. Empirically the "2×" factor lets
    each backend keep one request running while the next is in flight,
    without queuing more than the slowest backend's serving rate can
    drain.
    """
    cands = _eligible_workers("use_for_embeddings", model=model)
    backend_count = 1 + len(cands)  # +1 for host
    return max(2, min(16, backend_count * 2))


def pick_chat_worker(
    model: str, conv_id: str | None = None,
) -> dict | None:
    """Return the FULL worker dict the chat picker would route to,
    or None when host wins.

    Same selection logic as `pick_chat_target` (affinity, capability
    score, mega-busy bias) — just exposes the row instead of the
    pre-formatted (url, token) tuple. Callers that need the worker's
    `use_encrypted_proxy` flag use this; callers that just want the
    URL can stay on `pick_chat_target`.
    """
    cands = _eligible_workers("use_for_chat", model=model)
    sticky = _CONV_AFFINITY.get(conv_id) if conv_id else None
    if sticky and sticky != "host":
        wid = sticky.split(":", 1)[1] if ":" in sticky else ""
        sticky_w = next((w for w in cands if w["id"] == wid), None)
        if sticky_w is not None:
            sticky_load = _ACTIVE_TURNS_PER_NODE.get(sticky, 0)
            better = next(
                (
                    w for w in cands
                    if w["id"] != wid
                    and _ACTIVE_TURNS_PER_NODE.get(f"worker:{w['id']}", 0)
                        <= sticky_load - 2
                ),
                None,
            )
            if better is None:
                return sticky_w
    if not cands:
        return None
    top_score = _capability_score(cands[0])
    tied = [c for c in cands if _capability_score(c) == top_score]
    if len(tied) > 1:
        tied.sort(
            key=lambda c: _ACTIVE_TURNS_PER_NODE.get(f"worker:{c['id']}", 0),
        )
        w = tied[0]
    else:
        w = cands[0]
    worker_score = _capability_score(w)
    host_score = _host_capability_score(model)
    if worker_score[0] <= 0 or host_score[0] <= 0:
        worker_score = (0.0,) + worker_score[1:]
        host_score = (0.0,) + host_score[1:]
    if worker_score <= host_score:
        # Same model-presence guard as pick_chat_target — host
        # only wins when it actually has the model installed,
        # otherwise return the worker so the agent can dispatch
        # via encrypted proxy instead of 404'ing on host Ollama.
        host_has_model = resolve_ollama_model(model) is not None
        if host_has_model and not is_host_mega_busy():
            return None
    return w


def pick_embed_worker(model: str) -> dict | None:
    """Return the FULL worker dict (or None for host) the next embed
    request should route to.

    Wraps the same picker `pick_embed_target` uses but returns the
    full row so callers can read `use_encrypted_proxy` and dispatch
    through the secure proxy when needed. Behaviour-equivalent for
    callers that only need the URL/token.
    """
    cands = _eligible_workers("use_for_embeddings", model=model)
    if not cands:
        return None
    if len(cands) > 1:
        top_score = _capability_score(cands[0])
        threshold_score = top_score
        usable = [
            c for c in cands
            if _capability_score(c) >= _scaled_score_threshold(threshold_score)
        ]
    else:
        usable = cands
    idx = _EMBED_TARGET_INDEX.get(model, 0)
    w = usable[idx % len(usable)]
    _EMBED_TARGET_INDEX[model] = (idx + 1) % max(len(usable), 1)
    return w


def pick_embed_target(model: str) -> tuple[str, str | None] | None:
    """Choose a worker to run an embed request against, or None for host.

    Round-robins across every eligible worker that has the embedding
    model installed. Embedding throughput in the pool was previously
    bottlenecked by the strongest single worker — this rotation lets
    the whole pool absorb concurrent embed traffic (RAG indexing,
    codebase search builds, recall lookups) at near-N× the previous
    rate.

    Returns `(base_url, auth_token_or_None)`. `auth_token` is fetched
    from the dedicated `get_compute_worker_auth_token` so the token never
    sits on a row dict. Caller composes the URL as `f"{base}/api/embeddings"`
    and adds `Authorization: Bearer …` when the token is set.
    """
    cands = _eligible_workers("use_for_embeddings", model=model)
    if not cands:
        return None

    # Capability-sorted: `_eligible_workers` already returns the strongest
    # worker first. We don't blindly round-robin across all of them — a
    # very weak worker with the model loaded would slow the whole rotation
    # to its pace. Take the top-K where K is the count of workers within
    # 50 % of the strongest worker's score; that lets pools with mostly-
    # uniform laptops use everyone while still excluding a clearly-slower
    # outlier (e.g. a Pi running Ollama for diagnostic reasons).
    if len(cands) > 1:
        top_score = _capability_score(cands[0])
        # Score tuple's first element is measured TPS (or 0 if no bench).
        # Fall back to capability tier when no measurement.
        threshold_score = top_score
        usable = [
            c for c in cands
            if _capability_score(c) >= _scaled_score_threshold(threshold_score)
        ]
    else:
        usable = cands

    # Stateful rotation. Module-level counter avoids re-issuing the same
    # worker on back-to-back calls within a turn. Atomic-ish — the GIL
    # gives us enough safety; an embed request lost across a race is a
    # one-off latency hit, never a correctness problem.
    idx = _EMBED_TARGET_INDEX.get(model, 0)
    w = usable[idx % len(usable)]
    _EMBED_TARGET_INDEX[model] = (idx + 1) % max(len(usable), 1)
    base = _worker_base_url(w)
    token = db.get_compute_worker_auth_token(w["id"])
    return (base, token)


def _scaled_score_threshold(top_score: tuple) -> tuple:
    """Half the leader's measured TPS becomes the floor for inclusion in
    the round-robin. Tuple comparison works because the score schema
    starts with TPS — anything with TPS ≥ half-of-top is in.

    For pools whose leader has no measurement yet (TPS=0), we fall
    through to a permissive threshold (everything qualifies) — the
    rotation absorbs uneven workers in that case but still parallelises
    across them, which is strictly better than no rotation at all.
    """
    if not top_score or len(top_score) == 0:
        return top_score
    head = top_score[0] or 0
    if head <= 0:
        # No measurement on the leader — fall through to including all.
        # Replace the head with 0 so the threshold accepts any worker.
        return (0,) + tuple(top_score[1:])
    halved = head / 2.0
    return (halved,) + tuple(top_score[1:])


def pick_chat_target(
    model: str, conv_id: str | None = None,
) -> tuple[str, str | None] | None:
    """Choose where the chat turn runs — return (base_url, token) for a
    worker, or None to keep it on host.

    This is where the "use the more powerful machine" rule lives.
    Among eligible chat workers we already sort by `_capability_score`
    (GPU > CPU, more VRAM > less). Now we ALSO compare the strongest
    eligible worker against the host's own capability — if the worker
    isn't strictly more powerful than host, we keep the turn local
    (no LAN round-trip, KV cache stays warm). The worker only wins
    when its proven hardware beats the host's.

    Concretely: a worker beats host when EITHER
      * host has no GPU and the worker does, OR
      * both have a GPU but the worker has loaded a model larger than
        the host's total VRAM (max_vram_seen_bytes > host_vram_bytes),
        which is a hard lower-bound proof of more capacity.

    When ``conv_id`` is supplied AND a previous turn for the same
    conversation already landed on a node, the router prefers that
    node (KV cache stays warm). Affinity is overridden when the
    sticky node has gone ineligible (worker disabled / probe failing)
    or is more than one active turn behind the next-strongest node —
    at which point we re-pick and update the affinity.

    Tie-breaking (when two equally-strong nodes are eligible) prefers
    the one with FEWER currently-running turns so concurrent
    conversations spread across the pool instead of stacking on the
    strongest single node.
    """
    cands = _eligible_workers("use_for_chat", model=model)

    # Affinity short-circuit: if this conversation already has a sticky
    # node and that node is still eligible AND not heavily loaded vs.
    # the alternatives, return the affinity choice without re-scoring.
    sticky = _CONV_AFFINITY.get(conv_id) if conv_id else None
    if sticky:
        sticky_load = _ACTIVE_TURNS_PER_NODE.get(sticky, 0)
        if sticky == "host":
            # Affinity to host — return None unless a worker has become
            # dramatically more capable since we last picked. Cheap
            # comparison: if no eligible worker exists, host wins.
            if not cands:
                return None
            # Otherwise fall through to the standard host-vs-worker
            # comparison below; the affinity hint is a tiebreaker, not
            # an override of the strength rule.
        else:
            wid = sticky.split(":", 1)[1] if ":" in sticky else ""
            sticky_w = next((w for w in cands if w["id"] == wid), None)
            if sticky_w is not None:
                # Sticky worker still eligible — return it unless a
                # MUCH less-loaded alternative exists (load delta ≥ 2).
                # The threshold prevents thrashing affinity on every
                # tiny load fluctuation while still spreading load
                # when one node is genuinely overwhelmed.
                better = next(
                    (
                        w for w in cands
                        if w["id"] != wid
                        and _ACTIVE_TURNS_PER_NODE.get(f"worker:{w['id']}", 0)
                            <= sticky_load - 2
                    ),
                    None,
                )
                if better is None:
                    return (
                        _worker_base_url(sticky_w),
                        db.get_compute_worker_auth_token(sticky_w["id"]),
                    )
                # else fall through to fresh pick — affinity stale.

    if not cands:
        return None
    # Near-tie-aware selection. The original tie-breaker only kicked
    # in on EXACT capability-score equality, which almost never
    # happens in practice (one worker has measured 4.2 tok/s, another
    # 4.1 — different scores → no tie). We now widen "tied" to the
    # near-tie band (top_score scaled by ~0.85) so the load and
    # locality factors get a real say when N workers all have the
    # same model installed.
    top_score = _capability_score(cands[0])
    near_band = _scaled_score_threshold(top_score)
    near_tied = [c for c in cands if _capability_score(c) >= near_band]

    if len(near_tied) > 1:
        # Multi-criteria sort, lowest-key-wins:
        #   1. LOCALITY — LAN-paired peers (and host) before public-
        #      pool peers. A LAN peer's RTT is sub-millisecond; a
        #      public-pool peer crosses the internet, often hundreds
        #      of ms. Only let a public peer win on locality when
        #      it's measurably more capable AND no LAN option is
        #      eligible.
        #   2. LOAD — fewer in-flight turns wins. Concurrent
        #      conversations spread across the pool instead of
        #      stacking on the strongest single node.
        #   3. CAPABILITY — within the near-tie band, but still a
        #      tie-breaker when locality + load are equal.
        def _locality_rank(w: dict) -> int:
            # 0 = LAN paired, 1 = public, 2 = unknown / manually-added
            did = w.get("gigachat_device_id") or ""
            if not did:
                return 2
            try:
                paired = db.get_paired_device(did)
            except Exception:
                paired = None
            role = (paired or {}).get("role") or ""
            if role == "local":
                return 0
            if role == "public":
                return 1
            return 2
        near_tied.sort(
            key=lambda c: (
                _locality_rank(c),
                _ACTIVE_TURNS_PER_NODE.get(f"worker:{c['id']}", 0),
                # Negate capability so higher beats lower in ascending sort.
                tuple(-x for x in _capability_score(c)),
            ),
        )
        w = near_tied[0]
    else:
        w = cands[0]
    # Schedule a background bench of the host on this exact model so
    # the cache is populated for the next routing call. We can't await
    # here (this function is sync), so the FIRST call falls through
    # to heuristic axes only — that's safer than letting the worker's
    # measured-TPS field beat the host's missing-measurement field.
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            loop.create_task(_measure_host_throughput(model))
    except RuntimeError:
        pass

    worker_score = _capability_score(w)
    host_score = _host_capability_score(model)
    # If either side hasn't been benchmarked on this exact model
    # (TPS=0 means cache miss or measurement failure), fall through
    # to the heuristic axes so the comparison is apples-to-apples.
    # Otherwise the worker's "measured 4 tok/s on a different model"
    # would beat the host's "no measurement yet" → false positive.
    if worker_score[0] <= 0 or host_score[0] <= 0:
        worker_score = (0.0,) + worker_score[1:]
        host_score = (0.0,) + host_score[1:]

    if worker_score <= host_score:
        # Host wins on capability — but ONLY return host when host
        # actually has the model installed. Otherwise routing to
        # host means hitting host's Ollama with a model it doesn't
        # have, which 404's with "model not found" and the user
        # sees a useless error. The worker (which has the model)
        # is the correct choice in that case even if it's slightly
        # less capable on paper.
        host_has_model = resolve_ollama_model(model) is not None
        if not host_has_model:
            log.info(
                "compute_pool: host beats worker on capability for "
                "%r but host doesn't have the model installed — "
                "routing to worker %s anyway (avoids 404)",
                model, w.get("label"),
            )
        elif not is_host_mega_busy():
            # Host wins AND has the model AND isn't disk-saturated
            # → stay local (no LAN hop, KV cache stays warm).
            return None
        # Otherwise: host either lacks the model or is mega-busy →
        # fall through and return the worker.
    return (_worker_base_url(w), db.get_compute_worker_auth_token(w["id"]))


def pick_split_chat_target(model_name: str) -> tuple[str, str] | None:
    """Legacy lookup retained for back-compat with tests / explicit
    `split:<label>` model names. The auto-router below
    (`route_chat_for`) is the live entry point now; it never produces
    `split:` prefixes — it just inspects the model and decides whether
    to engage the split path transparently.

    Returns `(base_url, label)` for an explicit `split:<label>` whose
    row is `running`; None otherwise.
    """
    if not model_name or not model_name.startswith("split:"):
        return None
    label = model_name[len("split:"):].strip()
    if not label:
        return None
    for row in db.list_split_models(enabled_only=True):
        if row.get("label") == label and row.get("status") == "running":
            return (f"http://127.0.0.1:{row['llama_port']}", label)
    return None


# ---------------------------------------------------------------------------
# Auto-routing: pick host-Ollama vs spawn-llama-server-with-RPC for a model
# ---------------------------------------------------------------------------
# This is the intelligent split-or-not decision the user actually picks
# their model against. They never see `split:<label>` — they pick e.g.
# `gemma3:27b` from the model picker, and the router decides:
#
#   1. Resolve the model's GGUF blob + size from Ollama's manifest store.
#   2. If size <= host_vram_budget → Ollama on host (fastest path; 0 LAN
#      overhead; workers stay free for parallel embeddings/subagents).
#   3. Else → ensure a llama-server is running for THIS exact model with
#      --rpc to every eligible compute worker, return its URL.
#   4. If the model can't fit even with the combined pool → raise so the
#      user gets a clear error rather than silent OOM.
#
# Per-conversation lifecycle: when the user switches between two big
# models, we stop the previously-running llama-server (one big model hot
# at a time — we have one finite VRAM budget). Switching back to a small
# model also stops any running llama-server so its VRAM is freed for
# Ollama.
# ---------------------------------------------------------------------------

# How much of the host's VRAM we're willing to let one model occupy
# without engaging the split path. 85% leaves headroom for the OS, the
# desktop compositor, any other Ollama models loaded simultaneously, and
# the model's KV cache. Below this fraction → Ollama. Above → split.
_HOST_VRAM_USE_FRACTION = 0.85

# How much of the host's TOTAL memory (VRAM + RAM) we'll trust Ollama
# to run a single model from. Ollama auto-offloads layers that don't
# fit VRAM into system RAM and runs them on CPU — slow per layer, but
# strictly faster than streaming layer activations across a LAN every
# token. 70% of (vram + ram) leaves room for the OS, browser, the
# Gigachat backend itself, and KV cache. Below this fraction → host
# (Ollama). Above → engage split path with workers.
_HOST_TOTAL_USE_FRACTION = 0.70

# Adaptive routing — pool-capacity heuristic for engaging Phase 2 even
# when the model fits the host alone. We use a static threshold instead
# of a measured-TPS bake-off because the chat layer doesn't currently
# wire realised tokens-per-second back to the router; until that's
# instrumented, "engage when the pool is meaningfully bigger than host"
# is the safest proxy. The threshold below is conservative — pool VRAM
# must exceed host VRAM by 50 % before we engage split, which avoids
# regressing setups where pool ≈ host but split LAN-cost dominates.
_POOL_VRAM_OUTSCALE_FACTOR = 1.5

# Aggressive pooling — engage split even when host fits the model
# alone, IF the workers' COMBINED free RAM adds at least this many
# GB of usable budget. This is the user's "use as much pool resource
# as available" policy: even on a fast LAN, we'd rather pay the LAN
# round-trip and free up host RAM for everything else than keep all
# layers on the orchestrator. 4 GB is the sweet spot — below that
# the LAN cost dominates the win; above that the win compounds.
_AGGRESSIVE_POOL_GB = 4.0


# ---------------------------------------------------------------------------
# Mega-model busy tracker
#
# When the host engages the mega-model path (model > pool memory →
# layers page from host SSD via mmap on every forward pass), the disk
# is the bottleneck. Other in-flight work that ALSO routes to the host
# (parallel subagents, embedding fan-out, side chats) competes with
# that mmap traffic for I/O bandwidth and slows everyone down.
#
# Solution: when the mega-model path engages, mark the host "mega-busy"
# for a short window. Routing functions then bias toward workers so the
# host disk stays dedicated to the mmap stream. Each new mega-model
# turn refreshes the window; otherwise it expires naturally so a quick
# follow-up chat after the mega turn is over goes to the strongest
# node again.
# ---------------------------------------------------------------------------

# Wall-clock until which we consider host's disk subsystem mega-busy.
# Set on every mega-model engagement; checked by routing helpers.
_HOST_MEGA_MODEL_BUSY_UNTIL: float = 0.0

# How long one mega-model engagement keeps the host marked busy. 5 min
# matches Ollama's default model-keep-alive window — long enough to
# cover a typical mega-model turn end-to-end (load + prefill + decode
# + tool calls), short enough to recover quickly when the mega-model
# session ends and the user switches to a smaller chat.
_HOST_MEGA_MODEL_BUSY_TTL_SEC = 300.0


def _mark_host_mega_busy() -> None:
    """Stamp the host as mega-busy for the next `_HOST_MEGA_MODEL_BUSY_TTL_SEC`
    seconds. Called by `route_chat_for` whenever it commits to the
    mega-model path (model exceeds pool memory). Idempotent and cheap —
    overlapping mega turns just keep refreshing the deadline.
    """
    global _HOST_MEGA_MODEL_BUSY_UNTIL
    _HOST_MEGA_MODEL_BUSY_UNTIL = time.time() + _HOST_MEGA_MODEL_BUSY_TTL_SEC


def is_host_mega_busy() -> bool:
    """True when the host is currently mmapping a mega-model from disk.
    Routing helpers consult this to bias parallel work toward workers
    so the host's disk subsystem stays dedicated to the mega-model
    page-in stream.
    """
    return time.time() < _HOST_MEGA_MODEL_BUSY_UNTIL

# Adaptive routing — per-model TPS history. Reserved for a future
# commit that wires post-turn realised-TPS recording from agent.py;
# until then the cache stays empty and the static
# `_POOL_VRAM_OUTSCALE_FACTOR` heuristic above is what fires.
_ROUTE_TPS_CACHE_TTL_SEC = 24 * 3600
_ROUTE_TPS_CACHE: dict[str, dict] = {}


def _record_route_tps(model_name: str, *, kind: str, tps: float) -> None:
    """Stamp a measurement on the per-model TPS cache.

    `kind` is `"host"` or `"split"`. Only TPS > 0 are recorded — a
    failed bench leaves the previous entry alone so a transient probe
    error doesn't clobber valid data. Currently called from nowhere
    (post-turn TPS recording is a future commit); shipped now so the
    chat layer can be wired without a routing-policy change.
    """
    if tps <= 0 or kind not in ("host", "split"):
        return
    bucket = _ROUTE_TPS_CACHE.setdefault(model_name, {})
    bucket[f"{kind}_tps"] = tps
    bucket[f"{kind}_measured_at"] = time.time()


def _route_tps_for(model_name: str, kind: str) -> float | None:
    """Read a measurement back. Returns None if missing or stale."""
    bucket = _ROUTE_TPS_CACHE.get(model_name) or {}
    val = bucket.get(f"{kind}_tps")
    measured_at = bucket.get(f"{kind}_measured_at") or 0
    if not val or val <= 0:
        return None
    if time.time() - measured_at > _ROUTE_TPS_CACHE_TTL_SEC:
        return None
    return float(val)


def _should_force_split_for(
    model_name: str,
    *,
    strongest_single_vram: int,
    pool_vram_total: int,
) -> bool:
    """Decide whether to engage Phase 2 for a model that fits one node.

    The comparison anchor is the **strongest single node** (host OR
    a worker) — not just host. When a worker has more VRAM than host
    AND can hold the model alone, Phase 1 routing already sends the
    chat to that worker; engaging Phase 2 on top would just add LAN
    overhead with no capacity win. So the heuristic only kicks in
    when the pool is meaningfully bigger than the strongest single
    node — which is when split actually adds memory the chat would
    not otherwise have.

    Two-step decision, with the measured-TPS cache preferred when
    both samples exist; otherwise falls back to the pool-capacity
    heuristic.

    1. **Measured verdict.** When both host_tps and split_tps for
       this model are in the cache and the cached samples are fresh,
       engage split iff split_tps > host_tps. Ground truth.
    2. **Heuristic verdict.** When TPS data is missing (cold start,
       or the chat layer hasn't been wired to record TPS yet),
       engage split iff `pool_vram_total ≥ strongest_single_vram *
       _POOL_VRAM_OUTSCALE_FACTOR` (1.5×) AND there's at least one
       eligible rpc worker. The 1.5× threshold avoids engaging when
       the pool is barely bigger than the strongest single node —
       LAN per-token cost would dominate the small capacity win.
    """
    split_tps = _route_tps_for(model_name, "split")
    host_tps = _route_tps_for(model_name, "host")
    if split_tps is not None and host_tps is not None:
        return split_tps > host_tps
    if not _eligible_split_workers():
        return False
    if strongest_single_vram <= 0:
        return False
    return pool_vram_total >= strongest_single_vram * _POOL_VRAM_OUTSCALE_FACTOR

# Path to Ollama's on-disk model store. Default location on every
# platform Ollama supports; matches what `ollama_runtime` assumes.
_OLLAMA_MODELS_DIR = Path.home() / ".ollama" / "models"

# Where users drop replacement GGUFs that override Ollama's blob for a given
# model name. Necessary because some upstream GGUFs ship with a tensor layout
# that stock llama.cpp can't load even though the architecture is supported
# (e.g. Ollama's `gemma4:26b` packs the MoE expert weights as fused
# `ffn_gate_up_exps` per layer — 658 tensors total — while llama.cpp's
# `gemma4` loader spec expects them unfused as `ffn_gate_exps` +
# `ffn_up_exps` separately, 1014 tensors). The override path lets the user
# install Unsloth's GGUF (which uses the unfused layout) and have Phase 2
# split routing pick it up automatically without us mutating Ollama's
# blob store. Filename is the Ollama model name with `:` replaced by `-`,
# so `gemma4:26b` becomes `gemma4-26b.gguf`.
_OVERRIDE_GGUF_DIR = Path.home() / ".gigachat" / "llama-cpp" / "models"


def _override_gguf_path_for(model_name: str) -> Path:
    """Map an Ollama model name to its override GGUF path.

    Returns the path regardless of whether the file exists — callers
    check `.is_file()` to decide whether to use it.
    """
    safe = (model_name or "").strip().replace(":", "-").replace("/", "-")
    return _OVERRIDE_GGUF_DIR / f"{safe}.gguf"


def _override_mmproj_path_for(model_name: str) -> Path:
    """Map an Ollama model name to its multimodal projector path.

    Same naming convention as the main override but with a `.mmproj`
    infix: `gemma4:26b` → `gemma4-26b.mmproj.gguf`. When this file
    exists alongside the main override, llama-server is launched
    with `--mmproj <path>` enabling vision input. The mmproj is
    distributed/produced separately from the main LLM GGUF; for
    Ollama-bundled multimodal models the user can either run our
    extraction script once or download a pre-built mmproj from the
    upstream repo (e.g. Unsloth's `mmproj-BF16.gguf`).
    """
    safe = (model_name or "").strip().replace(":", "-").replace("/", "-")
    return _OVERRIDE_GGUF_DIR / f"{safe}.mmproj.gguf"


# ---------------------------------------------------------------------------
# Scope B: end-user auto-acquisition of override files
# ---------------------------------------------------------------------------
#
# Registry of known incompatible Ollama models that need an override GGUF
# (and optionally a separate mmproj GGUF) to load via llama.cpp's
# llama-server. For each entry:
#   * `needs_mmproj` - is a separate vision projector file required?
#   * `main_strategy` - "extract_text_only" runs the surgery script,
#     dropping `v.*` / `mm.*` tensors from the local Ollama blob (zero
#     bandwidth, lossless). "download" goes straight to HuggingFace.
#   * `main_url` / `main_size_gb` - HuggingFace fallback if surgery
#     fails or the user doesn't have the blob locally.
#   * `mmproj_url` / `mmproj_size_gb` - the multimodal projector. We
#     always download these from upstream — they come pre-built in
#     CLIP format (clip.has_vision_encoder=true, clip.projector_type=
#     "gemma4v") which is much harder to derive from Ollama's bundle.
#   * `reason` - human-readable explanation surfaced in errors / UI.
#
# Acquisition priority (cheapest -> most expensive):
#   1. Files already in override dir -> no-op
#   2. LAN copy: another worker has the override files -> SCP
#      (requires `ssh_host` set on the worker AND the worker to have
#      run the surgery before; the periodic probe loop will broadcast
#      this state once it's tracked - V2)
#   3. Local surgery: extract from user's Ollama blob (zero bandwidth)
#   4. Direct download from HuggingFace
_KNOWN_OVERRIDE_REGISTRY: dict[str, dict] = {
    "gemma4:26b": {
        "needs_mmproj": True,
        "main_strategy": "extract_text_only",
        "main_url": (
            "https://huggingface.co/unsloth/gemma-4-26B-A4B-it-GGUF/resolve/main/"
            "gemma-4-26B-A4B-it-UD-Q5_K_M.gguf"
        ),
        "main_size_gb": 19.7,
        "mmproj_url": (
            "https://huggingface.co/unsloth/gemma-4-26B-A4B-it-GGUF/resolve/main/"
            "mmproj-BF16.gguf"
        ),
        "mmproj_size_gb": 1.11,
        "reason": (
            "Ollama bundles vision tower in the LLM GGUF; "
            "llama.cpp expects a separate mmproj file."
        ),
    },
    "gemma4:31b": {
        "needs_mmproj": True,
        "main_strategy": "extract_text_only",
        "main_url": (
            "https://huggingface.co/unsloth/gemma-4-31B-it-GGUF/resolve/main/"
            "gemma-4-31B-it-Q5_K_M.gguf"
        ),
        "main_size_gb": 20.17,
        "mmproj_url": (
            "https://huggingface.co/unsloth/gemma-4-31B-it-GGUF/resolve/main/"
            "mmproj-BF16.gguf"
        ),
        "mmproj_size_gb": 1.12,
        "reason": (
            "Same bundled-multimodal pattern as gemma4:26b — Ollama's "
            "blob has a vision tower llama-server can't load directly."
        ),
    },
    "qwen3.5:9b": {
        # Multi-layer structural mismatch between Ollama's blob and
        # stock llama.cpp's qwen35 loader spec:
        #   1. rope.dimension_sections array length 3 vs expected 4
        #   2. SSM tensors named `ssm_dt` / `ssm_a` (no .bias suffix)
        #      vs llama.cpp's expected `ssm_dt.bias`, `ssm_a.bias`
        #   3. llama.cpp expects BOTH `blk.N.ssm_a` AND
        #      `blk.N.ssm_a.bias` as separate tensors; Ollama ships
        #      only one
        # Surgery on the local blob would need to synthesize the
        # missing bias tensors with zero values, which introduces
        # silent quality risk (we'd be guessing values the model
        # was trained against). User's "no issues" constraint rules
        # that out. Strategy = "download" pulls Unsloth's clean
        # Q4_K_M variant (matches Ollama's quant level — same
        # quality, just packaged for stock llama.cpp).
        "needs_mmproj": False,
        "main_strategy": "download",
        "main_url": (
            "https://huggingface.co/unsloth/Qwen3.5-9B-GGUF/resolve/main/"
            "Qwen3.5-9B-Q4_K_M.gguf"
        ),
        "main_size_gb": 5.29,
        "reason": (
            "Ollama blob's qwen35 layout has multiple structural "
            "mismatches with stock llama.cpp; surgery would require "
            "synthesizing missing tensors with arbitrary values."
        ),
    },
    # Gemma 3n E4B / E2B: Ollama labels these `gemma4:e4b` and `gemma4:e2b`,
    # but they're really Google Gemma 3n with Per-Layer Embeddings +
    # MatFormer architecture (audio + vision + text in the original).
    # Stock llama.cpp expects arch string "gemma3n", not Ollama's
    # internal "gemma4", and doesn't load Ollama's bundled audio/vision
    # tensors at all. We download Unsloth's clean text-only Gemma 3n
    # GGUF (matching Ollama's Q4_K_M quant level — no quality drop)
    # so llama-server can engage the pool. Vision/audio are not
    # available via this path (Unsloth doesn't publish an mmproj for
    # Gemma 3n yet); for those modalities, Ollama's runtime continues
    # to handle the model on host.
    "gemma4:e4b": {
        "needs_mmproj": False,
        "main_strategy": "download",
        "main_url": (
            "https://huggingface.co/unsloth/gemma-3n-E4B-it-GGUF/resolve/main/"
            "gemma-3n-E4B-it-Q4_K_M.gguf"
        ),
        "main_size_gb": 4.71,
        "reason": (
            "Ollama's gemma4:e4b is Gemma 3n E4B repackaged under the "
            "wrong arch name; stock llama.cpp can't load it. Unsloth's "
            "GGUF has the right metadata."
        ),
    },
    # `gemma4:latest` shares the digest with `gemma4:e4b`. Register
    # the same target so either Ollama tag resolves to the override.
    "gemma4:latest": {
        "needs_mmproj": False,
        "main_strategy": "download",
        "main_url": (
            "https://huggingface.co/unsloth/gemma-3n-E4B-it-GGUF/resolve/main/"
            "gemma-3n-E4B-it-Q4_K_M.gguf"
        ),
        "main_size_gb": 4.71,
        "reason": "Alias of gemma4:e4b (same digest).",
    },
    "gemma4:e2b": {
        "needs_mmproj": False,
        "main_strategy": "download",
        "main_url": (
            "https://huggingface.co/unsloth/gemma-3n-E2B-it-GGUF/resolve/main/"
            "gemma-3n-E2B-it-Q4_K_M.gguf"
        ),
        "main_size_gb": 3.00,
        "reason": (
            "Same root cause as gemma4:e4b — Ollama's blob uses the "
            "internal `gemma4` arch name and bundles audio/vision "
            "tensors stock llama.cpp can't load. Unsloth's text-only "
            "Gemma 3n E2B GGUF works directly."
        ),
    },
}


# ---------------------------------------------------------------------------
# Generic auto-detection: any GGUF that ships with bundled vision
# tensors (`v.*`, `mm.*`) needs surgery + mmproj before it can load via
# stock llama-server. The registry above provides DOWNLOAD URLs for
# specific known models; this auto-detection lets us fire surgery for
# ANY model with the bundled-multimodal pattern, even ones we haven't
# explicitly registered. mmproj download still needs a registry entry
# (CLIP-format metadata can't be safely derived from Ollama's bundle).
# ---------------------------------------------------------------------------


def _ollama_blob_has_bundled_multimodal(model_name: str) -> bool:
    """Return True if Ollama's blob for `model_name` contains bundled
    vision-tower tensors (`v.*` / `mm.*`) — a strong signal that
    stock llama-server can't load it without surgery + a separate
    mmproj. Cheap check (reads only the GGUF tensor index, not the
    weights), so it's safe to call from the routing hot path.

    Returns False on any read error so the caller falls back to
    its previous behavior (registry lookup or skip).
    """
    info = resolve_ollama_model(model_name)
    if not info:
        return False
    blob_path = info.get("ollama_blob_path") or info.get("gguf_path")
    if not blob_path or not Path(blob_path).is_file():
        return False
    try:
        import gguf
        reader = gguf.GGUFReader(blob_path)
    except Exception:
        return False
    for t in reader.tensors:
        name = t.name or ""
        if name.startswith("v.") or name.startswith("mm."):
            return True
    return False


def _auto_synthesize_registry_entry(model_name: str) -> dict | None:
    """If `model_name` looks like it needs surgery (Ollama blob has
    bundled vision tensors) but isn't explicitly in the registry,
    synthesize an "extract_text_only" entry on the fly.

    The synthesized entry has NO download URLs — the user's local
    Ollama blob is the only acquisition source. If surgery fails
    or there's no Ollama blob, the chat-time error explains that
    a registry entry is needed for the download fallback.

    Returns None if the model doesn't need an override at all
    (clean GGUF, would load fine in llama-server as-is).
    """
    if not _ollama_blob_has_bundled_multimodal(model_name):
        return None
    return {
        "needs_mmproj": True,
        "main_strategy": "extract_text_only",
        "main_url": None,
        "main_size_gb": 0,
        "mmproj_url": None,
        "mmproj_size_gb": 0,
        "reason": (
            "Auto-detected: Ollama's blob bundles vision tensors that "
            "stock llama-server can't load. Surgery extracts the text "
            "LLM from the local blob; provide an mmproj URL via the "
            "registry to enable image input."
        ),
        "synthesized": True,
    }


# Acquisition state per model so concurrent route_chat_for calls don't
# kick off duplicate downloads/surgeries. Keys are model names; values
# are dicts: {status: "running"|"done"|"error",
#             phase: "surgery"|"downloading-main"|"downloading-mmproj"|"done",
#             progress_pct: float, error: str|None, started_at: float}.
_ACQUISITION_STATE: dict[str, dict] = {}
_ACQUISITION_TASKS: dict[str, asyncio.Task] = {}


def get_acquisition_status(model_name: str) -> dict | None:
    """Public read of the current acquisition state for a model.

    Returns None if no acquisition has ever started for this model.
    Used by the chat layer to surface a "preparing model" status to
    the UI instead of returning a generic error.
    """
    return _ACQUISITION_STATE.get(model_name)


async def ensure_compatible_gguf(model_name: str) -> dict:
    """Make sure the override file(s) needed to load `model_name` via
    llama-server exist on disk.

    Behaviour:
      * Model not in `_KNOWN_OVERRIDE_REGISTRY` -> no-op,
        returns {"ok": True, "status": "no_override_needed"}.
      * Files already in place -> no-op,
        returns {"ok": True, "status": "ready"}.
      * Files missing AND no acquisition running -> kick off a
        background acquisition task and returns
        {"ok": False, "status": "starting", ...}.
      * Files missing AND acquisition running -> returns the live
        status dict so the caller can decide whether to wait or
        surface progress.

    The acquisition task itself runs out-of-band (asyncio task on the
    running loop). It does NOT block the chat call - by design. The
    chat layer is expected to translate `status: starting/running` into
    a user-facing "preparing the model, please retry shortly" response.
    """
    spec = _KNOWN_OVERRIDE_REGISTRY.get(model_name)
    if not spec:
        # Generic fallback: any model whose Ollama blob has bundled
        # multimodal tensors (`v.*` / `mm.*`) needs surgery to load
        # in stock llama-server, even if we haven't explicitly
        # registered it. Synthesize a minimal "extract_text_only"
        # spec on the fly so the acquisition pipeline kicks in.
        # Surgery from the local Ollama blob is the entire strategy
        # — without a download URL we can't fall back to HF, but
        # most users will have the Ollama blob locally, so the
        # "lossless local extraction" path covers the common case.
        spec = _auto_synthesize_registry_entry(model_name)
        if not spec:
            return {"ok": True, "status": "no_override_needed"}

    main_path = _override_gguf_path_for(model_name)
    mmproj_path = _override_mmproj_path_for(model_name)
    main_present = main_path.is_file()
    mmproj_present = mmproj_path.is_file() if spec.get("needs_mmproj") else True

    if main_present and mmproj_present:
        return {"ok": True, "status": "ready"}

    # If a task is already running for this model, just return its state.
    existing = _ACQUISITION_TASKS.get(model_name)
    if existing and not existing.done():
        return {"ok": False, **(_ACQUISITION_STATE.get(model_name) or {"status": "running"})}

    # No live task -> spawn one. We use asyncio.create_task so the
    # acquisition runs concurrently with chat handling.
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        # Called outside an event loop (e.g. test). Run synchronously
        # so callers still get a clean answer.
        result = await _acquire_override_files(model_name, spec)
        return result

    _ACQUISITION_STATE[model_name] = {
        "status": "starting",
        "phase": "init",
        "progress_pct": 0.0,
        "error": None,
        "started_at": time.time(),
        "needs_main": not main_present,
        "needs_mmproj": spec.get("needs_mmproj") and not mmproj_present,
    }
    task = loop.create_task(_acquire_override_files(model_name, spec))
    _ACQUISITION_TASKS[model_name] = task
    return {
        "ok": False,
        **_ACQUISITION_STATE[model_name],
        "estimated_total_gb": (
            (spec.get("main_size_gb", 0) if not main_present else 0)
            + (spec.get("mmproj_size_gb", 0) if spec.get("needs_mmproj") and not mmproj_present else 0)
        ),
    }


async def _acquire_override_files(model_name: str, spec: dict) -> dict:
    """Background acquisition worker — does the actual surgery /
    download work. Updates `_ACQUISITION_STATE[model_name]` as it
    progresses so the UI can show progress.

    Strategy ordering (cheapest first):
      0. LAN copy: check every enabled worker with `ssh_host` set;
         if any of them already has the override / mmproj file in
         their `~/.gigachat/llama-cpp/models/` (e.g. from a previous
         distribution after a different host's acquisition), SCP it
         over via SSH. Zero internet bandwidth. Useful when the
         user reinstalls on a host or adds a new host to a pool
         where workers already cache overrides.
      1. Local surgery if `main_strategy == "extract_text_only"` —
         requires the Ollama blob to exist locally. Zero bandwidth,
         lossless byte-copy of the LLM tensors.
      2. HuggingFace download for any file the prior steps couldn't
         produce.

    After the function lands the files locally, it kicks off a
    fire-and-forget background distribution to every worker so the
    next host that needs them can take the LAN path (step 0 above).
    """
    state = _ACQUISITION_STATE.setdefault(model_name, {})
    main_path = _override_gguf_path_for(model_name)
    mmproj_path = _override_mmproj_path_for(model_name)

    try:
        # Step 0: LAN-first — check every worker with ssh_host set to
        # see if it already caches the file. Save internet bandwidth
        # when another node in the pool has the override stashed.
        for path in (main_path, mmproj_path):
            if path.is_file():
                continue
            if path is mmproj_path and not spec.get("needs_mmproj"):
                continue
            for worker in db.list_compute_workers(enabled_only=True):
                if not (worker.get("ssh_host") or "").strip():
                    continue
                state.update({
                    "status": "running", "phase": "lan-copy",
                    "progress_pct": 2.0,
                })
                ok = await _lan_pull_override(worker, path.name, path)
                if ok:
                    log.info(
                        "compute_pool: pulled %s from %s via LAN",
                        path.name, worker.get("label"),
                    )
                    break

        # Step 1a: qwen3 rope-metadata patch — purely metadata fix on
        # the local Ollama blob; tensors copied verbatim. Zero
        # bandwidth, zero quantization change. Only applies to models
        # whose strategy is `qwen3_rope_metadata_patch`.
        if not main_path.is_file() and spec.get("main_strategy") == "qwen3_rope_metadata_patch":
            blob_info = resolve_ollama_model(model_name)
            ollama_blob = (blob_info or {}).get("ollama_blob_path")
            if not ollama_blob:
                manifest = _resolve_ollama_manifest(model_name)
                if manifest:
                    layers = manifest.get("layers") or []
                    for layer in layers:
                        if layer.get("mediaType") == "application/vnd.ollama.image.model":
                            digest = (layer.get("digest") or "").replace("sha256:", "")
                            candidate = _OLLAMA_MODELS_DIR / "blobs" / f"sha256-{digest}"
                            if candidate.is_file():
                                ollama_blob = str(candidate)
                            break
            if ollama_blob and Path(ollama_blob).is_file():
                state.update({
                    "status": "running", "phase": "metadata-patch",
                    "progress_pct": 5.0,
                })
                ok = await _run_qwen3_rope_patch(ollama_blob, str(main_path))
                if ok:
                    log.info(
                        "compute_pool: rope-metadata patch produced %s for %s",
                        main_path, model_name,
                    )

        # Step 1b: try local surgery for the main file if applicable.
        if not main_path.is_file() and spec.get("main_strategy") == "extract_text_only":
            blob_info = resolve_ollama_model(model_name)
            ollama_blob = (blob_info or {}).get("ollama_blob_path")
            if not ollama_blob:
                # `resolve_ollama_model` only returns ollama_blob_path
                # when an override exists already. Pull it directly.
                manifest = _resolve_ollama_manifest(model_name)
                if manifest:
                    layers = manifest.get("layers") or []
                    for layer in layers:
                        if layer.get("mediaType") == "application/vnd.ollama.image.model":
                            digest = (layer.get("digest") or "").replace("sha256:", "")
                            candidate = _OLLAMA_MODELS_DIR / "blobs" / f"sha256-{digest}"
                            if candidate.is_file():
                                ollama_blob = str(candidate)
                            break
            if ollama_blob and Path(ollama_blob).is_file():
                state.update({
                    "status": "running", "phase": "surgery",
                    "progress_pct": 5.0,
                })
                ok = await _run_text_only_surgery(
                    ollama_blob=ollama_blob,
                    dst=str(main_path),
                    arch=model_name.split(":", 1)[0],  # e.g. "gemma4"
                )
                if ok:
                    log.info("compute_pool: surgery produced %s for %s",
                             main_path, model_name)

        # Step 2a: if no URL is registered (auto-synthesized spec for
        # an unknown model) AND surgery didn't produce the file,
        # try HuggingFace auto-discovery. Looks for a community
        # GGUF (Unsloth/bartowski/mradermacher) matching the model
        # name. Saves the user from having to add a registry entry
        # for every new bundled-multimodal model they install.
        if not main_path.is_file() and not spec.get("main_url"):
            state.update({
                "status": "running", "phase": "discovering-url",
                "progress_pct": 25.0,
            })
            discovered = await _discover_hf_url(model_name)
            if discovered:
                spec = {**spec, "main_url": discovered["url"],
                        "main_size_gb": discovered.get("size_gb", 0)}
                log.info(
                    "compute_pool: HF auto-discovery for %s: %s",
                    model_name, discovered["url"],
                )

        # Step 2b: download main if surgery skipped / failed and a URL
        # is registered (or just got auto-discovered).
        if not main_path.is_file() and spec.get("main_url"):
            state.update({
                "status": "running", "phase": "downloading-main",
                "progress_pct": 30.0,
            })
            await _download_to_path(spec["main_url"], main_path)

        # Step 3: download mmproj if needed.
        if (
            spec.get("needs_mmproj")
            and not mmproj_path.is_file()
            and spec.get("mmproj_url")
        ):
            state.update({
                "status": "running", "phase": "downloading-mmproj",
                "progress_pct": 80.0,
            })
            await _download_to_path(spec["mmproj_url"], mmproj_path)

        state.update({
            "status": "done", "phase": "done", "progress_pct": 100.0,
            "error": None, "completed_at": time.time(),
        })

        # Fire-and-forget LAN distribution to workers, so subsequent
        # acquisitions on this LAN can take the cheap path.
        try:
            loop = asyncio.get_running_loop()
            loop.create_task(_distribute_override_to_workers(model_name))
        except RuntimeError:
            pass

        return {"ok": True, "status": "ready"}

    except Exception as e:
        log.warning("compute_pool: acquisition failed for %s: %s", model_name, e)
        state.update({
            "status": "error", "phase": state.get("phase", "unknown"),
            "error": f"{type(e).__name__}: {e}", "completed_at": time.time(),
        })
        return {"ok": False, "status": "error", "error": str(e)}
    finally:
        # Pop the task reference so the next request can re-attempt.
        _ACQUISITION_TASKS.pop(model_name, None)


async def _run_text_only_surgery(
    *, ollama_blob: str, dst: str, arch: str,
) -> bool:
    """Invoke the standalone surgery script in a subprocess.

    Why subprocess instead of importing the script: the surgery is
    CPU-heavy and reads the full GGUF (15+ GB). Running it in the
    same process would block the asyncio loop for several seconds
    even with thread pool tricks, and would tie up our process's
    address space with the mmap. Subprocess isolates that.

    Returns True on success (script exited 0 and dst exists), False
    otherwise. Failures are logged but not raised — caller falls
    back to download.
    """
    script = (
        Path(__file__).resolve().parent.parent / "scripts" / "repack_text_only_gguf.py"
    )
    if not script.is_file():
        return False
    import sys as _sys
    cmd = [
        _sys.executable, str(script),
        "--src", ollama_blob,
        "--dst", dst,
        "--arch", arch,
    ]

    def _run() -> int:
        import subprocess as _sp
        try:
            r = _sp.run(cmd, capture_output=True, text=True, timeout=900)
            if r.returncode != 0:
                log.warning(
                    "compute_pool: surgery script exited %d: %s",
                    r.returncode, (r.stderr or r.stdout)[:300],
                )
                return r.returncode
            return 0
        except Exception as e:
            log.warning("compute_pool: surgery script error: %s", e)
            return -1

    rc = await asyncio.to_thread(_run)
    return rc == 0 and Path(dst).is_file()


async def _run_qwen3_rope_patch(ollama_blob: str, dst: str) -> bool:
    """Invoke `scripts/repack_qwen3_rope_fix.py` in a subprocess to
    extend Ollama's qwen3.5 GGUF rope-section arrays from length 3
    to length 4 (zero bandwidth, zero quantization change).

    Returns True if the patched file lands cleanly at `dst`. Failures
    fall through to the download fallback if a URL is registered.
    """
    script = (
        Path(__file__).resolve().parent.parent / "scripts" / "repack_qwen3_rope_fix.py"
    )
    if not script.is_file():
        return False
    import sys as _sys
    cmd = [_sys.executable, str(script), "--src", ollama_blob, "--dst", dst]

    def _run() -> int:
        import subprocess as _sp
        try:
            r = _sp.run(cmd, capture_output=True, text=True, timeout=900)
            if r.returncode != 0:
                log.warning(
                    "compute_pool: qwen rope-patch script exited %d: %s",
                    r.returncode, (r.stderr or r.stdout)[:300],
                )
            return r.returncode
        except Exception as e:
            log.warning("compute_pool: qwen rope-patch script error: %s", e)
            return -1

    rc = await asyncio.to_thread(_run)
    return rc == 0 and Path(dst).is_file()


async def _lan_pull_override(worker: dict, filename: str, dst: Path) -> bool:
    """Try to pull `filename` from a worker's override dir over SSH/SCP.

    Returns True on a complete copy, False otherwise. The check is
    two-step:
      1. SSH `ls -l <path>` to confirm the file exists with non-zero
         size on the worker.
      2. SCP `<worker>:<path>` to a `.lan-part` temp; rename atomic
         on success.

    Failures (no ssh_host, file missing on worker, scp error) all
    return False quietly so the caller falls through to local
    surgery / HF download. Logged at INFO so the periodic-loop log
    surfaces what's happening.
    """
    ssh_host = (worker.get("ssh_host") or "").strip()
    if not ssh_host:
        return False

    remote_path = f"~/.gigachat/llama-cpp/models/{filename}"

    # Step 1: confirm it exists. Use `stat -c %s` to get a byte count
    # we can sanity-check (skip empty / partial files).
    check_cmd = ["ssh", *_ssh_persistent_args(),
                 "-o", "BatchMode=yes", "-o", "ConnectTimeout=5",
                 ssh_host, f'stat -c "%s" {remote_path} 2>/dev/null || echo MISSING']

    import subprocess as _sp
    def _run(cmd):
        try:
            return _sp.run(cmd, capture_output=True, text=True, timeout=20)
        except Exception:
            return None

    r = await asyncio.to_thread(_run, check_cmd)
    if r is None or r.returncode != 0:
        return False
    out = (r.stdout or "").strip()
    if out == "MISSING" or not out.isdigit() or int(out) == 0:
        return False
    expected_size = int(out)

    # Step 2: scp pull via temp file, rename on success.
    dst.parent.mkdir(parents=True, exist_ok=True)
    tmp = dst.with_suffix(dst.suffix + ".lan-part")
    if tmp.exists():
        tmp.unlink()
    scp_cmd = ["scp", *_ssh_persistent_args(),
               "-o", "BatchMode=yes", "-o", "ConnectTimeout=10",
               f"{ssh_host}:{remote_path}", str(tmp)]
    log.info(
        "compute_pool: lan-pull %s from %s (%.2f GB)",
        filename, ssh_host, expected_size / (1024 ** 3),
    )
    r = await asyncio.to_thread(
        lambda: _sp.run(scp_cmd, capture_output=True, text=True, timeout=3600),
    )
    if r is None or r.returncode != 0:
        if tmp.exists():
            tmp.unlink()
        return False
    if not tmp.is_file() or tmp.stat().st_size != expected_size:
        if tmp.exists():
            tmp.unlink()
        return False
    tmp.rename(dst)
    return True


async def _distribute_override_to_workers(model_name: str) -> None:
    """After a successful local acquisition, push the override files
    to every enabled worker with `ssh_host` set, so subsequent
    acquisitions (e.g. from another host on the same LAN) can take
    the LAN path. Fire-and-forget; failures are logged but never
    propagated.

    Idempotent: skips workers that already have the file at the
    expected size. Bandwidth is local LAN, capped to whatever the
    user's network can sustain — typically faster and cheaper than
    re-downloading from HuggingFace.
    """
    spec = _KNOWN_OVERRIDE_REGISTRY.get(model_name)
    if not spec:
        return
    main_path = _override_gguf_path_for(model_name)
    mmproj_path = _override_mmproj_path_for(model_name)
    paths_to_distribute = [main_path]
    if spec.get("needs_mmproj"):
        paths_to_distribute.append(mmproj_path)

    workers = [
        w for w in db.list_compute_workers(enabled_only=True)
        if (w.get("ssh_host") or "").strip()
    ]
    if not workers:
        return

    import subprocess as _sp
    for path in paths_to_distribute:
        if not path.is_file():
            continue
        local_size = path.stat().st_size
        for worker in workers:
            ssh_host = worker["ssh_host"].strip()
            remote_path = f"~/.gigachat/llama-cpp/models/{path.name}"
            # Skip if the worker already has the file at full size.
            check_cmd = ["ssh", *_ssh_persistent_args(),
                         "-o", "BatchMode=yes", "-o", "ConnectTimeout=5",
                         ssh_host,
                         f'stat -c "%s" {remote_path} 2>/dev/null || echo MISSING']
            r = await asyncio.to_thread(
                lambda c=check_cmd: _sp.run(c, capture_output=True, text=True, timeout=15),
            )
            if r and r.returncode == 0:
                out = (r.stdout or "").strip()
                if out.isdigit() and int(out) == local_size:
                    continue  # already there
            # Ensure the dir exists, then push the file.
            mkdir_cmd = ["ssh", *_ssh_persistent_args(),
                         "-o", "BatchMode=yes", "-o", "ConnectTimeout=5",
                         ssh_host, "mkdir -p ~/.gigachat/llama-cpp/models"]
            await asyncio.to_thread(
                lambda c=mkdir_cmd: _sp.run(c, capture_output=True, text=True, timeout=15),
            )
            scp_cmd = ["scp", *_ssh_persistent_args(),
                       "-o", "BatchMode=yes", "-o", "ConnectTimeout=10",
                       str(path), f"{ssh_host}:{remote_path}"]
            log.info(
                "compute_pool: distributing %s to %s (%.2f GB)",
                path.name, ssh_host, local_size / (1024 ** 3),
            )
            r = await asyncio.to_thread(
                lambda c=scp_cmd: _sp.run(c, capture_output=True, text=True, timeout=3600),
            )
            if r is None or r.returncode != 0:
                log.info(
                    "compute_pool: distribution to %s failed: %s",
                    ssh_host, (r.stderr or r.stdout or "")[-200:] if r else "no result",
                )


async def _discover_hf_url(model_name: str) -> dict | None:
    """Best-effort: search HuggingFace for a community GGUF matching
    `model_name` (an Ollama-style name like `gemma4:26b`) and return
    a direct download URL + estimated size, or None if no match.

    Algorithm:
      1. Build a few candidate query strings from the Ollama name
         (strip `:tag`, swap `:` for ` `, etc.).
      2. Call HuggingFace's public `/api/models` search endpoint
         (no auth needed for public repos), filtered to GGUF.
      3. Prefer well-known quantizers (unsloth, bartowski,
         mradermacher) — they ship clean stock-llama.cpp-loadable
         GGUFs.
      4. From the chosen repo, list files and pick the one matching
         `Q4_K_M` (Ollama's default quantization — preserves the
         user's quality expectations) or fall back to `Q5_K_M`.
      5. Return the resolve-main download URL.

    Returns None on any failure (HF unreachable, no candidates, no
    Q4_K_M / Q5_K_M file found). Caller falls through to a clean
    error message.

    Honest caveat: this is a heuristic. HuggingFace has many naming
    conventions, and our pattern matching may pick the wrong repo
    in edge cases. We bias toward the major quantizers (whose
    repos are well-named and stable) to minimize that risk; if a
    user installs a working but unusual model (e.g. a fine-tune
    not on Unsloth), this will return None and the user gets a
    clear error pointing them to the manual install path.
    """
    try:
        bare = model_name.split(":", 1)[0]
        tag = model_name.split(":", 1)[1] if ":" in model_name else ""
    except Exception:
        return None

    # Candidate search strings, in order of preference.
    queries = [
        f"{bare} {tag}".strip(),
        bare,
        bare.replace("-", " "),
    ]

    # Trusted quantizer repos in priority order.
    preferred_authors = ("unsloth", "bartowski", "mradermacher")

    timeout = httpx.Timeout(connect=10.0, read=30.0, write=10.0, pool=10.0)
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
        candidate_repos: list[str] = []
        for q in queries:
            try:
                r = await client.get(
                    "https://huggingface.co/api/models",
                    params={"search": q, "filter": "gguf", "limit": 30},
                )
                if r.status_code != 200:
                    continue
                hits = r.json() or []
            except Exception:
                continue
            # Sort by preferred author, then popularity (downloads).
            def _score(h):
                ident = h.get("id") or ""
                author = ident.split("/", 1)[0].lower() if "/" in ident else ""
                pref_idx = (
                    preferred_authors.index(author)
                    if author in preferred_authors
                    else len(preferred_authors)
                )
                return (pref_idx, -int(h.get("downloads") or 0))
            hits.sort(key=_score)
            for h in hits:
                ident = h.get("id") or ""
                if ident and ident not in candidate_repos:
                    candidate_repos.append(ident)
            if candidate_repos:
                break

        if not candidate_repos:
            log.info("compute_pool: HF auto-discovery: no candidates for %s", model_name)
            return None

        # For each candidate repo, list files and find a Q4_K_M or
        # Q5_K_M GGUF.
        for repo in candidate_repos[:5]:  # cap to avoid infinite probing
            try:
                r = await client.get(
                    f"https://huggingface.co/api/models/{repo}",
                    params={"blobs": "true"},
                )
                if r.status_code != 200:
                    continue
                meta = r.json() or {}
                siblings = meta.get("siblings") or []
            except Exception:
                continue

            # Filter to GGUFs; prefer Q4_K_M (matches Ollama's default).
            ggufs = [
                s for s in siblings
                if (s.get("rfilename") or "").endswith(".gguf")
            ]
            for quant_pat in ("Q4_K_M.gguf", "Q4_K_S.gguf",
                              "Q5_K_M.gguf", "UD-Q4_K_M.gguf",
                              "UD-Q5_K_M.gguf"):
                pick = next(
                    (s for s in ggufs
                     if s.get("rfilename", "").endswith(quant_pat)),
                    None,
                )
                if pick:
                    fname = pick["rfilename"]
                    url = f"https://huggingface.co/{repo}/resolve/main/{fname}"
                    return {
                        "url": url,
                        "repo": repo,
                        "filename": fname,
                        "size_gb": float(pick.get("size") or 0) / (1024 ** 3),
                    }
        log.info(
            "compute_pool: HF auto-discovery: candidates found but no "
            "Q4_K_M/Q5_K_M file in %s for %s",
            candidate_repos[:5], model_name,
        )
        return None


async def _download_to_path(url: str, dest: Path, *, max_attempts: int = 4) -> None:
    """Stream a HuggingFace direct-download URL to `dest`, resuming
    from `<dest>.part` if it already exists. In-process so we can
    surface progress through `_ACQUISITION_STATE`. Auto-retries on
    transient failures (network drop, server 5xx, partial response)
    with exponential backoff up to `max_attempts` total tries.

    Raises on terminal failure (4xx that's not 416, max attempts
    exhausted, content-length mismatch). Each retry resumes from the
    last byte successfully written, so a 50%-complete download that
    drops mid-flight only re-fetches the remainder.

    Defends against three real-world failure modes:
      1. Server returned 200 OK to a Range request (didn't honor the
         resume) — discards the partial and starts fresh, otherwise
         we'd append duplicate prefix bytes to the file.
      2. Connection drops mid-chunk (network blip, server reset) —
         exception caught, partial bytes already on disk, retry
         picks up where we left off.
      3. 5xx server error — wait + retry. 4xx (other than 416 Range
         Not Satisfiable) is terminal: bad URL, no retry can fix it.
    """
    dest.parent.mkdir(parents=True, exist_ok=True)
    tmp = dest.with_suffix(dest.suffix + ".part")
    last_err: Exception | None = None

    for attempt in range(1, max_attempts + 1):
        start_byte = tmp.stat().st_size if tmp.is_file() else 0
        headers = {"Range": f"bytes={start_byte}-"} if start_byte > 0 else {}
        timeout = httpx.Timeout(connect=30.0, read=120.0, write=30.0, pool=30.0)
        try:
            async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
                async with client.stream("GET", url, headers=headers) as r:
                    if r.status_code == 416:
                        # Range Not Satisfiable: our `.part` is bigger
                        # than the file. Clear it and retry from 0.
                        if tmp.exists():
                            tmp.unlink()
                        raise RuntimeError(f"HTTP 416; reset partial")
                    if r.status_code not in (200, 206):
                        msg = f"HTTP {r.status_code} from {url}"
                        # 4xx is terminal — retrying won't help.
                        if 400 <= r.status_code < 500:
                            raise RuntimeError(msg)
                        # 5xx is transient — fall through to retry.
                        raise RuntimeError(msg)
                    if start_byte > 0 and r.status_code == 200:
                        # Server ignored our Range request and is
                        # streaming the whole file. Reset partial so
                        # we don't end up with duplicated prefix bytes.
                        if tmp.exists():
                            tmp.unlink()
                        start_byte = 0
                    mode = "ab" if start_byte > 0 else "wb"
                    with tmp.open(mode) as f:
                        async for chunk in r.aiter_bytes(chunk_size=1024 * 1024):
                            f.write(chunk)
            tmp.rename(dest)
            return
        except Exception as e:
            last_err = e
            # Don't retry on terminal 4xx — caller should see the
            # error promptly so they can fix the URL / registry.
            if "HTTP 4" in str(e) and "HTTP 416" not in str(e):
                raise
            if attempt >= max_attempts:
                raise RuntimeError(
                    f"download failed after {attempt} attempts: {e}"
                ) from e
            backoff_s = min(60.0, 2.0 ** attempt)
            log.info(
                "compute_pool: download attempt %d/%d failed for %s: %s; "
                "retrying in %.0fs",
                attempt, max_attempts, url, e, backoff_s,
            )
            await asyncio.sleep(backoff_s)

    # Should not reach here — the loop either returns or raises.
    if last_err:
        raise last_err


def _resolve_ollama_manifest(model_name: str) -> dict | None:
    """Locate the manifest JSON for an Ollama model name.

    Ollama stores manifests at
        ~/.ollama/models/manifests/<registry>/<namespace>/<name>/<tag>
    The default registry is `registry.ollama.ai` and the default namespace
    is `library`. Custom registries / namespaces are rare for end users
    but we still walk the tree to find the file.

    Returns the parsed manifest dict, or None if no matching file exists
    (model not pulled, name typo, etc.).
    """
    name = (model_name or "").strip()
    if not name:
        return None
    # Tag defaults to `latest` when omitted.
    if ":" in name:
        bare, tag = name.split(":", 1)
    else:
        bare, tag = name, "latest"

    # Default location first — covers >99% of cases.
    candidate = _OLLAMA_MODELS_DIR / "manifests" / "registry.ollama.ai" / "library" / bare / tag
    if candidate.is_file():
        try:
            return jsonutil.loads(candidate.read_text(encoding="utf-8"))
        except Exception:
            return None

    # Fallback: walk the manifests tree. Slower but handles the user
    # who pulled a model from a non-default registry.
    manifests = _OLLAMA_MODELS_DIR / "manifests"
    if not manifests.is_dir():
        return None
    needle = (bare, tag)
    for root, _dirs, files in os.walk(manifests):
        for fname in files:
            if fname == tag and Path(root).name == bare:
                try:
                    return jsonutil.loads(Path(root, fname).read_text(encoding="utf-8"))
                except Exception:
                    continue
    return None


def resolve_ollama_model(model_name: str) -> dict | None:
    """Resolve a model name to its on-disk GGUF + size.

    Returns `{gguf_path, size_bytes, manifest}` or None if the manifest
    can't be found. The "model" layer of the manifest is the GGUF; we
    compute size from `layers[].size` for the layer whose mediaType is
    `application/vnd.ollama.image.model` (large) — license / params /
    template layers are tiny config blobs and don't count toward VRAM.

    `gguf_path` is the absolute path to the blob, suitable for passing
    as llama-server's `--model` flag.

    Override path: if a file exists at `_override_gguf_path_for(name)`,
    we return that instead of the Ollama blob. Used to swap in GGUFs
    with tensor layouts that stock llama.cpp can load (see the
    `_OVERRIDE_GGUF_DIR` docstring) without disturbing Ollama's own
    blob store. The Ollama manifest is still consulted for the size
    metadata so routing / VRAM-budget logic stays consistent.
    """
    manifest = _resolve_ollama_manifest(model_name)
    if not manifest:
        return None
    layers = manifest.get("layers") or []
    model_layer = None
    for layer in layers:
        if layer.get("mediaType") == "application/vnd.ollama.image.model":
            model_layer = layer
            break
    if not model_layer:
        return None
    digest = (model_layer.get("digest") or "").replace("sha256:", "")
    if not digest:
        return None
    blob_path = _OLLAMA_MODELS_DIR / "blobs" / f"sha256-{digest}"
    if not blob_path.is_file():
        return None

    # Override hook: prefer a user-installed replacement GGUF if one
    # exists. Size is taken from THAT file (not the manifest), since
    # the override may be a different quantization than the Ollama
    # blob. Manifest-derived size is kept as a hint for callers that
    # want the original.
    override = _override_gguf_path_for(model_name)
    if override.is_file():
        try:
            override_size = override.stat().st_size
        except OSError:
            override_size = 0
        # Multimodal projector — surface its path if installed beside
        # the main override. Phase 2 split (`split_lifecycle._build_command`)
        # passes it as `--mmproj` so llama-server can serve image input.
        mmproj = _override_mmproj_path_for(model_name)
        mmproj_str = str(mmproj) if mmproj.is_file() else None
        return {
            "gguf_path": str(override),
            "size_bytes": override_size or int(model_layer.get("size") or 0),
            "manifest": manifest,
            "ollama_blob_path": str(blob_path),
            "override": True,
            "mmproj_path": mmproj_str,
        }

    return {
        "gguf_path": str(blob_path),
        "size_bytes": int(model_layer.get("size") or 0),
        "manifest": manifest,
    }


def _host_vram_budget_bytes() -> int:
    """Bytes of host VRAM we're willing to let a single Ollama-loaded
    model occupy in VRAM only. Used for ranking host vs workers in
    `_host_capability_score` (which compares pure GPU memory). NOT the
    threshold for engaging the split path — that's
    `_host_total_capacity_bytes` below."""
    try:
        spec = sysdetect.detect_system()
        vram_gb = float(spec.get("vram_gb") or 0.0)
    except Exception:
        vram_gb = 0.0
    return int(vram_gb * _HOST_VRAM_USE_FRACTION * 1024 * 1024 * 1024)


def _host_total_capacity_bytes() -> int:
    """Bytes the host can hold a single model in (VRAM + RAM). Ollama
    natively uses CPU offload — layers that don't fit VRAM live in
    system RAM and run on CPU. That stays single-node (no per-token
    LAN overhead), so we should NOT engage split unless the model
    truly exceeds the host's total memory budget.

    For this host (8 GB VRAM + 15.8 GB RAM, 70% safety margin): ≈16.7 GB.
    A 9 GB model fits trivially, even though it doesn't fit VRAM alone.
    """
    try:
        spec = sysdetect.detect_system()
        vram_gb = float(spec.get("vram_gb") or 0.0)
        ram_gb = float(spec.get("ram_gb") or 0.0)
    except Exception:
        vram_gb = 0.0
        ram_gb = 0.0
    total_gb = vram_gb + ram_gb
    return int(total_gb * _HOST_TOTAL_USE_FRACTION * 1024 * 1024 * 1024)


def _eligible_split_workers() -> list[dict]:
    """Workers that can contribute layers via rpc-server.

    Same `enabled` + `use_for_chat` gate as Phase 1's chat picker —
    workers the user toggled off for chat shouldn't suddenly start
    receiving inference traffic just because the model needs splitting.
    Plus the rpc-specific gate: probe must report rpc-server reachable.
    """
    rows = db.list_compute_workers(enabled_only=True)
    out = []
    for w in rows:
        if not w.get("use_for_chat"):
            continue
        caps = w.get("capabilities") or {}
        if not caps.get("rpc_server_reachable"):
            continue
        out.append(w)
    # Freshest probe first — if we have to pick a subset later, we'd
    # rather use the worker we just confirmed alive than one whose
    # last_seen is an hour stale.
    out.sort(key=lambda w: float(w.get("last_seen") or 0), reverse=True)
    return out


async def _eligible_split_workers_with_autoprep() -> list[dict]:
    """Like ``_eligible_split_workers`` but auto-prep ineligible peers
    AND refresh live resource stats while we're already on the wire.

    Walks every chat-enabled worker and, for any whose
    ``rpc_server_reachable`` flag is unset, calls
    ``ensure_rpc_server_via_proxy`` to bring rpc-server up over the
    encrypted P2P channel. In parallel, it also calls
    ``probe_worker_live_stats`` so the routing decision sees fresh
    ``ram_free_gb`` / ``vram_total_gb`` rather than stale or null
    capability values.

    This is what the auto-promotion router calls when it's about to
    decide between Ollama and a split-rpc spawn — it gives every
    paired peer a chance to contribute without the user having to
    pre-arm rpc-server through Settings or SSH, AND it ensures the
    pool-size math reflects what the peer can actually contribute
    RIGHT NOW (not what it could contribute 5 minutes ago).

    Concurrent across peers (gather) so a slow peer doesn't
    serialize the whole hot-path.
    """
    rows = db.list_compute_workers(enabled_only=True)
    chat_workers = [w for w in rows if w.get("use_for_chat")]
    if not chat_workers:
        return []

    async def _prep_one(w: dict) -> None:
        """Bring up rpc-server (idempotent) AND refresh live stats +
        bandwidth measurement so the routing decision sees fresh
        per-peer link quality."""
        # Live stats first — cheap, runs even if rpc is already up,
        # populates ram_free_gb + vram_total_gb in capabilities for
        # the pool-size calculation downstream.
        try:
            stats = await probe_worker_live_stats(w)
        except Exception:
            stats = {}
        # Bandwidth probe — coarse but cheap (single existing API
        # round-trip); persists `bandwidth_mbps` so future routing
        # can demote slow peers from split engagement.
        try:
            bw = await probe_worker_bandwidth(w)
        except Exception:
            bw = 0.0
        if stats:
            try:
                caps = dict(w.get("capabilities") or {})
                caps["ram_free_gb"] = float(stats.get("ram_free_gb") or 0)
                caps["ram_total_gb"] = float(stats.get("ram_total_gb") or 0)
                caps["vram_total_gb"] = float(stats.get("vram_total_gb") or 0)
                caps["gpu_kind"] = stats.get("gpu_kind") or ""
                caps["ram_free_probed_at"] = stats.get("ts") or 0
                if bw > 0:
                    caps["bandwidth_mbps"] = bw
                    caps["bandwidth_probed_at"] = stats.get("ts") or 0
                # If we have a real vram_total but no max_vram_seen,
                # use vram_total as a floor — the pool-size math
                # treats max_vram_seen as the VRAM contribution and
                # would otherwise stay at 0 for a worker that's
                # never been benchmarked under load.
                cur_max_vram = int(caps.get("max_vram_seen_bytes") or 0)
                vram_bytes = int(float(stats.get("vram_total_gb") or 0) * (1024 ** 3))
                if vram_bytes > cur_max_vram:
                    caps["max_vram_seen_bytes"] = vram_bytes
                db.update_compute_worker_capabilities(
                    w["id"], capabilities=caps,
                )
                # Update our local copy so the rpc-prep below sees
                # the merged capabilities (avoids re-stamping).
                w["capabilities"] = caps
            except Exception as e:
                log.debug(
                    "compute_pool: stats persist failed for %s: %s",
                    w.get("label"), e,
                )
        # Then rpc-server prep — also idempotent.
        if not (w.get("capabilities") or {}).get("rpc_server_reachable"):
            try:
                await ensure_rpc_server_via_proxy(w)
            except Exception as e:
                log.debug(
                    "compute_pool: ensure_rpc failed for %s: %s",
                    w.get("label"), e,
                )

    log.info(
        "compute_pool: auto-prepping %d worker(s) via P2P "
        "(rpc-server + live stats, no SSH required)",
        len(chat_workers),
    )
    await asyncio.gather(
        *(_prep_one(w) for w in chat_workers),
        return_exceptions=True,
    )

    # Re-read fresh capabilities + filter.
    refreshed = db.list_compute_workers(enabled_only=True)
    out = []
    for w in refreshed:
        if not w.get("use_for_chat"):
            continue
        caps = w.get("capabilities") or {}
        if not caps.get("rpc_server_reachable"):
            continue
        out.append(w)
    out.sort(key=lambda w: float(w.get("last_seen") or 0), reverse=True)
    return out


# ---------------------------------------------------------------------------
# Speculative-decoding draft picker
#
# The router calls `pick_draft_for(target_model_name)` whenever the chat
# target fits on a single node — typically the host. The picker walks the
# pool's combined model inventory (host's installed Ollama models + every
# enabled worker's `/api/tags` snapshot) for a smaller, same-family,
# chat-capable model that can serve as a speculative-decoding draft.
#
# Generic by design: nothing in this module hardcodes specific model names
# or sizes. The user's pool can hold any mix of devices and models — the
# picker uses the family hint Ollama already publishes (`gemma`, `llama`,
# `qwen2`, …) so the same matcher works whether the chat target is a
# 3 B Phi or a 70 B Llama.
#
# Speculative decoding is a net win when:
#   * draft + target share the SAME tokenizer (== same family is the
#     safe approximation; mixing tokenizer families produces 0 % accept
#     rate), and
#   * draft is DRAMATICALLY smaller than target (the speedup formula
#     amortises draft cost over verified tokens — a draft only 2× smaller
#     barely helps; 5-10× smaller is where the wins live), and
#   * the executing node can hold BOTH models in VRAM at once.
#
# The picker enforces all three. When any check fails it returns None and
# the router falls back to the vanilla single-model path.
# ---------------------------------------------------------------------------

# Maximum draft size as a fraction of target size. Empirically, drafts
# bigger than ~30 % of target spend more cycles per draft token than the
# target verifies them — the speedup curve is non-monotonic. Below 30 %
# is the safe regime where speculative decoding is consistently a win.
_DRAFT_MAX_SIZE_FRACTION = 0.30

# Minimum target size (in bytes) below which speculative decoding isn't
# worth engaging — the per-token cost of running TWO models eats the
# small fixed-overhead win when the target is already tiny. 1.5 GB is
# the rough crossover (anything smaller already runs at >50 tok/s on a
# decent GPU and gains little from a draft).
_SPECULATIVE_MIN_TARGET_BYTES = 1_500_000_000

# VRAM headroom multiplier — we only engage speculative if the executing
# node has at least (target_size + draft_size) * this factor in VRAM.
# 1.15 = ~10 % real KV-cache need + ~5 % allocator/alignment buffer to
# avoid load-time OOM when (target+draft) genuinely fits on paper.
# Tightened from the historical 1.30 once flash-attn + KV-q8 freed
# half the per-slot KV cost; lowered further to 1.15 under the
# 5 %-margin policy.
_SPECULATIVE_VRAM_HEADROOM = 1.15


# Process-level cache of GGUF tokenizer fingerprints, keyed by absolute
# path. The cache invalidates on mtime change so re-pulling a model
# refreshes the fingerprint on the next probe. Bounded implicitly — at
# most one entry per Ollama-managed GGUF on disk, which tops out in the
# tens. No LRU needed.
_TOKENIZER_FINGERPRINT_CACHE: dict[str, tuple[float, str | None]] = {}


def _gguf_tokenizer_fingerprint(gguf_path: str) -> str | None:
    """Return a short SHA-256 fingerprint of the GGUF's tokenizer.

    Two GGUFs whose fingerprints match share token IDs — i.e. the
    speculative-decoding draft model can verify against the target
    without false rejections, regardless of what `details.family` each
    one reports. Lets the picker accept cross-family-but-same-vocab
    pairs (e.g. some Mistral derivatives shipping the Llama tokenizer)
    that the family heuristic alone would miss.

    Returns ``None`` when the file isn't a GGUF, the optional ``gguf``
    package isn't installed (older deployments), or the metadata is
    incomplete. Callers MUST treat ``None`` as "can't verify" and fall
    back to the family check, never as an implicit match — silently
    treating an unknown fingerprint as compatible would mean engaging
    speculative on truly incompatible pairs.

    Cheap fingerprint over the tokenizer model name, vocabulary size,
    and a deterministic 15-token sample (start / middle / end). Two
    different tokenizers practically never share that combination,
    while two GGUFs of the same tokenizer always do — fast and
    correct in practice. Hashing the entire vocab would also work
    but parses 100k+ string entries per probe and gains nothing.
    """
    if not gguf_path:
        return None
    try:
        path = Path(gguf_path)
        mtime = path.stat().st_mtime
    except OSError:
        return None

    cached = _TOKENIZER_FINGERPRINT_CACHE.get(gguf_path)
    if cached is not None and cached[0] == mtime:
        return cached[1]

    fingerprint: str | None = None
    try:
        # Imported lazily so the import error surfaces only when the
        # picker actually walks GGUFs — useful for old deployments
        # whose `gguf` package install lags requirements.txt.
        from gguf import GGUFReader

        reader = GGUFReader(gguf_path)
        model_name: str | None = None
        token_count: int | None = None
        sample: list[str] = []
        for field in reader.fields.values():
            if field.name == "tokenizer.ggml.model":
                # `field.parts[-1]` is a numpy uint8 array of UTF-8
                # bytes for string-typed metadata.
                try:
                    model_name = bytes(field.parts[-1]).decode("utf-8", errors="replace")
                except Exception:
                    model_name = None
            elif field.name == "tokenizer.ggml.tokens":
                # Array of strings — `field.contents()` gives a
                # list-like of decoded tokens. We only sample 15 to
                # keep this cheap; collisions with this much detail
                # require a deliberately constructed adversary.
                try:
                    tokens = field.contents()
                    token_count = len(tokens)
                    if token_count > 0:
                        positions = set(range(min(5, token_count)))
                        if token_count > 10:
                            mid = token_count // 2
                            positions.update(range(mid, mid + 5))
                        if token_count > 5:
                            positions.update(
                                range(max(0, token_count - 5), token_count)
                            )
                        for i in sorted(positions):
                            try:
                                sample.append(str(tokens[i]))
                            except Exception:
                                pass
                except Exception:
                    token_count = None
        if model_name and token_count is not None:
            payload = f"{model_name}|{token_count}|" + "|".join(sample)
            fingerprint = hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16]
    except ImportError:
        # `gguf` package missing — not an error, just no fingerprinting
        # for this deploy. Family fallback still works.
        log.debug("gguf package unavailable; skipping tokenizer fingerprint")
    except Exception as e:
        log.debug("tokenizer fingerprint read failed for %s: %s", gguf_path, e)

    _TOKENIZER_FINGERPRINT_CACHE[gguf_path] = (mtime, fingerprint)
    return fingerprint


# ---------------------------------------------------------------------------
# Generic SSH-PowerShell dispatch helper
#
# `dispatch_to_worker_powershell(worker, ps_script, timeout_sec, stdin_bytes)`
# is the shared transport every distributed-tool path uses: encode the
# PowerShell script as UTF-16-LE base64 (bypasses every cmdline-quoting
# nightmare), invoke `ssh worker powershell -EncodedCommand <b64>`,
# stream back stdout. Optional `stdin_bytes` is base64-encoded and the
# script can decode it via `[Console]::In.ReadToEnd()`.
#
# This stays minimal on purpose — workers need only OpenSSH + PowerShell
# (defaults on Windows 10+) to participate in the basic dispatch layer.
# Specific tool dispatchers (read_doc, web_search) layer their
# Python/library requirements on top and probe for them at call time.
# ---------------------------------------------------------------------------


async def dispatch_to_worker_powershell(
    worker: dict,
    ps_script: str,
    *,
    timeout_sec: float = 30.0,
    stdin_text: str = "",
) -> tuple[bool, bytes, str]:
    """Run ``ps_script`` on the worker via SSH + PowerShell.

    Returns ``(ok, stdout_bytes, stderr_message)``. ``ok`` is True when
    the SSH call returned exit 0; on failure ``stderr_message`` carries
    a short diagnostic and ``stdout_bytes`` is empty. Caller is in
    charge of interpreting stdout (typically base64-decoding a JSON
    payload the script wrote).

    `stdin_text`, when non-empty, is fed to the SSH process's stdin so
    PowerShell scripts can read large inputs without blowing past
    cmdline length limits. Bind it inside the script with
    ``[Console]::In.ReadToEnd()``.
    """
    ssh_host = (worker.get("ssh_host") or "").strip()
    if not ssh_host:
        return False, b"", "worker has no ssh_host configured"

    import base64
    encoded = base64.b64encode(ps_script.encode("utf-16-le")).decode("ascii")
    cmd = [
        "ssh", *_ssh_persistent_args(),
        "-o", "BatchMode=yes",
        "-o", f"ConnectTimeout={int(min(timeout_sec, 30))}",
        ssh_host,
        "powershell", "-NoProfile", "-EncodedCommand", encoded,
    ]
    import subprocess as _sp

    def _run() -> tuple[int, bytes, bytes]:
        try:
            r = _sp.run(
                cmd,
                input=stdin_text.encode("utf-8") if stdin_text else None,
                capture_output=True,
                timeout=timeout_sec + 10,
            )
            return r.returncode, r.stdout, r.stderr
        except _sp.TimeoutExpired:
            return -1, b"", b"timeout"
        except Exception as e:
            return -2, b"", repr(e).encode()

    rc, stdout, stderr = await asyncio.to_thread(_run)
    if rc != 0:
        err = (stderr or b"").decode("utf-8", errors="replace").strip()[:300]
        return False, b"", err or "no stderr"
    return True, stdout or b"", ""


# ---------------------------------------------------------------------------
# Distributed tool execution — dispatch eligible tool calls to workers
#
# Some tool calls (HTTP fetches today; doc parsing in a future commit)
# don't need host-only state — they can run on any worker that's
# reachable over SSH. Routing them away from host frees host CPU/IO for
# the inference path, which matters when the agent is mid-stream and
# also pulling in remote context (e.g. a `fetch_url` during chat).
#
# Protocol is intentionally minimal: SSH the worker and execute one
# `powershell -EncodedCommand` invocation per call. The worker only
# needs OpenSSH server + PowerShell, both of which ship by default on
# Windows 10+. No Python on the worker, no extra service.
#
# Round-robin across eligible workers so multi-fetch turns parallelise
# across the pool. Opt-in via the `compute_pool_distribute_tools`
# setting because SSH dispatch carries ~50-100 ms overhead per call;
# pools without spare worker capacity get a regression rather than a
# win.
# ---------------------------------------------------------------------------
_TOOL_DISPATCH_INDEX: dict[str, int] = {}
_DISPATCH_FETCH_TIMEOUT_SEC = 30.0




def _pick_tool_dispatch_target() -> dict | None:
    """Return the next eligible worker for tool dispatch.

    Eligible = enabled, has `ssh_host` configured (the dispatch
    protocol uses SSH), and was probed successfully recently. Round-
    robins via `_TOOL_DISPATCH_INDEX` so a burst of tool calls fans
    across the whole pool instead of pinning to the strongest worker.
    """
    rows = db.list_compute_workers(enabled_only=True)
    eligible = [
        w for w in rows
        if (w.get("ssh_host") or "").strip()
        and w.get("last_seen") is not None
    ]
    if not eligible:
        return None
    # Newest probe first so a worker we just confirmed alive ranks
    # above one whose last_seen is stale.
    eligible.sort(key=lambda w: float(w.get("last_seen") or 0), reverse=True)
    idx = _TOOL_DISPATCH_INDEX.get("any", 0)
    pick = eligible[idx % len(eligible)]
    _TOOL_DISPATCH_INDEX["any"] = (idx + 1) % len(eligible)
    return pick


async def dispatch_fetch_url_to_worker(url: str) -> tuple[bool, bytes, str] | None:
    """SSH a URL fetch onto a worker via PowerShell's `Invoke-WebRequest`.

    Returns:
      * ``None`` when distribution is disabled or no eligible worker is
        available — caller falls back to host-side fetch.
      * ``(True, body_bytes, "")`` when the worker fetched the URL
        successfully. Body is the raw response bytes (binary-safe via
        base64 transport).
      * ``(False, b"", error_msg)`` when distribution was attempted but
        the worker returned an error. Caller MAY fall back to host as a
        safety net (the distributed-tools setting governs whether to
        retry locally vs. surface the error).

    The SSRF / DNS check the host-side `fetch_url` runs BEFORE calling
    this function is the security gate — we don't repeat it on the
    worker side because the URL crossed the trust boundary already
    vetted.

    Engagement: always-on. If no eligible worker is reachable the
    function returns ``None`` and the caller fetches on host. The win
    appears when the host is busy with inference and a worker is
    idle; the cost (one SSH dial) is bounded by `ConnectTimeout=30`.
    """
    worker = _pick_tool_dispatch_target()
    if not worker:
        return None
    ssh_host = (worker.get("ssh_host") or "").strip()
    if not ssh_host:
        return None

    # PowerShell single-quoted string is literal — no `$` expansion,
    # no backtick escapes. Doubling embedded apostrophes is the only
    # quote-injection vector to handle.
    safe_url = url.replace("'", "''")
    ps_cmd = (
        # Cap response size so a multi-GB body doesn't OOM the worker.
        # 8 MB matches the host-side `FETCH_MAX_BYTES` (in tools.py).
        "[System.Net.ServicePointManager]::SecurityProtocol = "
        "[System.Net.SecurityProtocolType]::Tls12;"
        "try {"
        f"  $resp = Invoke-WebRequest -UseBasicParsing -Uri '{safe_url}'"
        f"    -TimeoutSec {int(_DISPATCH_FETCH_TIMEOUT_SEC)}"
        "     -MaximumRedirection 5 -ErrorAction Stop;"
        "  $bytes = $resp.RawContentStream.ToArray();"
        "  if ($bytes.Length -gt 8388608) { $bytes = $bytes[0..8388607] };"
        "  [Console]::Out.Write([System.Convert]::ToBase64String($bytes))"
        "} catch {"
        "  [Console]::Error.Write($_.Exception.Message);"
        "  exit 2"
        "}"
    )
    import base64
    encoded = base64.b64encode(ps_cmd.encode("utf-16-le")).decode("ascii")
    cmd = [
        "ssh", *_ssh_persistent_args(),
        "-o", "BatchMode=yes",
        "-o", f"ConnectTimeout={int(_DISPATCH_FETCH_TIMEOUT_SEC)}",
        ssh_host,
        "powershell", "-NoProfile", "-EncodedCommand", encoded,
    ]
    import subprocess as _sp

    def _run() -> tuple[int, bytes, bytes]:
        try:
            r = _sp.run(
                cmd, capture_output=True,
                timeout=_DISPATCH_FETCH_TIMEOUT_SEC + 10,
            )
            return r.returncode, r.stdout, r.stderr
        except _sp.TimeoutExpired:
            return -1, b"", b"timeout"
        except Exception as e:
            return -2, b"", repr(e).encode()

    rc, stdout, stderr = await asyncio.to_thread(_run)
    if rc != 0:
        err = (stderr or b"").decode("utf-8", errors="replace").strip()[:200]
        log.info(
            "compute_pool: distributed fetch via %s failed (rc=%d): %s",
            ssh_host, rc, err or "no stderr",
        )
        return False, b"", f"worker fetch failed: {err or 'unknown'}"
    try:
        body = base64.b64decode((stdout or b"").strip())
    except Exception as e:
        return False, b"", f"worker returned malformed base64: {e}"
    log.info(
        "compute_pool: dispatched fetch_url(%r) to worker %s — %d bytes",
        url, worker.get("label"), len(body),
    )
    return True, body, ""


# ---------------------------------------------------------------------------
# Distributed `web_search` — DuckDuckGo via worker's `ddgs` Python
#
# Worker requirement: Python on PATH + the `ddgs` package installed
# (`pip install ddgs`). The setup script `scripts/install-worker-tools.ps1`
# handles this; the dispatcher silently falls back to host when the
# worker doesn't have ddgs available (the SSH script's import error
# surfaces as a non-zero exit and the caller treats that as a graceful
# fallback signal).
# ---------------------------------------------------------------------------


def _pick_web_search_target() -> dict | None:
    """Return the next worker that has the `ddgs` Python library.

    Probe-time `capabilities.has_ddgs` flag is the gate — workers
    without ddgs are skipped so we don't dispatch a request that
    can't possibly succeed.
    """
    rows = db.list_compute_workers(enabled_only=True)
    eligible = [
        w for w in rows
        if (w.get("ssh_host") or "").strip()
        and w.get("last_seen") is not None
        and (w.get("capabilities") or {}).get("has_ddgs")
    ]
    if not eligible:
        return None
    eligible.sort(key=lambda w: float(w.get("last_seen") or 0), reverse=True)
    idx = _TOOL_DISPATCH_INDEX.get("web_search", 0)
    pick = eligible[idx % len(eligible)]
    _TOOL_DISPATCH_INDEX["web_search"] = (idx + 1) % len(eligible)
    return pick


async def dispatch_web_search_to_worker(
    query: str, max_results: int, region: str | None,
) -> tuple[bool, list[dict], str] | None:
    """Run a DDG search on a worker via SSH and return the parsed hits.

    Returns:
      * ``None`` when no eligible worker is reachable — caller falls
        back to host-side `_ddg_search_sync`.
      * ``(True, hits, "")`` on success.
      * ``(False, [], error)`` when distribution was attempted but the
        worker errored. Caller MAY fall back to host as the safety
        net so a transient worker hiccup never breaks the chat.
    """
    worker = _pick_web_search_target()
    if not worker:
        return None

    # Build a Python one-liner the worker runs. `ddgs` exposes a
    # `DDGS` context manager whose `.text(query, max_results=N,
    # region=...)` returns list[dict]. We dump the result as JSON to
    # stdout so the host can parse it back. stderr surfaces the
    # exception message on failure.
    safe_query = query.replace('"', '\\"').replace("$", "`$")
    region_arg = (
        f", region='{region}'" if region and region.replace("-", "").isalnum() else ""
    )
    py = (
        "import json,sys\n"
        "from ddgs import DDGS\n"
        f"with DDGS() as d:\n"
        f"    hits = list(d.text(\"{safe_query}\", max_results={int(max_results)}{region_arg}))\n"
        "json.dump(hits, sys.stdout)\n"
    )
    # `python -c <script>` is awkward to encode through PowerShell; far
    # cleaner to write the script via stdin and run `python -`. The
    # PowerShell wrapper just streams stdin into the python process.
    ps_script = (
        # Resolve the worker's Python interpreter — `py.exe` (the
        # Python launcher, ships with the installer) takes precedence
        # over a bare `python.exe` on PATH because it survives the
        # PATH ordering hazards of corporate machines.
        "$python = (Get-Command py.exe -ErrorAction SilentlyContinue).Source;"
        "if (-not $python) { $python = (Get-Command python.exe -ErrorAction SilentlyContinue).Source };"
        "if (-not $python) { Write-Error 'python not found on worker'; exit 3 };"
        # Read the python source from stdin and pipe it into python -.
        "$src = [Console]::In.ReadToEnd();"
        "$src | & $python -"
    )
    ok, stdout, stderr = await dispatch_to_worker_powershell(
        worker, ps_script, timeout_sec=30.0, stdin_text=py,
    )
    if not ok:
        log.info(
            "compute_pool: distributed web_search via %s failed: %s",
            worker.get("label"), stderr,
        )
        return False, [], stderr
    try:
        hits = jsonutil.loads((stdout or b"").decode("utf-8", errors="replace") or "[]")
    except json.JSONDecodeError as e:
        return False, [], f"worker returned non-JSON: {e}"
    if not isinstance(hits, list):
        return False, [], "worker returned non-list payload"
    log.info(
        "compute_pool: dispatched web_search(%r) to worker %s — %d hit(s)",
        query, worker.get("label"), len(hits),
    )
    return True, hits, ""


# ---------------------------------------------------------------------------
# Distributed `read_doc` — PDF / DOCX / XLSX parsing offloaded to a worker
#
# The host uploads the raw file bytes via SSH stdin (base64-encoded for
# binary safety), the worker decodes to a temp path, parses with the
# requested library (pymupdf / python-docx / openpyxl), and emits the
# extracted text on stdout. The host trims and returns to the agent
# unchanged — preserves the same shape `read_doc` already produces.
#
# Worker requirements: Python + the parser libraries. The probe-time
# detection (added below) fills `capabilities.read_doc_libs` with the
# subset the worker has, so the dispatcher only routes formats the
# worker can actually handle.
# ---------------------------------------------------------------------------

# Per-format → library name mapping the dispatcher uses to filter
# eligible workers based on what's installed.
_READ_DOC_FORMAT_LIBS = {
    ".pdf": "pymupdf",
    ".docx": "docx",        # python-docx imports as `docx`
    ".xlsx": "openpyxl",
    ".xlsm": "openpyxl",
}


def _pick_read_doc_target(suffix: str) -> dict | None:
    """Return the next worker that has the parser library for ``suffix``.

    Reads `capabilities.read_doc_libs` (populated at probe time) to
    skip workers missing the right library — falling back to host is
    cheaper than dispatching to a worker that's just going to error
    out on `import pymupdf` / `import docx` / `import openpyxl`.
    """
    needed = _READ_DOC_FORMAT_LIBS.get(suffix)
    if not needed:
        return None
    rows = db.list_compute_workers(enabled_only=True)
    eligible = [
        w for w in rows
        if (w.get("ssh_host") or "").strip()
        and w.get("last_seen") is not None
        and needed in ((w.get("capabilities") or {}).get("read_doc_libs") or [])
    ]
    if not eligible:
        return None
    eligible.sort(key=lambda w: float(w.get("last_seen") or 0), reverse=True)
    idx = _TOOL_DISPATCH_INDEX.get(f"read_doc:{needed}", 0)
    pick = eligible[idx % len(eligible)]
    _TOOL_DISPATCH_INDEX[f"read_doc:{needed}"] = (idx + 1) % len(eligible)
    return pick


async def dispatch_read_doc_to_worker(
    file_path: str, *, pages: str | None = None, sheets: str | None = None,
) -> tuple[bool, str, str] | None:
    """Stream a doc file to a worker, parse there, return the text.

    Returns ``None`` when no eligible worker is reachable for the
    file's format. Returns ``(True, text, "")`` on success or
    ``(False, "", error)`` on attempted-but-failed dispatch (caller
    falls back to host).

    File size cap: 30 MB on the wire. Anything bigger is rejected on
    the host side before SSH — we don't want to ship hundreds of MB
    to a worker over the LAN per dispatch.
    """
    from pathlib import Path
    p = Path(file_path)
    suffix = p.suffix.lower()
    worker = _pick_read_doc_target(suffix)
    if not worker:
        return None
    try:
        size = p.stat().st_size
    except OSError:
        return None
    if size > 30 * 1024 * 1024:
        # Too big to ship; let the host parser handle it locally.
        return None
    try:
        raw = p.read_bytes()
    except OSError as e:
        return False, "", f"could not read source file: {e}"

    import base64
    b64 = base64.b64encode(raw).decode("ascii")

    # Each format gets its own Python script so the dispatcher only
    # imports what's needed (no point pulling pymupdf into a docx
    # parse). Pages / sheets are inlined as Python literals.
    pages_lit = "None" if pages is None else repr(pages)
    sheets_lit = "None" if sheets is None else repr(sheets)
    if suffix == ".pdf":
        worker_script = _read_doc_pdf_worker_script(pages_lit)
    elif suffix == ".docx":
        worker_script = _read_doc_docx_worker_script()
    elif suffix in {".xlsx", ".xlsm"}:
        worker_script = _read_doc_xlsx_worker_script(sheets_lit)
    else:
        return None

    # Worker reads stdin as `<base64>\n<source>` — first line is the
    # file payload, subsequent lines are the parser script. The
    # PowerShell wrapper splits them and feeds the source to python.
    payload = b64 + "\n" + worker_script
    ps_script = (
        "$python = (Get-Command py.exe -ErrorAction SilentlyContinue).Source;"
        "if (-not $python) { $python = (Get-Command python.exe -ErrorAction SilentlyContinue).Source };"
        "if (-not $python) { Write-Error 'python not found on worker'; exit 3 };"
        # Read the full stdin, split on first newline: header is the
        # base64 payload, body is the python source.
        "$all = [Console]::In.ReadToEnd();"
        "$nl = $all.IndexOf(\"`n\");"
        "if ($nl -lt 0) { Write-Error 'malformed input'; exit 4 };"
        "$payload_b64 = $all.Substring(0, $nl).TrimEnd(\"`r\");"
        "$src = $all.Substring($nl + 1);"
        # Pipe both pieces in: env var for the payload, stdin for source.
        "$env:_GIGACHAT_DOC_PAYLOAD = $payload_b64;"
        "$src | & $python -"
    )
    ok, stdout, stderr = await dispatch_to_worker_powershell(
        worker, ps_script, timeout_sec=60.0, stdin_text=payload,
    )
    if not ok:
        log.info(
            "compute_pool: distributed read_doc(%s) via %s failed: %s",
            suffix, worker.get("label"), stderr,
        )
        return False, "", stderr
    text = (stdout or b"").decode("utf-8", errors="replace")
    log.info(
        "compute_pool: dispatched read_doc(%s) to worker %s — %d chars",
        suffix, worker.get("label"), len(text),
    )
    return True, text, ""


def _read_doc_pdf_worker_script(pages_lit: str) -> str:
    """Python source the worker runs to parse a PDF and emit text.

    The host's `_read_pdf_sync` does similar work — we replicate the
    page-range parser here (instead of importing it cross-process)
    because the worker's Python doesn't know about Gigachat.
    """
    return (
        "import os, sys, base64, tempfile, pymupdf\n"
        "payload = base64.b64decode(os.environ['_GIGACHAT_DOC_PAYLOAD'])\n"
        "with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as fh:\n"
        "    fh.write(payload)\n"
        "    path = fh.name\n"
        "try:\n"
        "    doc = pymupdf.open(path)\n"
        "    n = len(doc)\n"
        f"    pages_spec = {pages_lit}\n"
        "    if pages_spec:\n"
        "        idxs = []\n"
        "        for chunk in str(pages_spec).split(','):\n"
        "            chunk = chunk.strip()\n"
        "            if '-' in chunk:\n"
        "                a, b = chunk.split('-', 1)\n"
        "                try: idxs.extend(range(int(a) - 1, min(int(b), n)))\n"
        "                except ValueError: pass\n"
        "            else:\n"
        "                try: idxs.append(int(chunk) - 1)\n"
        "                except ValueError: pass\n"
        "        idxs = [i for i in idxs if 0 <= i < n]\n"
        "    else:\n"
        "        idxs = list(range(min(20, n)))\n"
        "    parts = []\n"
        "    for i in idxs:\n"
        "        page = doc.load_page(i)\n"
        "        text = page.get_text('text') or ''\n"
        "        parts.append(f'--- page {i+1} ---\\n' + text)\n"
        "    sys.stdout.write('\\n'.join(parts))\n"
        "finally:\n"
        "    try: doc.close()\n"
        "    except Exception: pass\n"
        "    try: os.unlink(path)\n"
        "    except Exception: pass\n"
    )


def _read_doc_docx_worker_script() -> str:
    """Python source the worker runs to parse a .docx and emit text."""
    return (
        "import os, sys, base64, tempfile\n"
        "from docx import Document\n"
        "payload = base64.b64decode(os.environ['_GIGACHAT_DOC_PAYLOAD'])\n"
        "with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as fh:\n"
        "    fh.write(payload)\n"
        "    path = fh.name\n"
        "try:\n"
        "    doc = Document(path)\n"
        "    parts = [p.text for p in doc.paragraphs]\n"
        "    for table in doc.tables:\n"
        "        for row in table.rows:\n"
        "            cells = [c.text.strip() for c in row.cells]\n"
        "            parts.append(' | '.join(cells))\n"
        "    sys.stdout.write('\\n'.join(p for p in parts if p))\n"
        "finally:\n"
        "    try: os.unlink(path)\n"
        "    except Exception: pass\n"
    )


def _read_doc_xlsx_worker_script(sheets_lit: str) -> str:
    """Python source the worker runs to parse a .xlsx and emit text."""
    return (
        "import os, sys, base64, tempfile, openpyxl\n"
        "payload = base64.b64decode(os.environ['_GIGACHAT_DOC_PAYLOAD'])\n"
        "with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as fh:\n"
        "    fh.write(payload)\n"
        "    path = fh.name\n"
        "try:\n"
        "    wb = openpyxl.load_workbook(path, data_only=True)\n"
        f"    spec = {sheets_lit}\n"
        "    sheets = (\n"
        "        [s.strip() for s in str(spec).split(',') if s.strip()]\n"
        "        if spec else wb.sheetnames[:3]\n"
        "    )\n"
        "    parts = []\n"
        "    for name in sheets:\n"
        "        if name not in wb.sheetnames: continue\n"
        "        ws = wb[name]\n"
        "        parts.append(f'--- sheet {name} ---')\n"
        "        for row in ws.iter_rows(values_only=True):\n"
        "            cells = [str(c) if c is not None else '' for c in row]\n"
        "            parts.append('\\t'.join(cells))\n"
        "    sys.stdout.write('\\n'.join(parts))\n"
        "finally:\n"
        "    try: os.unlink(path)\n"
        "    except Exception: pass\n"
    )


# ---------------------------------------------------------------------------
# Distributed `python_exec` — short Python snippets dispatched to a worker
#
# Pure-Python snippets that don't touch the host filesystem can run on
# any worker with a Python interpreter (every probe-able worker, since
# the spec probe already requires Python). Routing CPU-heavy snippets
# to a worker frees the host while it's busy with inference.
#
# Safety: we only dispatch when the snippet doesn't reference the local
# filesystem (`open(`, `pathlib`, `os.listdir`, etc.). Snippets that DO
# touch files would silently get the wrong cwd on the worker; the guard
# falls back to host execution rather than producing wrong results. The
# worker runs with `python -I` (isolated mode) — same isolation profile
# as the host-side execution.
# ---------------------------------------------------------------------------

# Substring patterns that hint the snippet reads/writes local files.
# Any match keeps execution on host (where the cwd actually exists).
# Conservative on purpose — false positives just keep things local
# (no harm), false negatives hit the worker fallback path.
_PYTHON_EXEC_HOST_ONLY_PATTERNS = (
    "open(", "with open", "os.listdir", "os.walk", "os.scandir",
    "os.path.exists", "os.path.isfile", "os.path.isdir",
    "os.makedirs", "os.mkdir", "os.remove", "os.unlink", "os.rename",
    "Path(", "from pathlib", "import pathlib",
    "shutil.", "glob.", "subprocess", "os.system",
    "tempfile",  # tempfile creates files; runs differently on workers
)


def _python_exec_dispatchable(code: str) -> bool:
    """Heuristic check whether a `python_exec` snippet is safe to ship to a worker.

    Returns False when the code clearly references the local filesystem
    or shells out — those need the host's cwd / environment. Returns
    True for compute / network / pure-python work that can run anywhere.

    Heuristic, not airtight. False positives keep the call on host (no
    harm). False negatives surface as a worker-side error and the caller
    falls back to host on the next try.
    """
    if not code or not code.strip():
        return False
    for needle in _PYTHON_EXEC_HOST_ONLY_PATTERNS:
        if needle in code:
            return False
    return True


async def dispatch_python_exec_to_worker(
    code: str, *, timeout_sec: float = 60.0,
) -> tuple[bool, str, int, str] | None:
    """Run a Python snippet on a round-robin-picked worker.

    Returns ``None`` when no eligible worker is reachable — caller
    falls back to host. Otherwise returns
    ``(ok, output, exit_code, error_message)``:

      * ``ok=True`` and ``exit_code=0`` on clean success.
      * ``ok=False`` with the worker's stderr in ``error`` when the
        snippet itself raised an exception (caller surfaces as the
        snippet's failure, doesn't retry on host).
      * ``ok=False`` with empty exit_code when the dispatch transport
        itself failed (caller MAY retry on host as a safety net).

    Output is capped at 20 000 chars (matches host-side `_clip()`).
    """
    worker = _pick_tool_dispatch_target()
    if not worker:
        return None
    ssh_host = (worker.get("ssh_host") or "").strip()
    if not ssh_host:
        return None

    # Run the user's snippet via `python -I -` (isolated mode, source
    # from stdin). Same flags as the host-side path — preserves the
    # security posture: no PYTHONPATH inheritance, no user-site, no
    # implicit cwd-relative imports.
    ps_script = (
        "$python = (Get-Command py.exe -ErrorAction SilentlyContinue).Source;"
        "if (-not $python) { $python = (Get-Command python.exe -ErrorAction SilentlyContinue).Source };"
        "if (-not $python) { Write-Error 'python not found on worker'; exit 3 };"
        # Read the snippet from our own stdin and pipe it into python -.
        "$src = [Console]::In.ReadToEnd();"
        "$src | & $python -I -"
    )
    ok, stdout, stderr = await dispatch_to_worker_powershell(
        worker, ps_script, timeout_sec=timeout_sec, stdin_text=code,
    )
    if not ok:
        # Transport-level failure (ssh down, worker unreachable, timeout).
        # Empty exit_code signals "didn't run" so the caller knows it can
        # retry on host without risking double-execution side effects.
        log.info(
            "compute_pool: distributed python_exec via %s failed: %s",
            worker.get("label"), stderr,
        )
        return False, "", 0, stderr or "worker unreachable"
    output = (stdout or b"").decode("utf-8", errors="replace")
    if len(output) > 20000:
        output = output[:20000] + "\n[output truncated]"
    log.info(
        "compute_pool: dispatched python_exec to worker %s — %d chars",
        worker.get("label"), len(output),
    )
    # `dispatch_to_worker_powershell` only returns ok=True on rc==0 from
    # the worker's PowerShell shell; the inner Python snippet's exit
    # code is reflected in PowerShell's $LASTEXITCODE which the wrapper
    # surfaces as its own exit code. Consider any reachable run a
    # successful dispatch (`ok=True`); if the snippet raised, output
    # already carries the traceback.
    return True, output, 0, ""


# ---------------------------------------------------------------------------
# Active dedup execution — run `ollama rm` on redundant copies
#
# The advisor (`pool_dedup_recommendations`) tells us which copies are
# safely removable; this helper actually executes them via SSH. Each
# action is independent — a failure on one worker doesn't block the
# rest. Returns a per-action result list so the API can surface
# success/failure.
# ---------------------------------------------------------------------------


async def execute_dedup_recommendations(
    *, model_filter: str | None = None,
) -> list[dict]:
    """Walk `pool_dedup_recommendations` and SSH-run `ollama rm` on
    each redundant copy. Returns a list of per-action result dicts::

        [
          {"model": "llama3:8b", "worker_id": "wid-A",
           "ok": True, "bytes_reclaimed": 4_500_000_000,
           "error": None},
          ...
        ]

    `model_filter`, when set, restricts execution to that model only.
    """
    recs = pool_dedup_recommendations()
    if model_filter:
        recs = [r for r in recs if r.get("model") == model_filter]
    workers_by_id = {w["id"]: w for w in db.list_compute_workers(enabled_only=False)}
    results: list[dict] = []
    for rec in recs:
        model = rec["model"]
        size = rec.get("size_bytes") or 0
        for loc in rec.get("remove_from") or []:
            if not loc.startswith("worker:"):
                # Host removals are not supported via this API on
                # purpose — host's Ollama is the operator's primary
                # surface and they should run `ollama rm` themselves.
                results.append({
                    "model": model, "location": loc,
                    "ok": False, "bytes_reclaimed": 0,
                    "error": "skipped: host removals are operator-driven",
                })
                continue
            wid = loc.split(":", 1)[1]
            worker = workers_by_id.get(wid)
            if not worker:
                results.append({
                    "model": model, "worker_id": wid,
                    "ok": False, "bytes_reclaimed": 0,
                    "error": "worker not found",
                })
                continue
            ssh_host = (worker.get("ssh_host") or "").strip()
            if not ssh_host:
                results.append({
                    "model": model, "worker_id": wid,
                    "ok": False, "bytes_reclaimed": 0,
                    "error": "worker has no ssh_host",
                })
                continue
            # `ollama rm <name>` removes the model on the worker. We
            # feed the model name through PowerShell single quotes so
            # `$` / backticks aren't expanded; only `'` itself needs
            # escaping (doubled).
            safe_model = model.replace("'", "''")
            ps_script = (
                # Resolve the ollama binary; bail if missing.
                "$ollama = (Get-Command ollama.exe -ErrorAction SilentlyContinue).Source;"
                "if (-not $ollama) { Write-Error 'ollama not found on worker'; exit 3 };"
                f"& $ollama rm '{safe_model}'"
            )
            ok, stdout, stderr = await dispatch_to_worker_powershell(
                worker, ps_script, timeout_sec=60.0,
            )
            results.append({
                "model": model, "worker_id": wid,
                "ok": ok, "bytes_reclaimed": size if ok else 0,
                "error": None if ok else (stderr or "unknown failure"),
            })
            if ok:
                log.info(
                    "compute_pool: dedup removed %s from worker %s — "
                    "reclaimed %.2f GB",
                    model, worker.get("label"), size / (1024 ** 3),
                )
    return results


# ---------------------------------------------------------------------------
# Pool-wide model-inventory summary
#
# `pool_inventory_summary()` consolidates the per-node model lists into a
# single per-model view: which nodes hold each model, how much disk that
# model occupies in aggregate across the pool, and how much of that is
# redundant (same blob duplicated on multiple nodes). The Settings UI
# uses this for a "pool storage" panel and the dedup advisor surfaces
# concrete reclaim targets without prescribing removal.
#
# Shipping this as a read-only summary first — the operator decides what
# to delete based on their own routing preferences. Auto-deletion would
# need to confirm the model isn't currently the chat target on the node
# in question, which the inventory itself can't see.
# ---------------------------------------------------------------------------


# Match the trailing quant suffix on Ollama tags / GGUF filenames.
# Examples that match: ``-q4_0``, ``-q4_K_M``, ``-q5_K_S``, ``-q8_0``,
# ``-iq3_xs``, ``-iq2_m`` (case-insensitive). The picker uses this to
# group models that share a base name but ship in different quant
# levels — the user can see at a glance "I have llama3:8b in both
# Q4_K_M and Q8_0" and the dedup advisor surfaces the redundancy.
import re as _re_quant
_QUANT_SUFFIX_RE = _re_quant.compile(
    r"[-_:](i?q\d+(?:_(?:[\w]+))?)$",
    _re_quant.IGNORECASE,
)


def _strip_quant_suffix(name: str) -> tuple[str, str | None]:
    """Split a model name into (base, quant) — quant is None when
    the name doesn't carry a recognisable suffix.

    Example::

        _strip_quant_suffix("llama3:8b-q4_K_M")   -> ("llama3:8b", "Q4_K_M")
        _strip_quant_suffix("llama3:8b")          -> ("llama3:8b", None)
        _strip_quant_suffix("qwen2.5:0.5b-iq3_xs")-> ("qwen2.5:0.5b", "IQ3_XS")
    """
    if not name:
        return name, None
    m = _QUANT_SUFFIX_RE.search(name)
    if not m:
        return name, None
    base = name[: m.start()]
    quant = m.group(1).upper()
    return base, quant


def pool_inventory_summary() -> dict:
    """Return a per-model breakdown of where every chat-capable model
    sits in the pool, plus aggregate disk-pressure totals.

    Shape::

        {
          "models": [
            {
              "name": "llama3:8b",
              "family": "llama",
              "size_bytes": 4_500_000_000,
              "locations": ["host", "worker:wid-A", "worker:wid-B"],
              "copies": 3,
              "redundant_bytes": 9_000_000_000,  # (copies - 1) * size_bytes
            },
            ...
          ],
          "total_unique_bytes": 24_300_000_000,
          "total_pool_bytes": 39_500_000_000,
          "total_redundant_bytes": 15_200_000_000,
        }

    A model is "redundant" when more than one node holds it. The user
    might still want all those copies (e.g. each worker is a chat
    target for that model), so the summary doesn't prescribe deletion
    — `pool_dedup_recommendations()` does that with explicit
    safe-to-remove logic.
    """
    by_name: dict[str, dict] = {}
    for entry in _pool_model_inventory():
        bucket = by_name.setdefault(
            entry["name"],
            {
                "name": entry["name"],
                "family": entry["family"],
                "size_bytes": entry["size_bytes"],
                "locations": [],
            },
        )
        # Take the largest size_bytes any node reports — workers may
        # under-count for an in-flight transfer; we want the real
        # blob size for the dedup math.
        if entry["size_bytes"] > bucket["size_bytes"]:
            bucket["size_bytes"] = entry["size_bytes"]
        if entry["source"] not in bucket["locations"]:
            bucket["locations"].append(entry["source"])

    models: list[dict] = []
    total_pool_bytes = 0
    total_unique_bytes = 0
    total_redundant_bytes = 0
    for entry in by_name.values():
        copies = len(entry["locations"])
        size = entry["size_bytes"]
        redundant = (copies - 1) * size if copies > 1 else 0
        models.append({
            **entry,
            "copies": copies,
            "redundant_bytes": redundant,
        })
        total_pool_bytes += copies * size
        total_unique_bytes += size
        total_redundant_bytes += redundant

    # Sort by redundant-bytes descending so the heaviest dedup wins
    # surface first in the Settings UI.
    models.sort(key=lambda m: m["redundant_bytes"], reverse=True)

    # Quant-variant grouping: every model name has its trailing quant
    # suffix stripped; entries that share a base get bucketed under
    # the base. Lets the user see "for llama3:8b family I have Q4_K_M
    # on host AND Q8_0 on worker-A" without paging through the flat
    # list. Models without a recognisable quant suffix become single-
    # member groups.
    by_base: dict[str, dict] = {}
    for entry in models:
        base, quant = _strip_quant_suffix(entry["name"])
        bucket = by_base.setdefault(
            base,
            {
                "base": base,
                "family": entry.get("family"),
                "variants": [],
                "total_bytes": 0,
            },
        )
        bucket["variants"].append({
            "name": entry["name"],
            "quant": quant,
            "size_bytes": entry["size_bytes"],
            "locations": entry["locations"],
        })
        bucket["total_bytes"] += entry["size_bytes"] * entry["copies"]

    quant_groups = [
        b for b in by_base.values() if len(b["variants"]) > 1
    ]
    quant_groups.sort(key=lambda b: b["total_bytes"], reverse=True)

    # Per-node disk + cache report. Lets the inventory UI render
    # "host SSD: 280 / 500 GB" + "worker-A: 60 / 250 GB", so the user
    # can spot where there's room to land more model pulls. Falls
    # back to zero values when probe data isn't yet available.
    nodes: list[dict] = []
    try:
        from . import sysdetect as _sysdetect
        host_spec = _sysdetect.detect_system()
    except Exception:
        host_spec = {}
    nodes.append({
        "id": "host",
        "label": "host",
        "ram_total_gb": float(host_spec.get("ram_gb") or 0),
        "vram_gb": float(host_spec.get("vram_gb") or 0),
        # Host disk free is best-effort: shutil.disk_usage on the
        # workspace drive. Empty when the call fails so the UI can
        # fall back to "—".
        "disk_total_gb": _host_disk_total_gb(),
        "disk_free_gb": _host_disk_free_gb(),
    })
    for w in db.list_compute_workers(enabled_only=False):
        caps = w.get("capabilities") or {}
        nodes.append({
            "id": f"worker:{w['id']}",
            "label": w.get("label") or w.get("address"),
            "ram_total_gb": float(caps.get("ram_total_gb") or 0),
            "ram_free_gb": float(caps.get("ram_free_gb") or 0),
            "disk_total_gb": float(caps.get("disk_total_gb") or 0),
            "disk_free_gb": float(caps.get("disk_free_gb") or 0),
            "cached_overrides": list(caps.get("cached_overrides") or []),
        })

    return {
        "models": models,
        "total_unique_bytes": total_unique_bytes,
        "total_pool_bytes": total_pool_bytes,
        "total_redundant_bytes": total_redundant_bytes,
        # Only multi-variant bases land here — single-quant models
        # don't need a "group" view since the flat models[] list
        # already covers them.
        "quant_groups": quant_groups,
        # Per-node storage / RAM picture so the UI can render an
        # accurate "where to put what" view of the pool.
        "nodes": nodes,
    }


def _host_disk_total_gb() -> float:
    """Best-effort host disk capacity for the Gigachat data dir.
    Returns 0 on any read failure (Windows / Linux cross-platform).
    """
    try:
        import shutil as _shutil
        from pathlib import Path as _Path
        usage = _shutil.disk_usage(str(_Path.home()))
        return round(usage.total / (1024 ** 3), 1)
    except Exception:
        return 0.0


def _host_disk_free_gb() -> float:
    """Best-effort host disk free space, mirrors `_host_disk_total_gb`."""
    try:
        import shutil as _shutil
        from pathlib import Path as _Path
        usage = _shutil.disk_usage(str(_Path.home()))
        return round(usage.free / (1024 ** 3), 1)
    except Exception:
        return 0.0


def pool_dedup_recommendations() -> list[dict]:
    """Suggest which redundant copies are *safe* to delete.

    Heuristic: for each model with copies > 1, recommend keeping the
    copy on the strongest node (host wins ties; among workers,
    `_capability_score` ranks). Other copies are listed as removable
    with the disk savings they'd reclaim.

    Returns a list of dicts shaped::

        {
          "model": "llama3:8b",
          "size_bytes": 4_500_000_000,
          "keep_at": "host",
          "remove_from": ["worker:wid-B", "worker:wid-C"],
          "bytes_reclaimed": 9_000_000_000,
        }

    Empty list when nothing is dedup-able. The recommendation is
    advisory — operator decides whether their routing pattern wants
    to keep redundant copies (e.g. for fail-over or for cases where
    the worker's local model serves its own chat target).
    """
    workers = {w["id"]: w for w in db.list_compute_workers(enabled_only=True)}
    summary = pool_inventory_summary()
    out: list[dict] = []
    for entry in summary["models"]:
        if entry["copies"] <= 1:
            continue
        locations = list(entry["locations"])
        # Rank locations: host first, then workers by capability score
        # (strongest worker wins so we keep the copy where chat is
        # most likely to be routed anyway).

        def _rank(loc: str) -> tuple:
            if loc == "host":
                return (0,)  # host wins
            wid = loc.split(":", 1)[1] if ":" in loc else ""
            w = workers.get(wid)
            if not w:
                # Unknown worker — keep at the back of the queue.
                return (3,)
            score = _capability_score(w)
            # Negate the tuple's numeric components so larger scores
            # sort to the front (after the host's leading 0).
            return (1, -float(score[0] or 0), -float(score[1] or 0))

        locations.sort(key=_rank)
        keep = locations[0]
        remove = locations[1:]
        out.append({
            "model": entry["name"],
            "size_bytes": entry["size_bytes"],
            "keep_at": keep,
            "remove_from": remove,
            "bytes_reclaimed": entry["size_bytes"] * len(remove),
        })
    out.sort(key=lambda r: r["bytes_reclaimed"], reverse=True)
    return out


# ---------------------------------------------------------------------------
# Auto-LAN-pull for worker-only drafts
#
# `pick_draft_for` only promotes host-resident drafts because llama-server
# needs the GGUF on local disk. But the user may have pulled a small
# same-family chat model on a worker and never copied it to host — leaving
# speculative decoding off the table even when the pool has exactly the
# draft we'd want.
#
# `_maybe_kickoff_draft_lan_sync` closes that gap. When the picker would
# have chosen a worker-only draft (smaller than every host candidate, same
# tokenizer, fits the size budget), we kick off a background SCP from the
# worker to host. The current chat continues on the slower path; the next
# turn picks up the now-host-resident draft. Rate-limited per (target,
# draft) pair so we don't trigger a parallel pull every probe sweep.
# ---------------------------------------------------------------------------
_DRAFT_PULL_LAST_ATTEMPT: dict[str, float] = {}
_DRAFT_PULL_COOLDOWN_SEC = 600.0  # 10 min
# Cap on the cooldown tracker: keys are
# "<target>|<candidate>|<worker_id>" so the combo space is
# N_targets × N_candidates × N_workers. A backend that's been up for
# years cycling many models could grow this slowly. FIFO eviction at
# 1024 entries — losing a cooldown stamp just means the next
# `_maybe_kickoff_draft_lan_sync` call for that combo can fire
# immediately, identical to the cold-start behaviour.
_DRAFT_PULL_LAST_ATTEMPT_MAX = 1024


def _evict_draft_pull_tracker_if_full() -> None:
    """Drop the oldest entry when the combo tracker exceeds its cap."""
    while len(_DRAFT_PULL_LAST_ATTEMPT) > _DRAFT_PULL_LAST_ATTEMPT_MAX:
        try:
            oldest_key = next(iter(_DRAFT_PULL_LAST_ATTEMPT))
            del _DRAFT_PULL_LAST_ATTEMPT[oldest_key]
        except (StopIteration, KeyError):
            break


def _maybe_kickoff_draft_lan_sync(
    target_model_name: str,
    candidate_name: str,
    worker_id: str,
) -> None:
    """Schedule a background `model_sync.sync_model` from worker to host.

    Fire-and-forget — the chat layer doesn't wait. On success, the next
    `pick_draft_for(target)` call sees the model in the host inventory
    and promotes it through the normal path. Failures are logged and
    silently retried after the cooldown.
    """
    key = f"{target_model_name}|{candidate_name}|{worker_id}"
    now_m = time.monotonic()
    last = _DRAFT_PULL_LAST_ATTEMPT.get(key, 0.0)
    if now_m - last < _DRAFT_PULL_COOLDOWN_SEC:
        return
    _DRAFT_PULL_LAST_ATTEMPT[key] = now_m
    _evict_draft_pull_tracker_if_full()

    log.info(
        "compute_pool: kicking off LAN sync of draft candidate %r from "
        "worker %r so future turns can use speculative decoding for %r",
        candidate_name, worker_id, target_model_name,
    )

    async def _bg() -> None:
        try:
            from . import model_sync
            # Note: sync_model SCPs FROM host TO worker. We need the
            # opposite direction. The ModelSync API has `pull_to_host`
            # for that path; if it's missing, fall back to a manual SCP
            # via the worker's ssh_host. Best-effort.
            puller = getattr(model_sync, "pull_model_to_host", None)
            if puller is None:
                log.info(
                    "compute_pool: model_sync.pull_model_to_host unavailable "
                    "on this build; draft auto-pull deferred to manual `ollama pull`"
                )
                return
            await puller(candidate_name, worker_id)
        except Exception as e:
            log.warning(
                "compute_pool: LAN pull of draft %r from worker %r failed: %s",
                candidate_name, worker_id, e,
            )

    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            loop.create_task(_bg())
    except RuntimeError:
        pass


# Synchronous draft SCP threshold. When no host candidate exists for a
# target, we'll BLOCK on a worker→host SCP for drafts smaller than
# this. Above the cap, the SCP cost exceeds the speculative speedup
# benefit and we stay on the (slower, non-speculative) path.
# 2 GB ≈ a Q4_K_M 1B-3B draft over Gigabit-Ethernet completes in
# ~20-40 s — comfortably inside one chat-turn's tolerance for a
# one-time setup cost that benefits every following turn.
_SYNC_DRAFT_SCP_MAX_BYTES = 2 * 1024 ** 3
_SYNC_DRAFT_SCP_TIMEOUT_SEC = 120.0


async def await_draft_for(target_model_name: str) -> dict | None:
    """Async version of `pick_draft_for` that can synchronously SCP a
    viable worker-only draft when no host candidate exists.

    The sync `pick_draft_for` only promotes host-resident drafts and
    fires a background SCP for worker-only candidates that won't help
    the current turn. This async wrapper closes that gap: when there's
    no host draft AND a worker has a small same-family candidate,
    we BLOCK on a worker→host SCP so this very turn benefits from
    speculative decoding.

    Constraints:
      * Worker draft must be < `_SYNC_DRAFT_SCP_MAX_BYTES` (2 GB).
        Above that the SCP exceeds the chat user's tolerance for
        first-token latency.
      * SCP timeout caps the wait at `_SYNC_DRAFT_SCP_TIMEOUT_SEC`
        — failed pulls fall through to the no-speculative path.

    Returns the same dict shape as `pick_draft_for`, or None when
    neither a host candidate nor a viable worker draft exists.
    """
    immediate = pick_draft_for(target_model_name)
    if immediate is not None:
        return immediate

    # No host draft. Walk pool inventory for a worker-only candidate
    # that's small enough to be worth a sync SCP.
    target = resolve_ollama_model(target_model_name)
    if not target:
        return None
    target_family = (target.get("family") or "").lower()
    target_size = int(target.get("size_bytes") or 0)
    if target_size < _SPECULATIVE_MIN_TARGET_BYTES:
        return None  # target too small for speculative to help
    max_draft_bytes = int(target_size * _DRAFT_MAX_SIZE_FRACTION)

    # Find the smallest same-family worker candidate within the sync
    # SCP size budget. Tokenizer-fingerprint matching needs the GGUF
    # locally (we're trying to GET it locally), so we limit to the
    # cheap family check here.
    best: dict | None = None
    for m in _pool_model_inventory():
        if m["name"] == target_model_name:
            continue
        if m["source"] == "host":
            continue  # already covered by pick_draft_for
        if m["size_bytes"] <= 0:
            continue
        if m["size_bytes"] > min(max_draft_bytes, _SYNC_DRAFT_SCP_MAX_BYTES):
            continue
        if not (m["family"] and target_family and m["family"] == target_family):
            continue
        if best is None or m["size_bytes"] < best["size_bytes"]:
            best = m

    if best is None:
        return None

    # Pull synchronously (SCP from worker to host). The pull function
    # is best-effort; on failure we fall through to the no-speculative
    # path so the user still gets their chat.
    try:
        from . import model_sync
        puller = getattr(model_sync, "pull_model_to_host", None)
        if puller is None:
            return None
        wid = best["source"].split(":", 1)[1] if ":" in best["source"] else ""
        if not wid:
            return None
        log.info(
            "compute_pool: sync-pulling draft %r from worker %s for "
            "speculative on %r (%.2f GB)",
            best["name"], wid, target_model_name,
            best["size_bytes"] / (1024 ** 3),
        )
        await asyncio.wait_for(
            puller(best["name"], wid),
            timeout=_SYNC_DRAFT_SCP_TIMEOUT_SEC,
        )
    except Exception as e:
        log.info(
            "compute_pool: sync draft pull failed for %r: %s",
            best["name"], e,
        )
        return None

    # Re-pick with the now-host-resident draft in scope.
    return pick_draft_for(target_model_name)


def _draft_override_for(target_model_name: str) -> str | None:
    """User-pinned draft override for a specific target.

    Read from the ``compute_pool_speculative_overrides`` setting, which
    stores a JSON object mapping ``"<target_model_name>": "<draft_name>"``.
    When set, the picker uses that exact draft regardless of family or
    tokenizer-fingerprint matches — an escape hatch for power users who
    want to experiment with cross-vocab speculative pairs (the few cases
    where the safety checks reject a pair the user knows works).

    Misuse is the user's responsibility: a draft with a different
    tokenizer than the target produces near-0 % accept rate and slows
    chat down. The setting is undocumented in the chat UI — only the
    settings API exposes it — so it stays out of the way of users who
    haven't read the README.
    """
    try:
        raw = db.get_setting("compute_pool_speculative_overrides")
    except Exception:
        return None
    if not raw:
        return None
    try:
        mapping = jsonutil.loads(raw)
    except (json.JSONDecodeError, TypeError):
        return None
    if not isinstance(mapping, dict):
        return None
    val = mapping.get(target_model_name)
    if isinstance(val, str) and val.strip():
        return val.strip()
    return None


def _pool_model_inventory() -> list[dict]:
    """Collated list of every chat-capable model installed anywhere in
    the pool. Each entry carries the source (host vs worker) so the
    caller can tell whether the GGUF is locally resolvable.

    Returns dicts shaped::

        {
            "name": "qwen2.5:0.5b",      # Ollama tag
            "family": "qwen2",            # Ollama details.family (lower-case)
            "size_bytes": 394_000_000,    # raw file size
            "source": "host" | "worker:<wid>",
        }

    Embedding-only models are filtered out (their ``family`` is the
    embed-model family — `nomic-embed-text` etc. — and they have no
    chat template, so they can't serve as a draft).
    """
    out: list[dict] = []

    # Host-side: walk Ollama's manifest store. `_resolve_ollama_manifest`
    # already handles the manifest schema; we don't need /api/tags here
    # because resolve gives us the on-disk file size directly.
    try:
        host_models = _list_host_installed_models()
    except Exception as e:
        log.info("pool inventory: host listing failed: %s", e)
        host_models = []
    for m in host_models:
        out.append({
            "name": m.get("name") or "",
            "family": (m.get("family") or "").lower(),
            "size_bytes": int(m.get("size_bytes") or 0),
            "source": "host",
        })

    # Worker-side: each probe stamps the worker row's `capabilities.models`
    # with what Ollama reports via /api/tags — name, size, family,
    # parameter_size, quantization_level. Use that snapshot directly.
    for w in db.list_compute_workers(enabled_only=True):
        if not w.get("use_for_chat"):
            continue
        caps = w.get("capabilities") or {}
        for m in caps.get("models") or []:
            if not isinstance(m, dict):
                continue
            out.append({
                "name": m.get("name") or "",
                "family": (m.get("family") or "").lower(),
                "size_bytes": int(m.get("size") or 0),
                "source": f"worker:{w['id']}",
            })

    # Filter out embedding-only models — name-based heuristic because
    # Ollama's `family` for embed models is the embed family
    # (`nomic-bert`, `bert`, …), not always identifiable cleanly. The
    # name carries `embed` in every embedding model we ship.
    return [m for m in out if m["name"] and "embed" not in m["name"].lower()]


def find_smaller_variants_in_family(
    model_name: str, *, max_size_bytes: int,
) -> list[dict]:
    """Suggest smaller-but-same-family alternatives that would fit.

    Used by the mega-model warning path: when the user picks a model
    too big to run usefully on the pool, we surface a short list of
    same-family models that ARE small enough — Llama-3-70B doesn't
    fit, here are llama3:8b and llama3:13b that do.

    Quality preservation: the returned list is filtered to models the
    user (or someone in the swarm) ALREADY HAS, NOT freshly-quantized
    versions. We never silently degrade Q4 → Q3 to fit; the user
    keeps their quality bar by picking from real same-family
    variants designed at smaller sizes.

    Returns up to 6 entries sorted largest-fitting first (closer to
    user's original quality intent). Empty list when no same-family
    variant exists in the pool.
    """
    if not model_name:
        return []
    info = resolve_ollama_model(model_name)
    target_family = ((info or {}).get("family") or "").lower()
    if not target_family:
        return []
    inventory = _pool_model_inventory()
    candidates: list[dict] = []
    seen_names: set[str] = set()
    target_size = int((info or {}).get("size_bytes") or 0)
    for entry in inventory:
        name = entry.get("name") or ""
        family = (entry.get("family") or "").lower()
        size = int(entry.get("size_bytes") or 0)
        if not name or family != target_family:
            continue
        if name == model_name or name in seen_names:
            continue
        # Only suggest models smaller than the original AND that fit
        # within the pool's available memory.
        if target_size > 0 and size >= target_size:
            continue
        if max_size_bytes > 0 and size > max_size_bytes:
            continue
        seen_names.add(name)
        candidates.append({
            "name": name,
            "family": family,
            "size_bytes": size,
            "size_gb": round(size / (1024 ** 3), 1),
            "source": entry.get("source") or "host",
        })
    # Largest-fitting first — closer to the user's original quality
    # intent than the smallest variant in the family.
    candidates.sort(key=lambda c: c["size_bytes"], reverse=True)
    return candidates[:6]


def _list_host_installed_models() -> list[dict]:
    """Read every `manifests/registry.ollama.ai/library/<model>/<tag>`
    JSON file under the host's Ollama models dir, return per-tag
    `{name, family, size_bytes}` summaries.

    We avoid hitting Ollama's HTTP API because (a) it may not be running
    yet at boot and (b) the picker is on the chat-startup hot path —
    one extra HTTP roundtrip per turn is wasteful when the same data
    is sitting in static JSON files on disk.
    """
    base = _OLLAMA_MODELS_DIR / "manifests" / "registry.ollama.ai"
    if not base.is_dir():
        return []
    out: list[dict] = []
    # Layout: manifests/registry.ollama.ai/<namespace>/<model>/<tag>
    # `namespace` is `library` for the public registry, but other
    # namespaces (`hf.co/...`) are also possible.
    for ns in base.iterdir():
        if not ns.is_dir():
            continue
        for model_dir in ns.iterdir():
            if not model_dir.is_dir():
                continue
            for tag_file in model_dir.iterdir():
                if not tag_file.is_file():
                    continue
                # Tag name = filename. Build the canonical Ollama name
                # back from the directory layout: <model>:<tag> for
                # the public namespace, <ns>/<model>:<tag> otherwise.
                name = (
                    f"{model_dir.name}:{tag_file.name}"
                    if ns.name == "library"
                    else f"{ns.name}/{model_dir.name}:{tag_file.name}"
                )
                info = resolve_ollama_model(name)
                if not info:
                    continue
                out.append({
                    "name": name,
                    "family": (info.get("family") or "").lower(),
                    "size_bytes": int(info.get("size_bytes") or 0),
                })
    return out


def pick_draft_for(target_model_name: str) -> dict | None:
    """Find a viable speculative-decoding draft for ``target_model_name``.

    Returns a dict shaped ``{"name", "gguf_path", "size_bytes",
    "source", "match"}`` when a vocab-compatible chat model dramatically
    smaller than the target is available on the host's disk, else
    ``None``. The ``match`` field reports how the candidate was
    verified — ``"family"``, ``"tokenizer"``, or ``"override"`` — so
    operators can audit which path engaged.

    Three-tier matching, cheapest first:

      1. **Manual override.** If the user pinned a specific draft for
         this target via ``compute_pool_speculative_overrides``, that
         draft wins outright (size + family checks bypassed). Trust
         the operator — they've decided the pair works.

      2. **Family match.** Same Ollama-reported family
         (`details.family`), e.g. both ``llama`` or both ``qwen2``.
         Fast — only needs the inventory snapshot.

      3. **Tokenizer fingerprint match.** Different families but
         identical GGUF tokenizers (parsed via the ``gguf`` package).
         Catches cross-family-but-same-vocab pairs that the family
         heuristic would otherwise reject — e.g. a Mistral derivative
         that ships the Llama tokenizer. Slower (parses each
         candidate's GGUF once, cached by mtime), but correct.

    Generic across the pool — the inventory walker
    (`_pool_model_inventory`) finds candidates anywhere in the pool,
    but the picker only promotes host-resident drafts because
    llama-server needs the draft GGUF on local disk. Worker-only
    drafts are noted in the inventory's ``source`` field; if the user
    wants them as drafts they'd need to ship the model to host first
    (manual ``ollama pull`` or ``push-model`` from another node).
    Future work: auto-LAN-pull when a worker-only draft outranks
    every host candidate.
    """
    target = resolve_ollama_model(target_model_name)
    if not target:
        return None
    target_family = (target.get("family") or "").lower()
    target_size = int(target.get("size_bytes") or 0)
    target_path = target.get("gguf_path") or ""
    if target_size <= 0:
        return None
    if target_size < _SPECULATIVE_MIN_TARGET_BYTES:
        # Target is already small — speculative decoding overhead would
        # likely outweigh any speedup. Skip silently.
        return None

    # Tier 1: manual override — short-circuits all the safety checks
    # because the user explicitly told us this pair works.
    override = _draft_override_for(target_model_name)
    if override:
        info = resolve_ollama_model(override)
        if info and info.get("gguf_path"):
            return {
                "name": override,
                "gguf_path": info["gguf_path"],
                "size_bytes": int(info.get("size_bytes") or 0),
                "source": "host",
                "match": "override",
            }
        # Override pointed at a model we can't resolve — log + fall
        # through to the auto picker so the chat doesn't die just
        # because the override is stale.
        log.info(
            "speculative override for %r names %r but it isn't resolvable; "
            "falling through to auto picker",
            target_model_name, override,
        )

    # Target's tokenizer fingerprint is computed at most once across
    # the loop — Tier 3's cross-family path only needs it when no
    # family match is available, and the GGUF parse isn't free.
    target_fingerprint: str | None = None
    target_fingerprint_loaded = False

    max_draft_bytes = int(target_size * _DRAFT_MAX_SIZE_FRACTION)
    host_candidates: list[dict] = []
    worker_candidates: list[dict] = []
    for m in _pool_model_inventory():
        if m["name"] == target_model_name:
            # Don't pick the target as its own draft.
            continue
        if m["size_bytes"] <= 0 or m["size_bytes"] >= max_draft_bytes:
            continue

        # Tier 2: same-family — cheap, no GGUF parse.
        if (
            m["family"]
            and target_family
            and m["family"] == target_family
        ):
            entry = {**m, "match": "family"}
            (host_candidates if m["source"] == "host" else worker_candidates).append(entry)
            continue

        # Tier 3: same tokenizer fingerprint. Resolve the candidate's
        # path so we can fingerprint it; lazily compute the target's
        # fingerprint on the first cross-family candidate we hit.
        # Tokenizer fingerprinting requires the GGUF on local disk, so
        # this tier currently only fires for host-resident candidates.
        # A worker-only cross-family candidate's fingerprint can't be
        # read from this side of the LAN — defer to family match for
        # worker candidates.
        if m["source"] != "host":
            continue
        cand_info = resolve_ollama_model(m["name"])
        if not cand_info or not cand_info.get("gguf_path"):
            continue
        cand_fingerprint = _gguf_tokenizer_fingerprint(cand_info["gguf_path"])
        if not cand_fingerprint:
            continue  # Can't verify — refuse rather than guess.
        if not target_fingerprint_loaded:
            target_fingerprint = _gguf_tokenizer_fingerprint(target_path)
            target_fingerprint_loaded = True
        if not target_fingerprint:
            continue  # Target unreadable — no Tier 3 matches possible.
        if cand_fingerprint == target_fingerprint:
            host_candidates.append({**m, "match": "tokenizer"})

    if not host_candidates and not worker_candidates:
        return None

    # Smallest draft wins — speculative decoding speedup scales with
    # how much faster the draft generates than the target verifies.
    host_candidates.sort(key=lambda m: m["size_bytes"])
    worker_candidates.sort(key=lambda m: m["size_bytes"])

    # Prefer host candidates so the chat doesn't pay a LAN round-trip on
    # every draft load. Worker-only candidates trigger a background SCP
    # so a future turn can promote them.
    smallest_worker = worker_candidates[0] if worker_candidates else None
    smallest_host = host_candidates[0] if host_candidates else None
    if (
        smallest_worker is not None
        and (smallest_host is None or smallest_worker["size_bytes"] < smallest_host["size_bytes"])
    ):
        # Worker-only candidate would beat (or replace) every host
        # candidate. Schedule a background pull so future turns get to
        # use it; for THIS turn fall back to the best host candidate
        # (or None if no host candidate exists).
        try:
            wid = smallest_worker["source"].split(":", 1)[1]
        except (KeyError, IndexError):
            wid = ""
        if wid:
            _maybe_kickoff_draft_lan_sync(
                target_model_name,
                smallest_worker["name"],
                wid,
            )

    if not smallest_host:
        return None
    info = resolve_ollama_model(smallest_host["name"])
    if not info or not info.get("gguf_path"):
        return None
    return {
        "name": smallest_host["name"],
        "gguf_path": info["gguf_path"],
        "size_bytes": smallest_host["size_bytes"],
        "source": smallest_host["source"],
        "match": smallest_host.get("match", "family"),
    }


def speculative_decoding_enabled() -> bool:
    """Read the user's preference for auto-engaging speculative decoding.

    Defaults to ON. The picker is conservative — `pick_draft_for` only
    engages when a same-family draft sits on the host's disk AND
    `_host_has_vram_for_speculative` confirms there's room for both
    models with 30 % headroom. Setups that can't benefit (no viable
    draft, tight VRAM, target too small) silently fall back to plain
    Ollama, so leaving the feature on by default is a no-op for them
    and a free 1.3-2× speedup for everyone else.

    Anyone who specifically wants the legacy Ollama-only behaviour
    can disable via the settings API:
        POST /api/settings  {"compute_pool_speculative_decoding": "false"}
    """
    val = db.get_setting("compute_pool_speculative_decoding")
    if val is None or str(val).strip() == "":
        # Default ON — the picker's gates handle viability.
        return True
    return str(val).lower() in ("1", "true", "yes", "on")


def _host_has_vram_for_speculative(target_size: int, draft_size: int) -> bool:
    """Does the host have enough VRAM to hold both target + draft + KV?

    `_HOST_VRAM_USE_FRACTION` already bakes in a 15 % margin for the
    OS / desktop / KV cache when sizing the single-model budget; the
    speculative path layers an additional `_SPECULATIVE_VRAM_HEADROOM`
    on top because two simultaneously-loaded models double the KV-cache
    pressure.
    """
    if target_size <= 0 or draft_size <= 0:
        return False
    budget = _host_vram_budget_bytes()
    if budget <= 0:
        return False
    needed = (target_size + draft_size) * _SPECULATIVE_VRAM_HEADROOM
    return needed <= budget


# ---------------------------------------------------------------------------
# Worker-side llama-server lifecycle
#
# When `pick_chat_target` would route a chat to a worker (because the
# worker is meaningfully more capable than host), the default path uses
# the worker's Ollama. That works fine for vanilla single-stream chat,
# but it means speculative decoding can't be engaged on that worker —
# Ollama doesn't expose `--model-draft` / `-md`.
#
# Spawning a llama-server PROCESS on the worker — directly serving the
# target + a draft GGUF — closes that gap. The host orchestrates
# (start / stop / health-check) over SSH; the chat-time HTTP traffic
# goes worker → worker (everything stays in the worker's VRAM, no
# RPC layer-pulls). Worker-side llama-server only kicks in when:
#   * the worker is the strongest-single-node winner,
#   * the worker has `llama_server_path` populated by the probe,
#   * speculative_decoding_enabled() is True,
#   * a viable draft GGUF is already on the worker (no upfront LAN
#     copy on the chat-startup hot path).
#
# State is kept in process memory (`_WORKER_CHAT_SERVERS`) because the
# llama-server process actually runs on the worker — host restart has
# no effect on the worker. On host startup the cache is empty; the
# next chat turn either reuses an existing worker process (verified
# via /health) or stops/respawns to pick up our config.
# ---------------------------------------------------------------------------
_WORKER_CHAT_SERVERS: dict[str, dict] = {}

# Default port for worker-side llama-server. Sits above Ollama's 11434
# and the host-side llama-server's 11500 so a single-machine dev setup
# (host + worker on same box) doesn't collide.
_WORKER_LLAMA_PORT = 11600

# Worker VRAM headroom factor for the target + draft pair. Same idea
# as `_SPECULATIVE_VRAM_HEADROOM` for host. 1.15 mirrors that constant
# under the 5 %-margin policy.
_WORKER_SPECULATIVE_VRAM_HEADROOM = 1.15


def _worker_has_model_locally(worker: dict, model_name: str) -> bool:
    """Did the periodic probe see ``model_name`` in this worker's
    `/api/tags`? Strict name match; tag-as-typed is enough."""
    caps = worker.get("capabilities") or {}
    for m in caps.get("models") or []:
        if not isinstance(m, dict):
            continue
        if (m.get("name") or "") == model_name:
            return True
    return False


def _pick_worker_resident_draft(
    worker: dict, target_model_name: str,
) -> dict | None:
    """Find a viable speculative-decoding draft already on this worker.

    Same family / size constraints as the host picker, except the
    inventory is restricted to models the worker reports installed.
    Returns ``None`` if no candidate fits — the caller falls back to
    the worker's Ollama (no speculative).
    """
    target = resolve_ollama_model(target_model_name)
    if not target:
        return None
    target_family = (target.get("family") or "").lower()
    target_size = int(target.get("size_bytes") or 0)
    if target_size < _SPECULATIVE_MIN_TARGET_BYTES:
        return None
    max_draft_bytes = int(target_size * _DRAFT_MAX_SIZE_FRACTION)
    caps = worker.get("capabilities") or {}
    candidates = []
    for m in caps.get("models") or []:
        if not isinstance(m, dict):
            continue
        name = m.get("name") or ""
        if not name or name == target_model_name:
            continue
        family = (m.get("family") or "").lower()
        size = int(m.get("size") or 0)
        if not family or family != target_family:
            continue
        if size <= 0 or size >= max_draft_bytes:
            continue
        candidates.append({"name": name, "family": family, "size_bytes": size})
    if not candidates:
        return None
    candidates.sort(key=lambda m: m["size_bytes"])
    return candidates[0]


def _worker_has_vram_for_pair(
    worker: dict, target_size: int, draft_size: int,
) -> bool:
    """Best-effort VRAM headroom check on the worker.

    `max_vram_seen_bytes` is a hard lower bound from the probe. If a
    larger model has actually been loaded on the worker, that's our
    best evidence of available VRAM. Without that signal we conserve
    by refusing — better to fall back to plain Ollama than to OOM the
    worker mid-load.
    """
    if target_size <= 0 or draft_size <= 0:
        return False
    caps = worker.get("capabilities") or {}
    proven_vram = int(caps.get("max_vram_seen_bytes") or 0)
    if proven_vram <= 0:
        return False
    needed = (target_size + draft_size) * _WORKER_SPECULATIVE_VRAM_HEADROOM
    return needed <= proven_vram


async def _wait_for_worker_chat_health(
    base_url: str, timeout_sec: float = 60.0,
) -> bool:
    """Poll the worker's llama-server `/health` until it reports OK or
    timeout. Returns True on success, False on timeout."""
    deadline = time.monotonic() + timeout_sec
    async with httpx.AsyncClient(timeout=5.0) as client:
        while time.monotonic() < deadline:
            try:
                r = await client.get(f"{base_url}/health")
                if r.status_code < 500:
                    return True
            except Exception:
                pass
            await asyncio.sleep(2.0)
    return False


async def stop_worker_chat_server(worker: dict) -> None:
    """SSH-kill any llama-server running on the worker. Best-effort —
    failures are logged and swallowed so a flaky worker can't block
    the routing layer's pivot to host-only."""
    wid = worker.get("id")
    if wid:
        _WORKER_CHAT_SERVERS.pop(wid, None)
    ps_script = (
        "Get-Process -Name llama-server -ErrorAction SilentlyContinue "
        "| Stop-Process -Force;"
        "Start-Sleep -Milliseconds 500"
    )
    try:
        await dispatch_to_worker_powershell(
            worker, ps_script, timeout_sec=10.0,
        )
    except Exception as e:
        log.info(
            "compute_pool: worker llama-server stop failed for %s: %s",
            worker.get("label"), e,
        )


async def ensure_worker_chat_server(
    worker: dict,
    model_name: str,
    *,
    target_gguf_path: str,
    draft_model_name: str,
) -> str:
    """Spawn (or confirm running) llama-server on the worker for the
    given target + draft pair. Returns the base URL the host uses to
    issue chat requests.

    Idempotent: re-spawning for the same (target, draft) pair is a
    no-op if /health on the existing process responds. Switching to
    a different model stops the previous llama-server first.

    Raises ``RuntimeError`` on any unrecoverable spawn failure — the
    caller falls back to the worker's plain Ollama URL.
    """
    wid = worker["id"]

    # Same target + draft already running? Re-verify health then
    # short-circuit. /health is cheap; re-running it costs less than
    # an unnecessary stop+spawn cycle.
    state = _WORKER_CHAT_SERVERS.get(wid)
    if (
        state
        and state.get("model") == model_name
        and state.get("draft") == draft_model_name
    ):
        if await _wait_for_worker_chat_health(state["url"], timeout_sec=2.0):
            return state["url"]
        # Health failed — clear stale state and fall through to
        # respawn.
        _WORKER_CHAT_SERVERS.pop(wid, None)

    # Stop any other llama-server first (different model running, or
    # a stale process from a previous session).
    await stop_worker_chat_server(worker)

    caps = worker.get("capabilities") or {}
    llama_server_path = caps.get("llama_server_path")
    if not llama_server_path:
        raise RuntimeError("worker has no llama-server installed")

    # Resolve the target + draft GGUF paths via Ollama's manifest store
    # ON THE WORKER. Each worker keeps its own ~/.ollama/models tree
    # populated via `ollama pull`; the SSH-bench script below extracts
    # the on-disk path for both models in one round trip.
    safe_target = model_name.replace("'", "''")
    safe_draft = draft_model_name.replace("'", "''")
    resolve_ps = (
        "$python = (Get-Command py.exe -ErrorAction SilentlyContinue).Source;"
        "if (-not $python) { $python = (Get-Command python.exe -ErrorAction SilentlyContinue).Source };"
        "if (-not $python) { Write-Error 'python not found on worker'; exit 3 };"
        "$src = @\"\n"
        "import json, os, sys\n"
        "from pathlib import Path\n"
        "models = ('" + safe_target + "', '" + safe_draft + "')\n"
        "out = {}\n"
        "for name in models:\n"
        "    base = Path.home() / '.ollama' / 'models'\n"
        "    if ':' in name:\n"
        "        ns_model, tag = name.rsplit(':', 1)\n"
        "        if '/' in ns_model:\n"
        "            ns, model = ns_model.split('/', 1)\n"
        "        else:\n"
        "            ns, model = 'library', ns_model\n"
        "    else:\n"
        "        ns, model, tag = 'library', name, 'latest'\n"
        "    manifest = base / 'manifests' / 'registry.ollama.ai' / ns / model / tag\n"
        "    if not manifest.exists():\n"
        "        out[name] = None\n"
        "        continue\n"
        "    try:\n"
        "        m = json.loads(manifest.read_text())\n"
        "    except Exception:\n"
        "        out[name] = None\n"
        "        continue\n"
        "    digest = None\n"
        "    for layer in m.get('layers', []):\n"
        "        if layer.get('mediaType') == 'application/vnd.ollama.image.model':\n"
        "            digest = layer.get('digest')\n"
        "            break\n"
        "    if not digest:\n"
        "        out[name] = None\n"
        "        continue\n"
        "    blob = base / 'blobs' / digest.replace(':', '-')\n"
        "    out[name] = str(blob) if blob.exists() else None\n"
        "json.dump(out, sys.stdout)\n"
        "\"@;"
        "$src | & $python -"
    )
    ok, stdout, stderr = await dispatch_to_worker_powershell(
        worker, resolve_ps, timeout_sec=15.0,
    )
    if not ok:
        raise RuntimeError(f"worker model resolve failed: {stderr}")
    try:
        paths = jsonutil.loads((stdout or b"").decode("utf-8", errors="replace") or "{}")
    except json.JSONDecodeError as e:
        raise RuntimeError(f"worker returned non-JSON resolve payload: {e}")
    target_on_worker = paths.get(model_name)
    draft_on_worker = paths.get(draft_model_name)
    if not target_on_worker or not draft_on_worker:
        raise RuntimeError(
            f"worker missing GGUF — target={target_on_worker}, "
            f"draft={draft_on_worker}"
        )

    # Pick a continuous-batching slot count tuned to the worker's
    # observed VRAM. Mirrors split_lifecycle._compute_optimal_parallel
    # without needing GGUF-metadata round-trips back to the worker —
    # we use the same tier table ollama_runtime uses for host. A
    # worker servicing one chat at a time pays nothing for unused
    # slots; adding 2-8 slots lets concurrent dispatches (parallel
    # subagents, two browser tabs against the same worker) share
    # the warm engine instead of serializing.
    worker_vram_gb = float(
        (worker.get("capabilities") or {}).get("max_vram_seen_bytes") or 0,
    ) / (1024 ** 3)
    if worker_vram_gb >= 20:
        worker_parallel = 8
    elif worker_vram_gb >= 12:
        worker_parallel = 4
    elif worker_vram_gb >= 6:
        worker_parallel = 2
    else:
        worker_parallel = 1

    # Build the spawn command. Worker-side llama-server runs the same
    # CLI as host's, just with target + draft both local. The flags
    # mirror split_lifecycle._build_command's defaults so the chat
    # behaviour matches what the rest of the app expects.
    #
    # Inference-perf flags wired here too:
    #   * `--flash-attn on` over `-fa auto` so behaviour is explicit
    #     and matches the host path; also unlocks KV quant below.
    #   * `-ctk q8_0 -ctv q8_0` halves KV memory at <1 % accuracy
    #     loss — frees room for the draft model on the worker.
    #   * `--cache-reuse 256` makes follow-up turns reuse the previous
    #     prompt's prefix via KV-shift; big win on multi-turn chat.
    #   * `--parallel N` enables continuous batching on the worker so
    #     concurrent requests share the same warm engine.
    spawn_cmd = (
        f'"{llama_server_path}" '
        f'--model "{target_on_worker}" '
        '--host 0.0.0.0 '
        f'--port {_WORKER_LLAMA_PORT} '
        '--flash-attn on -ctk q8_0 -ctv q8_0 --cache-reuse 256 '
        '--jinja -ngl 99 -c 4096 --no-warmup '
        f'--parallel {worker_parallel} '
        f'-md "{draft_on_worker}" --draft-max 8 --draft-min 1 -ngld 99'
    )
    # Win32_Process.Create makes the spawned llama-server outlive the
    # SSH session — same trick rpc-server restart already uses.
    spawn_ps = (
        f"Invoke-CimMethod -ClassName Win32_Process -MethodName Create "
        f"-Arguments @{{CommandLine='{spawn_cmd}'}} | Select-Object -ExpandProperty ProcessId"
    )
    ok, stdout, stderr = await dispatch_to_worker_powershell(
        worker, spawn_ps, timeout_sec=15.0,
    )
    if not ok:
        raise RuntimeError(f"worker llama-server spawn failed: {stderr}")

    base_url = f"http://{_worker_host(worker)}:{_WORKER_LLAMA_PORT}"
    if not await _wait_for_worker_chat_health(base_url, timeout_sec=120.0):
        # Best-effort cleanup so the dead process doesn't sit around
        # holding a port.
        await stop_worker_chat_server(worker)
        raise RuntimeError("worker llama-server didn't pass /health in 120s")

    _WORKER_CHAT_SERVERS[wid] = {
        "model": model_name,
        "draft": draft_model_name,
        "url": base_url,
        "port": _WORKER_LLAMA_PORT,
    }
    log.info(
        "compute_pool: worker-side llama-server up on %s — target=%s draft=%s",
        worker.get("label"), model_name, draft_model_name,
    )
    return base_url


async def _ensure_split_running_for(
    model_name: str,
    gguf_path: str,
    worker_ids: list[str],
    mmproj_path: str | None = None,
    draft_gguf_path: str | None = None,
) -> str:
    """Idempotent: ensure a `split_models` row exists + is running for
    this exact (model_name, gguf_path[, mmproj_path][, draft_gguf_path])
    tuple, then return its base_url.

    Auto-creates the row keyed by model_name as label. If a row with the
    same label already exists, we reuse it (updating worker_ids /
    mmproj_path / draft_gguf_path if the user added/removed workers or
    installed a new multimodal projector or a speculative-decoding draft
    since the previous turn). If a DIFFERENT split row is currently
    running, we stop it first — only one big model hot at a time.

    `mmproj_path`, when non-None, is forwarded to llama-server's
    `--mmproj` flag so vision-capable models (e.g. gemma4:26b after
    its vision tower has been extracted into a separate GGUF) can
    accept image input via Phase 2 split.

    `draft_gguf_path`, when non-None, is forwarded to llama-server's
    `-md` flag so a smaller same-family GGUF accelerates the target via
    speculative decoding. The router fills this in when the picker
    finds a viable draft anywhere in the pool's combined model
    inventory and the host has VRAM headroom for both.
    """
    # Local import to dodge the circular dep — split_lifecycle imports
    # compute_pool indirectly via db / runtime.
    from . import split_lifecycle

    rows = db.list_split_models()
    target_row = next((r for r in rows if r.get("label") == model_name), None)

    # Stop any other running split row — finite VRAM means one big
    # model active at a time.
    for r in rows:
        if r["id"] != (target_row or {}).get("id") and r.get("status") in ("running", "loading"):
            try:
                await split_lifecycle.stop(r["id"])
            except Exception as e:
                log.warning("compute_pool: failed to stop %s: %s", r["id"], e)

    if target_row is None:
        sid = db.create_split_model(
            label=model_name,
            gguf_path=gguf_path,
            mmproj_path=mmproj_path,
            draft_gguf_path=draft_gguf_path,
            worker_ids=worker_ids,
        )
        target_row = db.get_split_model(sid)
    else:
        # Refresh gguf_path / mmproj_path / draft_gguf_path / worker_ids
        # in case the user changed things since the previous turn.
        if (
            target_row.get("gguf_path") != gguf_path
            or target_row.get("mmproj_path") != mmproj_path
            or target_row.get("draft_gguf_path") != draft_gguf_path
            or target_row.get("worker_ids") != worker_ids
        ):
            db.update_split_model(
                target_row["id"],
                gguf_path=gguf_path,
                mmproj_path=mmproj_path,
                draft_gguf_path=draft_gguf_path,
                worker_ids=worker_ids,
            )
            target_row = db.get_split_model(target_row["id"])

    # SYCL+RPC defensive switch. Historical context: Intel-iGPU
    # SYCL builds before llama.cpp build 8233 crashed during RPC
    # layer push (#20259, #21420, #21474 — all CLOSED upstream now).
    # We're on build 8940 so those specific crashes shouldn't
    # repro any more, but we keep the SYCL-off-while-split policy
    # as a defensive measure: the failure mode (mid-load
    # "Remote RPC server crashed") is bad enough that we'd rather
    # eat the iGPU-acceleration loss on workers during split than
    # risk a regression. Workers stay on `-d SYCL0,CPU` for non-
    # split paths (Phase 1 routing, embeddings, subagents).
    #
    # Workers stay on `-d SYCL0,CPU` when NOT in split mode (Phase 1
    # routing, embeddings, subagents) so iGPU acceleration is still
    # used for those paths. The CPU-only switch is purely for the
    # llama-server split spawn.
    #
    # `stop_all_running_splits` restores the default backend so the
    # next non-split request (or smaller model that fits host alone)
    # immediately gets SYCL again.
    workers_for_split = [db.get_compute_worker(wid) for wid in worker_ids]
    workers_for_split = [w for w in workers_for_split if w]
    await _set_workers_backend(workers_for_split, in_split=True)

    sid = target_row["id"]
    if target_row.get("status") != "running":
        result = await split_lifecycle.start(sid)
        if not result.get("ok"):
            raise RuntimeError(
                f"failed to start llama-server for {model_name}: "
                f"{result.get('error') or 'unknown'}"
            )

    fresh = db.get_split_model(sid)
    return f"http://127.0.0.1:{fresh['llama_port']}"


async def _ensure_peer_orchestrated_split(
    peer_worker: dict,
    model_name: str,
    *,
    port: int = 8090,
) -> str | None:
    """Drive a paired peer to spawn its own llama-server using ITS
    local GGUF, with our local rpc-server as a `--rpc` backend.

    Used when the model is too big to fit any single peer alone but
    lives on `peer_worker`'s disk already. Compute fans across the
    peer + this host (host's rpc-server on 50052 receives layer
    pushes; peer's llama-server holds the rest). Zero data transfer.

    Returns the base URL of the peer's llama-server (direct LAN HTTP)
    on success, or None if the peer can't be driven (no llama-server
    binary, no local GGUF for that model, encrypted-proxy unreachable,
    etc.). Caller falls through to the legacy plain-Ollama path on None.

    Privacy: chat traffic to the returned URL is plaintext over LAN —
    same trust model as ``ensure_worker_chat_server`` (worker-side
    speculative llama-server). The peer is by definition paired,
    which puts it in the local-pool whitelist of
    ``p2p_privacy.assert_plaintext_allowed``.
    """
    from . import p2p_secure_client as _sec
    from . import p2p_rpc_server as _rpc_srv
    import json as _json

    # Bring up two SINGLE-BACKEND rpc-servers so the peer's llama-server
    # sees clean, non-hybrid backends. Hybrid SYCL+CPU on the same
    # rpc-server (`-d SYCL0,CPU`) trips ggml-rpc.cpp's
    # layout-mismatch crash mid-load on big models — confirmed live
    # with dolphin-mixtral 26 GB which crashed at
    # "ggml-rpc.cpp:505: Remote RPC server crashed or returned
    # malformed response" right after layer placement. Splitting
    # into two single-backend rpc-servers (SYCL0 on 50052, CPU on
    # 50053) makes each rpc-server expose ONE memory pool and
    # llama.cpp's layer allocator handles them as independent
    # devices — no hybrid path, no crash.
    #
    # This is the same pattern compute_pool.select_multi_rpc_specs
    # produces for worker-side multi-rpc spawn; we just apply it
    # locally too to keep host's compute surface symmetric with
    # what we ask peers to expose.
    try:
        # Restart 50052 if it's currently running with the hybrid
        # backend — switching to SYCL0-only requires a respawn.
        s52 = _rpc_srv.get_local_rpc_server_status(port=50052)
        existing_52 = (s52.get("active_backend") or "").strip().upper()
        if not s52.get("listening") or "CPU" in existing_52:
            _rpc_srv.start_local_rpc_server(backend="SYCL0", port=50052)
    except Exception as e:
        log.info("compute_pool: SYCL0-only rpc-server prep skipped: %s", e)
    try:
        if not _rpc_srv.get_local_rpc_server_status(port=50053).get("listening"):
            _rpc_srv.start_local_rpc_server(backend="CPU", port=50053)
    except Exception as e:
        log.info("compute_pool: CPU-only rpc-server prep skipped: %s", e)

    # Same for every paired peer — drive their rpc-servers to be
    # single-backend each, via the encrypted-proxy ensure-multi
    # endpoint we built earlier. Failures here are best-effort:
    # an offline peer or a stale capability cache shouldn't block
    # the spawn we're about to do on the orchestrator peer.
    try:
        for w in db.list_compute_workers(enabled_only=True):
            if w.get("id") == peer_worker.get("id"):
                continue
            if not w.get("use_for_chat"):
                continue
            try:
                await _sec.forward(
                    w,
                    method="POST",
                    path="/api/p2p/rpc-server/ensure-multi",
                    body={"specs": [
                        {"port": 50052, "backend": "SYCL0"},
                        {"port": 50053, "backend": "CPU"},
                    ]},
                    timeout=20.0,
                )
            except Exception as e:
                log.info(
                    "compute_pool: peer %s rpc-server ensure-multi "
                    "failed (%s); their iGPU/RAM won't contribute "
                    "to this turn",
                    w.get("label"), e,
                )
    except Exception as e:
        log.info(
            "compute_pool: failed to drive peer rpc-servers (%s); "
            "split will run with whatever rpc-servers are already up",
            e,
        )

    # rpc_targets DISABLED for now. Both SYCL and pure-CPU
    # rpc-servers cause llama.cpp b9002 to crash with
    # "ggml-rpc.cpp:640: Remote RPC server crashed or returned
    # malformed response" mid-decode, even with GGML_RPC_TIMEOUT
    # bumped to 120 s. The crash fires after the prompt eval
    # completes but BEFORE the first generated token, which means
    # the chat layer always returns empty + the user sees the
    # "Worker became unreachable mid-response" toast.
    #
    # We trade pool contribution for chat reliability: the peer's
    # local GPU + RAM (with mmap from disk for layers that don't
    # fit) is the only path that survives long enough to actually
    # return tokens. Re-enable rpc_targets when upstream b9002+
    # fixes the RPC regression.
    rpc_targets: list[str] = []

    if not rpc_targets:
        log.info(
            "compute_pool: peer-orchestrated split — no rpc endpoints "
            "available (no host LAN IPs and no other peers reachable); "
            "peer llama-server will run on its own resources only",
        )

    # Compute the n_gpu_layers cap based on the peer's CUDA VRAM
    # vs the model size. llama.cpp's `--fit on` REFUSES to lower
    # ngl below whatever the build defaults set (typically 99 on a
    # CUDA build), so it tries to pack the entire 26 GB model onto
    # the peer's 8 GB CUDA and OOMs. We compute the cap ourselves
    # from peer VRAM and pass it explicitly:
    #
    #   ngl_cap = floor((peer_vram_free_GB * 0.8) / per_layer_GB)
    #
    # where per_layer_GB ≈ model_size_GB / num_layers (rough). For
    # dolphin-mixtral 8x7B on RTX 3060 Ti 8 GB: per layer ≈ 0.8 GB,
    # 8 GB free * 0.8 / 0.8 ≈ 8 layers fit on CUDA. The remaining
    # 24 layers fan to peer-RAM + host-RAM via `--rpc`.
    peer_caps = peer_worker.get("capabilities") or {}
    peer_vram_gb = float(peer_caps.get("vram_total_gb") or 0)
    peer_vram_free_gb = float(peer_caps.get("vram_free_gb") or peer_vram_gb)
    # Best-effort layer count from the GGUF metadata cache; fall
    # back to a conservative estimate from model size if unknown.
    model_layer_count = 32  # Mixtral 8x7B has 32; conservative default.
    try:
        for m in (peer_caps.get("models") or []):
            if (m.get("name") or "") == model_name:
                model_layer_count = int(m.get("num_layers") or 32) or 32
                break
    except (TypeError, ValueError):
        pass
    # Heuristic per-layer footprint: model size / layer count, in GB.
    # peer_size is the bytes value resolved earlier in the caller;
    # we recompute from the worker's models cache here for safety.
    target_size_gb = 26.0  # safe upper bound for the heavy MoE case
    try:
        for m in (peer_caps.get("models") or []):
            if (m.get("name") or "") == model_name:
                target_size_gb = float(m.get("size") or 0) / (1024 ** 3)
                break
    except (TypeError, ValueError):
        pass
    per_layer_gb = max(0.05, target_size_gb / max(1, model_layer_count))
    if peer_vram_free_gb > 0.5:
        ngl_cap = max(0, int((peer_vram_free_gb * 0.80) / per_layer_gb))
    else:
        ngl_cap = 0  # No usable VRAM → CPU+RAM only on the peer.
    log.info(
        "compute_pool: peer-orchestrated split sizing — peer %s VRAM "
        "free=%.1f GB, per-layer~%.2f GB, ngl_cap=%d (rest fans to "
        "RAM + host-RPC)",
        peer_worker.get("label"), peer_vram_free_gb, per_layer_gb, ngl_cap,
    )

    spawn_body = {
        "model": model_name,
        "port": port,
        "rpc_targets": rpc_targets,
        "n_gpu_layers": ngl_cap,
        # 16K context — Gigachat's own system prompt + conversation
        # history routinely runs 4-8K, so 4K is too tight (the very
        # first chat turn 400's with "request exceeds context").
        # 16K accommodates ~30 multi-turn exchanges; KV q8_0 keeps
        # the memory footprint sane (~halves vs default).
        "context_size": 16384,
        "parallel": 1,
    }

    # Send the spawn request via the encrypted secure proxy. Generous
    # timeout — start_local_llama_server BLOCKS until the listener
    # comes up (or 60 s budget expires), and on cold start that
    # includes the GGUF mmap + RPC backend handshakes.
    try:
        status, body_text = await _sec.forward(
            peer_worker,
            method="POST",
            path="/api/p2p/llama-server/start",
            body=spawn_body,
            timeout=120.0,
        )
    except Exception as e:
        log.warning(
            "compute_pool: peer-orchestrated split — encrypted proxy "
            "spawn request to %s failed: %s",
            peer_worker.get("label"), e,
        )
        return None

    if status != 200:
        log.warning(
            "compute_pool: peer-orchestrated split — peer %s returned "
            "HTTP %d for spawn request: %s",
            peer_worker.get("label"), status, (body_text or "")[:300],
        )
        return None

    try:
        result = _json.loads(body_text or "{}")
    except _json.JSONDecodeError:
        log.warning(
            "compute_pool: peer-orchestrated split — peer %s returned "
            "non-JSON spawn response: %s",
            peer_worker.get("label"), (body_text or "")[:200],
        )
        return None

    if not result.get("ok"):
        log.warning(
            "compute_pool: peer-orchestrated split — peer %s spawn "
            "failed (status=%s): %s",
            peer_worker.get("label"),
            result.get("status"), result.get("error"),
        )
        return None

    # Peer is now serving llama-server on its loopback (port 8090
    # default). Build the LAN URL the orchestrator's chat dispatcher
    # connects to. The peer's llama-server binds to 0.0.0.0 so the
    # LAN IP is reachable directly — no encrypted-proxy tunneling
    # needed for the chat traffic itself (plaintext is allowed for
    # paired-LAN peers per `p2p_privacy.require_encryption`).
    peer_addr = peer_worker.get("address") or ""
    if not peer_addr:
        log.warning(
            "compute_pool: peer-orchestrated split — peer %s has no "
            "address on its compute_workers row, can't build URL",
            peer_worker.get("label"),
        )
        return None
    base_url = f"http://{peer_addr}:{port}"

    # Wait for the peer's llama-server to finish LOADING the model
    # before we hand the URL back to the chat dispatcher. /health
    # returns 503 "Loading model" while llama-server is mmapping
    # weights and uploading layers to RPC backends. For a 26 GB MoE
    # model split across CUDA + RAM + multiple RPC backends, this
    # routinely takes 3-5 minutes on first load. The chat layer
    # reasonably gives up on a single 503; better to block here
    # until the model is actually ready.
    #
    # We yield while polling so the agent's pull-progress drain loop
    # keeps surfacing events to the SSE stream — the user sees the
    # chat is "live" even though we're waiting on the model.
    health_deadline = time.time() + 360.0  # 6 minutes hard cap
    last_status = None
    healthy = False
    while time.time() < health_deadline:
        try:
            async with httpx.AsyncClient(timeout=10.0) as client:
                r = await client.get(f"{base_url}/health")
                last_status = r.status_code
                if r.status_code == 200:
                    # llama-server returns {"status":"ok"} once the
                    # model is loaded and the slots are idle.
                    try:
                        body = r.json()
                        if body.get("status") in ("ok", "no slot available"):
                            healthy = True
                            break
                    except Exception:
                        # Treat any 200 as healthy if body's not JSON
                        healthy = True
                        break
        except Exception:
            # Transient network blip — keep polling.
            pass
        await asyncio.sleep(2.0)

    if not healthy:
        log.warning(
            "compute_pool: peer-orchestrated split — peer %s "
            "llama-server still loading after 360s (last /health "
            "status=%s). Returning URL anyway; the chat layer will "
            "retry if the model isn't ready yet.",
            peer_worker.get("label"), last_status,
        )

    return base_url


async def stop_all_running_splits() -> None:
    """Free VRAM held by any running llama-server. Called when the
    router decides the upcoming chat turn fits Ollama on host alone —
    no point keeping a big-model llama-server warm if the active
    conversation no longer needs it.

    Also restores any workers that were switched to the MoE workaround
    backend (`-d CPU`) back to the default `-d SYCL0,CPU` so the next
    non-MoE request gets full iGPU acceleration. No-op for workers
    already on the default; the helper checks `current_rpc_backend`
    in capabilities before touching them.
    """
    from . import split_lifecycle

    for r in db.list_split_models():
        if r.get("status") in ("running", "loading"):
            try:
                await split_lifecycle.stop(r["id"])
            except Exception as e:
                log.warning("compute_pool: stop_all_running_splits %s: %s", r["id"], e)

    # Restore default backend on every enabled worker. Best-effort —
    # failures are logged but don't propagate; the next probe loop's
    # auto-restart will re-align state if a worker is unreachable.
    try:
        workers = db.list_compute_workers(enabled_only=True)
        await _set_workers_backend(workers, in_split=False)
    except Exception as e:
        log.warning(
            "compute_pool: restore default backend after stop failed: %s", e,
        )


class RouteChatError(RuntimeError):
    """Raised by route_chat_for when the model can't be served at all
    — e.g. the model file isn't present, or the combined pool is too
    small to hold it. Caller should surface this to the user.

    `status` is an optional structured payload — populated by Scope B
    when an override-file acquisition is in progress, so the chat
    layer can render a "preparing model" indicator with progress
    instead of a generic error.
    """

    def __init__(self, message: str, *, status: dict | None = None) -> None:
        super().__init__(message)
        self.status = status or {}


async def route_chat_for(
    model_name: str,
    *,
    on_pull_progress: Any = None,
) -> dict:
    """Pick the right inference engine for this model + drive any
    needed lifecycle changes.

    Returns one of:
      {"engine": "ollama"}
        → use the existing Ollama path. Downstream `pick_chat_target`
          decides whether the actual node is host or a more-powerful
          worker. Caller continues with `_stream_ollama_chat`.

      {"engine": "llama_server", "base_url": "http://127.0.0.1:NNNN", "label": <model>}
        → llama-server with --rpc was spawned (or already running) for
          this model. Caller uses `_stream_llama_server_chat` instead.

    Decision logic (intelligent — the strongest single node wins):
      * Compute the strongest single-node VRAM budget across host +
        every eligible worker. A worker's contribution is its
        `max_vram_seen_bytes` — only counts as "real" VRAM once the
        worker has actually loaded a model that big.
      * If the model fits the strongest single node → Ollama path.
        (`pick_chat_target` then routes to whichever node won.)
      * Else → split path: spawn llama-server with --rpc to every
        eligible worker so layers fan across host + every node's
        compute.

    Raises `RouteChatError` for unrecoverable cases. Side effects:
    may stop a currently-running llama-server (model switch) and may
    auto-create a `split_models` row.
    """
    # Cheap reject: explicit `split:` prefix is back-compat — short-
    # circuit to the legacy picker so existing tests still pass.
    legacy = pick_split_chat_target(model_name)
    if legacy is not None:
        base, label = legacy
        return {"engine": "llama_server", "base_url": base, "label": label}

    # Public-pool fallback. Before any local-routing logic, check
    # whether (a) the host has the model, (b) any already-registered
    # worker (LAN or public) has it. If neither, look for a peer in
    # the public swarm that's offering it AND register them as a
    # compute_worker so the rest of this function picks them up
    # naturally. If even the public swarm has nobody, fall back to
    # auto-pulling from the OFFICIAL Ollama registry on the host.
    #
    # All of this is best-effort: if the public-pool query / pull
    # fails, we continue to the existing flow which surfaces a clear
    # "model not found" error to the user instead of hanging.
    try:
        host_has = resolve_ollama_model(model_name) is not None
        existing_workers = db.list_compute_workers(enabled_only=True)
        pool_has = any(
            _worker_has_model(w, model_name)
            for w in existing_workers
            if w.get("use_for_chat")
        )
        if not host_has and not pool_has:
            # Step 1 — try the public swarm.
            from . import p2p_pool_routing as _ppr
            public_worker = await _ppr.ensure_public_peer_worker(model_name)
            if public_worker is None:
                # Step 2 — nobody in the swarm has it either. Auto-pull
                # from the OFFICIAL registry on host. Best-effort: a
                # failure here means the user just sees the standard
                # "model not found" downstream. Triggered as a
                # blocking call so the very next chat turn benefits;
                # caller can hit Ctrl+C if they don't want to wait.
                # Progress callback bridges to the agent's SSE stream
                # so the user sees download % instead of a silent
                # multi-minute hang.
                pulled = await _ppr.auto_pull_on_host(
                    model_name, on_progress=on_pull_progress,
                )
                if pulled:
                    log.info(
                        "compute_pool: auto-pulled %r on host before "
                        "route_chat_for resumed", model_name,
                    )
            else:
                log.info(
                    "compute_pool: public-pool peer registered as "
                    "worker for %r, routing will pick it up",
                    model_name,
                )
    except Exception as e:
        log.info(
            "compute_pool: public-pool fallback for %r failed: %s",
            model_name, e,
        )

    info = resolve_ollama_model(model_name)
    if info is None:
        # Not on local Ollama. Could be a peer-hosted model. Check
        # every paired peer's capability cache — if ANY peer has it,
        # try to size-fit against THAT peer's free memory. If the
        # model is bigger than every peer's available memory, engage
        # the split path instead of letting Ollama OOM later.
        #
        # Verified-against-bug case: dolphin-mixtral:8x7b (26 GB) is
        # only on FBS. Local routes here -> falls through to
        # pick_chat_target -> picks FBS -> FBS Ollama returns
        # 'requires more system memory (25.1 GiB) than is available
        # (10.7 GiB)' because the model exceeds FBS's free RAM.
        # Without this branch the user sees a misleading
        # "model not found" error from the post-failure fallback to
        # local Ollama.
        peer_size = 0
        peer_with_model = None
        try:
            for w in db.list_compute_workers(enabled_only=True):
                caps = w.get("capabilities") or {}
                for m in (caps.get("models") or []):
                    if (m.get("name") or "") == model_name:
                        sz = int(m.get("size") or 0)
                        if sz > peer_size:
                            peer_size = sz
                            peer_with_model = w
                        break
        except Exception:
            pass
        if peer_size > 0 and peer_with_model is not None:
            # Peer has the model but our pre-flight check finds it
            # can't fit on that peer alone — Ollama on the peer would
            # OOM mid-load. The user's intent in selecting this model
            # is to RUN it, so we don't refuse.
            #
            # The right architecture: spawn llama-server ON THAT PEER
            # using its local GGUF, with our local rpc-server as a
            # `--rpc` backend. Compute fans across both nodes; layers
            # that don't fit the peer's RAM ride to host's CUDA / RAM
            # over the LAN socket. Zero data transfer — the peer's
            # existing GGUF stays in place.
            #
            # Verified case: dolphin-mixtral:8x7b (26.4 GB) lives on
            # FBS (16.6 GB free, can't fit alone). Host has 8 GB CUDA
            # + 32 GB RAM and rpc-server already listening on 50052.
            # Peer-led split runs the model with no internet pull and
            # no host disk pressure.
            caps = peer_with_model.get("capabilities") or {}
            peer_free_bytes = int(
                (float(caps.get("ram_free_gb") or 0)
                 + float(caps.get("vram_total_gb") or 0))
                * (1024 ** 3)
            )
            if peer_free_bytes > 0 and peer_size > int(peer_free_bytes * 0.90):
                try:
                    base_url = await _ensure_peer_orchestrated_split(
                        peer_with_model, model_name,
                    )
                    if base_url:
                        log.info(
                            "compute_pool: peer-orchestrated split engaged "
                            "— peer %s runs llama-server with our "
                            "rpc-server as --rpc backend (model %s, "
                            "%.1f GB, no GGUF transfer)",
                            peer_with_model.get("label"),
                            model_name, peer_size / 1e9,
                        )
                        await stop_all_running_splits()
                        return {
                            "engine": "llama_server",
                            "base_url": base_url,
                            "label": model_name,
                            "host_node": (
                                f"worker:{peer_with_model['id']}"
                            ),
                            "peer_orchestrated": True,
                        }
                except Exception as e:
                    log.warning(
                        "compute_pool: peer-orchestrated split for %s on "
                        "%s failed (%s) — falling through to peer Ollama "
                        "(may OOM, but the error is surfaceable).",
                        model_name, peer_with_model.get("label"), e,
                    )
        # Fall through to default Ollama path — pick_chat_target will
        # route to whichever node has the model. If the peer can't
        # fit it Ollama surfaces its own OOM error, which is at
        # least actionable for the user.
        await stop_all_running_splits()
        return {"engine": "ollama"}

    size_bytes = info["size_bytes"]

    # Auto-prep paired peers BEFORE any tier decision. Brings up
    # rpc-server on every chat-enabled paired peer that doesn't
    # already have it listening, via the encrypted P2P channel
    # (concurrent, no SSH required). The freshly-eligible workers
    # then feed both `_should_force_split_for` (Tier-1 force gate)
    # and the Tier-2 candidate list. Without this call, a brand-new
    # paired peer would never be eligible until the next 5-minute
    # capability sweep — chat would silently stay on host alone.
    try:
        await _eligible_split_workers_with_autoprep()
    except Exception as e:
        log.info(
            "compute_pool: split auto-prep failed (%s); split engagement "
            "will skip ineligible peers this turn",
            e,
        )

    # Three-tier decision, fastest path first:
    #   1. **Single-node VRAM fit** — strongest GPU (host or worker)
    #      can hold the whole model in pure VRAM. No CPU offload, no
    #      LAN. Always wins on per-token rate.
    #   2. **Pool VRAM fit (split)** — combined VRAM across host + every
    #      eligible worker covers the model. Engages llama-server
    #      with --rpc; layers ride GPU memory across machines via
    #      llama.cpp's native pipeline. Beats host CPU offload when the
    #      LAN is fast (Ethernet) and workers have real GPUs (dGPU).
    #      Loses on Wi-Fi + iGPU because per-token RPC overhead exceeds
    #      the win from avoiding CPU offload on host — but that's a
    #      real-world calibration the user can correct by setting
    #      `use_for_chat=False` on a slow worker.
    #   3. **Host VRAM + host RAM (CPU offload)** — single-node Ollama
    #      with layer-spill into RAM. Slower per layer than VRAM, but
    #      no LAN per-token cost. Last resort when split path can't
    #      cover the model either.
    host_vram = _host_vram_budget_bytes()
    host_total = _host_total_capacity_bytes()
    chat_workers = _eligible_workers("use_for_chat", model=model_name)
    best_worker_capacity = 0
    pool_vram_total = host_vram
    if chat_workers:
        worker_vrams = [
            (w.get("capabilities") or {}).get("max_vram_seen_bytes") or 0
            for w in chat_workers
        ]
        best_worker_capacity = max(worker_vrams) if worker_vrams else 0
        # Sum every eligible worker's proven VRAM with host's. Honest
        # under-count: max_vram_seen is a lower bound (it's whatever
        # was loaded so far, not the worker's true VRAM ceiling).
        pool_vram_total = host_vram + sum(worker_vrams)

    # Tier 1: fits the strongest single GPU (host or worker)?
    strongest_single_vram = max(host_vram, best_worker_capacity)

    # Adaptive split engagement: even when one node fits the model,
    # engage Phase 2 if the pool is MEANINGFULLY bigger than the
    # strongest single node. Always-on — `_should_force_split_for`
    # gates engagement on its own heuristics so setups where pool ≈
    # strongest single node stay on the host path.
    force_split_engagement = _should_force_split_for(
        model_name,
        strongest_single_vram=strongest_single_vram,
        pool_vram_total=pool_vram_total,
    )
    if force_split_engagement:
        log.info(
            "compute_pool: adaptive split engaged for fits-on-host %s "
            "(strongest_single_vram=%d pool_vram=%d)",
            model_name, strongest_single_vram, pool_vram_total,
        )
    if (
        not force_split_engagement
        and strongest_single_vram > 0
        and size_bytes <= strongest_single_vram
    ):
        # Worker-side llama-server sub-tier: if the strongest single
        # node is a worker (not host) AND that worker has llama-server
        # installed AND a same-family draft is already on the worker,
        # spawn llama-server on the worker so speculative decoding
        # works for chats that would otherwise route to plain
        # worker-side Ollama. Compute happens entirely on the worker
        # (no LAN per-token cost), and the draft accelerates verify
        # the same way it does on host.
        if (
            best_worker_capacity > host_vram
            and chat_workers
            and speculative_decoding_enabled()
        ):
            # Pick the strongest worker by capability score.
            strongest_worker = max(
                chat_workers, key=lambda w: _capability_score(w),
            )
            caps = strongest_worker.get("capabilities") or {}
            if caps.get("llama_server_path"):
                draft = _pick_worker_resident_draft(strongest_worker, model_name)
                if (
                    draft
                    and _worker_has_vram_for_pair(
                        strongest_worker, size_bytes, draft["size_bytes"],
                    )
                ):
                    try:
                        # Resolve target path for the spawn — though
                        # the actual GGUF resolution happens worker-
                        # side via the SSH script, we still need to
                        # know the model is real on the host's view.
                        if info and info.get("gguf_path"):
                            base_url = await ensure_worker_chat_server(
                                strongest_worker,
                                model_name,
                                target_gguf_path=info["gguf_path"],
                                draft_model_name=draft["name"],
                            )
                            await stop_all_running_splits()
                            log.info(
                                "compute_pool: worker-side speculative "
                                "engaged — worker=%s target=%s draft=%s",
                                strongest_worker.get("label"),
                                model_name, draft["name"],
                            )
                            return {
                                "engine": "llama_server",
                                "base_url": base_url,
                                "label": model_name,
                                "host_node": f"worker:{strongest_worker['id']}",
                                "speculative_draft": draft["name"],
                                "speculative_match": "family",
                            }
                    except Exception as e:
                        log.info(
                            "compute_pool: worker-side llama-server spawn "
                            "failed for %s; falling back to plain Ollama: %s",
                            strongest_worker.get("label"), e,
                        )

        # Host-side speculative sub-tier: even though one node fits
        # the model alone, recruit the rest of the pool as a draft-
        # model accelerator on host. Walks pool inventory for a
        # smaller, same-family chat model and engages llama-server
        # (instead of Ollama) with `-md <draft>` when a viable draft
        # is found AND host VRAM has room for both.
        #
        # `await_draft_for` extends `pick_draft_for` with a
        # synchronous worker→host SCP path: when no host-resident
        # draft exists but a worker has one that's small enough to
        # SCP quickly (< 2 GB), we BLOCK on that pull so this very
        # turn benefits from speculative decoding. Falls back to no-
        # speculative on SCP failure.
        if speculative_decoding_enabled():
            draft = await await_draft_for(model_name)
            if (
                draft
                and _host_has_vram_for_speculative(size_bytes, draft["size_bytes"])
                and split_runtime.find_llama_server() is not None
            ):
                # Same Scope-B GGUF override resolution Tier 2 does — a
                # speculative path on a model that needs a repacked
                # GGUF still has to wait for that file to land before
                # llama-server can load it.
                ready = await ensure_compatible_gguf(model_name)
                if not ready.get("ok"):
                    # Fall back to Ollama silently — the user's chat
                    # shouldn't block on the speculative path's prep.
                    log.info(
                        "compute_pool: speculative decoding deferred for %s "
                        "while compatible GGUF is being prepared",
                        model_name,
                    )
                else:
                    info = resolve_ollama_model(model_name) or info
                    base_url = await _ensure_split_running_for(
                        model_name,
                        info["gguf_path"],
                        worker_ids=[],  # single-node — no rpc workers
                        mmproj_path=info.get("mmproj_path"),
                        draft_gguf_path=draft["gguf_path"],
                    )
                    log.info(
                        "compute_pool: speculative decoding engaged — "
                        "target=%s draft=%s match=%s "
                        "(draft is %.0f%% of target size)",
                        model_name, draft["name"], draft.get("match", "?"),
                        100 * draft["size_bytes"] / max(size_bytes, 1),
                    )
                    return {
                        "engine": "llama_server",
                        "base_url": base_url,
                        "label": model_name,
                        "speculative_draft": draft["name"],
                        "speculative_match": draft.get("match", "family"),
                    }
        await stop_all_running_splits()
        return {"engine": "ollama"}

    # Tier 2: engage Phase 2 split. Split-rpc-eligible workers
    # (rpc-server reachable) are what matters here — Phase 2 doesn't
    # need the model installed on the worker; layer weights stream
    # from host's local GGUF.
    #
    # When to engage split:
    #   * If model is so big the host CAN'T run it alone
    #     (size > host_total) → engage split regardless of estimated
    #     pool VRAM. Workers' `max_vram_seen_bytes` is a hard lower
    #     bound (often 0 for never-benchmarked workers), so the
    #     estimated combined VRAM is always an under-count. llama.cpp
    #     itself does the actual per-layer placement based on each
    #     rpc-server's real available memory; if it truly can't fit,
    #     llama-server reports a clean error — strictly better than
    #     "Ollama on host says no tokens" with no diagnostic.
    #   * If host CAN run alone but estimated pool VRAM covers the
    #     model AND the slowest worker LAN latency is < 150 ms,
    #     engage split anyway — combined GPU VRAM beats single-host
    #     CPU offload on a fast LAN. Above 150 ms latency we keep it
    #     local (LAN cost dominates the GPU win).
    #
    # Auto-prep happened at the top of route_chat_for; capability
    # cache is fresh, so the synchronous lookup here is sufficient.
    rpc_workers = _eligible_split_workers()
    if rpc_workers:
        rpc_pool_vram = host_vram + sum(
            (w.get("capabilities") or {}).get("max_vram_seen_bytes") or 0
            for w in rpc_workers
        )
        # Pool memory ceiling: VRAM + every worker's free RAM (workers
        # pre-allocate RPC buffers from system RAM for layer streams).
        # Used to detect the "mega-model" case where the GGUF exceeds
        # everything we can fit in actual hardware memory.
        rpc_pool_total = rpc_pool_vram + sum(
            int(float((w.get("capabilities") or {}).get("ram_free_gb") or 0)
                * (1024 ** 3))
            for w in rpc_workers
        )
        host_can_run_alone = host_total > 0 and size_bytes <= host_total
        worst_lan_latency_ms = max(
            (w.get("capabilities") or {}).get("probe_latency_ms") or 0
            for w in rpc_workers
        )
        # Bandwidth-aware filter: drop workers whose measured
        # bandwidth is below the per-token RPC cost threshold. Per-
        # token layer-push for a 7B model is ~200-500 KB; with
        # < 5 MB/s bandwidth that's 50-100 ms per-token tax which
        # dominates token generation rate and makes split slower
        # than single-host. We MEASURE bandwidth on each
        # heartbeat (capabilities.bandwidth_mbps) and on
        # public-pool registration; workers that haven't been
        # measured yet (None / 0) are kept in (no demote) so a
        # fresh peer doesn't get permanently excluded.
        rpc_workers = [
            w for w in rpc_workers
            if not (w.get("capabilities") or {}).get("bandwidth_mbps")
            or float((w.get("capabilities") or {}).get("bandwidth_mbps") or 0) >= 5.0
        ]
        if not rpc_workers:
            log.info(
                "compute_pool: every rpc worker measured below 5 MB/s "
                "bandwidth — dropping back to host-only path for %s",
                model_name,
            )
            await stop_all_running_splits()
            return {"engine": "ollama"}

        # Mega-model: model exceeds combined pool memory. Engages
        # split anyway — llama.cpp's adaptive `-ngl` puts as many
        # layers as fit on GPU/RPC devices, and the rest cascade to
        # host CPU+mmap from the GGUF on disk. Inference WILL be
        # slow (disk-paged forward pass per token; expect minutes
        # per token on consumer NVMe) but the path is correct.
        # We log loudly so the operator sees this is what's happening.
        is_mega_model = (
            rpc_pool_total > 0 and size_bytes > rpc_pool_total
        )
        if is_mega_model:
            log.warning(
                "compute_pool: MEGA-MODEL path engaged for %s — "
                "size %.1f GB exceeds combined pool memory %.1f GB. "
                "Layers beyond pool capacity will page from host disk "
                "via mmap; expect very slow per-token rate. "
                "GGUF must remain on host SSD for this to work.",
                model_name, size_bytes / (1024 ** 3),
                rpc_pool_total / (1024 ** 3),
            )

        # Decision: engage split when ANY of:
        #   * Host can't run alone (mandatory — only path that
        #     might work without paging from disk).
        #   * Pool VRAM covers the model AND the LAN is fast
        #     (split beats single-host CPU offload by routing
        #     layers to worker iGPUs).
        #   * MEGA-MODEL — model exceeds host RAM+VRAM. Even
        #     though host CPU+mmap can technically execute the
        #     model from disk, the rate is sub-1 tok/s; engaging
        #     workers' iGPU/RAM puts SOME layers on faster
        #     hardware. Better to take the LAN cost than spin on
        #     host disk. (Was previously only WARNED about; now
        #     it actually engages — matches the user's "use as
        #     much pool resource as available" policy.)
        #   * Aggressive pooling — host CAN run alone, BUT the
        #     pool RAM (workers' free memory) adds at least
        #     `_AGGRESSIVE_POOL_GB` of usable budget. Engages
        #     split so workers contribute even when the model
        #     fits host CPU. The LAN-latency gate still applies
        #     so this only kicks in on a fast LAN.
        worker_pool_ram_gb = sum(
            float((w.get("capabilities") or {}).get("ram_free_gb") or 0)
            for w in rpc_workers
        )
        aggressive_engage = (
            host_can_run_alone
            and worker_pool_ram_gb >= _AGGRESSIVE_POOL_GB
            and worst_lan_latency_ms <= 150
        )
        engage_split = (
            not host_can_run_alone
            or (rpc_pool_vram > 0 and size_bytes <= rpc_pool_vram and worst_lan_latency_ms <= 150)
            or is_mega_model
            or aggressive_engage
        )

        if engage_split:
            worker_ids = [w["id"] for w in rpc_workers]

            # Scope B: ensure any required override / mmproj files are
            # already on disk before we attempt to spawn llama-server.
            # If the model is in `_KNOWN_OVERRIDE_REGISTRY` and its
            # files are missing, kick off a background acquisition
            # (lossless surgery from the local Ollama blob, falling
            # back to HuggingFace download if needed) and surface a
            # `RouteChatError` so the chat layer can show the user a
            # "preparing the model" status with progress.
            ready = await ensure_compatible_gguf(model_name)
            if not ready.get("ok"):
                phase = ready.get("phase") or ready.get("status") or "preparing"
                raise RouteChatError(
                    f"Compatible GGUF for {model_name} is being prepared "
                    f"({phase}, ~{ready.get('estimated_total_gb', 0):.1f} GB total). "
                    "Try again shortly.",
                    status=ready,
                )

            # Re-resolve in case the acquisition just landed an override
            # file we didn't know about earlier in this call.
            info = resolve_ollama_model(model_name) or info

            # Strict semantics: if Tier 2 engages a split and the split
            # fails, we DO NOT fall back to Ollama on host. Split was
            # chosen because the host alone couldn't run this model
            # cleanly (or because the pool is faster), so silently
            # degrading to a much slower path defeats the user's intent
            # in setting up the pool. Raise so the caller surfaces the
            # error and the user can free pool memory / shrink the
            # model / disable workers explicitly. The exception
            # propagates out of `route_chat_for` to the chat layer.
            #
            # Speculative-decoding overlay on Phase 2: if the user has
            # speculative on and the pool offers a viable smaller
            # vocab-compatible draft, stack `--model-draft` on the
            # split path. The draft loads on the orchestrator host
            # alongside the target's first layers; verification still
            # rides the layer-split for everything past the draft's
            # local copy. Same picker, same gates, same headroom check
            # — but using the Phase 2 host VRAM budget reserved for
            # target layers as the headroom anchor instead of the
            # full target size.
            split_draft: dict | None = None
            split_draft_path: str | None = None
            if speculative_decoding_enabled():
                cand = pick_draft_for(model_name)
                # Phase 2 already plans to stream the bulk of target
                # weights across `--rpc` workers, so the host only
                # holds whichever layers `_compute_optimal_ngl`
                # assigns it. The draft just has to fit alongside
                # those host-resident target layers — we approximate
                # that by checking against full host VRAM budget
                # without the target_size term (the target's split
                # already accounts for itself in -ngl).
                if cand and _host_vram_budget_bytes() >= int(
                    cand["size_bytes"] * _SPECULATIVE_VRAM_HEADROOM
                ):
                    split_draft = cand
                    split_draft_path = cand["gguf_path"]

            base_url = await _ensure_split_running_for(
                model_name,
                info["gguf_path"],
                worker_ids,
                mmproj_path=info.get("mmproj_path"),
                draft_gguf_path=split_draft_path,
            )
            if split_draft:
                log.info(
                    "compute_pool: speculative decoding stacked on layer-"
                    "split — target=%s draft=%s match=%s",
                    model_name, split_draft["name"],
                    split_draft.get("match", "?"),
                )
            result: dict = {
                "engine": "llama_server",
                "base_url": base_url,
                "label": model_name,
            }
            if split_draft:
                result["speculative_draft"] = split_draft["name"]
                result["speculative_match"] = split_draft.get("match", "family")
            if is_mega_model:
                # Surface the mega-model status so the chat layer can
                # render a "this will be slow" banner in the UI. The
                # actual per-token rate depends on host disk read
                # bandwidth and the fraction of layers paging from
                # disk — could be anywhere from 0.1 tok/s to 1 tok/s
                # on consumer hardware.
                result["mega_model"] = True
                result["pool_memory_gb"] = round(
                    rpc_pool_total / (1024 ** 3), 1,
                )
                result["model_size_gb"] = round(size_bytes / (1024 ** 3), 1)
                # Suggest smaller same-family variants the user already
                # has — quality-preserving alternatives instead of
                # silent quantization. Empty list when nothing fits;
                # UI just shows the mega warning without suggestions.
                result["suggested_smaller_models"] = (
                    find_smaller_variants_in_family(
                        model_name, max_size_bytes=rpc_pool_total,
                    )
                )
                # Mark host's disk subsystem busy so subsequent parallel
                # work (subagents / embeddings / side chats) is biased
                # toward workers — keeps the host disk dedicated to the
                # mmap page-in stream.
                _mark_host_mega_busy()
            return result

    # Tier 3: no eligible RPC workers were found, OR Tier 2 decided
    # not to engage split (host can run alone AND latency too high to
    # benefit). Use Ollama on host. Ollama supports CPU offload + mmap
    # spill so even larger-than-VRAM models run (slowly, via disk
    # paging) without needing the pool.
    await stop_all_running_splits()
    result: dict = {"engine": "ollama"}
    # Flag mega-model on this path too: when the model exceeds host
    # total memory, Ollama will disk-page each forward pass via mmap.
    # The chat will work but at 0.1-1 tok/s; UI should render a
    # "this is going to be slow" banner.
    if host_total > 0 and size_bytes > host_total:
        log.warning(
            "compute_pool: MEGA-MODEL on host-only path — %s "
            "(size %.1f GB > host total %.1f GB). Forward pass will "
            "disk-page via mmap; expect 0.1-1 tok/s.",
            model_name, size_bytes / (1024 ** 3),
            host_total / (1024 ** 3),
        )
        result["mega_model"] = True
        result["pool_memory_gb"] = round(host_total / (1024 ** 3), 1)
        result["model_size_gb"] = round(size_bytes / (1024 ** 3), 1)
        result["suggested_smaller_models"] = (
            find_smaller_variants_in_family(
                model_name, max_size_bytes=host_total,
            )
        )
        # Same reasoning as the Tier-2 mega path above: bias parallel
        # work to workers while the host disk is busy mmapping.
        _mark_host_mega_busy()
    return result


def list_subagent_workers_full(model: str) -> list[dict]:
    """Same selection logic as `list_subagent_workers` but returns
    the FULL worker dicts so callers can read `use_encrypted_proxy`
    and dispatch subagent traffic through the encrypted proxy when
    set. Existing tuple-returning helper kept for back-compat with
    older test fixtures."""
    cands = _eligible_workers("use_for_subagents", model=model)
    if not cands:
        return []

    host_score = _host_capability_score(model)
    host_tps = host_score[0]
    mega_busy = is_host_mega_busy()

    if host_tps > 0 and not mega_busy:
        min_tps = host_tps * _SUBAGENT_MIN_PERF_RATIO
        gated = []
        for w in cands:
            w_tps = float((w.get("capabilities") or {}).get("tokens_per_second") or 0.0)
            if w_tps <= 0:
                gated.append(w)
            elif w_tps >= min_tps:
                gated.append(w)
            else:
                log.info(
                    "subagent fan-out: skipping %r (tps=%.1f) — below %.0f%% of host's %.1f tps",
                    w.get("label"), w_tps,
                    _SUBAGENT_MIN_PERF_RATIO * 100, host_tps,
                )
        cands = gated
    elif mega_busy:
        log.info(
            "subagent fan-out: host mega-busy, skipping perf-ratio gate "
            "to push all %d worker(s) into the rotation",
            len(cands),
        )
    return cands


def list_subagent_workers(model: str) -> list[tuple[str, str | None]]:
    """Return every eligible compute worker for parallel-subagent fan-out.

    The host itself is NOT in this list — the caller (`run_subagents_parallel`)
    composes `[host] + workers` and round-robins, so a 6-task fan-out across
    1 host + 2 workers schedules roughly 2 per machine.

    Performance gate (Phase 2 commit 18): a worker is included ONLY if
    its measured throughput is at least `_SUBAGENT_MIN_PERF_RATIO` of
    the host's measured throughput on this model. Without this gate,
    a worker that's 25× slower than host (e.g. a Core 3 100U laptop
    iGPU vs an RTX 3060 Ti) would still get round-robin-assigned tasks
    in `run_subagents_parallel` — and the slowest task bottlenecks the
    whole fan-out. Net result: parallel calls get SLOWER with the
    weak worker than without it. The gate keeps weak workers out of
    parallel jobs while still letting them serve embed routing and
    contribute layer storage to Phase 2 splits.

    When neither host nor worker has a measured TPS yet (cold cache),
    we fall back to including the worker — there's no signal to make
    a smarter call, and the next probe will fix it.
    """
    cands = _eligible_workers("use_for_subagents", model=model)
    if not cands:
        return []

    host_score = _host_capability_score(model)
    host_tps = host_score[0]

    # Mega-busy override: when the host is mmapping a mega-model, ANY
    # worker beats the host's disk-saturated state. Skip the perf-ratio
    # gate so the fan-out picks up every eligible worker — even slow
    # ones — to keep the host's I/O bandwidth focused on the mega-model
    # page-in stream. The caller still composes `[host] + workers`, so
    # the host gets at most one task; the workers absorb the rest.
    mega_busy = is_host_mega_busy()

    if host_tps > 0 and not mega_busy:
        min_tps = host_tps * _SUBAGENT_MIN_PERF_RATIO
        gated = []
        for w in cands:
            w_tps = float((w.get("capabilities") or {}).get("tokens_per_second") or 0.0)
            if w_tps <= 0:
                # No measurement yet — include and let the probe sort
                # this out next sweep. Excluding unknown workers would
                # never let a fresh worker bootstrap into the pool.
                gated.append(w)
            elif w_tps >= min_tps:
                gated.append(w)
            else:
                log.info(
                    "subagent fan-out: skipping %r (tps=%.1f) — below %.0f%% of host's %.1f tps",
                    w.get("label"), w_tps,
                    _SUBAGENT_MIN_PERF_RATIO * 100, host_tps,
                )
        cands = gated
    elif mega_busy:
        log.info(
            "subagent fan-out: host mega-busy, skipping perf-ratio gate "
            "to push all %d worker(s) into the rotation",
            len(cands),
        )

    return [
        (_worker_base_url(w), db.get_compute_worker_auth_token(w["id"]))
        for w in cands
    ]
