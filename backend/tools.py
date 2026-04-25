"""Tool implementations the agent can call.

Each tool returns a dict with keys:
  - ok (bool)           — whether it succeeded
  - output (str)        — plain-text output for the model (truncated)
  - error (str, opt.)   — error message when ok=False
  - exit_code (int, opt.) — bash return code
  - image_path (str, opt.) — path (relative to SCREENSHOT_DIR) of a PNG the
                             model should "see" as part of the tool result.
                             Only set for screenshot / computer-use tools.

The agent loop (see agent.py) reads `image_path` and, for multimodal models,
injects a follow-up user message with the image attached so the model can
actually look at the screen.
"""

from __future__ import annotations

import asyncio
import datetime
import difflib
import fnmatch
import ipaddress
import json
import os
import re
import shutil
import socket
import subprocess
import sys
import threading
import time
import uuid
from collections import OrderedDict
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import httpx

from . import db
from . import retention
from .telemetry import timed_tool

# --------------------------------------------------------------------------
# Windows DPI awareness.
#
# pyautogui's screenshots and mouse coordinates both live in *physical* pixels,
# but a high-DPI Windows app that hasn't declared DPI-awareness receives
# virtualised coordinates from the OS. That means pyautogui.size() can report
# a smaller screen than reality, screenshots come out blurry upscaled by the
# compositor, and click coordinates land on the wrong spot. Calling
# SetProcessDPIAware *before* importing pyautogui fixes all three. This is a
# no-op on non-Windows platforms.
# --------------------------------------------------------------------------
if sys.platform == "win32":
    try:
        import ctypes
        ctypes.windll.user32.SetProcessDPIAware()
    except Exception:
        # On some Windows editions / when run inside certain IDEs this can
        # throw; we fall back to pyautogui's default behavior.
        pass

# pyautogui + Pillow are imported lazily inside the tool functions so that
# import errors (e.g. on a headless server) don't break the whole backend.
# For screenshots we need Pillow too, which pyautogui already depends on.

MAX_OUTPUT_CHARS = 20000
DEFAULT_TIMEOUT_SEC = 120

# Screenshots are written here and served by the backend at /api/screenshots/*.
# Directory kept next to the SQLite DB so it travels with user data and is
# git-ignored via the `data/` entry in .gitignore.
SCREENSHOT_DIR = Path(__file__).resolve().parent.parent / "data" / "screenshots"

# Per-conversation file checkpoints live here. Every write_file / edit_file
# call snapshots the target file (if it existed) to this directory so the
# user can later restore by message id. Structured as:
#   CHECKPOINT_DIR / <conv_id> / <message_id> / <sha_of_resolved_path>.bin
CHECKPOINT_DIR = Path(__file__).resolve().parent.parent / "data" / "checkpoints"

# User-pasted images (drag-drop or clipboard paste from the UI) are stored
# here and attached to the next outbound user message as multimodal input.
UPLOAD_DIR = Path(__file__).resolve().parent.parent / "data" / "uploads"

# Persistent per-conversation memory notes. Each conversation gets its own
# plain-text file keyed by conversation id. The contents are injected into
# the system prompt on every turn, so anything the agent `remember`s here
# survives auto-compaction, server restarts, and long idle gaps. The format
# is human-readable markdown so the user can open + edit the file directly.
MEMORY_DIR = Path(__file__).resolve().parent.parent / "data" / "memory"

# Hard cap on how large any single memory file can grow. Prevents a runaway
# agent from ballooning the system prompt every turn. When exceeded, the
# oldest lines are trimmed.
MEMORY_MAX_CHARS = 16000

# Max dimension we feed to the model. We keep screenshots at native
# resolution whenever possible so the model's click coordinates map 1:1 to
# the real screen — eliminating a whole class of coord-space bugs. The cap
# only kicks in for very-high-DPI monitors (above 2560 long edge, i.e. 4K
# and larger), where the image would otherwise cost too many vision tokens.
# 2560 = native for 1080p and 1440p displays, which is what most users have.
MAX_SCREENSHOT_EDGE = 2560


def _clip(s: str) -> str:
    """Truncate long strings for display/return to the model."""
    if len(s) <= MAX_OUTPUT_CHARS:
        return s
    head = s[: MAX_OUTPUT_CHARS // 2]
    tail = s[-MAX_OUTPUT_CHARS // 2 :]
    return f"{head}\n\n... [output truncated, {len(s) - MAX_OUTPUT_CHARS} chars omitted] ...\n\n{tail}"


def _resolve(cwd: str, path: str, conv_id: str | None = None) -> Path:
    """Resolve `path` relative to the effective shell directory.

    When `conv_id` is provided AND bash has `cd`-ed somewhere inside the
    workspace during this conversation, relative paths resolve against that
    directory instead of the original `cwd`. This keeps the model's mental
    model consistent — if it ran `cd myapp` in bash, then
    `write_file("src/App.jsx", ...)` lands inside `myapp/src/`, not
    back at the workspace root.

    Absolute paths always win over both. Passing `conv_id=None` preserves the
    old "always relative to `cwd`" behaviour — used by tests / callers that
    don't want the bash-cwd behaviour to leak in.
    """
    p = Path(path)
    if not p.is_absolute():
        base = _effective_bash_cwd(cwd, conv_id) if conv_id else cwd
        p = Path(base) / p
    return p.resolve()


# ---------------------------------------------------------------------------
# File / shell tools
# ---------------------------------------------------------------------------
# Sentinel used to trail the shell's final $PWD after the user's command.
# We pick a long, improbable literal so it's vanishingly unlikely to collide
# with anything the command itself emits.
_BASH_CWD_MARKER = "__CCHAT_BASH_CWD_2E7F91A4__"


def _kill_process_tree(proc) -> None:
    """Kill a subprocess AND every descendant it spawned.

    `proc.kill()` on its own only signals the direct child. If that child
    forked (bash → npm → node → vite, cmd.exe → python, docker compose →
    multiple containers, …), the grandchildren inherit stdout/stderr and
    keep the pipes open — `proc.wait()` then hangs forever waiting for
    EOF that never comes, stranding the conversation.

    Windows: `taskkill /F /T` walks the process tree and terminates each
    process by pid. Much more reliable than chasing descendants manually.

    POSIX: if the process was spawned in its own group (not the default),
    `os.killpg` kills the whole group in one syscall. We don't currently
    set a new pgid up-front (would need `preexec_fn=os.setsid`), so we
    fall back to `proc.kill()` — fine for most foreground commands, and
    tree-escaping cases on POSIX are much rarer than on Windows because
    PIPE FDs close cleanly when the owner exits.
    """
    pid = getattr(proc, "pid", None)
    if sys.platform == "win32" and pid:
        try:
            # /F force, /T whole tree. Suppress all output because a
            # "process not found" is expected for fast-exit commands.
            subprocess.run(
                ["taskkill", "/F", "/T", "/PID", str(pid)],
                capture_output=True,
                timeout=5,
                creationflags=0x08000000,  # CREATE_NO_WINDOW
            )
        except Exception:
            # Fall back to plain kill — at least the direct child dies.
            try:
                proc.kill()
            except Exception:
                pass
    else:
        try:
            proc.kill()
        except Exception:
            pass


def _non_interactive_env() -> dict:
    """Return a copy of os.environ with CI/non-interactive flags set.

    The bash subprocess has no TTY, so any command that prompts will either
    hang until timeout or bail with "Operation cancelled". Most scaffolders
    and package managers check one of a handful of env vars to decide
    whether to prompt — setting them up-front picks the safe defaults so
    we don't have to remember the right CLI flag for every single tool.

    Keep this to well-known, broadly respected signals. Anything tool-
    specific should stay in the system-prompt guidance (so the model can
    override with explicit flags where it needs to).
    """
    env = os.environ.copy()
    env.setdefault("CI", "true")
    # Many interactive npm/pnpm/yarn prompts respect this.
    env.setdefault("npm_config_yes", "true")
    # Suppress the "Ok to proceed? (y)" confirm in npx when installing.
    env.setdefault("NPM_CONFIG_YES", "true")
    # apt-get, dpkg, debconf.
    env.setdefault("DEBIAN_FRONTEND", "noninteractive")
    # pip: skip the "new version available" nag that sometimes waits on tty.
    env.setdefault("PIP_DISABLE_PIP_VERSION_CHECK", "1")
    env.setdefault("PIP_NO_INPUT", "1")
    # Make git's pager / editor non-interactive if a command would invoke
    # one — `cat` won't page; `true` means any editor invocation exits
    # cleanly without opening anything.
    env.setdefault("GIT_PAGER", "cat")
    env.setdefault("PAGER", "cat")
    env.setdefault("GIT_EDITOR", "true")
    env.setdefault("EDITOR", "true")
    return env


def _bash_executable() -> str | None:
    """Locate a real bash binary so `subprocess_shell` stops defaulting to
    cmd.exe on Windows.

    The system prompt already tells the model to write bash syntax (`Shell:
    bash (Git Bash on Windows)`), but `asyncio.create_subprocess_shell`
    honours `COMSPEC` on Windows and falls back to cmd.exe — which means
    bash-only constructs (`$PWD`, `$?`, chained `cd a && cd b`) silently
    fail. We look up bash on PATH first, then fall back to the standard
    Git-for-Windows install locations so uvicorn launched from a plain
    cmd window (whose PATH has `Git\\cmd` for git.exe but NOT `Git\\bin`
    for bash.exe) still finds a shell. Without this fallback every `bash`
    call silently ran through cmd.exe, and `cd` / `$PWD` / the marker
    wrapper would all be interpreted as literal cmd syntax — blowing up
    cwd persistence and interactive-prompt handling.

    Returns None when no bash is found — callers fall back to the platform
    default so the tool still works on a minimal image.
    """
    import shutil
    found = shutil.which("bash")
    if found:
        return found
    # Known Git-for-Windows install locations. `usr/bin/bash.exe` is the
    # MSYS2 bash; `bin/bash.exe` is the wrapper that sets up env vars.
    # Prefer `bin/bash.exe` because that's what an interactive Git-Bash
    # session uses, but either works for our non-interactive wrapping.
    if sys.platform == "win32":
        for candidate in (
            r"C:\Program Files\Git\bin\bash.exe",
            r"C:\Program Files\Git\usr\bin\bash.exe",
            r"C:\Program Files (x86)\Git\bin\bash.exe",
            r"C:\Program Files (x86)\Git\usr\bin\bash.exe",
        ):
            if os.path.isfile(candidate):
                return candidate
    return None


def _wrap_with_cwd_marker(command: str) -> str:
    """Wrap a bash command so the shell prints its final `$PWD` on stdout.

    We run the user's command, stash its exit code, print the marker line,
    then re-exit with the original code. This lets us capture any `cd`
    effect (including chained `cd a && cd b`) without parsing the command
    ourselves — the shell does the work and we just read the marker back.

    On Git-bash / MSYS (Windows) the `$PWD` variable uses MSYS paths like
    `/c/Users/…` that Python's stdlib can't use as a cwd. We prefer
    `cygpath -w` when available (converts to `C:\\Users\\…`), falling back
    to a raw `$PWD` on POSIX where the value is already usable.
    """
    return (
        f"{command}\n"
        f"__CCHAT_RC=$?\n"
        f"__CCHAT_PWD=\"$(cygpath -w \"$PWD\" 2>/dev/null || printf %s \"$PWD\")\"\n"
        f"printf '\\n%s:%s\\n' '{_BASH_CWD_MARKER}' \"$__CCHAT_PWD\"\n"
        f"exit $__CCHAT_RC\n"
    )


def _split_cwd_marker(output: str) -> tuple[str, str | None]:
    """Extract the marker line from `output` and return (cleaned_output, cwd).

    Returns (original_output, None) if the marker isn't present — the wrapper
    didn't run (e.g. shell doesn't understand the syntax, or the command
    `exec`'d away before reaching the trailer).
    """
    needle = "\n" + _BASH_CWD_MARKER + ":"
    idx = output.rfind(needle)
    if idx < 0:
        # Also check for first-line placement (no preceding newline).
        head = _BASH_CWD_MARKER + ":"
        if output.startswith(head):
            rest = output[len(head):]
            nl = rest.find("\n")
            cwd = (rest if nl < 0 else rest[:nl]).rstrip()
            remainder = "" if nl < 0 else rest[nl + 1 :]
            return remainder, (cwd or None)
        return output, None
    head = output[:idx]
    rest = output[idx + len(needle):]
    nl = rest.find("\n")
    cwd = (rest if nl < 0 else rest[:nl]).rstrip()
    tail = "" if nl < 0 else rest[nl + 1 :]
    cleaned = head + (("\n" + tail) if tail else "")
    return cleaned, (cwd or None)


def _run_bash_blocking(
    cwd: str,
    command: str,
    timeout: int,
    bash_exe: str | None = None,
) -> dict:
    """Synchronous subprocess fallback used when the running event loop can't
    spawn subprocesses (e.g. WindowsSelectorEventLoop, which raises
    NotImplementedError from asyncio.create_subprocess_shell).

    Runs inside asyncio.to_thread so it doesn't block the event loop — we pay
    a thread hop but gain subprocess support on every loop type. When
    `bash_exe` is provided we invoke it via `[bash_exe, "-c", command]`
    (shell=False) so the async path and this fallback agree on which shell
    is parsing the script.
    """
    try:
        env = _non_interactive_env()
        # Use Popen manually so we can kill the whole process tree on
        # timeout — subprocess.run's built-in timeout handler only kills
        # the direct child, which hangs forever when that child has
        # spawned grandchildren holding the stdout pipe (npm→node→vite).
        if bash_exe:
            proc = subprocess.Popen(
                [bash_exe, "-c", command],
                cwd=cwd,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                stdin=subprocess.DEVNULL,
                env=env,
            )
        else:
            proc = subprocess.Popen(
                command,
                cwd=cwd,
                shell=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                stdin=subprocess.DEVNULL,
                env=env,
            )
        try:
            stdout, _ = proc.communicate(timeout=timeout)
        except subprocess.TimeoutExpired:
            _kill_process_tree(proc)
            try:
                proc.wait(timeout=2)
            except subprocess.TimeoutExpired:
                pass
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"command timed out after {timeout}s. If this is a "
                    "long-running server or watcher, use `bash_bg` "
                    "instead of `bash`."
                ),
            }
        return {
            "ok": proc.returncode == 0,
            "output": _clip((stdout or b"").decode("utf-8", errors="replace")),
            "exit_code": proc.returncode,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _augment_bash_error(command: str, output: str, exit_code: int | None) -> str:
    """Attach a short, platform-general hint when we detect a well-known
    pattern in a bash tool result.

    Purpose: small models don't reliably remember per-OS binary names and
    invocation quirks (e.g. `google-chrome` vs Windows' Start-menu 'Chrome',
    or `start` vs `open -a`). Rather than dumping platform trivia into the
    system prompt, we surface the hint *only on the error turn*, so the
    model learns the right tool choice from the result itself and only
    pays the token cost when it actually mattered.

    The hint steers the model toward `open_app` (for GUI apps) or the
    correct command form, without naming any particular app.
    """
    if exit_code == 0 or not output:
        return output

    # Patterns are intentionally broad so they catch variants across cmd.exe,
    # PowerShell, bash/zsh, and fish. We match case-insensitively on lowered
    # output.
    lowered = output.lower()

    # Command-not-found family — the single most common bash failure for a
    # model that misremembers a binary name or uses the wrong OS's name.
    command_not_found_markers = (
        "is not recognized as an internal or external command",  # cmd.exe
        "is not recognized as the name of a cmdlet",             # PowerShell
        "command not found",                                     # bash / zsh
        "no such file or directory",                             # posix exec
        "unknown command",                                       # fish
    )
    if any(m in lowered for m in command_not_found_markers):
        hint = (
            "\n\n[hint from the tool] This command was not found on PATH. "
            "If you were trying to launch a GUI application, DON'T retry "
            "with a different guess at the binary name — use the "
            "`open_app` tool instead. For plain launches: "
            "`open_app({\"name\": \"<display name>\"})`. For launches "
            "that need flags/args (guest mode, incognito, opening a file, "
            "new window): `open_app({\"name\": \"...\", \"args\": [\"--flag\"]})`. "
            "`open_app` handles per-OS invocation (Start / Spotlight / exec) "
            "and can locate installed apps by their display name, so you "
            "don't need to remember platform-specific binary names."
        )
        return output + hint

    # Interactive-prompt markers: the command wanted a TTY answer and bailed
    # because our subprocess has no stdin. Nudge the model to re-run with
    # non-interactive flags instead of trying to drive the CLI with
    # computer_key / Enter presses (which target whatever window is in
    # front, not this subprocess).
    interactive_markers = (
        "operation cancelled",
        "aborted by user",
        "user force closed the prompt",
        "? select ",
        "? choose ",
        "ok to proceed? (y)",
        "inquirerabort",
        "no tty",
        "not a tty",
        "input device is not a tty",
        "stdin is not a tty",
        "cannot prompt in non-interactive mode",
    )
    if any(m in lowered for m in interactive_markers):
        hint = (
            "\n\n[hint from the tool] The command tried to prompt for input, "
            "but `bash` runs without a TTY so there's no way to answer. "
            "Re-run with non-interactive flags instead — e.g. `-y` / `--yes` "
            "for npm/apt, `--template <x>` and a project name for `npm "
            "create vite`, `npx --yes <pkg>` to skip npx's install confirm, "
            "or pipe answers via a here-doc. Do NOT try to answer the "
            "prompt with `computer_key` — the prompt lives in this "
            "subprocess, which has no window, so keypresses go elsewhere."
        )
        return output + hint

    # Permission denied: remind the model that sudo / admin prompts can't
    # be answered from our sandbox.
    if "permission denied" in lowered or "access is denied" in lowered:
        hint = (
            "\n\n[hint from the tool] The command was denied by the OS. "
            "We can't satisfy interactive sudo / UAC prompts. Either ask "
            "the user to run the elevated command themselves, or choose "
            "a path that doesn't require elevation."
        )
        return output + hint

    return output


def _effective_bash_cwd(cwd: str, conv_id: str | None) -> str:
    """Resolve the shell's starting directory: persistent bash_cwd if set and
    still exists, otherwise the conversation's fixed workspace root.

    The fallback matters if the user deletes the directory the model had
    `cd`-ed into — we silently drop back to the workspace root instead of
    crashing every subsequent bash call with ENOENT.
    """
    if not conv_id:
        return cwd
    try:
        stored = db.get_bash_cwd(conv_id)
    except Exception:
        stored = None
    if stored and os.path.isdir(stored):
        return stored
    return cwd


def _persist_new_cwd(conv_id: str | None, new_cwd: str | None, workspace_cwd: str) -> None:
    """Write the captured pwd back to the conversation row, or clear it when
    the shell ended at the workspace root (so the DB doesn't carry a
    redundant override that would bias future lookups).
    """
    if not conv_id or not new_cwd:
        return
    try:
        norm_new = os.path.normcase(os.path.normpath(new_cwd))
        norm_base = os.path.normcase(os.path.normpath(workspace_cwd))
        db.set_bash_cwd(conv_id, None if norm_new == norm_base else new_cwd)
    except Exception:
        pass


async def run_bash(
    cwd: str,
    command: str,
    timeout: int = DEFAULT_TIMEOUT_SEC,
    conv_id: str | None = None,
) -> dict:
    """Run a shell command with a hard timeout.

    When `conv_id` is provided, the shell's starting directory is seeded
    from the conversation's persistent `bash_cwd` (so `cd subdir` in a
    previous call carries over), and the final `$PWD` is captured and
    written back so future calls continue from wherever this one ended.
    The caller sees clean output — the marker line is stripped before
    return.

    Uses asyncio's subprocess machinery on loops that support it (ProactorEventLoop
    on Windows, any POSIX loop), falling back to blocking subprocess + thread on
    loops that don't (WindowsSelectorEventLoop). The fallback matters because
    uvicorn sometimes picks SelectorEventLoop on Windows, which silently breaks
    every `bash` call with a bare `NotImplementedError` — and the model has no
    way to recover from that.
    """
    if not command or not command.strip():
        return {
            "ok": False,
            "output": "",
            "error": (
                "bash requires the `command` field, but you didn't pass one. "
                "Example: `bash({\"command\": \"cd myapp && ls\", "
                "\"reason\": \"...\"})`. The `reason` field on its own is not "
                "enough — `command` carries the actual shell text to run."
            ),
        }
    effective_cwd = _effective_bash_cwd(cwd, conv_id)
    # Only wrap when we actually have bash AND a conversation to persist
    # the cwd against — otherwise the marker line would just be noise
    # for callers that don't care (subagent, hook invocation, …) or
    # produce literal `$PWD` when the shell is cmd.exe.
    bash_exe = _bash_executable()
    can_wrap = conv_id is not None and bash_exe is not None
    wrapped = _wrap_with_cwd_marker(command) if can_wrap else command
    try:
        env = _non_interactive_env()
        if bash_exe:
            # Explicit `bash -c <script>` via exec bypasses the
            # platform default shell (cmd.exe on Windows) so bash
            # syntax — $PWD, $?, chained `cd` — works reliably.
            proc = await asyncio.create_subprocess_exec(
                bash_exe, "-c", wrapped,
                cwd=effective_cwd,
                stdin=asyncio.subprocess.DEVNULL,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.STDOUT,
                env=env,
            )
        else:
            proc = await asyncio.create_subprocess_shell(
                wrapped,
                cwd=effective_cwd,
                stdin=asyncio.subprocess.DEVNULL,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.STDOUT,
                env=env,
            )
        try:
            stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=timeout)
        except asyncio.TimeoutError:
            # Kill the whole process tree. `proc.kill()` only signals the
            # direct child (bash.exe on Windows, /bin/sh on POSIX) — any
            # grandchild that inherited stdout (vite, webpack, uvicorn,
            # docker, anything that forks) keeps the pipe open and
            # `proc.wait()` hangs forever waiting for EOF. `taskkill /T`
            # on Windows and POSIX process-groups on Linux/macOS tear
            # down the full tree so the call actually returns.
            _kill_process_tree(proc)
            try:
                await asyncio.wait_for(proc.wait(), timeout=2.0)
            except asyncio.TimeoutError:
                # Parent still hasn't reaped — abandon it. Python will
                # clean up eventually; we don't block the model on it.
                pass
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"command timed out after {timeout}s. If this is a "
                    "long-running server (dev server, file watcher, "
                    "docker compose up, etc.), use `bash_bg` instead of "
                    "`bash` — it returns immediately with a shell_id you "
                    "can poll via `bash_output` and stop with `kill_shell`."
                ),
            }
        out = stdout.decode("utf-8", errors="replace") if stdout else ""
        if can_wrap:
            out, new_cwd = _split_cwd_marker(out)
            _persist_new_cwd(conv_id, new_cwd, cwd)
        out = _augment_bash_error(command, out, proc.returncode)
        return {
            "ok": proc.returncode == 0,
            "output": _clip(out),
            "exit_code": proc.returncode,
        }
    except NotImplementedError:
        # Selector loop on Windows: transparently retry via blocking subprocess
        # in a worker thread. The model never sees this detour.
        result = await asyncio.to_thread(
            _run_bash_blocking, effective_cwd, wrapped, timeout, bash_exe,
        )
        if result.get("output") is not None and can_wrap:
            cleaned, new_cwd = _split_cwd_marker(result["output"])
            result["output"] = cleaned
            _persist_new_cwd(conv_id, new_cwd, cwd)
        if result.get("output") is not None:
            result["output"] = _augment_bash_error(
                command, result["output"], result.get("exit_code")
            )
        return result
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def python_exec(cwd: str, code: str, timeout: int = 60) -> dict:
    """Run a Python snippet in a subprocess and capture its stdout+stderr.

    Same shape as `run_bash` on purpose — the agent dispatches to either
    shell depending on the task. Using a subprocess (not eval/exec in-process)
    gives us:
      - a hard timeout the caller can rely on,
      - isolation from the backend's own globals/imports,
      - the same stdout/stderr combined stream model the model already knows.

    The script runs with the backend's Python interpreter so the user gets
    whatever libraries they pip-installed (pandas, numpy, requests, …)
    without having to configure a separate environment. It inherits no
    special privileges — security posture matches `bash`.
    """
    if not code or not code.strip():
        return {"ok": False, "output": "", "error": "empty code"}
    try:
        proc = await asyncio.create_subprocess_exec(
            sys.executable,
            "-I",  # isolated mode: ignore PYTHON* env vars, no user-site
            "-c",
            code,
            cwd=cwd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )
        try:
            stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=timeout)
        except asyncio.TimeoutError:
            proc.kill()
            await proc.wait()
            return {
                "ok": False,
                "output": "",
                "error": f"python_exec timed out after {timeout}s",
            }
        out = stdout.decode("utf-8", errors="replace") if stdout else ""
        return {
            "ok": proc.returncode == 0,
            "output": _clip(out),
            "exit_code": proc.returncode,
        }
    except NotImplementedError:
        # Selector loop on Windows: fall back to blocking subprocess run.
        def _blocking():
            try:
                cp = subprocess.run(
                    [sys.executable, "-I", "-c", code],
                    cwd=cwd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    timeout=timeout,
                )
                out = cp.stdout.decode("utf-8", errors="replace") if cp.stdout else ""
                return {
                    "ok": cp.returncode == 0,
                    "output": _clip(out),
                    "exit_code": cp.returncode,
                }
            except subprocess.TimeoutExpired:
                return {
                    "ok": False,
                    "output": "",
                    "error": f"python_exec timed out after {timeout}s",
                }
            except Exception as e:
                return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}
        return await asyncio.to_thread(_blocking)
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# --- Read-first guard for file-edit tools ---------------------------------
#
# The agent shouldn't blindly edit or overwrite an existing file it has
# never actually looked at — that's how hallucinated "here's what I think
# is in foo.py" replacements end up silently clobbering real code. We
# track which paths the model has read per conversation, and edit_file /
# write_file (when overwriting) refuse to proceed without a prior read.
#
# Design notes:
#   * The set is keyed by `conv_id` so two concurrent conversations don't
#     pollute each other. `conv_id=None` (tests, one-off CLI calls) skips
#     the guard entirely — there's no session to bind to.
#   * We record the RESOLVED absolute path so `./foo.py`, `foo.py`, and
#     `/abs/foo.py` all hash to the same entry.
#   * `write_file` updates the set after a successful write, so a
#     read → write → edit sequence works without re-reading in between.
#   * A `threading.Lock` guards the dict — dispatch runs on the asyncio
#     loop but file I/O may be `asyncio.to_thread`'d and a sibling tool
#     could mutate the set concurrently.
#   * Memory shape: one set of path strings per conversation. Cleared
#     explicitly when a conversation is deleted (see
#     `clear_read_state_for_conversation`). Process restart clears
#     everything; that's fine because fresh turns re-read anyway.
_READ_FILES_BY_CONV: dict[str, set[str]] = {}
_READ_FILES_LOCK = threading.Lock()


def _record_read(conv_id: str | None, path: Path) -> None:
    """Note that `path` has been read / written in `conv_id`."""
    if not conv_id:
        return
    key = str(path.resolve())
    with _READ_FILES_LOCK:
        _READ_FILES_BY_CONV.setdefault(conv_id, set()).add(key)


def _require_prior_read(conv_id: str | None, path: Path) -> str | None:
    """Return a helpful error message if `path` must-be-read-first.

    Returns None when the guard is satisfied (or skipped because
    `conv_id` is None, e.g. tests). The error string is phrased so the
    model knows exactly which tool to call next.
    """
    if not conv_id:
        return None
    key = str(path.resolve())
    with _READ_FILES_LOCK:
        seen = _READ_FILES_BY_CONV.get(conv_id)
        if seen and key in seen:
            return None
    return (
        f"file has not been read in this conversation: {path}\n"
        "Call `read_file` on this exact path first so the edit is based "
        "on the file's actual contents, not assumed contents."
    )


def clear_read_state_for_conversation(conv_id: str) -> None:
    """Drop the read-tracking set for a conversation (called on delete)."""
    with _READ_FILES_LOCK:
        _READ_FILES_BY_CONV.pop(conv_id, None)


async def read_file(
    cwd: str,
    path: str,
    max_bytes: int = 200_000,
    conv_id: str | None = None,
) -> dict:
    """Read a UTF-8 text file (truncated if huge).

    Also records the read against `conv_id` so subsequent `edit_file` /
    `write_file` calls on the same path are allowed through the Read-
    first guard.
    """
    try:
        p = _resolve(cwd, path, conv_id)
        if not p.exists():
            return {"ok": False, "output": "", "error": f"file not found: {p}"}
        if p.is_dir():
            return {"ok": False, "output": "", "error": f"is a directory: {p}"}
        data = p.read_bytes()[:max_bytes]
        text = data.decode("utf-8", errors="replace")
        _record_read(conv_id, p)
        return {"ok": True, "output": _clip(text), "path": str(p)}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def write_file(
    cwd: str,
    path: str,
    content: str,
    conv_id: str | None = None,
) -> dict:
    """Write (overwrite) a file and auto-create parent directories.

    Writes via `write_bytes` so Windows doesn't silently translate `\\n` into
    `\\r\\n`. The model's content is taken verbatim — if the caller wanted
    CRLF they'd author CRLF, and we shouldn't be rewriting every line
    ending on Windows as a side-effect of the platform's text-mode default.

    When `path` already exists, the Read-first guard is enforced: the
    model must have read the target in this conversation before
    overwriting it. Creating a brand-new file is exempt — there's
    nothing to have read.
    """
    try:
        p = _resolve(cwd, path, conv_id)
        if p.exists() and p.is_file():
            guard = _require_prior_read(conv_id, p)
            if guard:
                return {"ok": False, "output": "", "error": guard}
        p.parent.mkdir(parents=True, exist_ok=True)
        # Write bytes, not text: `write_text` on Windows text mode translates
        # `\n` → `\r\n` silently. That turns a "write this file" request into
        # "write this file with Windows line endings" — bad for shell scripts,
        # Dockerfiles, config files generated by the model, and anything round-
        # tripped through read_file (which uses read_bytes + decode, preserving
        # whatever's on disk).
        p.write_bytes(content.encode("utf-8"))
        # Post-write, the model has effectively "seen" this file — any follow-
        # up edit in the same conversation can proceed without a re-read.
        _record_read(conv_id, p)
        return {"ok": True, "output": f"wrote {len(content)} chars to {p}", "path": str(p)}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def list_dir(cwd: str, path: str = ".", conv_id: str | None = None) -> dict:
    """List entries in a directory as `type  size  name` lines.

    `conv_id` lets relative `path` values follow wherever bash `cd`-ed to
    in this conversation — so `list_dir({"path": "."})` after a
    `cd myapp` lists the project dir, not the workspace root.
    """
    try:
        p = _resolve(cwd, path, conv_id)
        if not p.exists():
            return {"ok": False, "output": "", "error": f"not found: {p}"}
        if not p.is_dir():
            return {"ok": False, "output": "", "error": f"not a directory: {p}"}
        entries = []
        for item in sorted(p.iterdir()):
            kind = "dir" if item.is_dir() else "file"
            try:
                size = item.stat().st_size if item.is_file() else 0
            except Exception:
                size = 0
            entries.append(f"{kind:4s}  {size:>10d}  {item.name}")
        return {"ok": True, "output": "\n".join(entries) if entries else "(empty)"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Computer-use tools
#
# These let the model drive the user's desktop. Every action except
# `mouse_move` implicitly captures a follow-up screenshot, because otherwise
# the model cannot observe the consequence of what it just did. All tool
# calls still go through the agent's approval layer (manual-mode gating +
# per-call Approve/Reject UI).
# ---------------------------------------------------------------------------
def _ensure_screenshot_dir() -> Path:
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    return SCREENSHOT_DIR


# Most recent screenshot's scale factor (real_px / displayed_px), set every
# time `_capture_screenshot_sync` runs. The model only ever sees the DISPLAYED
# (possibly downscaled) image, so whatever (x, y) it picks lives in that
# coordinate space. The click/scroll/move tools multiply by this factor so the
# click lands at the right spot on the real screen. Without this, on a
# 1920x1080 monitor that's been downscaled to 1568x882, a click the model aims
# at the taskbar lands ~210 px above it — exactly the bug we hit trying to
# open Microsoft Store via the Start button.
_LAST_SHOT_SCALE: tuple[float, float] = (1.0, 1.0)
# Pixel dimensions of the displayed screenshot (what the model sees). Used to
# bounds-check the coordinates the model picks, so we can return a loud error
# when it hallucinates a y past the bottom of the image instead of silently
# clamping and clicking the wrong thing.
_LAST_SHOT_DISPLAY: tuple[int, int] = (0, 0)
# (x_offset, y_offset) of the last captured monitor's top-left in the global
# virtual-screen coordinate space. This is 0,0 for the primary monitor, but
# non-zero for secondary monitors (e.g. a second display to the right of the
# primary might start at (1920, 0)). Clicks must add this back so the click
# lands on the monitor the model is looking at, not wherever the same (x, y)
# happens to fall on the primary display.
_LAST_SHOT_ORIGIN: tuple[int, int] = (0, 0)

# ---------------------------------------------------------------------------
# Change-detection + last-click marker
#
# Two cheap "did it work?" signals attached to every screenshot the model
# sees:
#
#   1. Pixel signature + window-title set → one-line diff summary like
#      "Δ 8% pixels; new window 'Save As' appeared". The model otherwise
#      has to eyeball two screenshots and reason about what changed —
#      a hard task for small VLMs and one that often produces hallucinated
#      progress claims ("the dialog opened" when it didn't).
#   2. Last click coordinate → red dot painted on the NEXT screenshot.
#      Lets the model literally see where it just clicked, which is the
#      single best feedback signal for "I aimed wrong" — without it, a
#      missed click looks identical to a click that hit a no-op control.
#
# Both are best-effort: any failure here just means the model gets the
# screenshot without the extra signal. Never raise from these helpers.
# ---------------------------------------------------------------------------
# 64x36 grayscale signature of the most recent pre-annotation screenshot.
# Compared against the current frame to compute the percent-pixels-changed
# diff. None means "no baseline yet" → first screenshot reports a fresh
# capture rather than a (meaningless) 0% diff.
_LAST_PIXEL_SIG: bytes | None = None
# Set of visible top-level window titles at the time of the last screenshot.
# Used to surface "new window 'X' appeared" / "window 'Y' disappeared"
# notes alongside the pixel diff, which catches state changes that
# pixel-diff alone misses (a tiny dialog popping up over a busy app).
_LAST_WINDOW_SET: frozenset[str] | None = None
# Real-screen coords (after _map_to_screen) of the most recent click, set
# by every click code path. Cleared after the next screenshot draws the
# marker — we deliberately show the dot ONCE so it doesn't linger and
# confuse the model into thinking it clicked there again.
_LAST_CLICK_SCREEN_POS: tuple[int, int] | None = None


def _record_click_pos(screen_x: int, screen_y: int) -> None:
    """Remember a real-screen click coord for the next screenshot to mark.

    Called by every click code path (computer_click, click_element,
    click_element_id, computer_drag, computer_batch click branches) so
    the next screenshot the model sees has a red dot at the exact pixel
    that was clicked. The dot is the simplest possible "did you aim at
    the right thing?" signal.
    """
    global _LAST_CLICK_SCREEN_POS
    try:
        _LAST_CLICK_SCREEN_POS = (int(screen_x), int(screen_y))
    except Exception:
        # Best-effort — coords arrive from external callers, never let a
        # bad type tank a click.
        _LAST_CLICK_SCREEN_POS = None


def _compute_pixel_sig(img) -> bytes:
    """Cheap, stable fingerprint of an image for percent-changed diffs.

    Downsamples to 64x36 grayscale and returns the raw bytes. Same
    resolution + threshold tuning as the `pixel_change` ui_wait kind so
    the two systems give consistent answers. ~2.3 KB per snapshot,
    constant time per comparison, and the downsample washes out the
    sub-pixel rendering noise that would otherwise fire false positives.

    Must be called BEFORE any overlay annotations (grid lines, last-click
    marker, set-of-mark badges) — otherwise the diff measures our own
    overlays, not real UI change.
    """
    try:
        g = img.convert("L").resize(_PIXEL_DIFF_GRID)
        return g.tobytes()
    except Exception:
        return b""


def _enumerate_window_titles_fast() -> frozenset[str] | None:
    """Snapshot the set of visible top-level window titles.

    Used for the "new/removed window" half of the change-detection
    summary. Windows-only — returns None on other platforms or when the
    UIA package isn't installed, in which case the diff falls back to
    pixel-only.

    Skips zero-titled / zero-area windows (those are usually system
    message-only windows and would just churn the diff). Capped at 100
    entries so a runaway desktop with hundreds of ghost windows can't
    spend visible time enumerating.
    """
    if sys.platform != "win32":
        return None
    try:
        import uiautomation as auto
    except ImportError:
        return None
    titles: set[str] = []
    titles_set: set[str] = set()
    try:
        for w in auto.GetRootControl().GetChildren():
            try:
                title = w.Name or ""
                if not title:
                    continue
                r = w.BoundingRectangle
                if not r:
                    continue
                if (r.right - r.left) <= 0 or (r.bottom - r.top) <= 0:
                    continue
                titles_set.add(title)
                if len(titles_set) >= 100:
                    break
            except Exception:
                continue
    except Exception:
        return None
    return frozenset(titles_set)


def _get_focused_control_name() -> str | None:
    """Return the accessible name of the currently keyboard-focused control.

    Used for the focus-drift detection on `computer_type` / `computer_key`:
    we snapshot the focused control name BEFORE the keystrokes go out and
    again AFTER, then warn the model if focus moved between the two reads.
    Catches the most common "I typed into the wrong field" failure where
    a popup, notification, or background process stole focus mid-action.

    Returns the trimmed accessible name, or None when:
      - we're not on Windows (UIA isn't available)
      - the `uiautomation` package isn't installed
      - no control is focused right now
      - the focused control's accessible name is empty (most edit fields
        do have a name; a blank name is real signal that focus is in
        nowhere-useful and the model should re-focus)
      - any UIA call fails (always treat focus as unknown rather than
        guessing — a wrong name would be more harmful than no name)

    Cheap (~5-30 ms) so it's safe to call inline before/after keystrokes.
    Names are truncated to 80 chars so a verbose "RichEdit" placeholder
    can't blow out the result string.
    """
    if sys.platform != "win32":
        return None
    try:
        import uiautomation as auto
    except ImportError:
        return None
    try:
        ctrl = auto.GetFocusedControl()
        if ctrl is None:
            return None
        name = (ctrl.Name or "").strip()
        if not name:
            return None
        return name[:80]
    except Exception:
        return None


def _capture_status_context() -> dict | None:
    """Snapshot a small dict of "what is the desktop doing right now?".

    Returns a dict with three best-effort fields:

      - foreground: the title of the currently-foreground top-level
        window (truncated to 80 chars). None when we can't read it.
      - focused: the accessible name of the keyboard-focused control
        within that window (truncated to 80 chars). None when no
        control is focused or its name is blank.
      - cursor: (x, y) tuple of the OS cursor's current screen position.

    These three signals together let the model orient itself in ZERO
    vision tokens — instead of having to read the full screenshot to
    figure out which app is in front, it sees the answer in text. They
    travel as a `[ctx: ...]` line on every screenshot result.

    All three fields are independently best-effort: if UIA fails, we
    still try to read the cursor; if the cursor read throws, we still
    return whatever UIA gave us. Returns None only if every signal
    failed AND we're not on Windows (no useful info at all).
    """
    foreground: str | None = None
    focused: str | None = None
    cursor: tuple[int, int] | None = None

    if sys.platform == "win32":
        try:
            import uiautomation as auto
            try:
                fg = auto.GetForegroundControl()
                if fg is not None:
                    title = (fg.Name or "").strip()
                    if title:
                        foreground = title[:80]
            except Exception:
                pass
        except ImportError:
            pass
        focused = _get_focused_control_name()

    try:
        import pyautogui
        pos = pyautogui.position()
        cursor = (int(pos.x), int(pos.y))
    except Exception:
        cursor = None

    if foreground is None and focused is None and cursor is None:
        return None
    return {"foreground": foreground, "focused": focused, "cursor": cursor}


def _format_focus_drift(before: str | None, after: str | None) -> str | None:
    """Format a focus-drift note for `computer_type` / `computer_key`.

    Compares the focused-control names captured before and after a
    keystroke action and returns one of:

      - None, when there's nothing useful to say (both sides are None
        because UIA isn't available, OR both sides match — no drift).
      - "[focus drifted: 'Search bar' → 'Address bar']" when focus
        moved between two known controls.
      - "[focus lost: 'Search bar' → none]" when focus was somewhere
        before but is now nowhere (popup closed, app crashed).
      - "[focus arrived: none → 'Address bar']" when focus was nowhere
        before but landed somewhere (newly-spawned dialog grabbed it).

    The arrows + bracket framing match the `[ctx: ...]` and `[Δ: ...]`
    style so the model parses them consistently.
    """
    # Nothing to compare — UIA unavailable or every read failed.
    if before is None and after is None:
        return None
    if before == after:
        return None
    if before is not None and after is not None:
        return f"[focus drifted: {before!r} → {after!r}]"
    if before is not None and after is None:
        return f"[focus lost: {before!r} → none]"
    return f"[focus arrived: none → {after!r}]"


def _format_status_context(ctx: dict | None) -> str | None:
    """Format the dict from `_capture_status_context` as a one-line string
    suitable for splicing into a `[ctx: ...]` tag on a screenshot output.

    Skips fields that are None so the line stays terse — e.g. on Linux
    the foreground/focused fields are absent and only the cursor shows
    up. Returns None when the input is None or every field is empty.
    """
    if not ctx:
        return None
    parts: list[str] = []
    fg = ctx.get("foreground")
    if fg:
        parts.append(f"foreground={fg!r}")
    foc = ctx.get("focused")
    if foc:
        parts.append(f"focused={foc!r}")
    cur = ctx.get("cursor")
    if cur:
        parts.append(f"cursor=({cur[0]}, {cur[1]})")
    if not parts:
        return None
    return "; ".join(parts)


def _compute_screenshot_change(
    new_sig: bytes,
    new_titles: frozenset[str] | None,
) -> str | None:
    """Compare the new screenshot against the previous baseline.

    Returns a short human-readable string describing what changed, or
    None when this is the first screenshot in the session (no baseline
    to diff against — saying "0% changed" would be misleading).

    Output examples:
      - "first screenshot — no prior frame to diff against"
      - "no visible change since last screenshot"
      - "Δ 8% pixels; new window 'Save As' appeared"
      - "Δ 42% pixels; gone: 'Notepad'; new: 'Untitled - Notepad'"

    Window add/remove notes are capped to keep the line short. Pixel %
    is computed from the 64x36 grayscale signature so it's monitor-size
    independent.
    """
    global _LAST_PIXEL_SIG, _LAST_WINDOW_SET
    prev_sig = _LAST_PIXEL_SIG
    prev_titles = _LAST_WINDOW_SET
    # Update baselines BEFORE returning so the next screenshot diffs
    # against this one, not the one before it.
    _LAST_PIXEL_SIG = new_sig if new_sig else _LAST_PIXEL_SIG
    if new_titles is not None:
        _LAST_WINDOW_SET = new_titles

    if prev_sig is None or not new_sig or len(prev_sig) != len(new_sig):
        return "first screenshot — no prior frame to diff against"

    # Percent of pixels that differ at all (any luminance delta). For the
    # 64x36 grid this is 2304 comparisons — trivial.
    differing = sum(1 for a, b in zip(prev_sig, new_sig) if a != b)
    pct = (differing / len(new_sig)) * 100.0 if new_sig else 0.0

    parts: list[str] = []
    if pct < 0.5:
        parts.append("no visible change since last screenshot")
    else:
        parts.append(f"Δ {pct:.0f}% pixels")

    # Window-set diff — only when both snapshots have a title set (UIA
    # available + Windows). Cap the listed titles so a launcher opening
    # 20 helper windows doesn't blow the line out.
    if prev_titles is not None and new_titles is not None:
        added = list(new_titles - prev_titles)
        removed = list(prev_titles - new_titles)
        if added:
            preview = ", ".join(repr(t[:50]) for t in added[:2])
            extra = f" (+{len(added) - 2} more)" if len(added) > 2 else ""
            parts.append(f"new: {preview}{extra}")
        if removed:
            preview = ", ".join(repr(t[:50]) for t in removed[:2])
            extra = f" (+{len(removed) - 2} more)" if len(removed) > 2 else ""
            parts.append(f"gone: {preview}{extra}")

    return "; ".join(parts)


def _draw_last_click_marker(img, scale: tuple[float, float], origin: tuple[int, int]) -> bool:
    """Paint a red dot at the last click's image-space position.

    Returns True when a marker was drawn (so callers can mention it in
    their output text), False when there's nothing to mark or the click
    was on a different region than this screenshot covers.

    The dot is drawn ONCE per click and then `_LAST_CLICK_SCREEN_POS` is
    cleared — otherwise the marker would persist across multiple
    screenshots and the model would see it linger like a stale UI
    artefact.

    The marker is a small filled circle (8 px radius) with a white
    outline so it's visible on both light and dark UIs. Designed to be
    impossible to miss but small enough not to obscure the control it
    points at.
    """
    global _LAST_CLICK_SCREEN_POS
    pos = _LAST_CLICK_SCREEN_POS
    if pos is None:
        return False
    sx, sy = pos
    ox, oy = origin
    fx, fy = scale
    if not fx or not fy:
        _LAST_CLICK_SCREEN_POS = None
        return False
    # Translate screen coords → image coords. Inverse of _map_to_screen:
    # screen = image * scale + origin → image = (screen - origin) / scale.
    ix = int(round((sx - ox) / fx))
    iy = int(round((sy - oy) / fy))
    w, h = img.size
    if ix < 0 or iy < 0 or ix >= w or iy >= h:
        # Click was off this screenshot's region (e.g. clicked on monitor 2
        # then took a screenshot of monitor 1, or window-cropped past it).
        # Don't draw; clear so the next matching screenshot can mark it.
        _LAST_CLICK_SCREEN_POS = None
        return False
    try:
        from PIL import ImageDraw
        draw = ImageDraw.Draw(img, "RGBA")
        # White outer ring for contrast on dark UIs, then red filled dot.
        draw.ellipse([(ix - 9, iy - 9), (ix + 9, iy + 9)], fill=(255, 255, 255, 220))
        draw.ellipse([(ix - 7, iy - 7), (ix + 7, iy + 7)], fill=(220, 30, 30, 255))
        # Tiny crosshair for sub-pixel precision feedback.
        draw.line([(ix - 11, iy), (ix - 4, iy)], fill=(255, 255, 255, 230), width=2)
        draw.line([(ix + 4, iy), (ix + 11, iy)], fill=(255, 255, 255, 230), width=2)
        draw.line([(ix, iy - 11), (ix, iy - 4)], fill=(255, 255, 255, 230), width=2)
        draw.line([(ix, iy + 4), (ix, iy + 11)], fill=(255, 255, 255, 230), width=2)
    except Exception:
        _LAST_CLICK_SCREEN_POS = None
        return False
    _LAST_CLICK_SCREEN_POS = None
    return True


def _pick_grid_step(img_size: tuple[int, int]) -> int:
    """Choose a coordinate-grid spacing that scales with the image.

    The original 100-px grid was tuned for full-monitor screenshots
    (~1500-1900 px wide). On a small `screenshot_window` capture
    (e.g. a 320x240 dialog), 100-px lines waste vision tokens — only
    2-3 lines fit and they crowd the actual UI without giving the
    model new positional info.

    Heuristic: aim for ~10-16 grid cells along the shorter edge so
    the grid is dense enough to read positions accurately but sparse
    enough that the lines themselves are mostly background. Snap to
    a small set of round numbers (25, 50, 100, 150) so the model
    sees consistent values across screenshots of similar sizes
    instead of jittery odd numbers.
    """
    w, h = img_size
    short = min(w, h) if w and h else 600
    # Target ~12 cells along the short edge.
    raw = max(20, short // 12)
    # Snap to a friendly set so labels read cleanly (multiples of 50
    # are easy for the model to do mental arithmetic with).
    for candidate in (25, 50, 100, 150, 200):
        if raw <= candidate:
            return candidate
    return 200


def _draw_coordinate_grid(img, step: int = 100) -> None:
    """Draw a light coordinate grid over the screenshot in-place.

    Small vision models (like Gemma 4 e4b) are poor at localizing arbitrary
    pixel positions, so they fall back on spatial priors from training
    ("taskbar is near the bottom") which are directionally correct but off
    by hundreds of pixels. Overlaying a labeled grid lets the model READ the
    y coordinate of the taskbar directly off the image instead of guessing —
    e.g. it sees the label "1000" next to a grid line near the taskbar and
    uses y=1040 instead of hallucinating y=95.

    Lines are drawn at every `step` pixels and labeled at every `step * 2`
    (so labels don't clutter the image). Grid is semi-transparent gray so
    the UI underneath stays legible.
    """
    from PIL import ImageDraw, ImageFont

    draw = ImageDraw.Draw(img, "RGBA")
    w, h = img.size
    line_color = (255, 255, 0, 110)  # yellow, semi-transparent → visible on light AND dark UIs
    label_bg = (0, 0, 0, 170)
    label_fg = (255, 255, 0, 255)

    # Pillow's default bitmap font is tiny but always available; try a TTF for
    # readability, fall back gracefully when none is installed.
    font = None
    for candidate in ("arial.ttf", "DejaVuSans.ttf", "Helvetica.ttc"):
        try:
            font = ImageFont.truetype(candidate, 14)
            break
        except Exception:
            continue
    if font is None:
        font = ImageFont.load_default()

    # Vertical lines + X labels along the top edge.
    for x in range(0, w, step):
        draw.line([(x, 0), (x, h)], fill=line_color, width=1)
        if x % (step * 2) == 0 and x > 0:
            txt = str(x)
            tb = draw.textbbox((0, 0), txt, font=font)
            tw, th = tb[2] - tb[0], tb[3] - tb[1]
            draw.rectangle([(x + 2, 2), (x + 2 + tw + 4, 2 + th + 4)], fill=label_bg)
            draw.text((x + 4, 3), txt, fill=label_fg, font=font)

    # Horizontal lines + Y labels along the left edge.
    for y in range(0, h, step):
        draw.line([(0, y), (w, y)], fill=line_color, width=1)
        if y % (step * 2) == 0 and y > 0:
            txt = str(y)
            tb = draw.textbbox((0, 0), txt, font=font)
            tw, th = tb[2] - tb[0], tb[3] - tb[1]
            draw.rectangle([(2, y + 2), (2 + tw + 4, y + 2 + th + 4)], fill=label_bg)
            draw.text((4, y + 3), txt, fill=label_fg, font=font)


def _list_monitors_sync() -> list[dict]:
    """Enumerate attached monitors via mss.

    Index 0 in mss is the synthetic "all monitors" rectangle; 1..N are real
    physical displays. Our public API uses 1-based indexing to match that —
    monitor 1 is the primary, 2 is the second, etc. Monitor 0 means "stitch
    all monitors together into one image" (useful for an overview shot).
    Returns the list in mss's native order with origin + size so the model
    can decide which one to screenshot/click.
    """
    import mss
    out = []
    try:
        with mss.mss() as sct:
            for idx, m in enumerate(sct.monitors):
                out.append({
                    "index": idx,
                    "left": int(m.get("left", 0)),
                    "top": int(m.get("top", 0)),
                    "width": int(m.get("width", 0)),
                    "height": int(m.get("height", 0)),
                    "is_virtual": idx == 0,
                })
    except Exception:
        # mss import failure → fall back to pyautogui primary.
        try:
            import pyautogui
            sz = pyautogui.size()
            out = [
                {"index": 0, "left": 0, "top": 0, "width": int(sz.width), "height": int(sz.height), "is_virtual": True},
                {"index": 1, "left": 0, "top": 0, "width": int(sz.width), "height": int(sz.height), "is_virtual": False},
            ]
        except Exception:
            pass
    return out


async def list_monitors() -> dict:
    """Return the list of attached monitors so the model can pick one.

    Index convention matches the `monitor` param in `screenshot` and
    `computer_click`: 0 = virtual "all screens" rectangle; 1 = primary;
    2..N = secondary displays. Every entry also reports its origin and
    size in the global virtual-screen coordinate space.
    """
    try:
        mons = await asyncio.to_thread(_list_monitors_sync)
        if not mons:
            return {"ok": False, "output": "", "error": "could not enumerate monitors"}
        summary = "; ".join(
            f"#{m['index']}: {m['width']}x{m['height']} at ({m['left']}, {m['top']})"
            + (" [virtual]" if m["is_virtual"] else "")
            for m in mons
        )
        return {"ok": True, "output": summary, "monitors": mons}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _capture_screenshot_sync(monitor: int | None = None) -> dict:
    """Blocking screenshot — called via asyncio.to_thread by the async wrappers.

    Saves a PNG inside SCREENSHOT_DIR, downscaled so its long edge is at most
    MAX_SCREENSHOT_EDGE px. Also records the scale factor so subsequent click
    coordinates (which the model picks from the downscaled image) can be
    mapped back to real screen pixels. A labeled coordinate grid is drawn
    over the image before saving so the model can read pixel positions
    directly instead of guessing from visual priors.

    `monitor` is 1-based (1=primary, 2=secondary, ...); 0 = all monitors
    stitched together; None = primary, same as 1 (preserves old behaviour).
    """
    import pyautogui
    from PIL import Image  # noqa: F401  (imported to surface clearer errors)

    origin = (0, 0)
    img = None
    if monitor is not None:
        try:
            import mss
            from PIL import Image as _PILImage
            with mss.mss() as sct:
                mons = sct.monitors
                if monitor < 0 or monitor >= len(mons):
                    raise ValueError(
                        f"monitor index {monitor} out of range; "
                        f"valid indices are 0..{len(mons) - 1} "
                        f"(call `list_monitors` to enumerate)"
                    )
                m = mons[monitor]
                raw = sct.grab(m)
                img = _PILImage.frombytes("RGB", raw.size, raw.rgb)
                origin = (int(m.get("left", 0)), int(m.get("top", 0)))
        except ValueError:
            raise
        except Exception:
            # mss failure — fall through to pyautogui below so we at least
            # return the primary screen instead of a hard error.
            img = None

    if img is None:
        img = pyautogui.screenshot()  # PIL.Image, RGB — primary display only
    orig_w, orig_h = img.size
    long_edge = max(orig_w, orig_h)
    if long_edge > MAX_SCREENSHOT_EDGE:
        scale = MAX_SCREENSHOT_EDGE / float(long_edge)
        new_size = (int(orig_w * scale), int(orig_h * scale))
        # LANCZOS is the highest quality downscaler Pillow ships.
        img = img.resize(new_size, Image.LANCZOS)

    # Compute pixel signature + window-title set BEFORE we overlay anything
    # — otherwise our own grid lines + click marker would dominate the diff
    # and the change-detection summary would always read "lots changed".
    sig = _compute_pixel_sig(img)
    titles = _enumerate_window_titles_fast()

    # Pre-compute scale (real_px / displayed_px) so the click-marker drawer
    # can translate screen coords back into the displayed image. Needs to
    # match what we'll store in `_LAST_SHOT_SCALE` after the save.
    pre_w, pre_h = img.size
    pre_scale = (
        orig_w / pre_w if pre_w else 1.0,
        orig_h / pre_h if pre_h else 1.0,
    )

    # Need RGBA for semi-transparent grid lines; convert back to RGB before
    # saving so the PNG stays small. Grid step adapts to image size — full
    # monitor shots get a 100-px grid, small `screenshot_window` crops get
    # 25-50 px, so the grid stays useful at every scale.
    img = img.convert("RGBA")
    grid_step = _pick_grid_step(img.size)
    _draw_coordinate_grid(img, step=grid_step)
    # Last-click marker — paints a red dot at the previous click's location
    # so the model can see where it just aimed. Returns True if drawn so we
    # can mention it in the output text.
    click_marked = _draw_last_click_marker(img, pre_scale, origin)
    img = img.convert("RGB")

    name = f"{uuid.uuid4().hex}.png"
    path = _ensure_screenshot_dir() / name
    # optimize=True keeps files small (matters when they're streamed back to
    # the browser as part of the chat history).
    img.save(path, format="PNG", optimize=True)
    w, h = img.size
    # Remember scale + displayed dims + origin. Guard against divide-by-zero
    # even though PIL shouldn't produce a zero-dim image in practice.
    global _LAST_SHOT_SCALE, _LAST_SHOT_DISPLAY, _LAST_SHOT_ORIGIN
    _LAST_SHOT_SCALE = (
        orig_w / w if w else 1.0,
        orig_h / h if h else 1.0,
    )
    _LAST_SHOT_DISPLAY = (w, h)
    _LAST_SHOT_ORIGIN = origin
    # Compute the change summary AFTER updating baselines via _compute_screenshot_change.
    change_summary = _compute_screenshot_change(sig, titles)
    # Snapshot foreground/focused/cursor at capture time. Cheaper than asking
    # the model to read it off the image, and stays in text so it doesn't
    # consume vision tokens.
    status_ctx = _capture_status_context()
    return {
        "name": name,
        "path": str(path),
        "width": w,
        "height": h,
        "original_width": orig_w,
        "original_height": orig_h,
        "monitor_origin": origin,
        "change_summary": change_summary,
        "click_marked": click_marked,
        "status_context": status_ctx,
    }


async def _capture_screenshot(monitor: int | None = None) -> dict:
    """Async wrapper so the blocking capture + PNG encode don't stall the loop."""
    return await asyncio.to_thread(_capture_screenshot_sync, monitor)


def _attach_shot_feedback(output: str, shot: dict) -> str:
    """Append short feedback tags to a screenshot tool's output.

    Centralises the formatting so every screenshot-returning tool surfaces
    the same cheap signals to the model:

      - `[ctx: foreground='...'; focused='...'; cursor=(x, y)]` — sourced
        from `status_context`. Tells the model in text which window is
        in front, which control has the keyboard caret, and where the
        OS cursor is pointing — answers it would otherwise have to read
        out of the screenshot pixels.
      - `[Δ: ...]` — sourced from `change_summary` + `click_marked`.
        One line describing what changed since the last screenshot
        (pixel %, new/removed windows) and a note when the red dot is
        painted on top of the previous click point.

    Tag order is `[ctx]` before `[Δ]` so the model reads "what's the
    state right now" before "what changed since last time" — the order
    a human would scan them in.

    Returns the output string unchanged when no signals are present
    (e.g. shot dict from a non-instrumented code path), so the helper
    is safe to wrap around any output without conditionals.
    """
    if not isinstance(shot, dict):
        return output

    extras: list[str] = []

    # [ctx: ...] block — the "where am I?" answer in text form.
    ctx_str = _format_status_context(shot.get("status_context"))
    if ctx_str:
        extras.append(f"[ctx: {ctx_str}]")

    # [Δ: ...] block — the "what changed since last time?" answer.
    bits: list[str] = []
    summary = shot.get("change_summary")
    if summary:
        bits.append(summary)
    if shot.get("click_marked"):
        bits.append("red dot marks the last click")
    if bits:
        extras.append(f"[Δ: {'; '.join(bits)}]")

    if not extras:
        return output
    sep = "" if output.endswith(("\n", " ")) else " "
    return f"{output}{sep}{' '.join(extras)}"


async def take_screenshot(
    monitor: int | None = None,
    with_elements: bool = False,
) -> dict:
    """Grab the current screen and make it available to the model.

    We deliberately report ONLY the displayed image size to the model, not the
    real monitor resolution or the downscale factor. That keeps the model in a
    single coordinate space — the image it sees — and the tool layer silently
    translates to real-screen pixels on click. Exposing two coordinate spaces
    invites the model to mix them up (we saw it pick y=910 on an 882-tall
    image because it was mentally reasoning in 1080p taskbar coords).

    `monitor` is a 1-based display index; 0 captures all monitors at once;
    None defaults to the primary display (same as omitting the arg). Use
    `list_monitors` to enumerate.

    When `with_elements=True`, the response also includes an `elements`
    list with `{id, role, name, bbox, enabled}` for every clickable
    control in the foreground window — ids are cached the same way
    `inspect_window` caches them, so the model can pass them straight
    to `click_element_id` without a second round-trip. Windows-only;
    on other platforms the field is simply omitted rather than erroring.
    """
    try:
        shot = await _capture_screenshot(monitor)
        suffix = ""
        if monitor is not None:
            suffix = f" (monitor {monitor})"
        # Mirror whatever spacing _pick_grid_step chose for this image so
        # the description doesn't lie about the overlay (we used to always
        # say "every 100 px" even on 320-px windows that got 25-px lines).
        grid_step = _pick_grid_step((shot["width"], shot["height"]))
        base = (
            f"Screenshot captured{suffix}. Image size: {shot['width']}x{shot['height']} pixels. "
            f"A yellow coordinate grid is overlaid every {grid_step} px with labels every "
            f"{grid_step * 2} px — read the nearest grid label to pick accurate click coordinates."
        )
        result = {
            "ok": True,
            "output": _attach_shot_feedback(base, shot),
            "image_path": shot["name"],
        }
        # --- optional accessibility-tree enrichment -----------------------
        # Runs on a worker thread because the UIA walk is blocking and
        # can take 50-300ms on heavy windows. Collapses the common
        # "screenshot, then inspect" two-call pattern into one call.
        if with_elements:
            elems = await asyncio.to_thread(_enumerate_elements_for_screenshot)
            if elems.get("ok"):
                result["elements"] = elems["elements"]
                result["window_title"] = elems["window_title"]
                result["output"] = _attach_shot_feedback(
                    (
                        f"{base} {len(elems['elements'])} clickable element"
                        f"{'s' if len(elems['elements']) != 1 else ''} "
                        f"cached from foreground window "
                        f"{elems['window_title']!r} — click any with "
                        f"`click_element_id({{\"id\": \"elN\"}})`."
                    ),
                    shot,
                )
            else:
                # Don't fail the whole screenshot on an enrichment miss.
                result["elements"] = []
                result["elements_error"] = elems.get("error")
        return result
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _screen_size() -> tuple[int, int]:
    import pyautogui
    size = pyautogui.size()
    return int(size.width), int(size.height)


def _map_to_screen(x: int | float, y: int | float) -> tuple[int, int]:
    """Translate model-supplied coords (in the last screenshot's pixel space)
    to real screen pixels, then clamp inside the virtual screen.

    The model only sees the DISPLAYED image, which may be smaller than the
    real monitor (we downscale to MAX_SCREENSHOT_EDGE to keep tokens cheap).
    Without this mapping, every click on a >1568px monitor lands in the wrong
    place proportional to the downscale factor.

    If the last screenshot was of a non-primary monitor, `_LAST_SHOT_ORIGIN`
    carries the monitor's top-left in the global virtual-screen coordinate
    system. We add that back so the click lands on the monitor the model is
    actually looking at, not on the primary display.
    """
    sx, sy = _LAST_SHOT_SCALE
    ox, oy = _LAST_SHOT_ORIGIN
    rx = float(x) * sx + ox
    ry = float(y) * sy + oy
    # Clamp inside the virtual screen bounds. With multi-monitor we can't use
    # pyautogui.size() because it only reports the primary dims; fall back to
    # mss's virtual rect when we have an origin offset set.
    vw, vh, vx0, vy0 = _virtual_screen_bounds()
    cx = max(vx0 + 1, min(int(round(rx)), vx0 + vw - 2))
    cy = max(vy0 + 1, min(int(round(ry)), vy0 + vh - 2))
    return cx, cy


def _virtual_screen_bounds() -> tuple[int, int, int, int]:
    """Return (width, height, left, top) of the full virtual screen — the
    bounding rectangle of every attached monitor. Falls back to the primary
    monitor's size when mss is unavailable.
    """
    try:
        import mss
        with mss.mss() as sct:
            v = sct.monitors[0]  # index 0 = virtual "all screens"
            return (
                int(v.get("width", 0)),
                int(v.get("height", 0)),
                int(v.get("left", 0)),
                int(v.get("top", 0)),
            )
    except Exception:
        import pyautogui
        s = pyautogui.size()
        return int(s.width), int(s.height), 0, 0


# Back-compat alias — kept so anything still calling the old name works, but
# routes through the new scaling path.
_clamp_point = _map_to_screen


def _clamp_to_image(x: int | float, y: int | float) -> tuple[int, int, str]:
    """Clamp (x, y) to the last screenshot's pixel bounds.

    Returns (clamped_x, clamped_y, notice). `notice` is an empty string when
    the input was already inside the image, otherwise a short human-readable
    note describing the adjustment so the model can see in its next turn that
    its coordinate was out of bounds and learn from it.

    Small 4B-class vision models (e.g. Gemma 4 e4b) often miss pixel coordinates
    by a few percent because they use spatial priors ("taskbar is near the
    bottom", "close button is top-right") rather than reading exact pixel
    positions. A hard reject on every overshoot stalls the conversation; a
    silent clamp hides genuine hallucinations. Clamp-with-notice gives the
    best of both: small overshoots succeed at the expected UI region, and the
    note tells the model what happened.

    Skipped when no screenshot has been taken yet — there's no reference
    coord space to clamp against.
    """
    ix, iy = int(round(float(x))), int(round(float(y)))
    w, h = _LAST_SHOT_DISPLAY
    if w <= 0 or h <= 0:
        return ix, iy, ""
    cx = max(0, min(ix, w - 1))
    cy = max(0, min(iy, h - 1))
    if (cx, cy) == (ix, iy):
        return cx, cy, ""
    return cx, cy, (
        f"note: you asked for ({ix}, {iy}) but the screenshot is only "
        f"{w}x{h}, so the point was clamped to ({cx}, {cy}). Next time, "
        f"pick coordinates inside the image to aim precisely."
    )


def _format_click_output(action: str, ix: int, iy: int, notice: str) -> str:
    """Human-readable result for the model, in image pixel-space."""
    base = f"{action} at ({ix}, {iy})."
    return f"{base} {notice}" if notice else base


def _computer_click_sync(x: int, y: int, button: str, clicks: int) -> tuple[int, int]:
    import pyautogui
    cx, cy = _map_to_screen(x, y)
    pyautogui.click(x=cx, y=cy, button=button, clicks=clicks, interval=0.05)
    # Record the click so the next screenshot paints a red dot here — the
    # single best feedback signal for "did I aim at the right control?".
    _record_click_pos(cx, cy)
    return cx, cy


async def computer_click(
    x: int | float,
    y: int | float,
    button: str = "left",
    double: bool = False,
) -> dict:
    """Move + click at pixel coordinates. Captures a screenshot right after.

    button: 'left' | 'right' | 'middle'.
    double: if true, issues two clicks with a short interval (double-click).
    """
    try:
        btn = (button or "left").lower()
        if btn not in {"left", "right", "middle"}:
            return {"ok": False, "output": "", "error": f"invalid button: {button!r}"}
        ix, iy, notice = _clamp_to_image(x, y)
        clicks = 2 if double else 1
        await asyncio.to_thread(_computer_click_sync, ix, iy, btn, clicks)
        # Brief pause so the UI has time to react before we grab the screen.
        await asyncio.sleep(0.25)
        shot = await _capture_screenshot()
        suffix = "double-clicked" if double else "clicked"
        return {
            "ok": True,
            "output": _attach_shot_feedback(
                _format_click_output(f"{btn}-{suffix}", ix, iy, notice), shot
            ),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _computer_type_sync(text: str, interval: float) -> None:
    import pyautogui
    # pyautogui.typewrite handles ASCII only; `write` is the newer alias but
    # still ASCII. For unicode (e.g. accents, emoji) we fall back to the
    # clipboard via pyperclip + ctrl+v. Detect by attempting ASCII encoding.
    try:
        text.encode("ascii")
        pyautogui.typewrite(text, interval=interval)
    except UnicodeEncodeError:
        import pyperclip
        pyperclip.copy(text)
        pyautogui.hotkey("ctrl", "v")


async def computer_type(text: str, interval: float = 0.02) -> dict:
    """Type literal text at the current cursor position."""
    try:
        if text is None:
            return {"ok": False, "output": "", "error": "text is required"}
        # Clamp interval to a sensible range so the model can't stall the app.
        iv = max(0.0, min(float(interval or 0.02), 0.2))
        # Focus-drift detection: snapshot the focused control name BEFORE
        # keystrokes go out, then again after a brief settle. If the names
        # differ, something stole focus mid-typing (notification popup,
        # background app, etc.) and the keystrokes likely landed somewhere
        # unexpected. Surface it so the model retries instead of trusting
        # a silent miss.
        focus_before = await asyncio.to_thread(_get_focused_control_name)
        await asyncio.to_thread(_computer_type_sync, str(text), iv)
        await asyncio.sleep(0.15)
        focus_after = await asyncio.to_thread(_get_focused_control_name)
        shot = await _capture_screenshot()
        preview = text if len(text) <= 80 else (text[:80] + "…")
        base = f"typed {len(text)} chars: {preview!r}"
        drift_note = _format_focus_drift(focus_before, focus_after)
        if drift_note:
            base = f"{base} {drift_note}"
        return {
            "ok": True,
            "output": _attach_shot_feedback(base, shot),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _computer_key_sync(keys: list[str]) -> None:
    import pyautogui
    if len(keys) == 1:
        pyautogui.press(keys[0])
    else:
        # hotkey() holds each modifier in order, presses the final key, then
        # releases in reverse — the right behavior for Ctrl+Shift+T etc.
        pyautogui.hotkey(*keys)


async def computer_key(keys: str | list[str]) -> dict:
    """Press a key or key combination.

    `keys` may be a single string ('enter'), a plus-joined combo
    ('ctrl+shift+t'), or a list of key names. pyautogui's accepted key names
    include: enter, tab, esc, space, backspace, delete, up, down, left,
    right, home, end, pageup, pagedown, f1..f12, ctrl, shift, alt, win,
    cmd, and any single printable character.
    """
    try:
        if isinstance(keys, list):
            seq = [str(k).strip().lower() for k in keys if str(k).strip()]
        else:
            seq = [p.strip().lower() for p in str(keys).split("+") if p.strip()]
        if not seq:
            return {"ok": False, "output": "", "error": "no keys provided"}
        # Same focus-drift dance as `computer_type` — record focus before
        # and after the key press so a popup-stealing-focus mid-press is
        # surfaced to the model. Especially useful for navigation keys
        # (Enter, Tab) where "did Enter go to the right field?" is a
        # common silent-failure mode.
        focus_before = await asyncio.to_thread(_get_focused_control_name)
        await asyncio.to_thread(_computer_key_sync, seq)
        await asyncio.sleep(0.2)
        focus_after = await asyncio.to_thread(_get_focused_control_name)
        shot = await _capture_screenshot()
        base = f"pressed: {'+'.join(seq)}"
        drift_note = _format_focus_drift(focus_before, focus_after)
        if drift_note:
            base = f"{base} {drift_note}"
        return {
            "ok": True,
            "output": _attach_shot_feedback(base, shot),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _open_app_sync(name: str) -> None:
    """Drive the OS's native search-and-launch: open search → type → Enter.

    Why this approach? Launching by URI scheme (`ms-windows-store:`) or shell
    folder (`shell:AppsFolder\\...`) requires the model to know per-app
    magic strings — knowledge small 4B models don't reliably have. The
    search-menu path is what a human would do: it works for every installed
    app, handles Win32, UWP, Store, and classic apps uniformly, and needs
    only the app's display name.

    Per-platform shortcut rationale:
      - Windows: `Win+S` opens the Search pane with focus already in the
        search box. Plain `Win` opens the Start menu and *usually* focuses
        the search field too, but on Windows 11 the focus can land on the
        pinned-apps grid instead — typing then goes nowhere. `Win+S` is
        more deterministic.
      - macOS: `Cmd+Space` summons Spotlight with focus in the input.
      - Linux: Super opens the activities/search view on GNOME and KDE;
        most DEs auto-focus the search field.
    """
    import pyautogui

    if sys.platform == "win32":
        # Win+S = Search pane (focus is guaranteed to be in the search box).
        pyautogui.hotkey("win", "s")
    elif sys.platform == "darwin":
        pyautogui.hotkey("command", "space")
    else:
        pyautogui.press("winleft")
    # Longer settle than the previous 0.6s because Windows 11's Search
    # overlay can take ~800ms to become keyboard-receptive on cold start.
    time.sleep(1.0)
    # Slightly slower per-keystroke interval — very fast typing can drop
    # characters when a just-opened search box is still hooking up input.
    pyautogui.typewrite(name, interval=0.05)
    # Give the search index a full beat to resolve and highlight the top
    # hit before we press Enter.
    time.sleep(1.2)
    pyautogui.press("enter")
    # App cold-start can be several seconds for Store apps. The fresh
    # screenshot is for verification, so it's worth waiting long enough
    # that the window is actually drawn.
    time.sleep(2.5)


def _spawn_app_with_args_sync(name: str, args: list[str]) -> tuple[bool, str]:
    """Launch an app *with command-line arguments* via the OS-appropriate
    spawner, so the caller doesn't have to know per-platform invocation.

    Returns (ok, detail). Never raises; turns subprocess failures into a
    string so the async wrapper can shape a normal tool response.

    Per-platform dispatch:
      - **Windows**: uses `cmd /c start "" <name> <args...>`. The empty
        first quoted arg is the well-known `start` quirk — it's the
        window title, not part of the command — without it `start` will
        treat `<name>` as the title when `<name>` is quoted. The `start`
        builtin resolves short names (`chrome`, `notepad`, `code`) via
        HKLM App Paths, so `chrome --guest` works without a full path.
      - **macOS**: `open -na <app> --args <args...>`. `-n` opens a new
        instance; `--args` passes the remaining argv through to the app.
      - **Linux**: direct exec via `shutil.which`. GNOME/KDE don't have
        a universal launch-with-args helper; PATH lookup is the most
        reliable thing we can do without desktop-specific code.

    We pass arguments as a list (shell=False) to avoid any shell
    injection through the model-supplied args.
    """
    import shlex
    import shutil

    if sys.platform == "win32":
        cmd = ["cmd", "/c", "start", "", name, *args]
    elif sys.platform == "darwin":
        cmd = ["open", "-na", name, "--args", *args]
    else:
        exe = shutil.which(name)
        if not exe:
            return (
                False,
                f"could not locate executable for {name!r} on PATH",
            )
        cmd = [exe, *args]

    # Pre-snapshot all top-level windows whose title already contains the
    # app name. We use this AFTER spawning to identify the newly-opened
    # window by HWND diff — critical when the app already has other
    # windows open (e.g. Chrome is already running the agent's own UI, so
    # a naive substring match for 'chrome' would focus the *existing*
    # window instead of the newly-launched guest one).
    pre_hwnds: set[int] = set()
    needle = name.strip().lower()
    if sys.platform == "win32":
        try:
            import uiautomation as auto
            for w in auto.GetRootControl().GetChildren():
                try:
                    title = (w.Name or "").lower()
                    if title and needle in title:
                        pre_hwnds.add(int(w.NativeWindowHandle))
                except Exception:
                    continue
        except Exception:
            # uiautomation import failure — skip the diff; the post-spawn
            # block below will just no-op on Windows without it.
            pass

    try:
        # Detach so the child outlives the request handler; we don't want
        # to block the agent loop waiting for a GUI process to exit.
        subprocess.Popen(
            cmd,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            stdin=subprocess.DEVNULL,
            close_fds=(sys.platform != "win32"),
        )
    except Exception as e:
        return False, f"spawn failed: {type(e).__name__}: {e}"

    # Wait for the new window to appear, then focus it. We poll every
    # 0.2s for up to 5s rather than sleeping a fixed amount, because the
    # new window's appearance is what actually matters — sleeping 2s
    # unconditionally either wastes time (if the window appears fast) or
    # loses the focus race (if it appears slow).
    matched_title = None
    if sys.platform == "win32":
        try:
            import uiautomation as auto
            deadline = time.monotonic() + 5.0
            while time.monotonic() < deadline:
                new_target = None
                for w in reversed(auto.GetRootControl().GetChildren()):
                    try:
                        title = (w.Name or "").lower()
                        if not title or needle not in title:
                            continue
                        hwnd = int(w.NativeWindowHandle)
                        if hwnd in pre_hwnds:
                            continue  # pre-existing window, skip
                        r = w.BoundingRectangle
                        if not r or (r.right - r.left) <= 0:
                            continue
                        new_target = w
                        break
                    except Exception:
                        continue
                if new_target is not None:
                    try:
                        new_target.SetActive()
                        new_target.SetFocus()
                        matched_title = new_target.Name
                    except Exception:
                        pass
                    break
                time.sleep(0.2)
        except Exception:
            pass
    # Give the UI one more beat to settle so the verification screenshot
    # below catches the final foreground state.
    time.sleep(1.0)

    detail_suffix = f" [focused new window: {matched_title!r}]" if matched_title else ""
    return True, "spawn: " + " ".join(shlex.quote(p) for p in cmd) + detail_suffix


async def open_app(name: str, args: list[str] | None = None) -> dict:
    """Launch a desktop application by display name via the OS launcher.

    Two modes, chosen by whether the caller passes ``args``:

    - **No args** → keyboard-driven search-and-launch (Win+S on Windows /
      Cmd+Space on macOS / Super on Linux). Handles Win32, UWP, Store,
      and classic apps uniformly, and needs only the display name.
    - **With args** → direct subprocess spawn. Use this when the user's
      request requires a flag or argument (guest mode, incognito, open a
      file in a specific app, new window, etc.) — the search-menu path
      cannot pass extra argv.

    The ``args`` list is passed to ``subprocess.Popen`` as a list
    (shell=False), so each element is treated as a single argv entry.
    No shell interpolation means no command injection even if the model
    emits a value with spaces or metacharacters.
    """
    try:
        if not name or not name.strip():
            return {"ok": False, "output": "", "error": "app name required"}
        clean = name.strip()
        # Guard against a runaway model handing us a giant string — the OS
        # search box cannot possibly need more than ~120 chars and anything
        # larger would just waste real time typing keystrokes.
        if len(clean) > 120:
            return {
                "ok": False,
                "output": "",
                "error": f"app name too long ({len(clean)} chars, max 120)",
            }

        # Coerce/validate args — must be a list of strings, nothing weird.
        clean_args: list[str] = []
        if args:
            if not isinstance(args, list):
                return {
                    "ok": False,
                    "output": "",
                    "error": "args must be a list of strings",
                }
            for a in args:
                if not isinstance(a, str):
                    return {
                        "ok": False,
                        "output": "",
                        "error": f"args entries must be strings, got {type(a).__name__}",
                    }
                if len(a) > 300:
                    return {
                        "ok": False,
                        "output": "",
                        "error": f"arg too long ({len(a)} chars, max 300)",
                    }
                clean_args.append(a)
            if len(clean_args) > 20:
                return {
                    "ok": False,
                    "output": "",
                    "error": f"too many args ({len(clean_args)}, max 20)",
                }

        if clean_args:
            ok, detail = await asyncio.to_thread(
                _spawn_app_with_args_sync, clean, clean_args
            )
            shot = await _capture_screenshot()
            if not ok:
                return {
                    "ok": False,
                    "output": "",
                    "error": _attach_shot_feedback(detail, shot),
                    "image_path": shot["name"],
                }
            return {
                "ok": True,
                "output": _attach_shot_feedback(
                    (
                        f"launched {clean!r} with args {clean_args!r}. Check the "
                        f"screenshot to confirm the window opened. ({detail})"
                    ),
                    shot,
                ),
                "image_path": shot["name"],
            }

        # No args → keyboard search path.
        await asyncio.to_thread(_open_app_sync, clean)
        shot = await _capture_screenshot()
        return {
            "ok": True,
            "output": _attach_shot_feedback(
                (
                    f"invoked OS launcher for {clean!r}. Check the screenshot "
                    f"to confirm the app window opened — if the launcher is still "
                    f"visible or the wrong app came up, the search didn't resolve "
                    f"as expected."
                ),
                shot,
            ),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _computer_scroll_sync(x: int, y: int, amount: int) -> tuple[int, int]:
    import pyautogui
    cx, cy = _map_to_screen(x, y)
    pyautogui.moveTo(cx, cy, duration=0.05)
    pyautogui.scroll(int(amount))
    return cx, cy


async def computer_scroll(
    x: int | float,
    y: int | float,
    direction: str = "down",
    amount: int = 5,
) -> dict:
    """Scroll the wheel at (x, y). `amount` is in 'clicks' of the wheel.

    direction='up' is positive scroll, 'down' is negative (matches pyautogui).
    """
    try:
        d = (direction or "down").lower()
        if d not in {"up", "down"}:
            return {"ok": False, "output": "", "error": f"invalid direction: {direction!r}"}
        # Clamp amount so an over-eager model doesn't scroll 9999 lines.
        a = max(1, min(int(amount or 5), 50))
        ticks = a if d == "up" else -a
        ix, iy, notice = _clamp_to_image(x, y)
        await asyncio.to_thread(_computer_scroll_sync, ix, iy, ticks)
        await asyncio.sleep(0.25)
        shot = await _capture_screenshot()
        base = f"scrolled {d} by {a} ticks at ({ix}, {iy})."
        text = f"{base} {notice}" if notice else base
        return {
            "ok": True,
            "output": _attach_shot_feedback(text, shot),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _computer_mouse_move_sync(x: int, y: int) -> tuple[int, int]:
    import pyautogui
    cx, cy = _map_to_screen(x, y)
    pyautogui.moveTo(cx, cy, duration=0.1)
    return cx, cy


async def computer_mouse_move(x: int | float, y: int | float) -> dict:
    """Move the cursor without clicking. Useful for hover states."""
    try:
        ix, iy, notice = _clamp_to_image(x, y)
        await asyncio.to_thread(_computer_mouse_move_sync, ix, iy)
        await asyncio.sleep(0.15)
        shot = await _capture_screenshot()
        return {
            "ok": True,
            "output": _attach_shot_feedback(
                _format_click_output("moved cursor to", ix, iy, notice), shot
            ),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Accessibility-tree targeted click
#
# Small vision-language models (Gemma 4 4B) are directionally OK but
# pixel-imprecise — they routinely miss click targets by 30-80px, especially
# on dense UIs. The fix is to stop relying on vision for *targeting* and use
# the OS accessibility tree instead: every UI control (button, link, menu
# item, list row) exposes a Name property and a bounding rectangle. We walk
# the tree, match by name, and click the rect's center — no pixel guessing.
#
# Windows: Microsoft UI Automation (via the `uiautomation` package). This is
# the same technology screen readers (NVDA, Narrator) use, so it covers
# Win32, UWP, WPF, and Electron apps that expose a11y (most of them do).
#
# macOS/Linux: the AX and AT-SPI APIs offer equivalent functionality but
# require different bindings — this tool returns a clear "not implemented"
# error on those platforms instead of silently falling back to pixel clicks.
# ---------------------------------------------------------------------------
def _click_element_sync(
    name: str,
    match: str,
    click_type: str,
    timeout: float,
) -> dict:
    """Synchronous worker — finds the control by accessible name and clicks
    its center. Returns a dict with ok/error and (when ok) the resolved
    screen coordinates, so the async wrapper can include them in the
    tool output for the model's benefit.
    """
    if sys.platform != "win32":
        return {
            "ok": False,
            "error": (
                f"click_element is currently only implemented on Windows; "
                f"this machine reports sys.platform={sys.platform!r}. Fall "
                f"back to `computer_click` with pixel coordinates."
            ),
        }

    try:
        import uiautomation as auto
    except ImportError:
        return {
            "ok": False,
            "error": (
                "the `uiautomation` package is not installed; run "
                "`pip install uiautomation` and restart the backend."
            ),
        }

    # Slow-down is disruptive for testing and for our own timing assumptions
    # — tell uiautomation to run at full speed.
    try:
        auto.SetGlobalSearchTimeout(timeout)
    except Exception:
        pass

    # Strategy: search inside the foreground window first (the one the user
    # is actually looking at), then fall back to a shallow desktop-wide
    # search. This keeps the common case fast while still finding popups /
    # dialogs that are technically separate top-level windows.
    needle = name.strip()

    def _matches(ctrl_name: str) -> bool:
        if not ctrl_name:
            return False
        if match == "exact":
            return ctrl_name == needle
        # Case-insensitive substring by default — small models rarely match
        # the exact casing of labels.
        return needle.lower() in ctrl_name.lower()

    def _iter_descendants(root, max_depth: int = 20):
        """BFS the a11y tree under `root`, yielding every control."""
        frontier = [(root, 0)]
        while frontier:
            ctrl, depth = frontier.pop(0)
            try:
                children = ctrl.GetChildren()
            except Exception:
                children = []
            for child in children:
                yield child
                if depth + 1 < max_depth:
                    frontier.append((child, depth + 1))

    candidates = []
    # Pass 1: foreground window.
    try:
        fg = auto.GetForegroundControl()
        if fg:
            for ctrl in _iter_descendants(fg, max_depth=20):
                try:
                    if _matches(ctrl.Name):
                        candidates.append(ctrl)
                        if len(candidates) >= 10:
                            break
                except Exception:
                    continue
    except Exception:
        pass

    # Pass 2: full desktop, shallower. Only if pass 1 found nothing.
    if not candidates:
        try:
            desktop = auto.GetRootControl()
            for ctrl in _iter_descendants(desktop, max_depth=8):
                try:
                    if _matches(ctrl.Name):
                        candidates.append(ctrl)
                        if len(candidates) >= 10:
                            break
                except Exception:
                    continue
        except Exception:
            pass

    if not candidates:
        return {
            "ok": False,
            "error": (
                f"no accessible element matches {needle!r} (match={match!r}). "
                f"The control may not expose a Name, may be inside a "
                f"webview that needs a different tool, or the spelling may "
                f"differ from what you see. Try `screenshot` to confirm the "
                f"label, or fall back to `computer_click`."
            ),
        }

    # Pick the best candidate: the first one whose bounding rect has nonzero
    # area (visible) and whose rect center lies on screen. Fall back to the
    # first candidate if nothing qualifies — better than nothing.
    chosen = None
    chosen_rect = None
    for c in candidates:
        try:
            r = c.BoundingRectangle
            if r and (r.right - r.left) > 0 and (r.bottom - r.top) > 0:
                chosen = c
                chosen_rect = r
                break
        except Exception:
            continue
    if chosen is None:
        chosen = candidates[0]
        try:
            chosen_rect = chosen.BoundingRectangle
        except Exception:
            chosen_rect = None
    if chosen_rect is None or (chosen_rect.right - chosen_rect.left) <= 0:
        return {
            "ok": False,
            "error": (
                f"found {len(candidates)} element(s) named {needle!r} but "
                f"none has a visible bounding rectangle — the control is "
                f"likely offscreen or collapsed."
            ),
        }

    cx = (chosen_rect.left + chosen_rect.right) // 2
    cy = (chosen_rect.top + chosen_rect.bottom) // 2

    # Actually click. pyautogui (not auto.Click) because it's what the rest
    # of the tool stack uses, so behaviour stays consistent with
    # `computer_click` — failsafe corner, same event timing, etc.
    import pyautogui
    pyautogui.moveTo(cx, cy, duration=0.1)
    button = {"right": "right", "middle": "middle"}.get(click_type, "left")
    if click_type == "double":
        pyautogui.doubleClick(cx, cy)
    else:
        pyautogui.click(x=cx, y=cy, button=button)
    # Mark the click so the post-click screenshot shows where it landed.
    _record_click_pos(cx, cy)

    try:
        matched_name = chosen.Name
    except Exception:
        matched_name = needle

    return {
        "ok": True,
        "x": cx,
        "y": cy,
        "matched_name": matched_name,
        "candidate_count": len(candidates),
    }


async def click_element(
    name: str,
    match: str = "contains",
    click_type: str = "left",
    timeout: float = 2.0,
) -> dict:
    """Click a UI control by its accessible name (via the OS a11y tree).

    This is far more reliable than `computer_click(x, y)` because it does
    not depend on the model's pixel-localization accuracy. The model just
    says the button's visible label and the tool looks it up in the
    accessibility tree — a screen-reader-grade API that knows exactly
    where each control is on screen.

    Args:
        name: The button/link/menu label as it appears on screen (or a
            substring; see `match`).
        match: "contains" (default, case-insensitive substring) or "exact".
        click_type: "left", "right", "middle", or "double".
        timeout: Seconds to let the a11y search run.
    """
    try:
        if not name or not str(name).strip():
            return {"ok": False, "output": "", "error": "name required"}
        clean = str(name).strip()
        if len(clean) > 200:
            return {
                "ok": False,
                "output": "",
                "error": f"name too long ({len(clean)} chars, max 200)",
            }
        if match not in ("contains", "exact"):
            return {
                "ok": False,
                "output": "",
                "error": f"match must be 'contains' or 'exact', got {match!r}",
            }
        if click_type not in ("left", "right", "middle", "double"):
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"click_type must be one of left/right/middle/double, "
                    f"got {click_type!r}"
                ),
            }
        try:
            timeout_f = max(0.1, min(float(timeout), 10.0))
        except (TypeError, ValueError):
            timeout_f = 2.0

        result = await asyncio.to_thread(
            _click_element_sync, clean, match, click_type, timeout_f
        )
        # Always return a fresh screenshot so the model can verify outcome
        # (same contract as computer_click).
        shot = await _capture_screenshot()
        if not result.get("ok"):
            return {
                "ok": False,
                "output": "",
                "error": result.get("error", "click_element failed"),
                "image_path": shot["name"],
            }
        matched = result.get("matched_name", clean)
        extra = ""
        if result.get("candidate_count", 1) > 1:
            extra = (
                f" ({result['candidate_count']} matches, clicked the first "
                f"visible one — pass match='exact' or a more specific name "
                f"if the wrong one was chosen)"
            )
        return {
            "ok": True,
            "output": _attach_shot_feedback(
                (
                    f"{click_type}-clicked accessible element {matched!r} at "
                    f"({result['x']}, {result['y']}){extra}."
                ),
                shot,
            ),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Window focus control
#
# The single biggest source of silent computer-use failures is focus loss:
# you launch Chrome, the user's attention (or Windows's foreground-stealing
# prevention) hands focus back to another app, and the model's next
# `computer_type` goes into the WRONG window. The model can't tell this
# from a screenshot alone — screenshots capture the *visible* pixels, not
# which window has keyboard focus.
#
# `focus_window(name)` lets the model re-assert focus on the window it
# intends to drive before typing. It walks the top-level windows via UIA
# and calls SetFocus on the best match.
# ---------------------------------------------------------------------------
def _focus_window_sync(name: str) -> dict:
    """Locate a top-level window whose title contains ``name`` (case-
    insensitive) and set it as the foreground window.

    Returns a dict describing what was matched (or why it failed) so the
    async wrapper can shape a useful tool result.
    """
    if sys.platform != "win32":
        return {
            "ok": False,
            "error": (
                f"focus_window is currently only implemented on Windows; "
                f"this machine reports sys.platform={sys.platform!r}."
            ),
        }
    try:
        import uiautomation as auto
    except ImportError:
        return {
            "ok": False,
            "error": (
                "the `uiautomation` package is not installed; run "
                "`pip install uiautomation` and restart the backend."
            ),
        }

    needle = name.strip().lower()
    try:
        desktop = auto.GetRootControl()
        children = desktop.GetChildren()
    except Exception as e:
        return {"ok": False, "error": f"UIA enumeration failed: {type(e).__name__}: {e}"}

    # Identify the currently-foreground window so we can de-prioritise it
    # when disambiguating. If the model is calling `focus_window`, it
    # almost always wants to switch TO a DIFFERENT window, not re-focus
    # the one that's already in front — so the foreground one is the
    # worst candidate to pick by default.
    fg_hwnd = None
    try:
        fg = auto.GetForegroundControl()
        if fg is not None:
            fg_hwnd = int(fg.NativeWindowHandle)
    except Exception:
        pass

    # Iterate in reverse order — Windows tends to list newer top-level
    # windows later, so reversed order approximates "most-recent first".
    # Keep only visible, non-zero-area windows whose title contains
    # `needle`. Record (window, hwnd, title) for each.
    matches: list[tuple] = []
    for w in reversed(children):
        try:
            title = (w.Name or "")
            low = title.lower()
            if not low or needle not in low:
                continue
            r = w.BoundingRectangle
            if not r or (r.right - r.left) <= 0 or (r.bottom - r.top) <= 0:
                continue
            hwnd = int(w.NativeWindowHandle)
            matches.append((w, hwnd, title))
        except Exception:
            continue

    if not matches:
        return {
            "ok": False,
            "error": (
                f"no visible top-level window title contains {needle!r}. "
                f"Make sure the app is actually running — if you just "
                f"launched it, wait a moment for its window to appear."
            ),
        }

    # Prefer a match that isn't the currently-foreground window. Falling
    # back to the foreground window when it's the only candidate is fine.
    non_fg = [m for m in matches if m[1] != fg_hwnd]
    ordered = non_fg if non_fg else matches
    target, target_hwnd, target_title = ordered[0]

    try:
        # SetActive + SetFocus covers most apps. SetFocus alone sometimes
        # fails to override Windows's foreground-stealing prevention; the
        # standard workaround is to also minimize-restore the target or
        # briefly attach our thread's input queue to the foreground
        # thread. uiautomation's SetActive wraps those tricks.
        target.SetActive()
        target.SetFocus()
    except Exception as e:
        return {
            "ok": False,
            "error": f"SetFocus failed on {target_title!r}: {type(e).__name__}: {e}",
        }

    # Include the full list of candidate titles when there was ambiguity,
    # so the model can retry with a more specific substring (e.g.
    # "Guest" instead of "Chrome") if we picked the wrong window.
    candidate_titles = [t for _, _, t in matches]
    return {
        "ok": True,
        "matched_name": target_title,
        "candidate_count": len(matches),
        "candidate_titles": candidate_titles,
    }


async def focus_window(name: str) -> dict:
    """Bring a window to the foreground by a substring of its title.

    Call this *immediately before* `computer_type` / `computer_key` when
    you've just launched an app or clicked something that might have
    shifted focus. Typing keystrokes go to whichever window is foreground
    at the time — there is no way to "type into window X" directly; you
    have to make X foreground first.
    """
    try:
        if not name or not str(name).strip():
            return {"ok": False, "output": "", "error": "name required"}
        clean = str(name).strip()
        if len(clean) > 200:
            return {
                "ok": False,
                "output": "",
                "error": f"name too long ({len(clean)} chars, max 200)",
            }
        result = await asyncio.to_thread(_focus_window_sync, clean)
        # Windows needs a short beat after SetForeground before keystrokes
        # reach the target reliably.
        await asyncio.sleep(0.4)
        shot = await _capture_screenshot()
        if not result.get("ok"):
            return {
                "ok": False,
                "output": "",
                "error": result.get("error", "focus_window failed"),
                "image_path": shot["name"],
            }
        # Build a one-line output. When there was ambiguity (more than one
        # window matched), list the other candidate titles so the model
        # can immediately retry with a more specific substring if we
        # picked the wrong one — crucial when the same app has several
        # windows open (e.g. two Chrome windows, one of which is the
        # agent's own UI).
        base = (
            f"focused window {result['matched_name']!r}. Keyboard input "
            f"now targets that window."
        )
        cnt = int(result.get("candidate_count") or 1)
        if cnt > 1:
            titles = result.get("candidate_titles") or []
            other = [t for t in titles if t != result["matched_name"]]
            if other:
                base += (
                    f" ({cnt} windows matched {name!r}; if this is the "
                    f"wrong one, retry with a more specific substring "
                    f"from one of these: {other!r})"
                )
        return {
            "ok": True,
            "output": _attach_shot_feedback(base, shot),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Web tools
#
# web_search  — DuckDuckGo (no API key). Returns title / url / snippet per hit.
# fetch_url   — Download + extract main article text (trafilatura) so the
#               model sees readable prose instead of raw HTML. Includes an
#               SSRF guard so a prompt-injected model can't be coerced into
#               fetching internal-network addresses or the loopback.
# ---------------------------------------------------------------------------
FETCH_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/124.0 Safari/537.36 Gigachat/1.0"
)
FETCH_MAX_BYTES = 2_000_000        # cap raw HTML at 2MB
FETCH_DEFAULT_MAX_CHARS = 15000    # cap extracted text returned to the model
FETCH_TIMEOUT_SEC = 30.0


def _resolves_to_public_ip(hostname: str) -> bool:
    """DNS-resolve `hostname` and return True iff every answer is a public IP.

    This is the SSRF backstop. Without it, an attacker could register a
    domain that resolves to 127.0.0.1 (or an RFC1918 host) and have the
    agent fetch data from localhost / the LAN.
    """
    try:
        infos = socket.getaddrinfo(hostname, None)
    except socket.gaierror:
        # DNS failed — let the actual fetch report the real error upstream.
        return True
    for *_rest, sockaddr in infos:
        addr = sockaddr[0]
        try:
            ip = ipaddress.ip_address(addr)
        except ValueError:
            continue
        if (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_multicast
            or ip.is_reserved
            or ip.is_unspecified
        ):
            return False
    return True


def _is_safe_url(url: str) -> tuple[bool, str]:
    """Return (ok, reason). Reject non-http(s) schemes and internal hosts."""
    try:
        u = urlparse(url)
    except Exception:
        return False, "unparseable URL"
    if u.scheme not in {"http", "https"}:
        return False, f"only http(s) allowed (got {u.scheme!r})"
    host = (u.hostname or "").strip()
    if not host:
        return False, "missing host"
    if host.lower() in {"localhost", "ip6-localhost", "ip6-loopback", "0.0.0.0"}:
        return False, "loopback hosts are blocked"
    # Literal-IP URLs: judge directly without a DNS round-trip.
    try:
        ip = ipaddress.ip_address(host)
        if (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_multicast
            or ip.is_reserved
            or ip.is_unspecified
        ):
            return False, f"{host} is not a public IP"
        return True, ""
    except ValueError:
        pass  # not an IP literal — fall through to DNS check below
    return (True, "") if True else (False, "")  # DNS check happens async


def _ddg_search_sync(query: str, max_results: int, region: str | None) -> list[dict]:
    """Blocking DuckDuckGo call; wrapped in asyncio.to_thread by the tool."""
    from ddgs import DDGS
    kwargs: dict[str, Any] = {"max_results": max_results}
    if region:
        kwargs["region"] = region
    with DDGS() as d:
        return list(d.text(query, **kwargs))


async def web_search(query: str, max_results: int = 5, region: str | None = None) -> dict:
    """Search the web via DuckDuckGo.

    Returns one numbered hit per line with title + URL + snippet so the
    model can pick a promising result and pass it to `fetch_url`.
    """
    q = (query or "").strip()
    if not q:
        return {"ok": False, "output": "", "error": "empty query"}
    # Clamp so the model can't pull hundreds of results into context.
    n = max(1, min(int(max_results or 5), 20))
    try:
        hits = await asyncio.to_thread(_ddg_search_sync, q, n, region)
    except Exception as e:
        return {"ok": False, "output": "", "error": f"search failed: {type(e).__name__}: {e}"}
    if not hits:
        return {"ok": True, "output": f"(no results for {q!r})"}
    lines: list[str] = []
    for i, r in enumerate(hits, 1):
        title = (r.get("title") or "").strip()
        url = (r.get("href") or r.get("url") or "").strip()
        snippet = (r.get("body") or "").strip()
        lines.append(f"{i}. {title}\n   {url}\n   {snippet}")
    return {"ok": True, "output": _clip("\n\n".join(lines))}


def _trafilatura_extract_sync(html: str, url: str) -> str:
    import trafilatura
    extracted = trafilatura.extract(
        html,
        url=url,
        include_comments=False,
        include_tables=True,
        favor_precision=True,
    )
    return extracted or ""


async def fetch_url(url: str, max_chars: int = FETCH_DEFAULT_MAX_CHARS) -> dict:
    """Fetch a URL and return its main readable text.

    Two-layer safety:
      - Scheme / host whitelist on the URL itself (rejects file://, private
        IP literals, localhost, etc.).
      - DNS resolution check on hostnames so a domain that cname's to a
        private IP is also rejected.

    Raw HTML is capped at FETCH_MAX_BYTES; the extracted prose is truncated
    to `max_chars` before being handed back to the model.
    """
    u = (url or "").strip()
    if not u:
        return {"ok": False, "output": "", "error": "empty url"}

    ok, reason = _is_safe_url(u)
    if not ok:
        return {"ok": False, "output": "", "error": f"blocked: {reason}"}

    hostname = urlparse(u).hostname or ""
    # For hostname URLs (not IP literals) we still need a DNS round-trip.
    try:
        ipaddress.ip_address(hostname)
    except ValueError:
        resolves_public = await asyncio.to_thread(_resolves_to_public_ip, hostname)
        if not resolves_public:
            return {"ok": False, "output": "", "error": f"blocked: {hostname} resolves to a non-public IP"}

    try:
        headers = {
            "User-Agent": FETCH_USER_AGENT,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.8",
        }
        async with httpx.AsyncClient(
            timeout=FETCH_TIMEOUT_SEC,
            follow_redirects=True,
            headers=headers,
        ) as client:
            # Stream so we can enforce a hard byte cap on pathological pages.
            async with client.stream("GET", u) as r:
                r.raise_for_status()
                ctype = r.headers.get("content-type", "").lower()
                chunks: list[bytes] = []
                total = 0
                async for chunk in r.aiter_bytes():
                    total += len(chunk)
                    if total > FETCH_MAX_BYTES:
                        break
                    chunks.append(chunk)
                raw = b"".join(chunks)
                final_url = str(r.url)
    except httpx.HTTPError as e:
        return {"ok": False, "output": "", "error": f"HTTP error: {e}"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}

    # Decode — httpx would do this for us, but we streamed raw bytes so we
    # handle encoding manually. UTF-8 with replacement is fine for the model.
    if "charset=" in ctype:
        # Best-effort encoding pull from the Content-Type header.
        enc = ctype.split("charset=", 1)[1].split(";", 1)[0].strip() or "utf-8"
    else:
        enc = "utf-8"
    try:
        html = raw.decode(enc, errors="replace")
    except LookupError:
        html = raw.decode("utf-8", errors="replace")

    # Plain-text / JSON / markdown — skip trafilatura, return as-is (capped).
    if not any(t in ctype for t in ("html", "xml")):
        body = html
    else:
        body = await asyncio.to_thread(_trafilatura_extract_sync, html, final_url)
        if not body:
            return {
                "ok": False,
                "output": "",
                "error": "could not extract readable content (site may be JS-rendered or blocked)",
            }

    cap = max(500, min(int(max_chars or FETCH_DEFAULT_MAX_CHARS), 50000))
    if len(body) > cap:
        body = body[:cap] + f"\n\n... [truncated, {len(body) - cap} chars omitted]"
    header_line = f"# Fetched: {final_url}\n\n"
    return {"ok": True, "output": header_line + body}


# ---------------------------------------------------------------------------
# http_request — generic HTTP client for REST APIs
# ---------------------------------------------------------------------------
# Unlike fetch_url (GET + HTML extraction only) this tool exposes the full
# HTTP surface so the agent can talk to arbitrary REST APIs: Slack, GitHub,
# Notion, OpenAI, a home router's admin API, anything. The shape matches
# what a human would type in curl.
#
# Secrets: the model never sees credential values. Instead it references
# them with a placeholder like `Bearer {{secret:GITHUB_TOKEN}}` in any
# header or body value. The backend calls ``db.get_secret_value`` right
# before sending, substitutes the resolved value, and logs the placeholder
# form (not the substituted form) in the tool-result output so secrets
# never leak into the conversation transcript.
HTTP_REQUEST_MAX_BODY_BYTES = 2_000_000        # raw response cap
HTTP_REQUEST_DEFAULT_OUTPUT_CHARS = 20_000     # string-form body cap in tool output
HTTP_REQUEST_TIMEOUT_SEC = 60.0
HTTP_REQUEST_MAX_TIMEOUT_SEC = 120.0
HTTP_REQUEST_ALLOWED_METHODS = frozenset(
    {"GET", "POST", "PUT", "PATCH", "DELETE", "HEAD", "OPTIONS"}
)
_SECRET_PLACEHOLDER_RE = re.compile(r"\{\{secret:([A-Za-z_][A-Za-z0-9_]{0,63})\}\}")


def _substitute_secrets(text: str) -> tuple[str, list[str], list[str]]:
    """Replace every ``{{secret:NAME}}`` with its stored value.

    Returns ``(resolved_text, resolved_names, missing_names)``. Callers abort
    the request when ``missing_names`` is non-empty so the model sees a clean
    error instead of a literal `{{secret:X}}` flying out to some API.
    """
    if not text or "{{secret:" not in text:
        return text, [], []
    resolved: list[str] = []
    missing: list[str] = []

    def _sub(m: re.Match) -> str:
        name = m.group(1)
        value = db.get_secret_value(name)
        if value is None:
            missing.append(name)
            return m.group(0)
        resolved.append(name)
        return value

    return _SECRET_PLACEHOLDER_RE.sub(_sub, text), resolved, missing


async def http_request(
    url: str,
    method: str = "GET",
    headers: dict | None = None,
    body: Any = None,
    query: dict | None = None,
    timeout: float = HTTP_REQUEST_TIMEOUT_SEC,
    allow_private: bool = False,
    max_output_chars: int = HTTP_REQUEST_DEFAULT_OUTPUT_CHARS,
) -> dict:
    """Make an arbitrary HTTP request and return status + headers + body.

    Parameters roughly mirror curl flags. ``body`` accepts a string (sent
    verbatim) or a dict/list (JSON-serialized; Content-Type is auto-set to
    application/json unless the caller already supplied one). ``headers``
    and any string value inside ``body`` may contain ``{{secret:NAME}}``
    placeholders — these are substituted from the SQLite secrets store
    just before sending and never echoed back in the tool output.

    By default the SSRF guard from ``fetch_url`` applies — loopback, RFC1918,
    link-local, and reserved IPs (including DNS-resolved hostnames) are
    refused. Set ``allow_private=True`` to bypass for legitimate LAN targets
    (home router, Home Assistant, local dev server). That still requires the
    usual write-class approval since http_request is gated by the permission
    mode.
    """
    u = (url or "").strip()
    if not u:
        return {"ok": False, "output": "", "error": "url is required"}
    m = (method or "GET").strip().upper()
    if m not in HTTP_REQUEST_ALLOWED_METHODS:
        return {
            "ok": False,
            "output": "",
            "error": f"unsupported method {m!r} (allowed: {sorted(HTTP_REQUEST_ALLOWED_METHODS)})",
        }

    # SSRF guard — same checks as fetch_url unless the user opted out.
    if not allow_private:
        ok, reason = _is_safe_url(u)
        if not ok:
            return {"ok": False, "output": "", "error": f"blocked URL: {reason}"}
        parsed = urlparse(u)
        hostname = (parsed.hostname or "").strip()
        if hostname:
            try:
                ipaddress.ip_address(hostname)  # literal IP: already vetted above
            except ValueError:
                resolves_public = await asyncio.to_thread(_resolves_to_public_ip, hostname)
                if not resolves_public:
                    return {
                        "ok": False,
                        "output": "",
                        "error": f"{hostname} resolves to a non-public IP",
                    }
    else:
        # Minimum scheme check even when private ranges are allowed —
        # `file://` / `javascript:` / `data:` are never acceptable.
        try:
            pu = urlparse(u)
        except Exception:
            return {"ok": False, "output": "", "error": "unparseable URL"}
        if pu.scheme not in {"http", "https"}:
            return {"ok": False, "output": "", "error": f"only http(s) allowed (got {pu.scheme!r})"}

    # Normalize + substitute secrets in headers. Placeholders resolve in both
    # keys and values so the model can hide a header name too if it wants
    # (rare, but not disallowed).
    in_headers = {str(k): str(v) for k, v in (headers or {}).items()}
    resolved_secrets: set[str] = set()
    missing_secrets: set[str] = set()
    out_headers: dict[str, str] = {}
    for k, v in in_headers.items():
        nk, r1, m1 = _substitute_secrets(k)
        nv, r2, m2 = _substitute_secrets(v)
        out_headers[nk] = nv
        resolved_secrets.update(r1)
        resolved_secrets.update(r2)
        missing_secrets.update(m1)
        missing_secrets.update(m2)

    # Body handling: dict/list → JSON, string → verbatim, None → no body.
    request_body: bytes | None = None
    body_preview_for_output: str = ""
    if body is None:
        pass
    elif isinstance(body, (dict, list)):
        try:
            body_str = json.dumps(body)
        except (TypeError, ValueError) as e:
            return {"ok": False, "output": "", "error": f"body not JSON-serialisable: {e}"}
        body_str, rb, mb = _substitute_secrets(body_str)
        resolved_secrets.update(rb)
        missing_secrets.update(mb)
        request_body = body_str.encode("utf-8")
        body_preview_for_output = body_str
        # Default Content-Type to application/json if the caller didn't set one.
        if not any(k.lower() == "content-type" for k in out_headers):
            out_headers["Content-Type"] = "application/json"
    elif isinstance(body, str):
        body_str, rb, mb = _substitute_secrets(body)
        resolved_secrets.update(rb)
        missing_secrets.update(mb)
        request_body = body_str.encode("utf-8")
        body_preview_for_output = body_str
    else:
        return {"ok": False, "output": "", "error": f"body must be string / dict / list / null (got {type(body).__name__})"}

    if missing_secrets:
        return {
            "ok": False,
            "output": "",
            "error": (
                f"unknown secret(s): {sorted(missing_secrets)}. Add them in "
                "Settings → Secrets first, or reference an existing name."
            ),
        }

    # Final sanity caps
    t = max(1.0, min(float(timeout or HTTP_REQUEST_TIMEOUT_SEC), HTTP_REQUEST_MAX_TIMEOUT_SEC))
    cap = max(500, min(int(max_output_chars or HTTP_REQUEST_DEFAULT_OUTPUT_CHARS), 100_000))

    try:
        async with httpx.AsyncClient(
            timeout=t,
            follow_redirects=True,
            max_redirects=5,
        ) as c:
            r = await c.request(
                m,
                u,
                headers=out_headers or None,
                params=query or None,
                content=request_body,
            )
            # Stream-free read because we already guarded the method and size
            # via the timeout; for really large responses we truncate below.
            raw = r.content
    except httpx.TimeoutException:
        return {"ok": False, "output": "", "error": f"request timed out after {t}s"}
    except httpx.HTTPError as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}
    except Exception as e:  # noqa: BLE001 — surface network exceptions cleanly
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}

    if len(raw) > HTTP_REQUEST_MAX_BODY_BYTES:
        raw = raw[:HTTP_REQUEST_MAX_BODY_BYTES]
        body_truncated_note = (
            f"\n\n... [response truncated at {HTTP_REQUEST_MAX_BODY_BYTES} bytes]"
        )
    else:
        body_truncated_note = ""

    # Decode for the text output. If the response is binary (images / zip /
    # etc.) fall back to a size note rather than mojibake. We key off the
    # actual Content-Type header, not the bytes.
    resp_ctype = r.headers.get("content-type", "").lower()
    looks_textual = (
        resp_ctype.startswith("text/")
        or "json" in resp_ctype
        or "xml" in resp_ctype
        or "javascript" in resp_ctype
        or "yaml" in resp_ctype
        or not resp_ctype  # unknown → best-effort decode
    )
    if looks_textual:
        try:
            body_text = raw.decode(r.encoding or "utf-8", errors="replace")
        except LookupError:
            body_text = raw.decode("utf-8", errors="replace")
    else:
        body_text = f"[binary response, {len(raw)} bytes, content-type={resp_ctype!r}]"

    if len(body_text) > cap:
        body_text = body_text[:cap] + f"\n\n... [output truncated, {len(body_text) - cap} chars omitted]"
    body_text += body_truncated_note

    # Redact the Authorization / X-API-Key / Cookie values in the echoed
    # request headers so a tool-result crawl doesn't expose the substituted
    # secret. Placeholder-form is already safe; plaintext credentials
    # (someone chose to paste the token directly into `headers`) get masked
    # using the same heuristic as HTTP clients' debug output.
    redacted_headers = {
        k: ("***" if k.lower() in {"authorization", "x-api-key", "cookie", "set-cookie", "proxy-authorization"} else v)
        for k, v in in_headers.items()
    }

    # Response headers echo — small allowlist + truncation so a chatty
    # server doesn't dump 40 Set-Cookie lines into the transcript.
    resp_headers_out: list[str] = []
    for key, value in list(r.headers.items())[:20]:
        v = value if len(value) <= 200 else value[:200] + "…"
        if key.lower() in {"set-cookie", "authorization"}:
            v = "***"
        resp_headers_out.append(f"  {key}: {v}")

    parts: list[str] = [
        f"HTTP {r.status_code} {r.reason_phrase}",
        f"{m} {u}",
    ]
    if redacted_headers:
        parts.append("Request headers:")
        for k, v in redacted_headers.items():
            parts.append(f"  {k}: {v}")
    if body_preview_for_output:
        preview = body_preview_for_output
        if len(preview) > 2000:
            preview = preview[:2000] + " …"
        parts.append(f"Request body:\n{preview}")
    if resolved_secrets:
        parts.append(f"(resolved secrets: {sorted(resolved_secrets)})")
    parts.append("")
    if resp_headers_out:
        parts.append("Response headers:")
        parts.extend(resp_headers_out)
        parts.append("")
    parts.append("Response body:")
    parts.append(body_text)

    output = "\n".join(parts)

    # Defence-in-depth: if we substituted any secrets into the outbound
    # request, scrub those raw values out of the final output before handing
    # it back to the model. Most servers don't echo credentials, but some
    # (httpbin.org, misconfigured reverse proxies, errant debug endpoints)
    # happily include them in the response body. We don't want the secret
    # landing in the model's context just because the other side was
    # careless. Longest-first replacement keeps substrings of one secret
    # from being partially masked by another.
    if resolved_secrets:
        for name in sorted(resolved_secrets, key=lambda n: len(db.get_secret_value(n) or ""), reverse=True):
            val = db.get_secret_value(name)
            if val and len(val) >= 4:  # tiny values are too likely to be benign false positives
                output = output.replace(val, "***")

    return {
        "ok": 200 <= r.status_code < 400,
        "output": output,
        "status": r.status_code,
    }


# ---------------------------------------------------------------------------
# File editing / search / listing tools
#
# edit_file   — surgical exact-string replacement (cheaper than read+write).
# grep_tool   — content search. Uses ripgrep if installed, falls back to
#               a Python-only implementation otherwise.
# glob_tool   — fast filename pattern match. Supports `**` for recursion.
# ---------------------------------------------------------------------------
def _checkpoint_file(conv_id: str | None, resolved_path: Path, current_bytes: bytes) -> None:
    """Snapshot `resolved_path` before a destructive write.

    Stores under CHECKPOINT_DIR / <conv_id> / <message_id_placeholder> / <hash>.bin
    so the user can restore a specific earlier state later. A missing
    `conv_id` disables the checkpoint (used by subagents / CLI tests).
    """
    if not conv_id:
        return
    try:
        # Use a filename derived from the real path so we can find it on
        # restore without needing to keep a separate index.
        import hashlib
        key = hashlib.sha1(str(resolved_path).encode("utf-8")).hexdigest()[:16]
        # Include microseconds in the stamp so lexicographic sort order ==
        # chronological order even when two edits happen inside the same
        # wall-clock second. Without the %f component, two rapid edits land
        # in stamps that differ only in the random uuid suffix, so
        # `sorted(stamps)` picks the "earlier" one at random — breaking any
        # caller (model, user, or test) that wants to restore to a specific
        # point in time. The uuid suffix is still appended to disambiguate
        # the rare sub-microsecond collision (two edits on different threads).
        stamp = datetime.datetime.now().strftime("%Y%m%dT%H%M%S_%f") + f"_{uuid.uuid4().hex[:4]}"
        dest = CHECKPOINT_DIR / conv_id / stamp / f"{key}.bin"
        dest.parent.mkdir(parents=True, exist_ok=True)
        # Store original path next to the content so restore() knows where
        # to put it. One tiny sidecar file keeps things dirt simple.
        (dest.parent / f"{key}.path").write_text(str(resolved_path), encoding="utf-8")
        dest.write_bytes(current_bytes)
        # JIT retention: keep only the most recent N stamps for this
        # conversation. Cheap (listdir + partial delete) and bounded; see
        # `retention.MAX_CHECKPOINTS_PER_CONV`. Periodic sweep in app.py
        # handles age-based expiry and orphan conv directories.
        retention.trim_conv_checkpoints(CHECKPOINT_DIR, conv_id)
    except Exception:
        # Checkpointing is a best-effort safety net — don't fail the actual
        # write just because we couldn't make a backup.
        pass


async def edit_file(
    cwd: str,
    path: str,
    old_string: str,
    new_string: str,
    replace_all: bool = False,
    conv_id: str | None = None,
) -> dict:
    """Replace `old_string` with `new_string` in `path`.

    If `replace_all` is False (default), `old_string` must appear exactly once
    — otherwise the call fails with a helpful error so the model can widen
    the context to disambiguate. This mirrors Claude Code's Edit tool and
    avoids the classic LLM pitfall of over-zealous find/replace.
    """
    try:
        if old_string == new_string:
            return {
                "ok": False,
                "output": "",
                "error": "old_string and new_string are identical — nothing to do",
            }
        p = _resolve(cwd, path, conv_id)
        if not p.exists():
            return {"ok": False, "output": "", "error": f"file not found: {p}"}
        if p.is_dir():
            return {"ok": False, "output": "", "error": f"is a directory: {p}"}
        # Read-first guard: the model must have read this file in this
        # conversation before editing, so it can't hallucinate the old_string
        # from training-data priors. The guard is a no-op when conv_id is
        # None (tests / CLI one-offs).
        guard = _require_prior_read(conv_id, p)
        if guard:
            return {"ok": False, "output": "", "error": guard}
        raw = p.read_bytes()
        text = raw.decode("utf-8", errors="replace")
        count = text.count(old_string)
        if count == 0:
            return {
                "ok": False,
                "output": "",
                "error": "old_string not found in file (check whitespace and line endings)",
            }
        if count > 1 and not replace_all:
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"old_string appears {count} times; either include more surrounding "
                    "context to make it unique, or pass replace_all=true"
                ),
            }
        new_text = (
            text.replace(old_string, new_string)
            if replace_all
            else text.replace(old_string, new_string, 1)
        )
        _checkpoint_file(conv_id, p, raw)
        # Bytes, not text: the file was read as bytes, and `text` preserves
        # whatever original line endings it had. `write_text` on Windows would
        # append a CR to every `\n` in the file — silently flipping a Unix-
        # formatted file to CRLF on the first edit and, worse, making the
        # next LF-authored `old_string` fail to match (since the file now
        # stores CRLF). `write_bytes` keeps the round-trip byte-identical
        # for unchanged regions.
        p.write_bytes(new_text.encode("utf-8"))
        # Return a tiny unified diff so the model can see what actually changed
        # without needing a follow-up read_file.
        diff = _compact_diff(text, new_text, str(p))
        suffix = "s" if count != 1 else ""
        return {
            "ok": True,
            "output": f"edited {p} ({count} replacement{suffix})\n\n{diff}",
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _compact_diff(old: str, new: str, path: str, max_lines: int = 40) -> str:
    """Render a small unified diff suitable for inline display."""
    diff_iter = difflib.unified_diff(
        old.splitlines(keepends=False),
        new.splitlines(keepends=False),
        fromfile=f"a/{path}",
        tofile=f"b/{path}",
        n=2,
    )
    lines = list(diff_iter)
    if len(lines) > max_lines:
        lines = lines[:max_lines] + [f"... [truncated, {len(lines) - max_lines} more lines]"]
    return "\n".join(lines)


def _rg_available() -> bool:
    """Cache-free check for ripgrep on PATH. Called once per grep call."""
    return shutil.which("rg") is not None


def _python_grep(
    base: Path,
    pattern: str,
    glob: str | None,
    case_insensitive: bool,
    output_mode: str,
    head_limit: int,
) -> str:
    """Pure-Python fallback when ripgrep isn't installed.

    Walks `base` recursively, reads each text file, and applies a compiled
    regex. Binary files are skipped by a cheap NUL-byte heuristic so we don't
    dump image/PDF bytes into the model's context.
    """
    flags = re.IGNORECASE if case_insensitive else 0
    try:
        rx = re.compile(pattern, flags)
    except re.error as e:
        raise ValueError(f"bad regex: {e}") from e
    hits: list[str] = []
    matched_files: set[str] = set()
    count_by_file: dict[str, int] = {}
    for root, dirs, files in os.walk(base):
        # Skip obvious non-source directories to keep this snappy.
        dirs[:] = [d for d in dirs if d not in {".git", "node_modules", ".venv", "__pycache__", "dist", "build"}]
        for fname in files:
            if glob and not fnmatch.fnmatch(fname, glob):
                continue
            fpath = Path(root) / fname
            try:
                # Quick binary check: peek first 1KB for NUL bytes.
                with fpath.open("rb") as fh:
                    head = fh.read(1024)
                if b"\x00" in head:
                    continue
                text = fpath.read_text(encoding="utf-8", errors="replace")
            except Exception:
                continue
            local_hits: list[str] = []
            for lineno, line in enumerate(text.splitlines(), 1):
                if rx.search(line):
                    local_hits.append(f"{fpath}:{lineno}:{line.rstrip()}")
                    count_by_file[str(fpath)] = count_by_file.get(str(fpath), 0) + 1
            if local_hits:
                matched_files.add(str(fpath))
                hits.extend(local_hits)
                if len(hits) >= head_limit * 10:  # plenty of slack; truncation below
                    break
        if len(hits) >= head_limit * 10:
            break
    if output_mode == "files_with_matches":
        out = sorted(matched_files)
    elif output_mode == "count":
        out = [f"{k}:{v}" for k, v in sorted(count_by_file.items())]
    else:  # content
        out = hits
    out = out[:head_limit]
    if not out:
        return "(no matches)"
    return "\n".join(out)


async def grep_tool(
    cwd: str,
    pattern: str,
    path: str = ".",
    glob: str | None = None,
    case_insensitive: bool = False,
    output_mode: str = "files_with_matches",
    head_limit: int = 100,
    conv_id: str | None = None,
) -> dict:
    """Search files for a regex, preferring ripgrep when available.

    `output_mode` controls the shape of the output:
      - "files_with_matches" (default): unique list of paths that contain a hit
      - "content": lines with filename:lineno prefix
      - "count": one `path:N` per file

    `conv_id` lets relative `path` values follow wherever bash `cd`-ed to
    in this conversation — same policy as `read_file` / `write_file`.
    """
    if not pattern:
        return {"ok": False, "output": "", "error": "empty pattern"}
    hl = max(1, min(int(head_limit or 100), 2000))
    base = _resolve(cwd, path, conv_id)
    if not base.exists():
        return {"ok": False, "output": "", "error": f"path not found: {base}"}
    om = output_mode if output_mode in {"files_with_matches", "content", "count"} else "files_with_matches"
    try:
        if _rg_available():
            # Let ripgrep do the heavy lifting — 50-100x faster than Python
            # for big repos, and it respects .gitignore by default.
            cmd = ["rg", "--no-config", "--color=never"]
            if case_insensitive:
                cmd.append("-i")
            if om == "files_with_matches":
                cmd.append("-l")
            elif om == "count":
                cmd.append("-c")
            else:
                cmd.append("-n")
            if glob:
                cmd += ["--glob", glob]
            cmd += ["-e", pattern, str(base)]
            result = await asyncio.to_thread(
                subprocess.run,
                cmd,
                capture_output=True,
                text=True,
                timeout=30,
            )
            lines = (result.stdout or "").splitlines()
            if not lines:
                return {"ok": True, "output": "(no matches)"}
            if len(lines) > hl:
                lines = lines[:hl] + [f"... [truncated, {len(lines) - hl} more]"]
            return {"ok": True, "output": _clip("\n".join(lines))}
        # Fallback: Python regex walker.
        out = await asyncio.to_thread(
            _python_grep, base, pattern, glob, case_insensitive, om, hl
        )
        return {"ok": True, "output": _clip(out)}
    except ValueError as e:
        return {"ok": False, "output": "", "error": str(e)}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def glob_tool(cwd: str, pattern: str, path: str = ".", conv_id: str | None = None) -> dict:
    """Find files by pathname pattern. Supports `**` for recursive match.

    Results are sorted by modification time (newest first) to surface
    recently-edited files — useful when the model is trying to locate the
    file it just changed without re-reading the entire tree.

    `conv_id` lets relative `path` values follow wherever bash `cd`-ed to
    in this conversation — same policy as `read_file` / `write_file`.
    """
    if not pattern:
        return {"ok": False, "output": "", "error": "empty pattern"}
    base = _resolve(cwd, path, conv_id)
    if not base.exists():
        return {"ok": False, "output": "", "error": f"path not found: {base}"}
    try:
        # pathlib's glob supports `**` natively on 3.12+.
        matches = list(base.glob(pattern))
        # Newest first so the model sees what's freshest.
        matches.sort(key=lambda p: p.stat().st_mtime if p.exists() else 0, reverse=True)
        if not matches:
            return {"ok": True, "output": "(no matches)"}
        out = "\n".join(str(p) for p in matches[:500])
        if len(matches) > 500:
            out += f"\n... [truncated, {len(matches) - 500} more]"
        return {"ok": True, "output": out}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Clipboard tools
#
# Give the model a way to move text between itself and the user's desktop
# without hallucinating a `clip.exe` invocation. pyperclip handles Win/Mac/X11.
# ---------------------------------------------------------------------------
async def clipboard_read() -> dict:
    try:
        import pyperclip
        text = pyperclip.paste() or ""
        return {"ok": True, "output": _clip(text) or "(clipboard is empty)"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def clipboard_write(text: str) -> dict:
    if text is None:
        return {"ok": False, "output": "", "error": "text is required"}
    try:
        import pyperclip
        pyperclip.copy(str(text))
        return {"ok": True, "output": f"copied {len(str(text))} chars to clipboard"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Background shell processes
#
# Long-lived commands (dev servers, builds, test watchers) don't fit the
# one-shot `bash` tool because that awaits completion. The trio below lets
# the agent start a process, periodically peek its accumulated output, and
# kill it when no longer needed.
#
# State lives in the module-level _BG_SHELLS dict. Each entry holds the
# subprocess + two async reader tasks that keep draining stdout/stderr into
# in-memory buffers. Buffers are cleared on each bash_output read, so the
# model only ever sees *new* lines — not the whole backlog from minute one.
# ---------------------------------------------------------------------------
class _PopenProcessAdapter:
    """asyncio.subprocess.Process-compatible facade around a blocking Popen.

    Only bash_bg's selector-event-loop fallback uses this — when
    `asyncio.create_subprocess_exec` raises NotImplementedError because the
    running event loop is the Windows Selector variant, we spawn the child
    with `subprocess.Popen` instead. bash_output / kill_shell still want to
    treat `_BgShell.proc` uniformly, so this shim exposes just the surface
    those call sites touch: `.returncode`, `.terminate()`, `.kill()`, and
    awaitable `.wait()`.
    """

    def __init__(self, popen: subprocess.Popen) -> None:
        self._popen = popen

    @property
    def returncode(self) -> int | None:
        # Popen.poll() reaps if the child has exited; returncode is None otherwise.
        return self._popen.poll()

    def terminate(self) -> None:
        try:
            self._popen.terminate()
        except Exception:
            pass

    def kill(self) -> None:
        try:
            self._popen.kill()
        except Exception:
            pass

    async def wait(self) -> int:
        # Popen.wait is blocking; off-load to a thread so we don't stall the loop.
        return await asyncio.to_thread(self._popen.wait)


class _BgShell:
    def __init__(
        self,
        proc: "asyncio.subprocess.Process | _PopenProcessAdapter",
        command: str,
        cwd: str,
    ) -> None:
        self.proc = proc
        self.command = command
        self.cwd = cwd
        self.start_time = time.time()
        self.stdout_buf: list[str] = []
        self.stderr_buf: list[str] = []
        # Readers are asyncio.Task on the Proactor path and threading.Thread on
        # the Popen-fallback path. kill_shell uses isinstance to branch.
        self._stdout_task: "asyncio.Task | threading.Thread | None" = None
        self._stderr_task: "asyncio.Task | threading.Thread | None" = None


_BG_SHELLS: dict[str, _BgShell] = {}
_BG_OUTPUT_CHAR_LIMIT = 50_000  # per-stream soft cap to prevent unbounded growth


def _append_to_bg_buf(buf: list[str], line: str) -> None:
    """Append a single decoded line to a bg-shell buffer, enforcing the cap.

    Extracted so the async and thread drains share identical trimming logic.
    """
    buf.append(line)
    if sum(len(s) for s in buf) > _BG_OUTPUT_CHAR_LIMIT:
        while buf and sum(len(s) for s in buf) > _BG_OUTPUT_CHAR_LIMIT // 2:
            buf.pop(0)


async def _drain(stream: asyncio.StreamReader | None, buf: list[str]) -> None:
    """Continuously append decoded lines from `stream` to `buf` until EOF."""
    if stream is None:
        return
    while True:
        try:
            line = await stream.readline()
        except Exception:
            break
        if not line:
            break
        _append_to_bg_buf(buf, line.decode("utf-8", errors="replace"))


def _drain_blocking(stream, buf: list[str]) -> None:
    """Thread-backed equivalent of `_drain` for the Popen fallback.

    Runs on a daemon thread so it dies with the process. Exits cleanly when
    the pipe closes (readline returns b'') or any read raises.
    """
    if stream is None:
        return
    try:
        for raw in iter(stream.readline, b""):
            if not raw:
                break
            _append_to_bg_buf(buf, raw.decode("utf-8", errors="replace"))
    except Exception:
        return


async def bash_bg(cwd: str, command: str, conv_id: str | None = None) -> dict:
    """Launch `command` in the background and return a handle.

    Use this for anything that doesn't terminate quickly: dev servers, file
    watchers, long builds. Poll with `bash_output(shell_id)` to read new
    output; terminate with `kill_shell(shell_id)`.

    Starts from the conversation's persistent bash cwd (same as foreground
    `bash`) when `conv_id` is provided, so the model can `cd subdir` in a
    prior turn and have `bash_bg: npm run dev` land in the right directory.
    We do NOT wrap with a cwd marker here — the process is expected to
    outlive the turn, so there's no clean point to capture its final PWD.
    """
    if not command or not command.strip():
        return {
            "ok": False,
            "output": "",
            "error": (
                "bash requires the `command` field, but you didn't pass one. "
                "Example: `bash({\"command\": \"cd myapp && ls\", "
                "\"reason\": \"...\"})`. The `reason` field on its own is not "
                "enough — `command` carries the actual shell text to run."
            ),
        }
    effective_cwd = _effective_bash_cwd(cwd, conv_id)
    bash_exe = _bash_executable()
    env = _non_interactive_env()
    try:
        if bash_exe:
            proc = await asyncio.create_subprocess_exec(
                bash_exe, "-c", command,
                cwd=effective_cwd,
                stdin=asyncio.subprocess.DEVNULL,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )
        else:
            proc = await asyncio.create_subprocess_shell(
                command,
                cwd=effective_cwd,
                stdin=asyncio.subprocess.DEVNULL,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )
    except NotImplementedError:
        # Selector loop on Windows: asyncio subprocess support isn't there.
        # Mirror the fallback `bash` already does — spawn via subprocess.Popen
        # on a worker thread, drain the pipes from daemon reader threads, and
        # wrap the Popen in an adapter so kill_shell/bash_output don't need
        # to branch on backing type. The model sees the same shell_id API.
        def _spawn_popen() -> subprocess.Popen:
            if bash_exe:
                return subprocess.Popen(
                    [bash_exe, "-c", command],
                    cwd=effective_cwd,
                    stdin=subprocess.DEVNULL,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    env=env,
                )
            return subprocess.Popen(
                command,
                shell=True,
                cwd=effective_cwd,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=env,
            )
        try:
            popen = await asyncio.to_thread(_spawn_popen)
        except Exception as e:
            return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}
        shell_id = f"bg_{uuid.uuid4().hex[:8]}"
        shell = _BgShell(_PopenProcessAdapter(popen), command, effective_cwd)
        t_out = threading.Thread(
            target=_drain_blocking, args=(popen.stdout, shell.stdout_buf), daemon=True
        )
        t_err = threading.Thread(
            target=_drain_blocking, args=(popen.stderr, shell.stderr_buf), daemon=True
        )
        t_out.start()
        t_err.start()
        shell._stdout_task = t_out
        shell._stderr_task = t_err
        _BG_SHELLS[shell_id] = shell
        return {
            "ok": True,
            "output": (
                f"started background shell {shell_id}\n"
                f"command: {command[:200]}\n"
                f"poll with bash_output({shell_id!r}); stop with kill_shell({shell_id!r})."
            ),
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}
    shell_id = f"bg_{uuid.uuid4().hex[:8]}"
    shell = _BgShell(proc, command, effective_cwd)
    shell._stdout_task = asyncio.create_task(_drain(proc.stdout, shell.stdout_buf))
    shell._stderr_task = asyncio.create_task(_drain(proc.stderr, shell.stderr_buf))
    _BG_SHELLS[shell_id] = shell
    return {
        "ok": True,
        "output": (
            f"started background shell {shell_id}\n"
            f"command: {command[:200]}\n"
            f"poll with bash_output({shell_id!r}); stop with kill_shell({shell_id!r})."
        ),
    }


async def bash_output(shell_id: str) -> dict:
    """Return and clear the newly buffered stdout/stderr for a background shell."""
    shell = _BG_SHELLS.get(shell_id)
    if not shell:
        return {"ok": False, "output": "", "error": f"no such background shell: {shell_id}"}
    stdout = "".join(shell.stdout_buf)
    stderr = "".join(shell.stderr_buf)
    shell.stdout_buf.clear()
    shell.stderr_buf.clear()
    rc = shell.proc.returncode
    status = f"exited ({rc})" if rc is not None else "running"
    parts = [f"[status: {status}, uptime: {time.time() - shell.start_time:.1f}s]"]
    if stdout:
        parts.append("--- stdout ---")
        parts.append(stdout.rstrip())
    if stderr:
        parts.append("--- stderr ---")
        parts.append(stderr.rstrip())
    if not stdout and not stderr:
        parts.append("(no new output)")
    return {"ok": True, "output": _clip("\n".join(parts))}


async def kill_shell(shell_id: str) -> dict:
    """Terminate a background shell and remove it from the registry."""
    shell = _BG_SHELLS.get(shell_id)
    if not shell:
        return {"ok": False, "output": "", "error": f"no such background shell: {shell_id}"}
    try:
        if shell.proc.returncode is None:
            shell.proc.terminate()
            try:
                await asyncio.wait_for(shell.proc.wait(), timeout=3.0)
            except asyncio.TimeoutError:
                shell.proc.kill()
                await shell.proc.wait()
        # Stop the readers so they don't block on a dead pipe forever.
        # Async-task readers we cancel; daemon threads self-exit on EOF once
        # terminate()/kill() closed the Popen pipes, so there's nothing to do.
        for t in (shell._stdout_task, shell._stderr_task):
            if isinstance(t, asyncio.Task) and not t.done():
                t.cancel()
        _BG_SHELLS.pop(shell_id, None)
        return {"ok": True, "output": f"killed {shell_id}"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Drag + window management
#
# Pixel-click + type + scroll cover 95% of desktop interaction, but a few
# common gestures still need dedicated primitives:
#   - `computer_drag` — drag from (x1, y1) to (x2, y2). Needed for reordering
#     tabs, resizing columns, drawing selection boxes, dragging files onto
#     targets, etc. pyautogui.dragTo handles the event timing.
#   - `window_action` — minimize / maximize / restore / close a window by a
#     substring of its title, without the model having to guess the pixel
#     position of the title-bar buttons.
#   - `window_bounds` — move + resize a window to a specific rectangle, or
#     just read its current rectangle. Useful for staging side-by-side
#     comparisons or pushing a noisy window off-screen.
# ---------------------------------------------------------------------------
def _computer_drag_sync(x1: int, y1: int, x2: int, y2: int, duration: float, button: str) -> None:
    import pyautogui
    # moveTo first so the starting position is stable before the mouseDown,
    # then dragTo to the destination. Works across monitors because both
    # points are in absolute virtual-screen coords.
    pyautogui.moveTo(x1, y1, duration=0.05)
    pyautogui.dragTo(x2, y2, duration=duration, button=button)
    # Mark the drag END (the spot where the mouse was actually released)
    # so the next screenshot's red dot points at the drop target — usually
    # what the model wants to verify.
    _record_click_pos(x2, y2)


async def computer_drag(
    x1: int | float,
    y1: int | float,
    x2: int | float,
    y2: int | float,
    duration: float = 0.4,
    button: str = "left",
) -> dict:
    """Press, move, and release — a full drag from one point to another.

    Both endpoints are in the last screenshot's pixel space (exactly like
    `computer_click`). `duration` is in seconds and caps the how-long-to-
    drag — too short and fast-path apps may miss the intermediate move
    events; too long wastes real time. `button` is 'left' (default),
    'right', or 'middle'.
    """
    try:
        btn = (button or "left").lower()
        if btn not in {"left", "right", "middle"}:
            return {"ok": False, "output": "", "error": f"invalid button: {button!r}"}
        # Clamp both endpoints to the last screenshot's bounds so the drag
        # stays on screen. We use _clamp_to_image → _map_to_screen like the
        # click path, so the (x, y) math is identical.
        ix1, iy1, _n1 = _clamp_to_image(x1, y1)
        ix2, iy2, _n2 = _clamp_to_image(x2, y2)
        rx1, ry1 = _map_to_screen(ix1, iy1)
        rx2, ry2 = _map_to_screen(ix2, iy2)
        dur = max(0.05, min(float(duration), 5.0))
        await asyncio.to_thread(_computer_drag_sync, rx1, ry1, rx2, ry2, dur, btn)
        await asyncio.sleep(0.25)
        shot = await _capture_screenshot()
        return {
            "ok": True,
            "output": _attach_shot_feedback(
                f"{btn}-dragged from ({ix1}, {iy1}) to ({ix2}, {iy2}) over {dur:.2f}s.",
                shot,
            ),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _find_window_sync(name: str):
    """Locate a visible top-level window by case-insensitive title substring.

    Returns the uiautomation WindowControl or None. Reused by the window
    management and inspect_window tools. Factored out because all of them
    need the same "find by fuzzy name, prefer non-foreground, skip zero-
    area" logic as `_focus_window_sync`.
    """
    try:
        import uiautomation as auto
    except ImportError:
        return None
    needle = name.strip().lower()
    try:
        fg_hwnd = None
        try:
            fg = auto.GetForegroundControl()
            if fg is not None:
                fg_hwnd = int(fg.NativeWindowHandle)
        except Exception:
            pass
        matches = []
        for w in reversed(auto.GetRootControl().GetChildren()):
            try:
                title = (w.Name or "").lower()
                if not title or needle not in title:
                    continue
                r = w.BoundingRectangle
                if not r or (r.right - r.left) <= 0 or (r.bottom - r.top) <= 0:
                    continue
                matches.append((w, int(w.NativeWindowHandle)))
            except Exception:
                continue
        if not matches:
            return None
        non_fg = [m for m in matches if m[1] != fg_hwnd]
        return (non_fg[0] if non_fg else matches[0])[0]
    except Exception:
        return None


def _window_action_sync(name: str, action: str) -> dict:
    """Perform one of minimize / maximize / restore / close on a window.

    Uses Win32 APIs via ctypes so we don't depend on uiautomation pattern
    support (which is spotty for non-native apps). Falls back to sending
    a WM_CLOSE for 'close' which is the polite way to ask any window to
    shut down (apps get to run their 'save before quit' dialog).
    """
    if sys.platform != "win32":
        return {"ok": False, "error": "window_action is Windows-only"}
    w = _find_window_sync(name)
    if w is None:
        return {"ok": False, "error": f"no window title contains {name!r}"}
    import ctypes
    user32 = ctypes.windll.user32
    hwnd = int(w.NativeWindowHandle)
    # SW_* constants — see ShowWindow docs.
    SW_MINIMIZE, SW_MAXIMIZE, SW_RESTORE = 6, 3, 9
    WM_CLOSE = 0x0010
    try:
        if action == "minimize":
            user32.ShowWindow(hwnd, SW_MINIMIZE)
        elif action == "maximize":
            user32.ShowWindow(hwnd, SW_MAXIMIZE)
        elif action == "restore":
            user32.ShowWindow(hwnd, SW_RESTORE)
        elif action == "close":
            user32.PostMessageW(hwnd, WM_CLOSE, 0, 0)
        else:
            return {"ok": False, "error": f"unknown action: {action!r}"}
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}
    return {"ok": True, "matched_name": w.Name or ""}


async def window_action(name: str, action: str) -> dict:
    """Minimize / maximize / restore / close a window by title substring.

    Much more reliable than pixel-clicking the title-bar buttons, whose
    coordinates depend on DPI, window chrome style, and whether the app
    has a custom title bar. `name` is a case-insensitive substring of the
    window's title; `action` is one of: 'minimize', 'maximize', 'restore',
    'close'.
    """
    try:
        if not name or not name.strip():
            return {"ok": False, "output": "", "error": "name required"}
        act = (action or "").strip().lower()
        if act not in {"minimize", "maximize", "restore", "close"}:
            return {
                "ok": False,
                "output": "",
                "error": f"action must be minimize|maximize|restore|close, got {action!r}",
            }
        result = await asyncio.to_thread(_window_action_sync, name.strip(), act)
        if not result.get("ok"):
            return {"ok": False, "output": "", "error": result.get("error") or "window_action failed"}
        await asyncio.sleep(0.25)
        shot = await _capture_screenshot()
        return {
            "ok": True,
            "output": _attach_shot_feedback(
                f"{act}d window {result.get('matched_name')!r}.", shot
            ),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _window_bounds_sync(name: str, x, y, width, height) -> dict:
    """Read or set a window's bounding rectangle via SetWindowPos.

    When every bound is None, we just read the current rect. Otherwise
    we move+resize. Falls back to None on unsupported OSes.
    """
    if sys.platform != "win32":
        return {"ok": False, "error": "window_bounds is Windows-only"}
    w = _find_window_sync(name)
    if w is None:
        return {"ok": False, "error": f"no window title contains {name!r}"}
    import ctypes
    user32 = ctypes.windll.user32
    hwnd = int(w.NativeWindowHandle)
    # Read current rect via GetWindowRect.
    class RECT(ctypes.Structure):
        _fields_ = [("left", ctypes.c_long), ("top", ctypes.c_long),
                    ("right", ctypes.c_long), ("bottom", ctypes.c_long)]
    rect = RECT()
    user32.GetWindowRect(hwnd, ctypes.byref(rect))
    current = {
        "left": rect.left, "top": rect.top,
        "width": rect.right - rect.left, "height": rect.bottom - rect.top,
    }
    if all(v is None for v in (x, y, width, height)):
        return {"ok": True, "matched_name": w.Name or "", "bounds": current}
    # SetWindowPos: uFlags=0x0004 is SWP_NOZORDER, 0x0040 is SWP_SHOWWINDOW.
    nx = int(x) if x is not None else current["left"]
    ny = int(y) if y is not None else current["top"]
    nw = int(width) if width is not None else current["width"]
    nh = int(height) if height is not None else current["height"]
    user32.SetWindowPos(hwnd, 0, nx, ny, nw, nh, 0x0044)
    return {
        "ok": True,
        "matched_name": w.Name or "",
        "bounds": {"left": nx, "top": ny, "width": nw, "height": nh},
    }


async def window_bounds(
    name: str,
    x: int | None = None,
    y: int | None = None,
    width: int | None = None,
    height: int | None = None,
) -> dict:
    """Read or set a window's position + size by title substring.

    Call with only `name` to READ current bounds. Pass any of x/y/width/height
    to move / resize — omitted fields are kept at their current value.
    """
    try:
        if not name or not name.strip():
            return {"ok": False, "output": "", "error": "name required"}
        result = await asyncio.to_thread(_window_bounds_sync, name.strip(), x, y, width, height)
        if not result.get("ok"):
            return {"ok": False, "output": "", "error": result.get("error") or "window_bounds failed"}
        b = result["bounds"]
        verb = "read" if all(v is None for v in (x, y, width, height)) else "set"
        return {
            "ok": True,
            "output": (
                f"{verb} window {result['matched_name']!r} bounds: "
                f"pos=({b['left']}, {b['top']}) size=({b['width']}x{b['height']})"
            ),
            "bounds": b,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Accessibility-tree dump — `inspect_window` + `click_element_id`
#
# `click_element` finds controls by accessible name, but the model often
# needs to FIRST see what names are available. `inspect_window` dumps the
# UIA tree of the foreground (or named) window as compact structured text
# so the model can pick the right control name before calling click_element.
#
# To eliminate the cost of re-walking the tree on every click, every
# inspect_window dump also assigns a stable identifier (`el42`, `el43`, …)
# to each visible control and caches the click-target coordinates module-
# locally. The model can then call `click_element_id("el42")` to click
# without paying for another tree walk OR risking a fuzzy-name mismatch
# when several controls share the same label.
#
# The cache is bounded (FIFO eviction) and protected by a lock because
# multiple sub-agents may inspect concurrently when running under
# `delegate_parallel`.
# ---------------------------------------------------------------------------
# Bounded LRU cache of element click targets. Entry value is a dict with
# (cx, cy, label, bbox, enabled, minted_at) — screen-space click point,
# short human-readable label for success messages, the screen-space
# bounding rect, the control's IsEnabled state at inspect time, and a
# monotonic timestamp for diagnostics.
#
# Eviction: LRU, not FIFO. `_element_cache_get` promotes the accessed
# entry to the most-recently-used end. Rationale: a chatty session that
# mints hundreds of new ids for each screenshot can otherwise evict its
# own just-minted ids before it clicks them — FIFO punishes "active"
# ids exactly as hard as stale ones. Promoting on read keeps the
# working set resident.
_ELEMENT_CACHE: "OrderedDict[str, dict]" = OrderedDict()
_ELEMENT_CACHE_MAX = 5000
_ELEMENT_NEXT_ID = 0
# Highest ID we've ever minted. Compared against the numeric suffix of
# a requested id so we can distinguish "evicted" from "never existed".
_ELEMENT_HIGH_WATER = 0
_ELEMENT_LOCK = threading.Lock()


# --------------------------------------------------------------------------
# Set-of-Mark overlay constants for inspect_window's annotated screenshot.
#
# Recent vision-LLM research (Set-of-Mark prompting) shows that overlaying
# numbered badges directly on the UI controls dramatically improves a small
# VLM's ability to "click that thing". Instead of asking the model to read
# coordinates off a grid, we paint each cached `[elN]` ID right on top of
# the matching control, so the next prompt can literally say "click el7".
#
# Caps:
# - _INSPECT_OVERLAY_MAX_BADGES — beyond ~80 badges the image becomes a
#   wall of yellow boxes that hurts more than it helps; the text dump still
#   lists every cached id so the model can fall back to id-by-name.
# - _BADGE_PADDING — minimum gap (px) between badges before we shift one.
# --------------------------------------------------------------------------
_INSPECT_OVERLAY_MAX_BADGES = 80
_BADGE_PADDING = 4


def _draw_inspect_overlay(
    window_rect: tuple[int, int, int, int],
    items: list[tuple[str, int, int, int, int, str]],
) -> dict | None:
    """Render an annotated screenshot of the inspected window with [elN] badges.

    `window_rect` is the screen-space (left, top, right, bottom) of the
    window we just inspected. `items` is a list of
    ``(eid, ele_left, ele_top, ele_right, ele_bottom, label)`` — one tuple
    per control that was cached in `_ELEMENT_CACHE`.

    Steps:
      1. Clip the window rect to the virtual screen so a partly-off-screen
         window doesn't crash mss / capture garbage.
      2. mss-grab the clipped region (no DPI re-scale — UIA is already DPI-
         aware on this code path because we set DPI awareness at startup).
      3. For each item, anchor a yellow badge at the control's top-left
         (so the badge does not cover the visible label of the control).
         A simple collision walker shifts colliding badges down then right
         so dense panels stay legible.
      4. Save the PNG via the same uuid-named path scheme as `screenshot`.

    Returns ``{"name", "path", "width", "height", "drawn"}`` on success, or
    ``None`` when the window is fully off-screen / has zero area / mss
    fails. The caller is expected to surface ``image_path`` only when
    something was actually rendered.
    """
    import mss
    from PIL import Image, ImageDraw, ImageFont

    wleft, wtop, wright, wbottom = window_rect
    if (wright - wleft) <= 0 or (wbottom - wtop) <= 0:
        return None

    # Clip to the union of all attached monitors. _virtual_screen_bounds
    # returns (width, height, left, top) — convert to right/bottom for
    # straightforward intersection.
    vw, vh, vx, vy = _virtual_screen_bounds()
    cleft = max(wleft, vx)
    ctop = max(wtop, vy)
    cright = min(wright, vx + vw)
    cbottom = min(wbottom, vy + vh)
    cw = cright - cleft
    ch = cbottom - ctop
    if cw <= 0 or ch <= 0:
        return None

    try:
        with mss.mss() as sct:
            raw = sct.grab({"left": cleft, "top": ctop, "width": cw, "height": ch})
            img = Image.frombytes("RGB", raw.size, raw.rgb)
    except Exception:
        return None

    # Cap really enormous windows so the PNG stays small + the model isn't
    # billed for thousands of vision tokens. Mirrors the cap used by the
    # primary screenshot path.
    orig_w, orig_h = img.size
    long_edge = max(orig_w, orig_h)
    scale = 1.0
    if long_edge > MAX_SCREENSHOT_EDGE:
        scale = MAX_SCREENSHOT_EDGE / float(long_edge)
        new_size = (int(orig_w * scale), int(orig_h * scale))
        img = img.resize(new_size, Image.LANCZOS)
    img_w, img_h = img.size

    img = img.convert("RGBA")
    draw = ImageDraw.Draw(img, "RGBA")

    # Pick a readable TTF; bitmap fallback when no system font is available.
    font = None
    for cand in ("arial.ttf", "DejaVuSans.ttf", "Helvetica.ttc"):
        try:
            font = ImageFont.truetype(cand, 13)
            break
        except Exception:
            continue
    if font is None:
        font = ImageFont.load_default()

    badge_fill = (255, 215, 0, 235)   # gold — readable on light AND dark UIs
    badge_edge = (60, 40, 0, 255)
    badge_text = (20, 20, 20, 255)

    drawn_rects: list[tuple[int, int, int, int]] = []
    drawn = 0
    for eid, eleft, etop, _eright, _ebottom, _label in items:
        if drawn >= _INSPECT_OVERLAY_MAX_BADGES:
            break
        # Translate from screen coords to image-local coords. Apply the
        # same downscale factor we used for the capture, so a badge
        # anchored at the control's top-left lands in the correct spot.
        bx_full = max(0, eleft - cleft)
        by_full = max(0, etop - ctop)
        bx = int(bx_full * scale)
        by = int(by_full * scale)
        # Off-image (e.g. element scrolled out of view) → skip.
        if bx >= img_w or by >= img_h:
            continue

        text = eid  # already prefixed 'el' (e.g. 'el42')
        tb = draw.textbbox((0, 0), text, font=font)
        tw, th = tb[2] - tb[0], tb[3] - tb[1]
        bw = tw + 8
        bh = th + 4

        # Resolve collisions: shift down then right when the candidate
        # rect overlaps an already-drawn badge. Bounded retries so we
        # don't loop forever on a fully packed UI.
        for _attempt in range(8):
            rect = (bx, by, bx + bw, by + bh)
            collide = False
            for r in drawn_rects:
                if not (rect[2] + _BADGE_PADDING < r[0]
                        or rect[0] > r[2] + _BADGE_PADDING
                        or rect[3] + _BADGE_PADDING < r[1]
                        or rect[1] > r[3] + _BADGE_PADDING):
                    collide = True
                    break
            if not collide:
                break
            by += bh + _BADGE_PADDING
            if by + bh > img_h:
                by = max(0, int((etop - ctop) * scale))
                bx += bw + _BADGE_PADDING
                if bx + bw > img_w:
                    break  # give up — overlay one anyway at original anchor

        rect = (bx, by, bx + bw, by + bh)
        draw.rectangle(rect, fill=badge_fill, outline=badge_edge, width=1)
        draw.text((bx + 4, by + 1), text, fill=badge_text, font=font)
        drawn_rects.append(rect)
        drawn += 1

    img = img.convert("RGB")
    name = f"{uuid.uuid4().hex}.png"
    path = _ensure_screenshot_dir() / name
    img.save(path, format="PNG", optimize=True)
    return {
        "name": name,
        "path": str(path),
        "width": img.size[0],
        "height": img.size[1],
        "drawn": drawn,
    }


def _element_cache_put(
    cx: int,
    cy: int,
    label: str,
    bbox: tuple[int, int, int, int] | None = None,
    enabled: bool | None = None,
) -> str:
    """Allocate a fresh element ID, store the click target, return the ID.

    Lock-protected because multiple inspect_window calls can run in
    parallel from sibling sub-agents under `delegate_parallel`.

    `bbox` is (left, top, right, bottom) in screen coordinates. Preserved
    so callers that want to describe or wait-for the control have its
    footprint, not just the centre point. `enabled` mirrors the UIA
    `IsEnabled` flag at inspect time — consumed by `ui_wait` kind
    `element_enabled` so the model can wait for a greyed-out button
    to become clickable without re-walking the tree every poll.
    """
    global _ELEMENT_NEXT_ID, _ELEMENT_HIGH_WATER
    with _ELEMENT_LOCK:
        _ELEMENT_NEXT_ID += 1
        eid = f"el{_ELEMENT_NEXT_ID}"
        _ELEMENT_HIGH_WATER = max(_ELEMENT_HIGH_WATER, _ELEMENT_NEXT_ID)
        _ELEMENT_CACHE[eid] = {
            "cx": int(cx),
            "cy": int(cy),
            "label": str(label),
            "bbox": tuple(int(x) for x in bbox) if bbox else None,
            "enabled": enabled,
            "minted_at": time.monotonic(),
        }
        # LRU eviction: oldest-touched entry drops when we're at capacity.
        # Stale entries are unavoidable (the UI might have moved since
        # the inspect call), but the model will see the click had no
        # effect via the post-click screenshot and can re-inspect.
        while len(_ELEMENT_CACHE) > _ELEMENT_CACHE_MAX:
            _ELEMENT_CACHE.popitem(last=False)
        return eid


def _element_cache_get(eid: str) -> dict | None:
    """Look up a cached entry by ID; None if unknown / evicted.

    On hit, promotes the entry to the most-recently-used end so active
    ids aren't evicted by a burst of new inspections. Returned value is
    a *copy* — callers mutate their own dict, not the cache entry.
    """
    with _ELEMENT_LOCK:
        entry = _ELEMENT_CACHE.get(eid)
        if entry is None:
            return None
        _ELEMENT_CACHE.move_to_end(eid, last=True)
        # Return a shallow copy so the caller can't accidentally mutate
        # the cached bbox tuple or label in place.
        return dict(entry)


def _element_id_status(eid: str) -> str:
    """Explain why a lookup missed — 'evicted', 'not minted', or 'bad format'.

    Surfaced to the model in click_element_id and ui_wait errors so it can
    choose the right recovery (re-inspect vs. check the id it typed).
    """
    s = str(eid or "").strip()
    if not s.startswith("el") or not s[2:].isdigit():
        return "bad format"
    try:
        n = int(s[2:])
    except ValueError:
        return "bad format"
    with _ELEMENT_LOCK:
        high = _ELEMENT_HIGH_WATER
    if n > high:
        return "not minted"
    return "evicted"


# Named types we keep even when their accessible name is empty. Anonymous
# structural Panes / Groups are discarded (noise). Extracted to module
# scope so the element-scan helper and inspect_window share one list.
_KEPT_CONTROL_TYPES = {
    "ButtonControl", "EditControl", "ComboBoxControl", "HyperlinkControl",
    "MenuItemControl", "CheckBoxControl", "RadioButtonControl",
    "ListItemControl", "TabItemControl", "TreeItemControl",
    "TextControl", "DocumentControl",
}


def _scan_window_elements_sync(
    root,
    max_depth: int,
    max_nodes: int,
) -> dict:
    """Walk the accessibility tree under `root`, mint element ids, and
    return everything callers (inspect_window, screenshot(with_elements))
    need to describe what's clickable.

    Returns a dict with:
      - `lines`       — indented text dump, one line per kept node
      - `overlay_items` — list of (eid, left, top, right, bottom, label)
                         for the Set-of-Mark renderer
      - `elements`    — structured list `[{id, role, name, bbox, enabled}]`
                         in SCREEN coordinates, ready for direct
                         consumption by `screenshot(with_elements=True)`
      - `node_count`  — nodes visited (<= max_nodes)

    Pure function: all UIA access happens here; no disk, no network.
    Raises nothing — per-node errors are swallowed, whole-call errors
    propagate to the caller.
    """
    lines: list[str] = []
    overlay_items: list[tuple[str, int, int, int, int, str]] = []
    elements: list[dict] = []
    queue = [(root, 0)]
    visited = 0
    while queue and visited < max_nodes:
        ctrl, depth = queue.pop(0)
        visited += 1
        try:
            cname = ctrl.Name or ""
            ctype = ctrl.ControlTypeName or "Unknown"
            rect = ctrl.BoundingRectangle
            if cname or ctype in _KEPT_CONTROL_TYPES:
                # Mint a stable ID for any node with a real bounding rect,
                # so the model can subsequently call click_element_id(eid)
                # without re-walking the tree. Off-screen / zero-area
                # controls don't get an ID — clicking them would be a
                # no-op and they'd just clutter the output.
                bbox = ""
                prefix = "       "  # 7 spaces lines up with the widest "[elNNN] "
                if rect and (rect.right - rect.left) > 0 and (rect.bottom - rect.top) > 0:
                    bbox = f" @({rect.left},{rect.top},{rect.right - rect.left}x{rect.bottom - rect.top})"
                    cx = (rect.left + rect.right) // 2
                    cy = (rect.top + rect.bottom) // 2
                    # UIA's IsEnabled reports whether the control is
                    # interactive right now (greyed-out Save buttons
                    # read False). We capture it so ui_wait can gate on
                    # it without re-walking the tree per poll.
                    try:
                        is_enabled = bool(ctrl.IsEnabled)
                    except Exception:
                        # Some controls raise when probed (protected
                        # frames, disposed handles) — treat as unknown.
                        is_enabled = None
                    screen_bbox = (
                        int(rect.left), int(rect.top),
                        int(rect.right), int(rect.bottom),
                    )
                    eid = _element_cache_put(
                        cx, cy, cname or ctype,
                        bbox=screen_bbox, enabled=is_enabled,
                    )
                    overlay_items.append((
                        eid,
                        screen_bbox[0], screen_bbox[1],
                        screen_bbox[2], screen_bbox[3],
                        cname or ctype,
                    ))
                    elements.append({
                        "id": eid,
                        "role": ctype,
                        "name": cname,
                        "bbox": list(screen_bbox),
                        "enabled": is_enabled,
                    })
                    prefix = f"[{eid}] "
                lines.append(f"{'  ' * depth}{prefix}{ctype}{bbox}  Name={cname!r}")
        except Exception:
            continue
        if depth + 1 > max_depth:
            continue
        try:
            for child in ctrl.GetChildren():
                queue.append((child, depth + 1))
        except Exception:
            continue
    return {
        "lines": lines,
        "overlay_items": overlay_items,
        "elements": elements,
        "node_count": visited,
    }


def _resolve_inspect_root_sync(name: str | None):
    """Shared helper: given an optional window-title substring, return the
    UIA root to inspect. Returns ``(root_or_None, error_string_or_None)``.

    Factored so inspect_window and screenshot(with_elements) agree on how
    to resolve "inspect *what*" — either the named window or the current
    foreground — without duplicating the ImportError / platform-guard
    handshake.
    """
    if sys.platform != "win32":
        return None, "Windows-only"
    try:
        import uiautomation as auto
    except ImportError:
        return None, "uiautomation package not installed"
    if name and name.strip():
        root = _find_window_sync(name.strip())
        if root is None:
            return None, f"no window title contains {name!r}"
        return root, None
    try:
        root = auto.GetForegroundControl()
    except Exception as e:
        return None, f"could not get foreground window: {e}"
    if root is None:
        return None, "no root window to inspect"
    return root, None


def _inspect_window_sync(
    name: str | None,
    max_depth: int,
    max_nodes: int,
    overlay: bool = True,
) -> dict:
    root, err = _resolve_inspect_root_sync(name)
    if root is None:
        # `Windows-only` becomes the full error here — match the
        # previous "inspect_window is Windows-only" wording for
        # callers that string-match on it.
        if err == "Windows-only":
            return {"ok": False, "error": "inspect_window is Windows-only"}
        return {"ok": False, "error": err or "inspect_window failed"}
    # Delegate the walk to the shared helper (also used by
    # screenshot(with_elements=True)) so the tree-filter rules and
    # cache-write behaviour stay in lockstep.
    scan = _scan_window_elements_sync(root, max_depth, max_nodes)

    # --- Set-of-Mark overlay -------------------------------------------------
    # Render an annotated screenshot of the inspected window. Best-effort:
    # any failure here (OS quirk, off-screen window, mss hiccup) just means
    # the model gets the text dump without an image — the inspection itself
    # still succeeded, so we don't surface the renderer error.
    image_payload: dict | None = None
    if overlay and scan["overlay_items"]:
        try:
            wrect = root.BoundingRectangle
            if wrect:
                image_payload = _draw_inspect_overlay(
                    (int(wrect.left), int(wrect.top), int(wrect.right), int(wrect.bottom)),
                    scan["overlay_items"],
                )
        except Exception:
            image_payload = None

    lines = scan["lines"]
    return {
        "ok": True,
        "output": "\n".join(lines) if lines else "(no nameable controls found)",
        "node_count": scan["node_count"],
        "window_title": (root.Name or ""),
        "id_count": len(scan["elements"]),
        "image_payload": image_payload,
    }


def _enumerate_elements_for_screenshot(
    name: str | None = None,
    max_depth: int = 12,
    max_nodes: int = 500,
) -> dict:
    """Return a structured element map for whatever window the
    screenshot is aimed at.

    This is the entry point used by `screenshot(with_elements=True)` and
    `screenshot_window(with_elements=True)`. It mints fresh ids and
    stashes them in the element cache, same as `inspect_window` would,
    so the model can immediately `click_element_id(...)` on any returned
    element. Graceful on non-Windows: returns ``{"ok": False, "error":
    "Windows-only", "elements": []}`` without raising.
    """
    root, err = _resolve_inspect_root_sync(name)
    if root is None:
        return {
            "ok": False,
            "error": err or "no window",
            "elements": [],
            "window_title": "",
        }
    try:
        scan = _scan_window_elements_sync(root, max_depth, max_nodes)
    except Exception as e:
        return {
            "ok": False,
            "error": f"{type(e).__name__}: {e}",
            "elements": [],
            "window_title": (getattr(root, "Name", None) or ""),
        }
    return {
        "ok": True,
        "elements": scan["elements"],
        "node_count": scan["node_count"],
        "window_title": (getattr(root, "Name", None) or ""),
    }


async def inspect_window(
    name: str | None = None,
    max_depth: int = 12,
    max_nodes: int = 500,
    overlay: bool = True,
) -> dict:
    """Dump the accessibility tree of a window as indented text.

    Without args, dumps the foreground window. Pass `name` to target a
    specific window by title substring. Output shows control type, bounding
    box, and accessible name. Each visible control is also tagged with a
    stable identifier like `[el42]` — pass that ID to `click_element_id`
    to click without re-walking the tree (faster, and it removes the
    fuzzy-name disambiguation step entirely).

    When `overlay=True` (default) the call also returns an annotated
    screenshot of the window with each `[elN]` badge painted on top of
    the matching control (Set-of-Mark prompting). The model can then
    literally see "el7 is the OK button" instead of correlating bbox
    coordinates against a separate screenshot. Pass `overlay=False` to
    skip the rendering cost when you only need the text dump.
    """
    try:
        md = max(1, min(int(max_depth or 12), 30))
        mn = max(10, min(int(max_nodes or 500), 5000))
        result = await asyncio.to_thread(
            _inspect_window_sync, name, md, mn, bool(overlay)
        )
        if not result.get("ok"):
            return {"ok": False, "output": "", "error": result.get("error") or "inspect_window failed"}
        ids = result.get("id_count") or 0
        image_payload = result.get("image_payload")
        # Surface a one-line note about the overlay so the model can tell
        # what it's looking at (and, when the overlay was suppressed for
        # reasons like off-screen windows, it knows the text is the only
        # source of truth).
        if image_payload:
            overlay_note = (
                f" Annotated screenshot attached "
                f"({image_payload.get('drawn', 0)} of {ids} ids badged)."
            )
        elif overlay:
            overlay_note = " (no annotated screenshot — window is off-screen or render failed)"
        else:
            overlay_note = ""
        out = {
            "ok": True,
            "output": (
                f"Window {result['window_title']!r} — "
                f"{result['node_count']} nodes scanned (max {mn}), "
                f"{ids} clickable id{'s' if ids != 1 else ''} cached. "
                f"Click any one with `click_element_id({{\"id\": \"elN\"}})`."
                f"{overlay_note}\n"
                + result["output"]
            ),
        }
        if image_payload:
            out["image_path"] = image_payload["name"]
        return out
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Click by cached element ID — the fast path after `inspect_window`.
#
# `click_element` walks the UIA tree on every call (slow on heavy apps and
# vulnerable to label collisions when N controls share a name). When the
# model has just inspected a window, it already has unambiguous IDs in
# hand — `click_element_id` skips the tree walk and clicks the cached
# screen-space center directly.
# ---------------------------------------------------------------------------
def _click_element_id_sync(eid: str, click_type: str) -> dict:
    cached = _element_cache_get(eid)
    if cached is None:
        # Distinguish "never existed" (bad id the model invented or
        # typo'd) from "evicted" (LRU pushed it out after a flurry of
        # inspections). The model's recovery differs: typo → re-read
        # the dump; evicted → re-inspect.
        reason = _element_id_status(eid)
        return {
            "ok": False,
            "error": (
                f"unknown element id {eid!r} ({reason}). IDs are minted "
                f"by `inspect_window` or `screenshot(with_elements=True)` "
                f"and live LRU in a {_ELEMENT_CACHE_MAX}-entry cache for "
                f"the duration of this backend process. Call inspect "
                f"again to mint a fresh id."
            ),
        }
    cx, cy, label = cached["cx"], cached["cy"], cached["label"]
    import pyautogui
    pyautogui.moveTo(cx, cy, duration=0.1)
    button = {"right": "right", "middle": "middle"}.get(click_type, "left")
    if click_type == "double":
        pyautogui.doubleClick(cx, cy)
    else:
        pyautogui.click(x=cx, y=cy, button=button)
    # Mark the click so the post-click screenshot shows where it landed.
    _record_click_pos(cx, cy)
    return {"ok": True, "x": cx, "y": cy, "label": label}


async def click_element_id(id: str, click_type: str = "left") -> dict:
    """Click a previously-discovered UI element by its cached ID.

    Run `inspect_window` first — every clickable control in the dump is
    tagged with `[elN]`. Pass that exact ID here to click. This is the
    preferred click path after an inspection because:
      - no tree walk (instant);
      - no fuzzy-name matching (the ID points at one specific control,
        even if N other controls share its label).

    `click_type` is one of: 'left' (default), 'right', 'middle', 'double'.
    """
    try:
        if not id or not str(id).strip():
            return {"ok": False, "output": "", "error": "id required"}
        clean_id = str(id).strip()
        if click_type not in ("left", "right", "middle", "double"):
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"click_type must be one of left/right/middle/double, "
                    f"got {click_type!r}"
                ),
            }
        result = await asyncio.to_thread(_click_element_id_sync, clean_id, click_type)
        if not result.get("ok"):
            # Fast-path error: no click happened, so there's nothing new
            # on screen worth screenshotting. Skipping the capture here
            # also means the "unknown id" error survives on headless /
            # no-display machines where a screenshot call would itself
            # crash with `KeyError: 'DISPLAY'` and drown out the real
            # reason the click failed.
            return {
                "ok": False,
                "output": "",
                "error": result.get("error", "click_element_id failed"),
            }
        # Settle so the post-click screenshot reflects the new state.
        await asyncio.sleep(0.15)
        shot = await _capture_screenshot()
        return {
            "ok": True,
            "output": _attach_shot_feedback(
                (
                    f"{click_type}-clicked {result.get('label', '?')!r} "
                    f"(id {clean_id}) at ({result['x']}, {result['y']})."
                ),
                shot,
            ),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Combo: focus + click + type via accessibility name — `type_into_element`.
#
# Filling out a form takes three round-trips today: `click_element` to focus
# the field, `computer_type` to enter the text, and `screenshot` to verify.
# That's two model turns and two screenshots before we even know whether the
# text landed in the right field. `type_into_element` collapses the first
# two into a single call by reusing `click_element`'s a11y-tree search to
# focus, then immediately typing — no pixel guessing, no second turn.
#
# Why not call `click_element` then `computer_type` from inside the agent
# loop? Because the model would have to OBSERVE the click via a screenshot
# before deciding to type, which is exactly the round-trip we're saving.
#
# The same UIA-name search that powers `click_element` is reused here, so
# behaviour stays consistent: case-insensitive substring match by default,
# `match='exact'` for ambiguous cases, foreground-window-first lookup that
# falls back to a shallow desktop walk for popups.
# ---------------------------------------------------------------------------
def _type_into_element_sync(
    name: str,
    text: str,
    match: str,
    clear: bool,
    interval: float,
) -> dict:
    """Synchronous worker — find the control by name, click to focus it,
    then type ``text``. Optionally clears the field first by selecting all
    + delete (Ctrl+A, Delete) so re-entering text doesn't append to what
    was already there.

    Returns a dict with ok / error and (when ok) the resolved click coords
    + the matched element's accessible name, so the async wrapper can
    surface them in the tool result.
    """
    # Reuse the click-element worker so all the search behaviour, candidate
    # ranking, and platform guards stay in one place. A failed click bails
    # out before any keystrokes go anywhere — important so we don't blast
    # text into the wrong window when focus didn't change as expected.
    click_result = _click_element_sync(name, match, "left", timeout=2.0)
    if not click_result.get("ok"):
        return click_result

    import pyautogui

    # The accessibility-tree click moved focus to the control, but Windows
    # sometimes needs a brief beat before keystrokes route to the newly-
    # focused widget — same reason `focus_window` sleeps 0.4s afterwards.
    # 80ms is enough for the input queue to drain on the test rigs we
    # tried; faster than focus_window because we only need one widget's
    # focus to settle, not a whole window's foreground change.
    import time as _time
    _time.sleep(0.08)

    if clear:
        # Select-all + Delete is the most portable "clear this field"
        # gesture: works for line edits, multi-line text areas, browser
        # form fields, IDE editors. Some apps treat Backspace differently
        # (e.g. character-by-character), so we deliberately use Delete on
        # a selected range instead.
        try:
            pyautogui.hotkey("ctrl", "a")
            _time.sleep(0.04)
            pyautogui.press("delete")
            _time.sleep(0.04)
        except Exception:
            # Best-effort — if the clear keystroke combo throws, fall
            # through to typing anyway. Worst case the new text appends
            # to the old, which the model can fix on its next turn.
            pass

    # Same ASCII-vs-unicode dispatch as `_computer_type_sync`. Keeping the
    # branching duplicated (rather than calling the other helper) so a
    # future change to typing strategy here doesn't surprise the simpler
    # standalone tool.
    try:
        text.encode("ascii")
        pyautogui.typewrite(text, interval=interval)
    except UnicodeEncodeError:
        import pyperclip
        pyperclip.copy(text)
        pyautogui.hotkey("ctrl", "v")

    return {
        "ok": True,
        "x": click_result.get("x"),
        "y": click_result.get("y"),
        "matched_name": click_result.get("matched_name"),
        "candidate_count": click_result.get("candidate_count", 1),
        "cleared": bool(clear),
    }


async def type_into_element(
    name: str,
    text: str,
    match: str = "contains",
    clear: bool = False,
    interval: float = 0.02,
) -> dict:
    """Find a UI control by accessible name, focus it, then type into it.

    One-shot replacement for the click-then-type pattern: where you'd
    normally call `click_element({"name": "Search"})` followed by
    `computer_type({"text": "..."})`, this does both in a single tool
    call. Halves the round-trip count for forms and search boxes.

    Args:
        name: The text-field's accessible name (substring or exact, see
            `match`). Same as `click_element` — usually the placeholder
            text, label, or aria-label of the field.
        text: The text to type into the focused control.
        match: "contains" (default, case-insensitive substring) or "exact".
        clear: When True, sends Ctrl+A + Delete BEFORE typing so the new
            text replaces whatever was there. Default False (appends).
        interval: Per-character typing delay in seconds. Defaults to a
            very fast 0.02s — bump it if a slow IME / web app drops keys.
    """
    try:
        if not name or not str(name).strip():
            return {"ok": False, "output": "", "error": "name required"}
        if text is None:
            return {"ok": False, "output": "", "error": "text is required"}
        clean_name = str(name).strip()
        if len(clean_name) > 200:
            return {
                "ok": False,
                "output": "",
                "error": f"name too long ({len(clean_name)} chars, max 200)",
            }
        text_str = str(text)
        # Same 10k cap as computer_type — anything bigger is almost
        # certainly a runaway / mistake, and Ollama context limits make
        # huge blobs a poor fit anyway.
        if len(text_str) > 10000:
            return {
                "ok": False,
                "output": "",
                "error": f"text too long ({len(text_str)} chars, max 10000)",
            }
        if match not in ("contains", "exact"):
            return {
                "ok": False,
                "output": "",
                "error": f"match must be 'contains' or 'exact', got {match!r}",
            }
        # Clamp interval — same range as computer_type so timing is
        # consistent between the two paths.
        try:
            iv = max(0.0, min(float(interval or 0.02), 0.2))
        except (TypeError, ValueError):
            iv = 0.02

        result = await asyncio.to_thread(
            _type_into_element_sync, clean_name, text_str, match, bool(clear), iv
        )
        # Settle so the post-action screenshot reflects the typed text.
        await asyncio.sleep(0.2)
        shot = await _capture_screenshot()
        if not result.get("ok"):
            return {
                "ok": False,
                "output": "",
                "error": result.get("error", "type_into_element failed"),
                "image_path": shot["name"],
            }
        matched = result.get("matched_name", clean_name)
        preview = text_str if len(text_str) <= 60 else (text_str[:60] + "…")
        clear_note = " (cleared first)" if result.get("cleared") else ""
        ambiguity_note = ""
        if result.get("candidate_count", 1) > 1:
            ambiguity_note = (
                f" — {result['candidate_count']} controls matched, used the "
                f"first; pass match='exact' or a more specific name if the "
                f"wrong one was chosen"
            )
        base = (
            f"typed {len(text_str)} char{'s' if len(text_str) != 1 else ''} "
            f"into {matched!r}{clear_note}: {preview!r}{ambiguity_note}."
        )
        return {
            "ok": True,
            "output": _attach_shot_feedback(base, shot),
            "image_path": shot["name"],
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Window-scoped screenshot + window enumeration.
#
# `take_screenshot` always grabs an entire monitor. For window-scoped tasks
# (one settings dialog, one app) that wastes vision tokens — most of the
# image is unrelated wallpaper / sibling apps the model doesn't need to
# reason about. `screenshot_window` crops to one window's bounding rect,
# typically 4-10x fewer tokens, and updates `_LAST_SHOT_*` so subsequent
# clicks land in screen coords just like a normal screenshot.
#
# `list_windows` is the "what's open right now?" introspection tool — much
# more reliable than asking the model to read window titles off a screenshot
# of the taskbar. The model picks a title here, then passes it to
# `focus_window` / `screenshot_window` / `inspect_window` / `window_action`.
# ---------------------------------------------------------------------------
def _screenshot_window_sync(name: str) -> dict:
    """Capture a screenshot cropped to the named window's bounding rect.

    Pipeline mirrors `_capture_screenshot_sync` so the model sees a
    consistent image (downscaled to MAX_SCREENSHOT_EDGE if huge, with the
    yellow coordinate grid overlaid). The big difference is the source
    region: we crop to one window's rect via mss instead of grabbing a
    whole monitor.

    Updates `_LAST_SHOT_*` globals so a subsequent computer_click(x, y)
    treats the model's image-space coords as window-local and translates
    them back to screen coords correctly.
    """
    if sys.platform != "win32":
        return {"ok": False, "error": "screenshot_window is Windows-only"}
    w = _find_window_sync(name)
    if w is None:
        return {"ok": False, "error": f"no window title contains {name!r}"}
    rect = w.BoundingRectangle
    if not rect:
        return {"ok": False, "error": "could not read window bounds"}
    wleft, wtop = int(rect.left), int(rect.top)
    ww = int(rect.right - rect.left)
    wh = int(rect.bottom - rect.top)
    if ww <= 0 or wh <= 0:
        return {
            "ok": False,
            "error": (
                "window has zero area (it's probably minimized — "
                "call window_action({\"action\": \"restore\"}) first)"
            ),
        }

    # Clip to virtual screen bounds so a window dragged half off-screen
    # doesn't crash mss. _virtual_screen_bounds returns (w, h, left, top).
    vw, vh, vx, vy = _virtual_screen_bounds()
    cleft = max(wleft, vx)
    ctop = max(wtop, vy)
    cright = min(wleft + ww, vx + vw)
    cbottom = min(wtop + wh, vy + vh)
    cw = cright - cleft
    ch = cbottom - ctop
    if cw <= 0 or ch <= 0:
        return {"ok": False, "error": "window is fully off-screen"}

    import mss
    from PIL import Image
    try:
        with mss.mss() as sct:
            raw = sct.grab({"left": cleft, "top": ctop, "width": cw, "height": ch})
            img = Image.frombytes("RGB", raw.size, raw.rgb)
    except Exception as e:
        return {"ok": False, "error": f"mss capture failed: {e}"}

    orig_w, orig_h = img.size
    long_edge = max(orig_w, orig_h)
    if long_edge > MAX_SCREENSHOT_EDGE:
        scale = MAX_SCREENSHOT_EDGE / float(long_edge)
        new_size = (int(orig_w * scale), int(orig_h * scale))
        img = img.resize(new_size, Image.LANCZOS)

    # Compute pixel signature + window-title set BEFORE annotation so the
    # diff measures real UI change, not our own overlays. Same contract as
    # `_capture_screenshot_sync` — see those comments for the rationale.
    sig = _compute_pixel_sig(img)
    titles = _enumerate_window_titles_fast()

    pre_w, pre_h = img.size
    pre_scale = (
        orig_w / pre_w if pre_w else 1.0,
        orig_h / pre_h if pre_h else 1.0,
    )

    # Adaptive grid overlay — small dialogs get a denser grid, full-window
    # captures get the same 100-px grid as a monitor screenshot. Picked by
    # `_pick_grid_step` from the image dimensions, so a 320x240 settings
    # popup shows readable 25-px lines instead of 2 useless 100-px ones.
    img = img.convert("RGBA")
    grid_step = _pick_grid_step(img.size)
    _draw_coordinate_grid(img, step=grid_step)
    click_marked = _draw_last_click_marker(img, pre_scale, (cleft, ctop))
    img = img.convert("RGB")

    name_png = f"{uuid.uuid4().hex}.png"
    path = _ensure_screenshot_dir() / name_png
    img.save(path, format="PNG", optimize=True)
    w_disp, h_disp = img.size

    # Update last-shot globals so subsequent computer_click coordinates
    # land on the right pixel of the right monitor. The origin is the
    # CLIPPED window top-left (so a partly off-screen window still has
    # its visible portion clickable).
    global _LAST_SHOT_SCALE, _LAST_SHOT_DISPLAY, _LAST_SHOT_ORIGIN
    _LAST_SHOT_SCALE = (
        orig_w / w_disp if w_disp else 1.0,
        orig_h / h_disp if h_disp else 1.0,
    )
    _LAST_SHOT_DISPLAY = (w_disp, h_disp)
    _LAST_SHOT_ORIGIN = (cleft, ctop)

    change_summary = _compute_screenshot_change(sig, titles)
    status_ctx = _capture_status_context()

    return {
        "ok": True,
        "name": name_png,
        "path": str(path),
        "width": w_disp,
        "height": h_disp,
        "matched_name": w.Name or "",
        "window_origin": (cleft, ctop),
        "source_size": (cw, ch),
        "change_summary": change_summary,
        "click_marked": click_marked,
        "status_context": status_ctx,
    }


async def screenshot_window(name: str, with_elements: bool = False) -> dict:
    """Take a screenshot cropped to one window's bounding rect.

    Cuts vision-token cost by 4-10x for window-scoped tasks (one app, one
    dialog, one settings panel) versus the full-monitor `screenshot`.
    Coordinates the model picks off the returned image are translated
    to screen pixels by `computer_click` & friends just like a normal
    screenshot — call this exactly the way you would call `screenshot`.

    Useful when:
      - the target app is small and surrounded by unrelated windows;
      - you want the model to focus on form fields without seeing the
        rest of the desktop;
      - vision-token budget is tight (long sessions).

    When `with_elements=True`, returns an `elements` list alongside the
    image — every clickable control inside that window with a cached
    id you can pass straight to `click_element_id`. Eliminates the
    screenshot→inspect_window two-call pattern for window-scoped tasks.

    Currently Windows-only.
    """
    try:
        if not name or not name.strip():
            return {"ok": False, "output": "", "error": "name required"}
        target = name.strip()
        result = await asyncio.to_thread(_screenshot_window_sync, target)
        if not result.get("ok"):
            return {
                "ok": False,
                "output": "",
                "error": result.get("error", "screenshot_window failed"),
            }
        ox, oy = result["window_origin"]
        sw, sh = result["source_size"]
        # `result` already carries `change_summary` and `click_marked` from
        # the sync helper, so feed it straight into _attach_shot_feedback
        # to surface the same one-line diff every other screenshot tool
        # gets.
        base = (
            f"Captured {result['matched_name']!r} — image "
            f"{result['width']}x{result['height']} from screen region "
            f"({ox},{oy}) {sw}x{sh}."
        )
        payload = {
            "ok": True,
            "output": _attach_shot_feedback(base, result),
            "image_path": result["name"],
        }
        if with_elements:
            # Use the matched window title from the sync helper so we
            # inspect exactly the same window we just captured (avoids
            # racing against a newly-foregrounded window of the same
            # prefix).
            elems = await asyncio.to_thread(
                _enumerate_elements_for_screenshot,
                result["matched_name"],
            )
            if elems.get("ok"):
                payload["elements"] = elems["elements"]
                payload["window_title"] = elems["window_title"]
                payload["output"] = _attach_shot_feedback(
                    (
                        f"{base} {len(elems['elements'])} clickable element"
                        f"{'s' if len(elems['elements']) != 1 else ''} "
                        f"cached — click any with "
                        f"`click_element_id({{\"id\": \"elN\"}})`."
                    ),
                    result,
                )
            else:
                payload["elements"] = []
                payload["elements_error"] = elems.get("error")
        return payload
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _list_windows_sync(max_count: int) -> dict:
    """Enumerate visible top-level windows.

    Returns one entry per window with title, screen bounds, foreground
    flag, and minimized flag. Filters out:
      - zero-titled windows (system message-only / hidden tool windows);
      - zero-area windows (collapsed system shells);
    Sorted: foreground first, then non-minimized, then enumeration order
    (which on Windows is roughly z-order top-down).
    """
    if sys.platform != "win32":
        return {"ok": False, "error": "list_windows is Windows-only"}
    try:
        import uiautomation as auto
    except ImportError:
        return {"ok": False, "error": "uiautomation package not installed"}
    import ctypes
    user32 = ctypes.windll.user32

    fg_hwnd = None
    try:
        fg = auto.GetForegroundControl()
        if fg is not None:
            fg_hwnd = int(fg.NativeWindowHandle)
    except Exception:
        pass

    out: list[dict] = []
    try:
        for w in auto.GetRootControl().GetChildren():
            try:
                title = w.Name or ""
                if not title:
                    continue
                r = w.BoundingRectangle
                if not r:
                    continue
                left = int(r.left)
                top = int(r.top)
                width = int(r.right - r.left)
                height = int(r.bottom - r.top)
                # Zero-area windows are usually invisible / collapsed —
                # they'd just clutter the list and can't be screenshotted.
                if width <= 0 or height <= 0:
                    continue
                hwnd = int(w.NativeWindowHandle)
                # IsIconic returns non-zero for minimized windows. The
                # call is cheap; per-window it's fine.
                try:
                    minimized = bool(user32.IsIconic(hwnd))
                except Exception:
                    minimized = False
                out.append({
                    "title": title,
                    "left": left,
                    "top": top,
                    "width": width,
                    "height": height,
                    "foreground": (hwnd == fg_hwnd),
                    "minimized": minimized,
                    "control_type": w.ControlTypeName or "",
                })
                if len(out) >= max_count:
                    break
            except Exception:
                # Per-window failure is non-fatal — many ghost / system
                # windows raise on .Name or .NativeWindowHandle access.
                continue
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}

    # Foreground first, then non-minimized, then keep the existing order.
    out.sort(key=lambda d: (not d["foreground"], d["minimized"]))
    return {"ok": True, "windows": out}


async def list_windows(max_count: int = 40) -> dict:
    """List visible top-level windows so the model can pick one.

    Returns each window's title, bounding rect, foreground flag, and
    minimized flag. Far more reliable than asking the model to read
    titles off a screenshot of the taskbar — the title strings are
    exact, and the bbox tells `screenshot_window` exactly where to
    crop. Once the model has a title, it can pass that to
    `focus_window`, `screenshot_window`, `inspect_window`, or
    `window_action`. Currently Windows-only.

    `max_count` is clamped to [1, 100] to keep output bounded.
    """
    try:
        cap = max(1, min(int(max_count or 40), 100))
        result = await asyncio.to_thread(_list_windows_sync, cap)
        if not result.get("ok"):
            return {
                "ok": False,
                "output": "",
                "error": result.get("error", "list_windows failed"),
            }
        wins = result["windows"]
        if not wins:
            return {
                "ok": True,
                "output": "(no nameable top-level windows found)",
                "windows": [],
            }
        lines = []
        for w in wins:
            tags = []
            if w["foreground"]:
                tags.append("foreground")
            if w["minimized"]:
                tags.append("minimized")
            tag_str = (" [" + ",".join(tags) + "]") if tags else ""
            lines.append(
                f"{w['title']!r}{tag_str} — {w['control_type']} "
                f"@({w['left']},{w['top']},{w['width']}x{w['height']})"
            )
        return {
            "ok": True,
            "output": (
                f"{len(wins)} window{'s' if len(wins) != 1 else ''} visible "
                f"(cap {cap}):\n" + "\n".join(lines)
            ),
            "windows": wins,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# UI-state polling — `ui_wait`
#
# The desktop is asynchronous: clicks open dialogs, navigations load pages,
# launchers hand off to apps. Without an explicit wait primitive the model
# falls into a screenshot-spam loop ("did the dialog appear yet? what
# about now? what about now?") that burns vision tokens and can race
# against the very thing it's waiting for.
#
# `ui_wait` polls one of four signals until it triggers (or a timeout
# elapses), then returns ONE screenshot. The `monitor` tool is the
# generic file/url/bash variant; this is its UI-specific sibling.
#
# kinds:
#   - window       — wait for a top-level window whose title contains `target`
#   - element      — wait for an a11y control whose accessible name contains
#                    `target` to become visible (Windows-only)
#   - text         — wait for OCR to find `target` somewhere on screen
#   - pixel_change — wait until the screen differs noticeably from the
#                    baseline taken at the start of the wait (good for
#                    animations / loads / paint-when-ready signals)
# ---------------------------------------------------------------------------
# Resolution we downsample screenshots to before pixel-diffing the
# baseline-vs-current frames. Tiny enough that compression noise mostly
# averages out, large enough that real UI changes (a dialog opening,
# a page rendering) flip plenty of pixels.
_PIXEL_DIFF_GRID = (64, 36)
# Fraction of pixels that must differ for `pixel_change` to fire. 5% is
# enough to catch a dialog opening anywhere on screen but not so low that
# subpixel anti-aliasing flicker triggers a false positive.
_PIXEL_DIFF_THRESHOLD = 0.05


async def _ui_wait_check(
    kind: str,
    target: str,
    baseline_pixels: list | None,
    require_enabled: bool = False,
) -> tuple[bool, str]:
    """One iteration of the `ui_wait` poll loop.

    Returns (matched, detail) where detail is a short status string
    surfaced in the timeout error so the model knows WHY the wait failed.

    `require_enabled` (used by kinds `element` and `element_enabled`)
    demands the matching control's IsEnabled == True — useful when
    waiting for a "Save" button to un-grey, not just to appear.
    """
    if kind == "window":
        if sys.platform != "win32":
            return False, "ui_wait kind='window' is Windows-only"
        win = await asyncio.to_thread(_find_window_sync, target)
        if win is None:
            return False, f"no window title contains {target!r} yet"
        try:
            return True, f"window {win.Name!r} is visible"
        except Exception:
            return True, "window appeared"

    if kind == "window_gone":
        # Inverse of `window` — succeeds when no matching window exists.
        # Useful for "wait until the progress dialog closes" without
        # polling pixel_change (which is noisy under video / ads).
        if sys.platform != "win32":
            return False, "ui_wait kind='window_gone' is Windows-only"
        win = await asyncio.to_thread(_find_window_sync, target)
        if win is None:
            return True, f"no window titled like {target!r} present"
        try:
            return False, f"window {win.Name!r} is still visible"
        except Exception:
            return False, f"matching window still visible"

    if kind in ("element", "element_enabled"):
        if sys.platform != "win32":
            return False, f"ui_wait kind={kind!r} is Windows-only"
        try:
            import uiautomation as auto  # noqa: F401  (probe for availability)
        except ImportError:
            return False, "uiautomation package not installed"
        # element_enabled is shorthand for `element` + require_enabled.
        enabled_only = require_enabled or (kind == "element_enabled")

        def _scan() -> tuple[str, bool | None]:
            # Bounded BFS in the foreground window — same shape as the
            # `click_element` search, just stop on the first visible
            # (and, if asked, enabled) match. Returns (name, is_enabled)
            # — is_enabled is None when probing the attribute raised.
            try:
                import uiautomation as auto2
                fg = auto2.GetForegroundControl()
            except Exception:
                return "", None
            if fg is None:
                return "", None
            needle = target.lower()
            queue = [(fg, 0)]
            visited = 0
            while queue and visited < 500:
                ctrl, depth = queue.pop(0)
                visited += 1
                try:
                    cname = ctrl.Name or ""
                    if cname and needle in cname.lower():
                        r = ctrl.BoundingRectangle
                        if r and (r.right - r.left) > 0 and (r.bottom - r.top) > 0:
                            try:
                                ce = bool(ctrl.IsEnabled)
                            except Exception:
                                ce = None
                            if enabled_only and ce is not True:
                                # Keep scanning — there may be another
                                # control with the same name substring
                                # that IS enabled (common when a dialog
                                # has both greyed-out and active copies
                                # of the same label in different panes).
                                pass
                            else:
                                return cname, ce
                except Exception:
                    pass
                if depth >= 12:
                    continue
                try:
                    for c in ctrl.GetChildren():
                        queue.append((c, depth + 1))
                except Exception:
                    continue
            return "", None

        match, is_enabled = await asyncio.to_thread(_scan)
        if match:
            if enabled_only and is_enabled is not True:
                return False, f"element {match!r} present but not yet enabled"
            return True, f"element {match!r} appeared" + (
                " (enabled)" if is_enabled else ""
            )
        return False, f"no element name contains {target!r} yet"

    if kind == "text":
        # Use the existing OCR pipeline. Capture a fresh screenshot so we
        # see whatever's on screen *right now*, then look for `target`
        # in the recognised words. The OCR call already supports a
        # `match` filter, so non-empty `words` after filtering means
        # the text is on screen.
        shot = await _capture_screenshot()
        ocr = await ocr_screenshot(shot["name"], match=target)
        if ocr.get("ok") and (ocr.get("words") or []):
            words = ocr["words"]
            sample = words[0]
            return True, (
                f"OCR matched {len(words)} word{'s' if len(words) != 1 else ''} "
                f"containing {target!r} (e.g. {sample.get('text')!r} @"
                f"({sample.get('x')}, {sample.get('y')}))"
            )
        return False, f"OCR has not found {target!r} yet"

    if kind == "pixel_change":
        # Compare the current frame's downsampled pixel histogram to the
        # baseline captured before polling started. Any meaningful change
        # in screen content trips this — a dialog opening, a page
        # rendering, a video advancing.
        shot = await _capture_screenshot()
        try:
            from PIL import Image
            with Image.open(shot["path"]) as im:
                cur = list(im.resize(_PIXEL_DIFF_GRID, Image.LANCZOS).getdata())
        except Exception as e:
            return False, f"could not read screenshot for diff: {e}"
        if baseline_pixels is None:
            return False, "no baseline (internal error)"
        diff = sum(1 for a, b in zip(baseline_pixels, cur) if a != b)
        ratio = diff / max(1, len(cur))
        if ratio >= _PIXEL_DIFF_THRESHOLD:
            return True, f"~{ratio * 100:.1f}% of pixels changed"
        return False, f"only ~{ratio * 100:.1f}% of pixels changed (threshold {_PIXEL_DIFF_THRESHOLD * 100:.0f}%)"

    return False, f"unknown kind {kind!r}"


async def ui_wait(
    kind: str,
    target: str = "",
    timeout_seconds: int = 15,
    interval_seconds: float = 1.0,
    require_enabled: bool = False,
) -> dict:
    """Poll a UI signal until it triggers (or a timeout elapses).

    `kind` is one of:
      - `window`          — `target` is a substring of the window title to wait for
      - `window_gone`     — wait for no window whose title contains `target` to remain
      - `element`         — `target` is a substring of the accessible name to wait for
      - `element_enabled` — as `element`, but only matches when the control is
                            interactive (UIA IsEnabled == True). Shorthand for
                            `element` + `require_enabled=True`.
      - `text`            — `target` is text to OCR-search for on screen
      - `pixel_change`    — wait for the screen to change noticeably (no `target` needed)

    `require_enabled` — when true, the `element` kind additionally demands
    the matching control's IsEnabled flag be True before triggering.
    `element_enabled` is a convenience alias that sets this for you.

    Capture-side, the loop runs at most every `interval_seconds`
    (clamped 0.25-5s) and gives up after `timeout_seconds` (clamped
    1-120s). Returns ONE screenshot when the condition triggers — the
    model sees the post-trigger state without paying for a screenshot
    on every poll iteration.
    """
    try:
        kind_norm = (kind or "").strip().lower()
        valid_kinds = {
            "window", "window_gone",
            "element", "element_enabled",
            "text", "pixel_change",
        }
        if kind_norm not in valid_kinds:
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"kind must be one of {sorted(valid_kinds)}, "
                    f"got {kind!r}"
                ),
            }
        target_clean = str(target or "").strip()
        if kind_norm != "pixel_change" and not target_clean:
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"target required for kind={kind_norm!r} — "
                    f"pass the substring you're waiting for"
                ),
            }
        try:
            timeout = max(1.0, min(float(timeout_seconds), 120.0))
        except (TypeError, ValueError):
            timeout = 15.0
        try:
            interval = max(0.25, min(float(interval_seconds), 5.0))
        except (TypeError, ValueError):
            interval = 1.0

        # Capture the baseline ONCE for pixel_change. The wait loop
        # then compares each fresh frame against this same baseline,
        # not against the previous frame, so a slow fade still trips
        # the check eventually.
        baseline_pixels = None
        if kind_norm == "pixel_change":
            try:
                shot = await _capture_screenshot()
                from PIL import Image
                with Image.open(shot["path"]) as im:
                    baseline_pixels = list(
                        im.resize(_PIXEL_DIFF_GRID, Image.LANCZOS).getdata()
                    )
            except Exception as e:
                return {
                    "ok": False,
                    "output": "",
                    "error": f"could not capture baseline: {type(e).__name__}: {e}",
                }

        start = time.monotonic()
        deadline = start + timeout
        attempts = 0
        last_status = ""
        while True:
            attempts += 1
            ok, status = await _ui_wait_check(
                kind_norm, target_clean, baseline_pixels,
                require_enabled=bool(require_enabled),
            )
            last_status = status
            if ok:
                shot = await _capture_screenshot()
                elapsed = time.monotonic() - start
                return {
                    "ok": True,
                    "output": _attach_shot_feedback(
                        (
                            f"ui_wait({kind_norm}, {target_clean!r}) triggered "
                            f"after {elapsed:.1f}s ({attempts} check{'s' if attempts != 1 else ''}): "
                            f"{status}"
                        ),
                        shot,
                    ),
                    "image_path": shot["name"],
                }
            if time.monotonic() >= deadline:
                shot = await _capture_screenshot()
                return {
                    "ok": False,
                    "output": "",
                    "error": _attach_shot_feedback(
                        (
                            f"ui_wait({kind_norm}, {target_clean!r}) timed out "
                            f"after {timeout:.0f}s ({attempts} checks). "
                            f"Last status: {status}"
                        ),
                        shot,
                    ),
                    "image_path": shot["name"],
                }
            await asyncio.sleep(interval)
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# `computer_batch` — sequence multiple desktop primitives in one tool call
#
# Form-filling is a brutal use of round-trips: every {click, type, key}
# pair is a separate model turn, with a screenshot round-tripped each
# time. `computer_batch` lets the model bundle those primitives together
# — one tool call, one screenshot at the end. For a 5-step login flow
# this halves end-to-end latency and a third of the vision tokens.
#
# Allowed step actions are LIMITED to the desktop primitives we actually
# trust to compose safely. Anything that branches based on result
# (read_file, web_search, bash) is deliberately excluded — those need
# the model in the loop. Anything destructive (delete, write, kill) is
# also excluded.
# ---------------------------------------------------------------------------
# Hard cap on steps per batch. Long chains get unreliable fast (early
# steps invalidate later coordinates), so prefer short batches.
_BATCH_MAX_STEPS = 20
# Inter-step settle. Long enough that consecutive clicks register on
# slow apps; short enough that a 10-step batch doesn't take 10 seconds.
_BATCH_STEP_SETTLE_SEC = 0.10
# Cap on per-step `wait_ms` so a malicious / confused sequence can't
# stall the agent for minutes.
_BATCH_MAX_WAIT_MS = 5000

_BATCH_ALLOWED_ACTIONS = {
    "click",
    "double_click",
    "right_click",
    "middle_click",
    "type",
    "key",
    "scroll",
    "mouse_move",
    "drag",
    "click_element",
    "click_element_id",
    "focus_window",
    "wait_ms",
}


async def _batch_run_step(step: dict, idx: int) -> tuple[bool, str]:
    """Execute one batch step. Returns (ok, summary).

    Each step is a dict with `action` plus the action-specific args.
    Any failure stops the batch — the caller still takes the trailing
    screenshot so the model can see the partial-execution state.
    """
    if not isinstance(step, dict):
        return False, f"step {idx}: not a dict"
    action = str(step.get("action") or "").strip().lower()
    if action not in _BATCH_ALLOWED_ACTIONS:
        return False, (
            f"step {idx}: action {action!r} not allowed; valid actions are "
            f"{sorted(_BATCH_ALLOWED_ACTIONS)}"
        )

    # Map mouse-click variants onto computer_click's button + double args.
    # We keep these as separate top-level actions in the schema for
    # readability, but they all reduce to the same primitive.
    if action in {"click", "double_click", "right_click", "middle_click"}:
        x = step.get("x")
        y = step.get("y")
        if x is None or y is None:
            return False, f"step {idx} ({action}): x and y required"
        button = {
            "click": "left",
            "double_click": "left",
            "right_click": "right",
            "middle_click": "middle",
        }[action]
        double = action == "double_click"
        try:
            ix, iy, _n = _clamp_to_image(x, y)
            rx, ry = _map_to_screen(ix, iy)
            clicks = 2 if double else 1
            await asyncio.to_thread(
                _computer_click_sync, rx, ry, button, clicks
            )
        except Exception as e:
            return False, f"step {idx} ({action}) failed: {type(e).__name__}: {e}"
        return True, f"{action}({ix},{iy})"

    if action == "type":
        text = step.get("text", "")
        if not isinstance(text, str):
            return False, f"step {idx} (type): text must be a string"
        if len(text) > 10000:
            return False, f"step {idx} (type): text exceeds 10k chars"
        try:
            interval = float(step.get("interval", 0.02))
        except (TypeError, ValueError):
            interval = 0.02
        interval = max(0.0, min(interval, 0.2))
        try:
            await asyncio.to_thread(_computer_type_sync, text, interval)
        except Exception as e:
            return False, f"step {idx} (type) failed: {type(e).__name__}: {e}"
        preview = text if len(text) <= 30 else (text[:30] + "…")
        return True, f"type({preview!r})"

    if action == "key":
        keys = step.get("keys", "")
        if not keys:
            return False, f"step {idx} (key): keys required"
        # Normalize to the same shape computer_key accepts.
        if isinstance(keys, str):
            seq = [k.strip() for k in keys.split("+") if k.strip()]
        elif isinstance(keys, list):
            seq = [str(k).strip() for k in keys if str(k).strip()]
        else:
            return False, f"step {idx} (key): keys must be string or list"
        if not seq:
            return False, f"step {idx} (key): no key names parsed from {keys!r}"
        try:
            await asyncio.to_thread(_computer_key_sync, seq)
        except Exception as e:
            return False, f"step {idx} (key) failed: {type(e).__name__}: {e}"
        return True, f"key({'+'.join(seq)})"

    if action == "scroll":
        x = step.get("x")
        y = step.get("y")
        if x is None or y is None:
            return False, f"step {idx} (scroll): x and y required"
        direction = (step.get("direction") or "down").lower()
        amount = int(step.get("amount", 5))
        ticks = amount if direction == "up" else -amount
        try:
            ix, iy, _n = _clamp_to_image(x, y)
            rx, ry = _map_to_screen(ix, iy)
            await asyncio.to_thread(_computer_scroll_sync, rx, ry, ticks)
        except Exception as e:
            return False, f"step {idx} (scroll) failed: {type(e).__name__}: {e}"
        return True, f"scroll({direction} x{amount} @ {ix},{iy})"

    if action == "mouse_move":
        x = step.get("x")
        y = step.get("y")
        if x is None or y is None:
            return False, f"step {idx} (mouse_move): x and y required"
        try:
            ix, iy, _n = _clamp_to_image(x, y)
            rx, ry = _map_to_screen(ix, iy)
            await asyncio.to_thread(_computer_mouse_move_sync, rx, ry)
        except Exception as e:
            return False, f"step {idx} (mouse_move) failed: {type(e).__name__}: {e}"
        return True, f"mouse_move({ix},{iy})"

    if action == "drag":
        for k in ("x1", "y1", "x2", "y2"):
            if step.get(k) is None:
                return False, f"step {idx} (drag): {k} required"
        try:
            duration = float(step.get("duration", 0.4))
        except (TypeError, ValueError):
            duration = 0.4
        duration = max(0.05, min(duration, 5.0))
        button = (step.get("button") or "left").lower()
        if button not in {"left", "right", "middle"}:
            return False, f"step {idx} (drag): invalid button {button!r}"
        try:
            ix1, iy1, _ = _clamp_to_image(step["x1"], step["y1"])
            ix2, iy2, _ = _clamp_to_image(step["x2"], step["y2"])
            rx1, ry1 = _map_to_screen(ix1, iy1)
            rx2, ry2 = _map_to_screen(ix2, iy2)
            await asyncio.to_thread(
                _computer_drag_sync, rx1, ry1, rx2, ry2, duration, button
            )
        except Exception as e:
            return False, f"step {idx} (drag) failed: {type(e).__name__}: {e}"
        return True, f"drag({ix1},{iy1}→{ix2},{iy2})"

    if action == "click_element":
        name = str(step.get("name") or "").strip()
        if not name:
            return False, f"step {idx} (click_element): name required"
        match = step.get("match") or "contains"
        click_type = step.get("click_type") or "left"
        try:
            timeout_f = float(step.get("timeout", 2.0))
        except (TypeError, ValueError):
            timeout_f = 2.0
        timeout_f = max(0.1, min(timeout_f, 10.0))
        result = await asyncio.to_thread(
            _click_element_sync, name, match, click_type, timeout_f
        )
        if not result.get("ok"):
            return False, f"step {idx} (click_element) failed: {result.get('error')}"
        return True, f"click_element({name!r})"

    if action == "click_element_id":
        eid = str(step.get("id") or "").strip()
        if not eid:
            return False, f"step {idx} (click_element_id): id required"
        click_type = step.get("click_type") or "left"
        if click_type not in ("left", "right", "middle", "double"):
            return False, (
                f"step {idx} (click_element_id): invalid click_type "
                f"{click_type!r}"
            )
        result = await asyncio.to_thread(
            _click_element_id_sync, eid, click_type
        )
        if not result.get("ok"):
            return False, f"step {idx} (click_element_id) failed: {result.get('error')}"
        return True, f"click_element_id({eid})"

    if action == "focus_window":
        name = str(step.get("name") or "").strip()
        if not name:
            return False, f"step {idx} (focus_window): name required"
        result = await asyncio.to_thread(_focus_window_sync, name)
        # Brief settle so subsequent type/key target the freshly-focused window.
        await asyncio.sleep(0.4)
        if not result.get("ok"):
            return False, f"step {idx} (focus_window) failed: {result.get('error')}"
        return True, f"focus_window({name!r}→{result.get('matched_name')!r})"

    if action == "wait_ms":
        try:
            ms = int(step.get("ms", 0))
        except (TypeError, ValueError):
            return False, f"step {idx} (wait_ms): ms must be an int"
        ms = max(0, min(ms, _BATCH_MAX_WAIT_MS))
        await asyncio.sleep(ms / 1000.0)
        return True, f"wait_ms({ms})"

    # Defensive — should be unreachable because of the allowlist check above.
    return False, f"step {idx}: action {action!r} reached default branch"


async def computer_batch(
    steps: list[dict] | None = None,
    screenshot: bool = True,
) -> dict:
    """Run a SHORT sequence of desktop primitives in one tool call.

    Each entry of `steps` is a dict with `action` plus action-specific
    fields. Allowed actions:

      - `{action: "click", x, y}` (also `double_click`, `right_click`, `middle_click`)
      - `{action: "type", text, interval?}`
      - `{action: "key", keys: "ctrl+l"}` or `{action: "key", keys: ["enter"]}`
      - `{action: "scroll", x, y, direction: "down"|"up", amount}`
      - `{action: "mouse_move", x, y}`
      - `{action: "drag", x1, y1, x2, y2, duration?, button?}`
      - `{action: "click_element", name, match?, click_type?}`
      - `{action: "click_element_id", id, click_type?}`
      - `{action: "focus_window", name}`
      - `{action: "wait_ms", ms}`

    Up to 20 steps per call. A small inter-step settle is inserted
    automatically. If any step fails, the batch stops and we still take
    a screenshot so the model can see the partial state. With
    `screenshot=true` (default) one screenshot is captured at the end —
    set false only if the next thing you do is a screenshot anyway.
    """
    if not isinstance(steps, list) or not steps:
        return {
            "ok": False,
            "output": "",
            "error": "steps must be a non-empty list",
        }
    if len(steps) > _BATCH_MAX_STEPS:
        return {
            "ok": False,
            "output": "",
            "error": (
                f"too many steps ({len(steps)}); cap is {_BATCH_MAX_STEPS}. "
                f"Split into multiple batches."
            ),
        }
    summaries: list[str] = []
    failed_at: int | None = None
    failure_msg = ""
    for i, step in enumerate(steps):
        try:
            ok, summary = await _batch_run_step(step, i + 1)
        except Exception as e:
            ok = False
            summary = f"step {i + 1} crashed: {type(e).__name__}: {e}"
        summaries.append(("OK " if ok else "ERR ") + summary)
        if not ok:
            failed_at = i + 1
            failure_msg = summary
            break
        # Inter-step settle to let the previous action's UI side-effects
        # land before the next step targets coordinates / focus.
        await asyncio.sleep(_BATCH_STEP_SETTLE_SEC)

    out_lines = "\n".join(f"  {s}" for s in summaries)
    if failed_at is not None:
        # Always include a screenshot on failure so the model sees the
        # partial-execution state and can pick the right recovery move.
        try:
            shot = await _capture_screenshot()
            return {
                "ok": False,
                "output": "",
                "error": _attach_shot_feedback(
                    (
                        f"computer_batch stopped at step {failed_at}/{len(steps)}: "
                        f"{failure_msg}\nSteps run:\n{out_lines}"
                    ),
                    shot,
                ),
                "image_path": shot["name"],
            }
        except Exception:
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"computer_batch stopped at step {failed_at}/{len(steps)}: "
                    f"{failure_msg}\nSteps run:\n{out_lines}"
                ),
            }

    base_output = (
        f"computer_batch ran {len(steps)} step{'s' if len(steps) != 1 else ''} "
        f"successfully:\n{out_lines}"
    )
    result: dict = {"ok": True, "output": base_output}
    if screenshot:
        try:
            shot = await _capture_screenshot()
            result["image_path"] = shot["name"]
            # Append the same change-summary feedback every other
            # screenshot tool surfaces, so a 5-step form fill ends with
            # an explicit "Δ 12% pixels; new window 'Confirm' appeared"
            # line instead of the model having to compare images.
            result["output"] = _attach_shot_feedback(base_output, shot)
        except Exception:
            # If the post-batch screenshot fails, the steps still ran
            # — surface success but without an image. The model can
            # always call `screenshot` next.
            pass
    return result


# ---------------------------------------------------------------------------
# Document readers — `read_doc`
#
# `read_file` is for text files. For PDFs, .docx, and .xlsx we need real
# parsers to extract prose. Each backend gets a tiny converter that returns
# plain text with soft section markers (page/sheet/paragraph) so the model
# can cite page numbers when summarising.
# ---------------------------------------------------------------------------
# Cap on characters we ship back. A big PDF trivially exceeds any model's
# context; the model can ask for a specific page range if it needs more.
_DOC_MAX_CHARS = 40000


def _read_pdf_sync(path: str, pages: str | None) -> dict:
    try:
        import fitz  # pymupdf
    except ImportError:
        return {"ok": False, "error": "pymupdf is not installed"}
    try:
        doc = fitz.open(path)
    except Exception as e:
        return {"ok": False, "error": f"could not open PDF: {e}"}
    try:
        n = doc.page_count
        # Parse `pages` like "1-5" or "3" or "1,5,8-10". Empty = first 20 pages.
        wanted: list[int] = []
        if pages:
            for part in pages.split(","):
                part = part.strip()
                if "-" in part:
                    a, b = part.split("-", 1)
                    try:
                        a, b = int(a), int(b)
                        wanted.extend(range(a, b + 1))
                    except Exception:
                        continue
                else:
                    try:
                        wanted.append(int(part))
                    except Exception:
                        continue
        else:
            wanted = list(range(1, min(n, 20) + 1))
        out: list[str] = []
        for p in wanted:
            if p < 1 or p > n:
                continue
            page = doc.load_page(p - 1)
            text = page.get_text("text") or ""
            out.append(f"--- page {p} ---\n{text.strip()}")
        full = "\n\n".join(out)
        truncated = False
        if len(full) > _DOC_MAX_CHARS:
            full = full[:_DOC_MAX_CHARS] + "\n\n...[truncated]"
            truncated = True
        return {
            "ok": True,
            "output": full,
            "page_count": n,
            "pages_returned": [p for p in wanted if 1 <= p <= n],
            "truncated": truncated,
        }
    finally:
        doc.close()


def _read_docx_sync(path: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"ok": False, "error": "python-docx is not installed"}
    try:
        d = Document(path)
    except Exception as e:
        return {"ok": False, "error": f"could not open docx: {e}"}
    parts: list[str] = []
    # Paragraphs in document order.
    for para in d.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)
    # Tables after paragraphs — crude flatten, one row per line.
    for ti, tbl in enumerate(d.tables):
        parts.append(f"--- table {ti + 1} ---")
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    full = "\n\n".join(parts)
    truncated = False
    if len(full) > _DOC_MAX_CHARS:
        full = full[:_DOC_MAX_CHARS] + "\n\n...[truncated]"
        truncated = True
    return {"ok": True, "output": full, "truncated": truncated}


def _read_xlsx_sync(path: str, sheets: str | None) -> dict:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"ok": False, "error": "openpyxl is not installed"}
    try:
        wb = load_workbook(path, data_only=True, read_only=True)
    except Exception as e:
        return {"ok": False, "error": f"could not open xlsx: {e}"}
    names = wb.sheetnames
    wanted = [s.strip() for s in (sheets.split(",") if sheets else []) if s.strip()]
    if not wanted:
        wanted = names[:3]  # first three sheets by default
    parts: list[str] = []
    for sn in wanted:
        if sn not in names:
            continue
        ws = wb[sn]
        parts.append(f"--- sheet {sn!r} ({ws.max_row} rows x {ws.max_column} cols) ---")
        # Cap rows to keep outputs manageable — model can ask for more.
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 200:
                parts.append(f"...[{ws.max_row - i} more rows]")
                break
            parts.append(" | ".join("" if v is None else str(v) for v in row))
    full = "\n".join(parts)
    truncated = False
    if len(full) > _DOC_MAX_CHARS:
        full = full[:_DOC_MAX_CHARS] + "\n\n...[truncated]"
        truncated = True
    return {"ok": True, "output": full, "sheets": names, "truncated": truncated}


async def read_doc(
    path: str,
    pages: str | None = None,
    sheets: str | None = None,
) -> dict:
    """Extract readable text from a PDF, .docx, or .xlsx file.

    `path` must exist and point at one of those three formats (detected by
    extension). `pages` is a page-range spec for PDFs (e.g. '1-5', '3,7,9');
    omit for the first 20 pages. `sheets` is a comma-separated list of sheet
    names for xlsx; omit for the first 3 sheets.
    """
    try:
        if not path or not str(path).strip():
            return {"ok": False, "output": "", "error": "path required"}
        # Resolve path relative to cwd for the same ergonomics as read_file.
        p = Path(path).expanduser()
        if not p.is_absolute():
            p = Path.cwd() / p
        if not p.is_file():
            return {"ok": False, "output": "", "error": f"file not found: {p}"}
        suffix = p.suffix.lower()
        if suffix == ".pdf":
            r = await asyncio.to_thread(_read_pdf_sync, str(p), pages)
        elif suffix == ".docx":
            r = await asyncio.to_thread(_read_docx_sync, str(p))
        elif suffix in {".xlsx", ".xlsm"}:
            r = await asyncio.to_thread(_read_xlsx_sync, str(p), sheets)
        else:
            return {
                "ok": False,
                "output": "",
                "error": (
                    f"unsupported extension {suffix!r}; read_doc handles "
                    f".pdf / .docx / .xlsx. For plain text use `read_file`."
                ),
            }
        if not r.get("ok"):
            return {"ok": False, "output": "", "error": r.get("error") or "read_doc failed"}
        return {"ok": True, "output": r.get("output") or "(empty document)"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# OCR — `ocr_screenshot`
#
# When `click_element` can't reach something (web content without a11y, custom-
# rendered apps, screenshots of remote desktops), the fallback is pixel clicks.
# OCR lets the model ground those clicks in *words*: instead of "click the
# pixel at (732, 918)", it can say "click the word 'Sign in' at its OCR
# bounding box centre". Two backends, tried in order:
#   1. winsdk.windows.media.ocr — built into every Win10+ machine, no binary.
#   2. pytesseract — cross-platform, but needs a Tesseract binary installed.
# ---------------------------------------------------------------------------
def _ocr_windows_sync(image_path: str) -> dict:
    try:
        import asyncio as _asyncio
        from winsdk.windows.media.ocr import OcrEngine
        from winsdk.windows.graphics.imaging import BitmapDecoder
        from winsdk.windows.storage import StorageFile
        from winsdk.windows.storage.streams import RandomAccessStreamReference
    except Exception as e:
        return {"ok": False, "error": f"winsdk not available: {e}"}

    async def _run():
        # Windows OCR takes a SoftwareBitmap from an async decoder pipeline.
        sf = await StorageFile.get_file_from_path_async(image_path)
        stream = await sf.open_async(0)  # 0 = Read
        decoder = await BitmapDecoder.create_async(stream)
        bmp = await decoder.get_software_bitmap_async()
        engine = OcrEngine.try_create_from_user_profile_languages()
        if engine is None:
            # User hasn't installed any OCR-capable language pack.
            return {"ok": False, "error": "no OCR language pack installed (Settings > Time & Language > Language)"}
        result = await engine.recognize_async(bmp)
        words = []
        for line in result.lines:
            for w in line.words:
                r = w.bounding_rect
                words.append({
                    "text": w.text,
                    "x": int(r.x),
                    "y": int(r.y),
                    "width": int(r.width),
                    "height": int(r.height),
                })
        return {"ok": True, "words": words, "backend": "windows"}

    # winsdk uses its own asyncio-compatible event loop. Run it blocking.
    try:
        return _asyncio.run(_run())
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}


def _ocr_tesseract_sync(image_path: str) -> dict:
    try:
        import pytesseract
        from PIL import Image as PILImage
    except Exception as e:
        return {"ok": False, "error": f"pytesseract not available: {e}"}
    try:
        img = PILImage.open(image_path)
    except Exception as e:
        return {"ok": False, "error": f"could not open image: {e}"}
    try:
        # image_to_data returns a dict with keys: text, left, top, width, height, conf, ...
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
    except pytesseract.TesseractNotFoundError:
        return {"ok": False, "error": "tesseract binary not found on PATH (install Tesseract)"}
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}
    words = []
    for i, txt in enumerate(data.get("text", [])):
        t = (txt or "").strip()
        if not t:
            continue
        words.append({
            "text": t,
            "x": int(data["left"][i]),
            "y": int(data["top"][i]),
            "width": int(data["width"][i]),
            "height": int(data["height"][i]),
        })
    return {"ok": True, "words": words, "backend": "tesseract"}


async def ocr_screenshot(
    image_path: str | None = None,
    match: str | None = None,
) -> dict:
    """Run OCR on a screenshot and return `[{text, x, y, width, height}]` for
    every recognised word, in reading order.

    If `image_path` is omitted, captures a fresh screenshot first (most common
    case). If `match` is provided, the output is filtered to only words whose
    text contains `match` (case-insensitive substring) — handy shortcut for
    "give me the bbox of 'Sign in' on screen right now".
    """
    try:
        # Resolve / capture the image.
        if image_path and image_path.strip():
            p = Path(image_path).expanduser()
            if not p.is_absolute():
                # The model normally receives screenshot names like
                # "abc123.png" and refers to them by that name; translate
                # back to the screenshots dir on disk.
                cand = _ensure_screenshot_dir() / p
                if cand.is_file():
                    p = cand
                else:
                    p = Path.cwd() / p
            if not p.is_file():
                return {"ok": False, "output": "", "error": f"image not found: {p}"}
            path_str = str(p)
        else:
            shot = await _capture_screenshot()
            path_str = shot["path"]
        # Try Windows OCR first, then Tesseract fallback.
        r = await asyncio.to_thread(_ocr_windows_sync, path_str)
        if not r.get("ok"):
            r = await asyncio.to_thread(_ocr_tesseract_sync, path_str)
        if not r.get("ok"):
            return {"ok": False, "output": "", "error": r.get("error") or "no OCR backend available"}
        words = r["words"]
        if match and match.strip():
            needle = match.strip().lower()
            words = [w for w in words if needle in w["text"].lower()]
        # Build compact printable summary — keep first 200 words so the
        # response stays within reasonable size. Full word list is in the
        # structured fields for programmatic use.
        printable = "\n".join(
            f"{i + 1}. {w['text']!r} @({w['x']}, {w['y']}, {w['width']}x{w['height']})"
            for i, w in enumerate(words[:200])
        )
        more = "" if len(words) <= 200 else f"\n...[{len(words) - 200} more words]"
        return {
            "ok": True,
            "output": (
                f"OCR backend: {r.get('backend')}; found {len(words)} word(s)"
                + (f" matching {match!r}" if match else "")
                + (":\n" + printable if printable else " (none)")
                + more
            ),
            "words": words,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Browser automation via Chrome DevTools Protocol — `browser_*`
#
# Pixel-clicking a browser is a vision gamble. If the user has Chrome running
# with `--remote-debugging-port=9222`, we can drive it via CDP: JSON-over-
# WebSocket commands that address DOM nodes by CSS selector. This entirely
# sidesteps the OCR / click_element / screenshot-reading pipeline for web
# tasks.
#
# The tools we expose:
#   - browser_tabs     — list every open tab (id + url + title)
#   - browser_goto     — navigate the active (or chosen) tab to a URL
#   - browser_click    — click a DOM node by CSS selector
#   - browser_type     — type into a focused element
#   - browser_text     — extract text content from a selector
#   - browser_eval     — run arbitrary JS in the page (escape hatch)
#
# The client is deliberately minimal — a single connection per call, talks
# JSON-RPC over websockets, no persistent session. Keeps the code small at
# the cost of reconnecting every turn, which is fine for agent usage.
# ---------------------------------------------------------------------------
_CDP_DEFAULT_PORT = 9222


async def _cdp_list_tabs(port: int) -> list[dict]:
    """Return the raw tab list from `/json` on the DevTools HTTP endpoint."""
    import httpx
    async with httpx.AsyncClient(timeout=5.0) as c:
        r = await c.get(f"http://127.0.0.1:{port}/json")
        r.raise_for_status()
        tabs = r.json()
    # Filter for actual pages (CDP also reports iframes + service workers).
    return [t for t in tabs if t.get("type") == "page"]


async def _cdp_send(ws, method: str, params: dict | None = None) -> dict:
    """Send one JSON-RPC command and return its result. Each call uses its
    own monotonic id so we correctly pair responses when multiple commands
    are in flight.
    """
    import json as _json
    _cdp_send.counter = getattr(_cdp_send, "counter", 0) + 1
    msg_id = _cdp_send.counter
    payload = {"id": msg_id, "method": method, "params": params or {}}
    await ws.send(_json.dumps(payload))
    # Drain messages until we see our response. CDP intersperses events
    # (Page.frameStoppedLoading, Runtime.executionContextCreated, etc.) in
    # the same stream, so we have to filter by id.
    import asyncio as _asyncio
    deadline = _asyncio.get_event_loop().time() + 15.0
    while True:
        remaining = deadline - _asyncio.get_event_loop().time()
        if remaining <= 0:
            raise TimeoutError(f"CDP {method} timed out")
        raw = await _asyncio.wait_for(ws.recv(), timeout=remaining)
        msg = _json.loads(raw)
        if msg.get("id") == msg_id:
            if "error" in msg:
                raise RuntimeError(f"CDP error: {msg['error']}")
            return msg.get("result", {})
        # else: it's an event, ignore.


async def _cdp_connect(port: int, tab_index: int | None) -> tuple[object, dict]:
    """Open a WebSocket to the requested tab's debugger endpoint.

    Picks the first tab by default, or the one at the given 0-based index.
    Returns (websocket, tab_dict). Caller is responsible for closing the ws.
    """
    import websockets as _ws_lib  # lazy import — not installed everywhere
    tabs = await _cdp_list_tabs(port)
    if not tabs:
        raise RuntimeError(
            f"no Chrome tabs visible on CDP port {port}. Make sure Chrome was "
            f"launched with --remote-debugging-port={port} (open_app can "
            f"pass it via args)."
        )
    idx = 0 if tab_index is None else int(tab_index)
    if idx < 0 or idx >= len(tabs):
        raise RuntimeError(
            f"tab_index {idx} out of range; {len(tabs)} tabs open (0..{len(tabs) - 1})"
        )
    tab = tabs[idx]
    ws_url = tab.get("webSocketDebuggerUrl")
    if not ws_url:
        raise RuntimeError("selected tab has no webSocketDebuggerUrl")
    ws = await _ws_lib.connect(ws_url, max_size=32 * 1024 * 1024)
    return ws, tab


async def browser_tabs(port: int = _CDP_DEFAULT_PORT) -> dict:
    """List every browser tab CDP can see, with index / title / URL."""
    try:
        tabs = await _cdp_list_tabs(port)
        lines = [
            f"{i}. {t.get('title', '(no title)')[:80]}  {t.get('url', '')[:120]}"
            for i, t in enumerate(tabs)
        ]
        return {
            "ok": True,
            "output": (
                f"{len(tabs)} tab(s) on CDP port {port}:\n" + "\n".join(lines)
                if tabs else
                f"No tabs visible on CDP port {port}."
            ),
            "tabs": tabs,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def browser_goto(
    url: str,
    tab_index: int | None = None,
    port: int = _CDP_DEFAULT_PORT,
) -> dict:
    """Navigate a tab to a URL and wait for it to finish loading."""
    try:
        if not url or not url.strip():
            return {"ok": False, "output": "", "error": "url required"}
        u = url.strip()
        # Scheme guard — CDP will happily navigate to file://, javascript:,
        # data:, chrome: URLs, which are all prompt-injection risks if the
        # URL came from a tool output that a hostile page can influence.
        # Parse properly instead of string-matching so `javascript:alert(1)`
        # (single colon, no double-slash) gets rejected too.
        parsed = urlparse(u)
        if parsed.scheme:
            if parsed.scheme.lower() not in {"http", "https"}:
                return {
                    "ok": False,
                    "output": "",
                    "error": f"only http/https schemes are allowed, got {parsed.scheme!r}",
                }
        else:
            # No scheme → assume https://
            u = "https://" + u
        ws, tab = await _cdp_connect(port, tab_index)
        try:
            await _cdp_send(ws, "Page.enable")
            await _cdp_send(ws, "Page.navigate", {"url": u})
            # Wait a short moment for the navigation; a full loadEventFired
            # wait would block on ad trackers etc. — for agent flow 2s is
            # enough to get the main document.
            import asyncio as _asyncio
            await _asyncio.sleep(2.0)
            r = await _cdp_send(ws, "Runtime.evaluate", {"expression": "document.title"})
            title = (r.get("result", {}) or {}).get("value") or ""
            r2 = await _cdp_send(ws, "Runtime.evaluate", {"expression": "location.href"})
            final = (r2.get("result", {}) or {}).get("value") or u
            return {
                "ok": True,
                "output": f"navigated to {final!r} (title: {title!r})",
                "url": final,
                "title": title,
            }
        finally:
            await ws.close()
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _js_literal(s: str) -> str:
    """Encode a Python string as a JSON-safe JS string literal."""
    import json as _json
    return _json.dumps(s)


async def browser_click(
    selector: str,
    tab_index: int | None = None,
    port: int = _CDP_DEFAULT_PORT,
) -> dict:
    """Click the first DOM element matching a CSS selector in the target tab."""
    try:
        if not selector or not selector.strip():
            return {"ok": False, "output": "", "error": "selector required"}
        ws, tab = await _cdp_connect(port, tab_index)
        try:
            # Use Runtime.evaluate so we don't have to set up DOM nodeIds.
            expr = (
                "(() => { const el = document.querySelector(" + _js_literal(selector) + "); "
                "if (!el) return {found: false}; "
                "el.scrollIntoView({block: 'center'}); "
                "el.click(); "
                "return {found: true, text: (el.innerText || el.value || '').slice(0, 200)}; })()"
            )
            r = await _cdp_send(ws, "Runtime.evaluate", {
                "expression": expr,
                "returnByValue": True,
            })
            val = (r.get("result", {}) or {}).get("value") or {}
            if not val.get("found"):
                return {"ok": False, "output": "", "error": f"no element matches {selector!r}"}
            return {
                "ok": True,
                "output": f"clicked {selector!r} (text: {val.get('text', '')!r})",
            }
        finally:
            await ws.close()
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def browser_type(
    selector: str,
    text: str,
    press_enter: bool = False,
    tab_index: int | None = None,
    port: int = _CDP_DEFAULT_PORT,
) -> dict:
    """Focus the element at `selector`, set its value, fire input events,
    and optionally press Enter. Works for <input>, <textarea>, and
    contenteditable elements.
    """
    try:
        if not selector or not selector.strip():
            return {"ok": False, "output": "", "error": "selector required"}
        ws, _tab = await _cdp_connect(port, tab_index)
        try:
            # Setting .value + dispatching 'input' handles React/Vue forms.
            expr = (
                "(() => { const el = document.querySelector(" + _js_literal(selector) + "); "
                "if (!el) return {found: false}; "
                "el.focus(); "
                "if ('value' in el) { el.value = " + _js_literal(text) + "; "
                "el.dispatchEvent(new Event('input', {bubbles: true})); "
                "el.dispatchEvent(new Event('change', {bubbles: true})); } "
                "else { el.innerText = " + _js_literal(text) + "; } "
                "return {found: true}; })()"
            )
            r = await _cdp_send(ws, "Runtime.evaluate", {
                "expression": expr,
                "returnByValue": True,
            })
            val = (r.get("result", {}) or {}).get("value") or {}
            if not val.get("found"):
                return {"ok": False, "output": "", "error": f"no element matches {selector!r}"}
            if press_enter:
                # Dispatch an Enter keypress at the focused element — this
                # triggers form submissions even when there's no submit button.
                for ev_type in ("keydown", "keypress", "keyup"):
                    await _cdp_send(ws, "Input.dispatchKeyEvent", {
                        "type": "rawKeyDown" if ev_type == "keydown" else (
                            "char" if ev_type == "keypress" else "keyUp"
                        ),
                        "key": "Enter",
                        "code": "Enter",
                        "windowsVirtualKeyCode": 13,
                        "nativeVirtualKeyCode": 13,
                    })
            return {"ok": True, "output": f"typed into {selector!r}" + (" and pressed Enter" if press_enter else "")}
        finally:
            await ws.close()
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def browser_text(
    selector: str = "body",
    tab_index: int | None = None,
    port: int = _CDP_DEFAULT_PORT,
    max_chars: int = 15000,
) -> dict:
    """Return the rendered text content of a CSS selector (default: whole
    page). Much more reliable than `fetch_url` for JS-rendered pages.
    """
    try:
        ws, _tab = await _cdp_connect(port, tab_index)
        try:
            expr = (
                "(() => { const el = document.querySelector(" + _js_literal(selector) + "); "
                "if (!el) return null; return el.innerText; })()"
            )
            r = await _cdp_send(ws, "Runtime.evaluate", {
                "expression": expr,
                "returnByValue": True,
            })
            val = (r.get("result", {}) or {}).get("value")
            if val is None:
                return {"ok": False, "output": "", "error": f"no element matches {selector!r}"}
            text = str(val)
            cap = max(500, min(int(max_chars or 15000), 50000))
            truncated = False
            if len(text) > cap:
                text = text[:cap] + "\n...[truncated]"
                truncated = True
            return {
                "ok": True,
                "output": text,
                "truncated": truncated,
            }
        finally:
            await ws.close()
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def browser_eval(
    expression: str,
    tab_index: int | None = None,
    port: int = _CDP_DEFAULT_PORT,
) -> dict:
    """Run an arbitrary JS expression in the page and return its value as JSON.

    Escape hatch when the other browser_* primitives aren't enough. Because
    this can do absolutely anything the page can do (including exfiltrate
    cookies), the model should only use it when a simpler browser tool can't.
    """
    try:
        if not expression or not expression.strip():
            return {"ok": False, "output": "", "error": "expression required"}
        ws, _tab = await _cdp_connect(port, tab_index)
        try:
            r = await _cdp_send(ws, "Runtime.evaluate", {
                "expression": expression,
                "returnByValue": True,
                "awaitPromise": True,
            })
            val = r.get("result", {})
            v = val.get("value")
            # Serialize for display — compact JSON for objects, raw for scalars.
            import json as _json
            if isinstance(v, (dict, list)):
                text = _json.dumps(v)[:8000]
            else:
                text = repr(v)[:8000]
            return {"ok": True, "output": text, "result": v}
        finally:
            await ws.close()
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Scheduled tasks — `schedule_task` / `list_scheduled_tasks` / `cancel_scheduled_task`
#
# Creates a persistent row in the scheduled_tasks SQLite table with a run-at
# timestamp (or cron-like recurrence) and an agent prompt to execute. A
# background daemon (started by backend.app on boot) polls the table every
# 30s, fires due tasks as new conversations, and updates next_run_at for
# recurring ones.
# ---------------------------------------------------------------------------
async def schedule_task(
    name: str,
    prompt: str,
    run_at: str | None = None,
    every_minutes: int | None = None,
    cwd: str | None = None,
) -> dict:
    """Queue a prompt to run in a fresh conversation at a specific time (or
    on a repeating interval). Exactly one of `run_at` (ISO 8601 datetime)
    or `every_minutes` (int, >=1) must be given. `cwd` overrides the
    working directory for the scheduled run; defaults to the current cwd.
    """
    try:
        from . import db as _db
        name = (name or "").strip()
        prompt = (prompt or "").strip()
        if not name:
            return {"ok": False, "output": "", "error": "name required"}
        if not prompt:
            return {"ok": False, "output": "", "error": "prompt required"}
        if len(prompt) > 8000:
            return {"ok": False, "output": "", "error": f"prompt too long ({len(prompt)} chars, max 8000)"}
        # Validate recurrence vs one-shot.
        if not run_at and not every_minutes:
            return {
                "ok": False, "output": "",
                "error": "must pass either run_at (ISO datetime) or every_minutes (int)",
            }
        if run_at and every_minutes:
            return {
                "ok": False, "output": "",
                "error": "pass only ONE of run_at / every_minutes, not both",
            }
        import datetime as _dt
        next_ts: float
        interval = None
        if run_at:
            try:
                dt = _dt.datetime.fromisoformat(run_at)
            except Exception:
                return {"ok": False, "output": "", "error": f"run_at must be ISO 8601, got {run_at!r}"}
            next_ts = dt.timestamp()
        else:
            n = int(every_minutes)
            if n < 1 or n > 60 * 24 * 30:
                return {"ok": False, "output": "", "error": f"every_minutes out of range (1..{60*24*30})"}
            interval = n * 60
            next_ts = _dt.datetime.now().timestamp() + interval
        task_id = _db.create_scheduled_task(
            name=name, prompt=prompt, next_run_at=next_ts,
            interval_seconds=interval, cwd=cwd or str(Path.cwd()),
        )
        when = _dt.datetime.fromtimestamp(next_ts).isoformat(timespec="seconds")
        return {
            "ok": True,
            "output": (
                f"scheduled {name!r} (id={task_id[:8]}) to run at {when}"
                + (f", repeating every {every_minutes} min" if every_minutes else "")
            ),
            "id": task_id,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def list_scheduled_tasks() -> dict:
    """Return every pending scheduled task (name, next_run_at, interval)."""
    try:
        from . import db as _db
        import datetime as _dt
        rows = _db.list_scheduled_tasks()
        if not rows:
            return {"ok": True, "output": "no scheduled tasks."}
        out = []
        for r in rows:
            when = _dt.datetime.fromtimestamp(r["next_run_at"]).isoformat(timespec="seconds")
            interval = f", every {r['interval_seconds'] // 60}min" if r.get("interval_seconds") else ""
            out.append(f"- {r['id'][:8]}: {r['name']!r} at {when}{interval}")
        return {"ok": True, "output": "\n".join(out), "tasks": rows}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def cancel_scheduled_task(id: str) -> dict:
    """Delete a scheduled task by (prefix of its) id."""
    try:
        from . import db as _db
        if not id or not id.strip():
            return {"ok": False, "output": "", "error": "id required"}
        removed = _db.cancel_scheduled_task(id.strip())
        if not removed:
            return {"ok": False, "output": "", "error": f"no scheduled task with id matching {id!r}"}
        return {"ok": True, "output": f"cancelled {removed} task(s)"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# schedule_wakeup — self-paced agent re-entry
#
# Unlike schedule_task, this inserts a one-shot row that RESUMES the current
# conversation (appends `note` as the next user turn) instead of starting a
# new chat. Use cases: "check this build in 4 min", "poke again if no user
# reply by X", background-polling patterns where the agent wants to come back
# into the same thread.
#
# Bounds:
#   delay_seconds: 60 .. 3600 (clamped)  — <60s fights the poll loop cadence
#                                          >1h crosses into "use schedule_task"
#   note:          <= 2000 chars          — a nudge, not a new task
# ---------------------------------------------------------------------------
_WAKEUP_MIN = 60
_WAKEUP_MAX = 3600
_WAKEUP_NOTE_MAX = 2000


async def schedule_wakeup(
    delay_seconds: int,
    note: str,
    conversation_id: str | None = None,
    cwd: str | None = None,
) -> dict:
    """Schedule yourself to wake up and continue THIS conversation after
    `delay_seconds` seconds. `note` becomes the user-turn text the daemon
    appends when it fires (use it to remind yourself what to do).

    `conversation_id` is injected by the dispatch layer — the model does not
    set it directly, and the tool rejects calls where it's missing (i.e. from
    a subagent loop that has no parent chat to re-enter).
    """
    try:
        from . import db as _db
        if not conversation_id:
            return {
                "ok": False, "output": "",
                "error": "schedule_wakeup can only be used inside a conversation",
            }
        try:
            n = int(delay_seconds)
        except Exception:
            return {"ok": False, "output": "", "error": "delay_seconds must be an integer"}
        if n < _WAKEUP_MIN or n > _WAKEUP_MAX:
            return {
                "ok": False, "output": "",
                "error": f"delay_seconds out of range ({_WAKEUP_MIN}..{_WAKEUP_MAX})",
            }
        note_clean = (note or "").strip()
        if not note_clean:
            return {"ok": False, "output": "", "error": "note is required — what should you do when you wake up?"}
        if len(note_clean) > _WAKEUP_NOTE_MAX:
            note_clean = note_clean[:_WAKEUP_NOTE_MAX] + "\n\n[note truncated]"
        import datetime as _dt
        next_ts = _dt.datetime.now().timestamp() + n
        task_id = _db.create_scheduled_task(
            name=f"wakeup: {note_clean[:40]}",
            prompt=note_clean,
            next_run_at=next_ts,
            interval_seconds=None,
            cwd=cwd or str(Path.cwd()),
            target_conversation_id=conversation_id,
        )
        when = _dt.datetime.fromtimestamp(next_ts).isoformat(timespec="seconds")
        return {
            "ok": True,
            "output": f"scheduled wakeup in {n}s (at {when}, id={task_id[:8]})",
            "id": task_id,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# start_loop / stop_loop — autonomous loop mode
#
# Same primitive as schedule_wakeup but recurring. `start_loop` inserts a row
# with kind='loop' and interval_seconds set, so the daemon re-fires the prompt
# into the same conversation at every interval until the agent (or user) calls
# stop_loop. One active loop per conversation — calling start_loop again
# reschedules instead of stacking loops (avoids a runaway double-fire).
#
# Bounds:
#   interval_seconds: 60 .. 3600  — same range as schedule_wakeup's delay;
#                                   anything tighter fights the 30s poll loop.
#   goal:             <= 4000 chars — a rolling instruction, not a full prompt.
# ---------------------------------------------------------------------------
_LOOP_MIN = 60
_LOOP_MAX = 3600
_LOOP_GOAL_MAX = 4000


async def start_loop(
    goal: str,
    interval_seconds: int = 300,
    conversation_id: str | None = None,
    cwd: str | None = None,
) -> dict:
    """Start (or reset) an autonomous loop on THIS conversation.

    Every `interval_seconds` the daemon will re-append `goal` as a user turn
    and trigger another agent pass. Use this to self-drive progress on a
    rolling objective — "keep checking the build", "keep refining the draft
    until I say stop". Call stop_loop (or the user clicks Stop loop in the
    UI) to end it. Calling start_loop while a loop is already active replaces
    the existing loop — use this to adjust the goal or interval mid-flight.
    """
    try:
        from . import db as _db
        if not conversation_id:
            return {
                "ok": False, "output": "",
                "error": "start_loop can only be used inside a conversation",
            }
        try:
            n = int(interval_seconds)
        except Exception:
            return {"ok": False, "output": "", "error": "interval_seconds must be an integer"}
        if n < _LOOP_MIN or n > _LOOP_MAX:
            return {
                "ok": False, "output": "",
                "error": f"interval_seconds out of range ({_LOOP_MIN}..{_LOOP_MAX})",
            }
        goal_clean = (goal or "").strip()
        if not goal_clean:
            return {"ok": False, "output": "", "error": "goal is required — what should the loop pursue?"}
        if len(goal_clean) > _LOOP_GOAL_MAX:
            goal_clean = goal_clean[:_LOOP_GOAL_MAX] + "\n\n[goal truncated]"
        # Replace any existing loop first so start_loop is idempotent — the
        # agent can safely call it on every pass without stacking ticks.
        _db.cancel_loops_for_conversation(conversation_id)
        import datetime as _dt
        next_ts = _dt.datetime.now().timestamp() + n
        task_id = _db.create_scheduled_task(
            name=f"loop: {goal_clean[:40]}",
            prompt=goal_clean,
            next_run_at=next_ts,
            interval_seconds=n,
            cwd=cwd or str(Path.cwd()),
            target_conversation_id=conversation_id,
            kind="loop",
        )
        return {
            "ok": True,
            "output": f"loop started, next tick in {n}s (id={task_id[:8]})",
            "id": task_id,
            "interval_seconds": n,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def stop_loop(conversation_id: str | None = None) -> dict:
    """Stop the autonomous loop on THIS conversation, if any is active.

    Idempotent — calling stop_loop when no loop is running is a no-op and
    returns ok=True with `stopped=0`. Use this once the rolling goal has been
    satisfied or the user explicitly wants out.
    """
    try:
        from . import db as _db
        if not conversation_id:
            return {
                "ok": False, "output": "",
                "error": "stop_loop can only be used inside a conversation",
            }
        n = _db.cancel_loops_for_conversation(conversation_id)
        return {
            "ok": True,
            "output": f"loop stopped (removed {n})" if n else "no loop was running",
            "stopped": n,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# spawn_task — flag a drive-by issue as a chip without derailing the turn
#
# This writes a row to `side_tasks` and emits a side_task_flagged SSE event
# (the agent loop picks that up from the tool result and forwards it). The
# frontend renders pending side tasks as chips under the assistant bubble;
# clicking Open calls POST /api/side-tasks/<id>/open which spins a new
# conversation. Clicking Dismiss calls the dismiss endpoint. Both transition
# the row to a terminal state.
# ---------------------------------------------------------------------------
_SPAWN_TITLE_MAX = 120
_SPAWN_TLDR_MAX = 400
_SPAWN_PROMPT_MAX = 8000


async def spawn_task(
    title: str,
    prompt: str,
    tldr: str = "",
    conversation_id: str | None = None,
) -> dict:
    """Flag an out-of-scope observation as a side task the user can open
    later without derailing the current turn.

    Use when you notice something worth fixing but it would bloat the current
    change: a stale README line, a dead config option, a missing test, a
    TODO you confirmed is real. The `prompt` must be self-contained — the
    spawned session has no memory of this conversation.
    """
    try:
        from . import db as _db
        if not conversation_id:
            return {
                "ok": False, "output": "",
                "error": "spawn_task can only be used inside a conversation",
            }
        t = (title or "").strip()
        p = (prompt or "").strip()
        l = (tldr or "").strip()
        if not t:
            return {"ok": False, "output": "", "error": "title is required"}
        if len(t) > _SPAWN_TITLE_MAX:
            return {"ok": False, "output": "", "error": f"title too long (max {_SPAWN_TITLE_MAX})"}
        if not p:
            return {"ok": False, "output": "", "error": "prompt is required — the spawned session has no context from this one"}
        if len(p) > _SPAWN_PROMPT_MAX:
            return {"ok": False, "output": "", "error": f"prompt too long (max {_SPAWN_PROMPT_MAX})"}
        if l and len(l) > _SPAWN_TLDR_MAX:
            l = l[:_SPAWN_TLDR_MAX]
        row = _db.create_side_task(
            source_conversation_id=conversation_id,
            title=t, prompt=p, tldr=l or None,
        )
        return {
            "ok": True,
            "output": f"flagged side task {row['id'][:8]}: {t}",
            "side_task": row,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# ask_user_question — structured multi-choice inline prompt
#
# Returns the chosen option as the tool result. The agent loop emits a
# special `await_user_answer` SSE event instead of running a handler — the
# frontend renders the buttons, the user clicks one, the backend resumes
# the future with the chosen value. See agent.py for the resume hook; this
# function body is never actually executed because dispatch() short-circuits
# the tool name (like `delegate`).
# ---------------------------------------------------------------------------
_ASK_MAX_OPTIONS = 6
_ASK_OPTION_LABEL_MAX = 80
_ASK_QUESTION_MAX = 500


async def ask_user_question(question: str, options: list) -> dict:
    """Stub — never called directly. dispatch() intercepts `ask_user_question`
    and handles it via the agent loop so the streaming SSE event can pause
    the turn until the user clicks one of the rendered buttons. The returned
    value from dispatch is a dict `{"ok": True, "output": "<chosen value>"}`.
    """
    return {
        "ok": False, "output": "",
        "error": "ask_user_question must be dispatched by the agent loop",
    }


def _validate_ask_question_args(question: str, options: list) -> str | None:
    q = (question or "").strip()
    if not q:
        return "question is required"
    if len(q) > _ASK_QUESTION_MAX:
        return f"question too long (max {_ASK_QUESTION_MAX})"
    if not isinstance(options, list) or not options:
        return "options must be a non-empty list"
    if len(options) > _ASK_MAX_OPTIONS:
        return f"too many options (max {_ASK_MAX_OPTIONS})"
    seen = set()
    for o in options:
        if not isinstance(o, dict):
            return "each option must be an object with `label` and `value`"
        label = (o.get("label") or "").strip()
        value = (o.get("value") or "").strip()
        if not label or not value:
            return "each option needs a non-empty `label` and `value`"
        if len(label) > _ASK_OPTION_LABEL_MAX:
            return f"option label too long (max {_ASK_OPTION_LABEL_MAX})"
        if value in seen:
            return f"duplicate option value: {value!r}"
        seen.add(value)
    return None


# ---------------------------------------------------------------------------
# Git worktree tools — create a throwaway worktree for risky edits
#
# These spawn `git worktree add` on the conversation's cwd, branch off HEAD
# (or whatever base_ref the caller chose), and return the new worktree's
# absolute path. The agent can then ask the user to switch cwd into it, do
# its work, and either merge the branch or drop the worktree.
#
# Safety rails:
#   * repo must be inside the conversation's cwd
#   * branch name is validated against a conservative regex
#   * base_ref is validated against the same regex (no shell metachars)
#   * worktree path lives under `<cwd>/.worktrees/<short-id>/` — we create
#     it ourselves, so a malicious branch name can't escape via a crafted
#     path argument.
# ---------------------------------------------------------------------------
# Branch and ref names are passed unquoted to `git worktree add` as argv
# elements. The regex locks them to shell-safe characters AND requires a
# non-dash first character, so a base_ref like "-fsomething" can never be
# interpreted by git as an option. Length caps bound the argv tail.
_GIT_REF_RE = re.compile(r"^[A-Za-z0-9._/][A-Za-z0-9._/\-]{0,199}$")
_BRANCH_NAME_RE = re.compile(r"^[A-Za-z0-9._/][A-Za-z0-9._/\-]{0,99}$")


def _is_git_repo(path: Path) -> bool:
    """True if `path` is inside a git repo (has a .git dir or file at root)."""
    try:
        p = path.resolve()
    except Exception:
        return False
    for parent in [p, *p.parents]:
        if (parent / ".git").exists():
            return True
    return False


async def create_worktree(
    branch: str,
    base_ref: str = "HEAD",
    conversation_id: str | None = None,
    cwd: str | None = None,
) -> dict:
    """Create a throwaway git worktree off `base_ref` on a new `branch`.

    Returns `{ok, output, worktree: {id, path, branch, base_ref}}`. The path
    is absolute and lives under `<cwd>/.worktrees/<short-id>/`. The caller
    (or the user) can then switch the conversation's cwd into that path to
    do risky work in isolation. Run `remove_worktree(<id>)` when done.
    """
    try:
        from . import db as _db
        if not conversation_id:
            return {"ok": False, "output": "", "error": "conversation_id required (internal)"}
        if not _BRANCH_NAME_RE.match(branch or ""):
            return {"ok": False, "output": "", "error": f"invalid branch name: {branch!r}"}
        if not _GIT_REF_RE.match(base_ref or ""):
            return {"ok": False, "output": "", "error": f"invalid base_ref: {base_ref!r}"}
        repo = Path(cwd or ".").resolve()
        if not _is_git_repo(repo):
            return {"ok": False, "output": "", "error": f"not a git repository: {repo}"}
        # Walk up to find the actual repo root (where .git lives) so the
        # worktree-list directory sits at the top level of the repo, not
        # inside some arbitrary subdirectory.
        repo_root = repo
        for parent in [repo, *repo.parents]:
            if (parent / ".git").exists():
                repo_root = parent
                break
        wt_parent = repo_root / ".worktrees"
        wt_parent.mkdir(exist_ok=True)
        short_id = uuid.uuid4().hex[:8]
        wt_path = wt_parent / short_id
        # `--` terminates option parsing so subsequent args are treated as
        # positional even if they happen to start with '-' (defense in depth
        # on top of the regex guard). `-b` consumes `branch` as its value so
        # it doesn't need to come after `--`.
        cmd = [
            "git", "-C", str(repo_root),
            "worktree", "add",
            "-b", branch,
            "--",
            str(wt_path), base_ref,
        ]
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        try:
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=60)
        except asyncio.TimeoutError:
            try:
                proc.kill()
            except Exception:
                pass
            return {"ok": False, "output": "", "error": "git worktree add timed out"}
        if proc.returncode != 0:
            return {
                "ok": False, "output": "",
                "error": f"git worktree add failed: {stderr.decode('utf-8', 'replace').strip()}",
            }
        row = _db.create_worktree_row(
            conversation_id=conversation_id,
            repo_path=str(repo_root),
            path=str(wt_path),
            branch=branch,
            base_ref=base_ref,
        )
        return {
            "ok": True,
            "output": (
                f"created worktree at {wt_path} on branch {branch!r}\n"
                f"switch the conversation's cwd into it to work in isolation."
            ),
            "worktree": row,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def list_worktrees(conversation_id: str | None = None) -> dict:
    """List git worktrees this conversation created (active + removed)."""
    try:
        from . import db as _db
        if not conversation_id:
            return {"ok": False, "output": "", "error": "conversation_id required (internal)"}
        rows = _db.list_worktrees_for_conversation(conversation_id)
        if not rows:
            return {"ok": True, "output": "no worktrees."}
        out = []
        for r in rows:
            badge = "✓" if r["status"] == "active" else "✗"
            out.append(f"- {badge} {r['id'][:8]} {r['branch']!r} @ {r['path']}")
        return {"ok": True, "output": "\n".join(out), "worktrees": rows}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def remove_worktree(id: str, conversation_id: str | None = None) -> dict:
    """Drop a worktree previously created via create_worktree.

    Runs `git worktree remove --force <path>` then marks the DB row as
    removed. `--force` is safe here because the worktree is a disposable
    branch we created ourselves; the caller is expected to have merged (or
    to not care about) any changes inside it.
    """
    try:
        from . import db as _db
        if not conversation_id:
            return {"ok": False, "output": "", "error": "conversation_id required (internal)"}
        # Allow matching by 8-char short id for convenience.
        rows = _db.list_active_worktrees_for_conversation(conversation_id)
        match = None
        for r in rows:
            if r["id"] == id or r["id"].startswith(id):
                match = r
                break
        if not match:
            return {"ok": False, "output": "", "error": f"no active worktree with id matching {id!r}"}
        cmd = [
            "git", "-C", match["repo_path"],
            "worktree", "remove", "--force", match["path"],
        ]
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        try:
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=60)
        except asyncio.TimeoutError:
            try:
                proc.kill()
            except Exception:
                pass
            return {"ok": False, "output": "", "error": "git worktree remove timed out"}
        if proc.returncode != 0:
            # Still mark the row removed — the worktree directory may be gone
            # already (manual rm), or the branch name may conflict. Either
            # way, the DB row should not linger as "active".
            _db.mark_worktree_removed(match["id"])
            return {
                "ok": False, "output": "",
                "error": f"git worktree remove failed: {stderr.decode('utf-8', 'replace').strip()}",
            }
        _db.mark_worktree_removed(match["id"])
        return {"ok": True, "output": f"removed worktree {match['id'][:8]} at {match['path']}"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Monitor tool — poll a target until a condition flips.
#
# This is the "wait until X happens" primitive. Three kinds of target are
# supported:
#   - file:<path>       — watch a path on disk
#   - url:<http(s)://…> — watch an HTTP endpoint (SSRF-guarded, same rules as
#                         fetch_url)
#   - bash:<command>    — watch the exit code / stdout of a shell command
#
# Conditions are a small DSL evaluated every tick:
#   exists              — file/URL is reachable (HTTP 2xx for URLs, path
#                         exists for files, exit-0 for bash)
#   missing             — inverse of exists (404/ENOENT/non-zero exit)
#   contains:<text>     — substring appears in the read content
#   not_contains:<text> — substring is ABSENT from the read content
#   changed             — content/mtime/status differs from the first tick
#   status:<int>        — HTTP status equals the given code (URL only)
#   exit_code:<int>     — bash exit code equals the given int (bash only)
#   regex:<pattern>     — Python regex matches the read content
#
# The tool blocks up to `timeout_seconds` and returns as soon as the condition
# holds. One "timed out" result when the deadline passes without the condition
# flipping. Output always includes the number of ticks polled so the agent
# can see how long the wait actually took.
# ---------------------------------------------------------------------------
_MONITOR_MIN_INTERVAL = 1
_MONITOR_MAX_INTERVAL = 60
_MONITOR_MIN_TIMEOUT = 1
_MONITOR_MAX_TIMEOUT = 60 * 30   # 30 minutes; matches other long-running bounds
_MONITOR_MAX_PREVIEW = 500       # chars of target content echoed back


def _monitor_parse_target(target: str) -> tuple[str, str] | tuple[None, str]:
    """Split a `kind:value` target string. Returns (kind, value) or (None, error)."""
    if not target or ":" not in target:
        return None, (
            "target must be 'file:<path>', 'url:<http(s)://...>', or 'bash:<command>'"
        )
    kind, _, value = target.partition(":")
    kind = kind.strip().lower()
    value = value.strip()
    if kind not in {"file", "url", "bash"}:
        return None, f"unknown target kind {kind!r}; use file / url / bash"
    if not value:
        return None, f"empty {kind} target"
    return kind, value


async def _monitor_tick_file(cwd: str, value: str) -> dict:
    """One observation of a file target. Never raises."""
    try:
        p = _resolve(cwd, value)
        if not p.exists():
            return {"ok": True, "exists": False, "content": "", "status": None, "exit": None, "tag": "missing"}
        # Use mtime + size as a cheap "changed?" fingerprint rather than
        # reading the whole file every tick. For the `contains`/`regex`
        # predicates we still need the content, so read a bounded slice.
        st = p.stat()
        content = ""
        if p.is_file() and st.st_size > 0:
            try:
                with p.open("rb") as f:
                    raw = f.read(200_000)
                content = raw.decode("utf-8", errors="replace")
            except Exception:
                content = ""
        tag = f"mtime={st.st_mtime:.3f},size={st.st_size}"
        return {"ok": True, "exists": True, "content": content, "status": None, "exit": None, "tag": tag}
    except Exception as e:
        return {"ok": False, "exists": False, "content": "", "status": None, "exit": None, "tag": f"error:{e}"}


async def _monitor_tick_url(value: str) -> dict:
    """One observation of a URL target. Applies the same SSRF guard as fetch_url."""
    ok, reason = _is_safe_url(value)
    if not ok:
        return {"ok": False, "exists": False, "content": "", "status": None, "exit": None, "tag": f"blocked:{reason}"}
    parsed = urlparse(value)
    if parsed.hostname and not _resolves_to_public_ip(parsed.hostname):
        return {"ok": False, "exists": False, "content": "", "status": None, "exit": None, "tag": "blocked:private IP"}
    try:
        # GET (not HEAD) so `contains` / `regex` predicates have body text.
        # Keep the response bounded — we don't need the whole page.
        async with httpx.AsyncClient(timeout=10.0, follow_redirects=True) as c:
            r = await c.get(value)
            body = r.text[:200_000] if r.text else ""
            status = r.status_code
        exists = 200 <= status < 300
        return {"ok": True, "exists": exists, "content": body, "status": status, "exit": None, "tag": f"status={status}"}
    except Exception as e:
        return {"ok": True, "exists": False, "content": "", "status": None, "exit": None, "tag": f"error:{type(e).__name__}"}


async def _monitor_tick_bash(cwd: str, value: str) -> dict:
    """One observation of a bash target. Runs the command with a short timeout."""
    # Reuse the existing run_bash plumbing with a tight timeout so a slow
    # command can't starve the whole polling loop.
    res = await run_bash(cwd, value, timeout=30)
    content = res.get("output", "")
    exit_code = res.get("exit_code")
    exists = exit_code == 0
    return {"ok": True, "exists": exists, "content": content, "status": None, "exit": exit_code, "tag": f"exit={exit_code}"}


def _monitor_check_condition(condition: str, obs: dict, first_tag: str) -> tuple[bool, str]:
    """Return (met, reason). `condition` is one of the DSL strings above."""
    cond = (condition or "").strip()
    if not cond:
        return False, ""
    # Bare predicates first.
    if cond == "exists":
        return (bool(obs.get("exists")), "target reachable" if obs.get("exists") else "")
    if cond == "missing":
        return (not obs.get("exists"), "target unreachable" if not obs.get("exists") else "")
    if cond == "changed":
        met = obs.get("tag") and obs.get("tag") != first_tag
        return (bool(met), f"target changed ({first_tag!r} → {obs.get('tag')!r})" if met else "")
    # Parametric predicates: "<op>:<arg>".
    if ":" not in cond:
        return False, f"unknown condition {cond!r}"
    op, _, arg = cond.partition(":")
    op = op.strip().lower()
    arg = arg  # preserve whitespace — important for `contains` on formatted output
    content = obs.get("content") or ""
    if op == "contains":
        return (arg in content, f"content contains {arg!r}" if arg in content else "")
    if op == "not_contains":
        absent = arg not in content
        return (absent, f"content does NOT contain {arg!r}" if absent else "")
    if op == "regex":
        try:
            if re.search(arg, content):
                return True, f"regex {arg!r} matched"
            return False, ""
        except re.error as e:
            return False, f"bad regex: {e}"
    if op == "status":
        try:
            target_status = int(arg)
        except ValueError:
            return False, f"status arg must be int, got {arg!r}"
        got = obs.get("status")
        return (got == target_status, f"HTTP status {got}" if got == target_status else "")
    if op == "exit_code":
        try:
            target_exit = int(arg)
        except ValueError:
            return False, f"exit_code arg must be int, got {arg!r}"
        got = obs.get("exit")
        return (got == target_exit, f"exit code {got}" if got == target_exit else "")
    return False, f"unknown condition operator {op!r}"


async def monitor(
    target: str,
    condition: str,
    interval_seconds: int = 5,
    timeout_seconds: int = 120,
    cwd: str = ".",
) -> dict:
    """Poll `target` on an interval until `condition` is met or `timeout_seconds` elapses.

    Returns early as soon as the condition holds. Designed for "wait until a
    build output file appears", "poll a health-check URL until it returns 200",
    "re-run `pytest` until it passes", and similar "watch for state change"
    patterns that would otherwise require a shell loop.

    All URL probes reuse the same SSRF guard as fetch_url so an attacker-controlled
    string can't be used to scan the LAN. Bash probes run with a 30 s per-tick
    cap so a hung command doesn't stall the loop.
    """
    kind_or_none, value_or_reason = _monitor_parse_target(target)
    if kind_or_none is None:
        return {"ok": False, "output": "", "error": value_or_reason}
    kind: str = kind_or_none
    value: str = value_or_reason

    if not condition or not condition.strip():
        return {"ok": False, "output": "", "error": "condition required"}

    # Clamp interval + timeout so a buggy call can't burn CPU forever.
    interval = max(_MONITOR_MIN_INTERVAL, min(int(interval_seconds or 5), _MONITOR_MAX_INTERVAL))
    timeout = max(_MONITOR_MIN_TIMEOUT, min(int(timeout_seconds or 120), _MONITOR_MAX_TIMEOUT))

    # First observation — used as the baseline for the `changed` predicate.
    # Done outside the loop so "condition already holds" returns immediately.
    if kind == "file":
        first = await _monitor_tick_file(cwd, value)
    elif kind == "url":
        first = await _monitor_tick_url(value)
    else:
        first = await _monitor_tick_bash(cwd, value)
    first_tag = first.get("tag") or ""
    ticks = 1
    start = time.time()

    met, reason = _monitor_check_condition(condition, first, first_tag)
    if met:
        preview = (first.get("content") or "")[:_MONITOR_MAX_PREVIEW]
        return {
            "ok": True,
            "output": (
                f"condition met on first check ({reason}). "
                f"ticks=1, elapsed=0.0s.\n\n"
                + (f"target preview:\n{preview}" if preview else "")
            ),
            "ticks": ticks,
            "elapsed_seconds": 0.0,
            "final_status": first.get("status"),
            "final_exit_code": first.get("exit"),
        }

    deadline = start + timeout
    while time.time() < deadline:
        # Sleep BEFORE the next tick so we poll at ~interval cadence rather
        # than hammering the target back-to-back.
        await asyncio.sleep(interval)
        ticks += 1
        if kind == "file":
            obs = await _monitor_tick_file(cwd, value)
        elif kind == "url":
            obs = await _monitor_tick_url(value)
        else:
            obs = await _monitor_tick_bash(cwd, value)
        met, reason = _monitor_check_condition(condition, obs, first_tag)
        if met:
            elapsed = time.time() - start
            preview = (obs.get("content") or "")[:_MONITOR_MAX_PREVIEW]
            return {
                "ok": True,
                "output": (
                    f"condition met ({reason}) after {elapsed:.1f}s / {ticks} ticks.\n\n"
                    + (f"target preview:\n{preview}" if preview else "")
                ),
                "ticks": ticks,
                "elapsed_seconds": elapsed,
                "final_status": obs.get("status"),
                "final_exit_code": obs.get("exit"),
            }

    # Fell through the deadline without the condition flipping.
    elapsed = time.time() - start
    return {
        "ok": False,
        "output": (
            f"timed out after {elapsed:.1f}s / {ticks} ticks without the "
            f"condition {condition!r} holding on target {target!r}."
        ),
        "error": "timeout",
        "ticks": ticks,
        "elapsed_seconds": elapsed,
    }


# ---------------------------------------------------------------------------
# Local document search — `doc_index` / `doc_search`
#
# Walks a directory, chunks each file, embeds every chunk via Ollama's
# `/api/embeddings` endpoint, and stores vectors in SQLite. `doc_search`
# embeds the query and returns top-k chunks by cosine similarity — scored
# in Python (no vector-ext needed). Model choice is configurable but we
# default to `nomic-embed-text`, which is a common small-footprint option
# available through Ollama.
# ---------------------------------------------------------------------------
_DEFAULT_EMBED_MODEL = "nomic-embed-text"
_DOC_INDEX_CHUNK_CHARS = 1200
_DOC_INDEX_CHUNK_OVERLAP = 200


def _chunk_text(text: str, size: int, overlap: int) -> list[str]:
    """Split `text` into overlapping windows of `size` chars.

    Overlap guards against a concept straddling the boundary — cosine search
    still finds it if any one chunk catches enough context. Both ends of the
    window land on non-whitespace where possible.
    """
    out: list[str] = []
    i = 0
    n = len(text)
    step = max(1, size - overlap)
    while i < n:
        chunk = text[i:i + size]
        if chunk.strip():
            out.append(chunk)
        i += step
    return out


async def _ollama_embed(client, text: str, model: str) -> list[float]:
    """Call Ollama's /api/embeddings and return the vector."""
    r = await client.post(
        "http://127.0.0.1:11434/api/embeddings",
        json={"model": model, "prompt": text},
    )
    r.raise_for_status()
    data = r.json()
    vec = data.get("embedding")
    if not vec:
        raise RuntimeError(f"ollama /api/embeddings returned no vector: {data}")
    return vec


async def doc_index(
    path: str,
    extensions: list[str] | None = None,
    model: str | None = None,
) -> dict:
    """Walk `path`, chunk every matching file, embed each chunk, and store
    the vectors in the `doc_chunks` table. Existing rows for the same file
    are replaced, so re-indexing the same directory is idempotent.

    `extensions` restricts which files get indexed (by lowercase suffix,
    e.g. ['.py', '.md']); default indexes common text formats.
    """
    try:
        import httpx
        from . import db as _db
        m = (model or _DEFAULT_EMBED_MODEL).strip()
        p = Path(path).expanduser()
        if not p.is_absolute():
            p = Path.cwd() / p
        if not p.exists():
            return {"ok": False, "output": "", "error": f"path not found: {p}"}
        exts = set((e or "").lower().strip() for e in (extensions or []) if e)
        if not exts:
            exts = {".md", ".txt", ".py", ".js", ".jsx", ".ts", ".tsx",
                    ".java", ".c", ".cpp", ".h", ".hpp", ".rs", ".go",
                    ".json", ".yaml", ".yml", ".html", ".css", ".scss"}
        # Collect candidate files. Cap the walk so we don't accidentally
        # embed a gigabyte of node_modules.
        files: list[Path] = []
        MAX_FILES = 500
        for f in p.rglob("*"):
            if len(files) >= MAX_FILES:
                break
            if not f.is_file():
                continue
            if f.suffix.lower() not in exts:
                continue
            # Skip obvious noise.
            if any(part in ("node_modules", ".git", "__pycache__", ".venv", "venv", "dist", "build")
                   for part in f.parts):
                continue
            try:
                if f.stat().st_size > 2 * 1024 * 1024:  # 2 MB ceiling
                    continue
            except Exception:
                continue
            files.append(f)
        if not files:
            return {"ok": True, "output": f"indexed 0 files under {p} (no matching files)"}
        chunks_total = 0
        async with httpx.AsyncClient(timeout=60.0) as client:
            for f in files:
                try:
                    text = f.read_text(encoding="utf-8", errors="replace")
                except Exception:
                    continue
                if not text.strip():
                    continue
                chunks = _chunk_text(text, _DOC_INDEX_CHUNK_CHARS, _DOC_INDEX_CHUNK_OVERLAP)
                # Ditch any previous rows for this file BEFORE embedding so
                # a partial failure leaves us in a clean state on retry.
                _db.delete_doc_chunks_for(str(f))
                for ci, chunk in enumerate(chunks):
                    try:
                        vec = await _ollama_embed(client, chunk, m)
                    except Exception:
                        continue
                    _db.insert_doc_chunk(
                        path=str(f), ordinal=ci, text=chunk, vector=vec, model=m,
                    )
                    chunks_total += 1
        return {
            "ok": True,
            "output": f"indexed {chunks_total} chunk(s) across {len(files)} file(s) under {p} (model={m})",
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def doc_search(
    query: str,
    top_k: int = 5,
    path_glob: str | None = None,
    model: str | None = None,
) -> dict:
    """Embed `query` via Ollama and return the top `top_k` indexed chunks
    by cosine similarity. `path_glob` optionally filters to rows whose
    path matches a substring/glob-like pattern (simple substring match).
    """
    try:
        import httpx
        from . import db as _db
        q = (query or "").strip()
        if not q:
            return {"ok": False, "output": "", "error": "query required"}
        k = max(1, min(int(top_k or 5), 30))
        m = (model or _DEFAULT_EMBED_MODEL).strip()
        async with httpx.AsyncClient(timeout=30.0) as client:
            qvec = await _ollama_embed(client, q, m)
        # Pull every candidate row and score in Python. At ~thousands of
        # chunks this is instant; for much larger indexes we'd want a
        # proper vector DB, but that's overkill for the 4B-model use case.
        rows = _db.all_doc_chunks(path_substr=(path_glob or None))
        if not rows:
            return {"ok": True, "output": "index is empty — call `doc_index` first."}
        # Cosine similarity. Using raw Python dot-products is fine at this
        # scale and avoids a numpy import.
        def cos(a, b):
            import math
            s = sum(x * y for x, y in zip(a, b))
            na = math.sqrt(sum(x * x for x in a)) or 1.0
            nb = math.sqrt(sum(x * x for x in b)) or 1.0
            return s / (na * nb)
        scored = []
        for r in rows:
            try:
                score = cos(qvec, r["vector"])
            except Exception:
                continue
            scored.append((score, r))
        scored.sort(key=lambda x: x[0], reverse=True)
        top = scored[:k]
        lines: list[str] = []
        for score, r in top:
            snippet = (r["text"] or "").strip().replace("\n", " ")[:200]
            lines.append(f"[{score:.3f}] {r['path']} #{r['ordinal']}: {snippet}")
        return {
            "ok": True,
            "output": (
                f"top {len(top)}/{len(rows)} chunks for {q!r}:\n" + "\n".join(lines)
                if top else "no matches"
            ),
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Codebase index — per-cwd semantic index that answers "where is X defined?"
#
# Reuses the doc_chunks store under the hood: every chunk's `path` is the
# absolute file path so we can both (a) track per-cwd rows via the
# `codebase_indexes` registry and (b) scope a search by path prefix. The
# indexer is gitignore-aware when the cwd is a git repo (via `git ls-files`);
# otherwise it falls back to the same noise-skip list doc_index uses.
# ---------------------------------------------------------------------------
_CODEBASE_INDEX_EXTS = {
    ".md", ".txt", ".py", ".js", ".jsx", ".ts", ".tsx",
    ".java", ".c", ".cpp", ".h", ".hpp", ".rs", ".go", ".rb",
    ".json", ".yaml", ".yml", ".html", ".css", ".scss",
    ".toml", ".ini", ".sh", ".ps1", ".php", ".swift", ".kt",
}
_CODEBASE_INDEX_MAX_FILES = 1500
_CODEBASE_INDEX_MAX_FILE_BYTES = 2 * 1024 * 1024


def _codebase_list_files(root: Path) -> list[Path]:
    """Return the files to index under `root`.

    Two strategies:
      1. If `root` is a git repo, shell out to `git ls-files -co --exclude-standard`
         — this is the same list the user sees in `git status`, so the agent
         naturally skips build artifacts, node_modules, secrets, etc.
      2. Otherwise, rglob + a hand-coded noise blacklist (same one used by
         doc_index). Less accurate but safe.
    Both paths honor the file-size ceiling and the allowed-extensions set so
    a runaway binary can't poison the index.
    """
    import subprocess
    files: list[Path] = []
    git_dir = root / ".git"
    if git_dir.exists():
        try:
            # -c = cached, -o = other/untracked, --exclude-standard = respect
            # .gitignore + global + info/exclude. -z gives NUL-separated output
            # so paths with newlines (rare but possible) still parse.
            proc = subprocess.run(
                ["git", "ls-files", "-co", "--exclude-standard", "-z"],
                cwd=str(root),
                capture_output=True,
                timeout=30,
                check=False,
            )
            if proc.returncode == 0:
                for rel in proc.stdout.split(b"\x00"):
                    if not rel:
                        continue
                    try:
                        rel_str = rel.decode("utf-8", errors="replace")
                    except Exception:
                        continue
                    p = root / rel_str
                    if not p.is_file():
                        continue
                    if p.suffix.lower() not in _CODEBASE_INDEX_EXTS:
                        continue
                    try:
                        if p.stat().st_size > _CODEBASE_INDEX_MAX_FILE_BYTES:
                            continue
                    except Exception:
                        continue
                    files.append(p)
                    if len(files) >= _CODEBASE_INDEX_MAX_FILES:
                        break
                return files
        except Exception:
            # git missing / timed out — fall through to the rglob path
            pass
    # Non-git fallback.
    NOISE = {"node_modules", ".git", "__pycache__", ".venv", "venv",
             "dist", "build", ".next", ".nuxt", "target", "out"}
    for f in root.rglob("*"):
        if len(files) >= _CODEBASE_INDEX_MAX_FILES:
            break
        if not f.is_file():
            continue
        if f.suffix.lower() not in _CODEBASE_INDEX_EXTS:
            continue
        if any(part in NOISE for part in f.parts):
            continue
        try:
            if f.stat().st_size > _CODEBASE_INDEX_MAX_FILE_BYTES:
                continue
        except Exception:
            continue
        files.append(f)
    return files


async def _codebase_index_cwd_impl(cwd: str, model: str | None = None) -> dict:
    """Walk `cwd`, chunk + embed every matching file, update the registry.

    Safe to call concurrently with the same cwd — the 'indexing' status in
    the registry is advisory only; duplicate work wastes CPU but never
    corrupts data (each file's chunks are deleted before re-insert).
    """
    import httpx
    from . import db as _db
    m = (model or _DEFAULT_EMBED_MODEL).strip()
    root = Path(cwd).expanduser().resolve()
    if not root.exists() or not root.is_dir():
        _db.upsert_codebase_index(
            str(root), status="error", error=f"cwd not found: {root}",
        )
        return {"ok": False, "error": f"cwd not found: {root}"}
    _db.upsert_codebase_index(str(root), status="indexing")
    try:
        files = _codebase_list_files(root)
        chunks_total = 0
        async with httpx.AsyncClient(timeout=60.0) as client:
            for f in files:
                try:
                    text = f.read_text(encoding="utf-8", errors="replace")
                except Exception:
                    continue
                if not text.strip():
                    continue
                chunks = _chunk_text(
                    text, _DOC_INDEX_CHUNK_CHARS, _DOC_INDEX_CHUNK_OVERLAP,
                )
                _db.delete_doc_chunks_for(str(f))
                for ci, chunk in enumerate(chunks):
                    try:
                        vec = await _ollama_embed(client, chunk, m)
                    except Exception:
                        continue
                    _db.insert_doc_chunk(
                        path=str(f), ordinal=ci, text=chunk, vector=vec, model=m,
                    )
                    chunks_total += 1
        import time as _time
        _db.upsert_codebase_index(
            str(root),
            status="ready",
            file_count=len(files),
            chunk_count=chunks_total,
            last_indexed_at=_time.time(),
        )
        return {"ok": True, "files": len(files), "chunks": chunks_total}
    except Exception as e:
        _db.upsert_codebase_index(
            str(root), status="error", error=f"{type(e).__name__}: {e}",
        )
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}


async def codebase_search(
    query: str,
    top_k: int = 8,
    conversation_id: str | None = None,
    cwd: str | None = None,
) -> dict:
    """Answer "where is X defined / used?" over the current conversation's cwd.

    This is a thin wrapper around the doc-chunk store: we embed `query` and
    return the top-k chunks whose `path` starts with the active cwd. The
    index is maintained automatically when a conversation's cwd is set (see
    app.py); if it isn't ready yet, we return a hint so the agent waits or
    falls back to grep/read.
    """
    try:
        import httpx
        from . import db as _db
        q = (query or "").strip()
        if not q:
            return {"ok": False, "output": "", "error": "query required"}
        target_cwd = (cwd or "").strip() or str(Path.cwd())
        root = str(Path(target_cwd).expanduser().resolve())
        idx = _db.get_codebase_index(root)
        if not idx or idx["status"] in ("pending", "indexing"):
            return {
                "ok": True,
                "output": (
                    f"codebase index for {root} is {idx['status'] if idx else 'not started'}; "
                    "no results yet — try again in a moment or fall back to grep / read."
                ),
                "status": idx["status"] if idx else "none",
            }
        if idx["status"] == "error":
            return {
                "ok": False, "output": "",
                "error": f"codebase index errored: {idx.get('error') or 'unknown'}",
            }
        k = max(1, min(int(top_k or 8), 30))
        m = _DEFAULT_EMBED_MODEL
        async with httpx.AsyncClient(timeout=30.0) as client:
            qvec = await _ollama_embed(client, q, m)
        # Scope the candidate rows to files under the cwd (both separators —
        # Windows paths use backslashes, POSIX forward slashes).
        prefix = root.rstrip("/\\")
        # all_doc_chunks takes a substring filter; use the cwd as the filter
        # and additionally post-filter on the prefix separator to avoid
        # sibling-directory false positives.
        rows = _db.all_doc_chunks(path_substr=prefix)
        def under_root(p: str) -> bool:
            return p.startswith(prefix + "/") or p.startswith(prefix + "\\")
        rows = [r for r in rows if under_root(r["path"])]
        if not rows:
            return {"ok": True, "output": f"index is empty for {root}"}
        def cos(a, b):
            import math
            s = sum(x * y for x, y in zip(a, b))
            na = math.sqrt(sum(x * x for x in a)) or 1.0
            nb = math.sqrt(sum(x * x for x in b)) or 1.0
            return s / (na * nb)
        scored: list[tuple[float, dict]] = []
        for r in rows:
            try:
                scored.append((cos(qvec, r["vector"]), r))
            except Exception:
                continue
        scored.sort(key=lambda x: x[0], reverse=True)
        top = scored[:k]
        lines: list[str] = []
        for score, r in top:
            snippet = (r["text"] or "").strip().replace("\n", " ")[:240]
            # Relative paths are easier on the eye; fall back to absolute if
            # the chunk somehow sits outside the cwd (shouldn't happen after
            # the prefix filter, but defensive coding costs nothing).
            try:
                rel = str(Path(r["path"]).relative_to(root))
            except Exception:
                rel = r["path"]
            lines.append(f"[{score:.3f}] {rel} #{r['ordinal']}: {snippet}")
        return {
            "ok": True,
            "output": (
                f"top {len(top)}/{len(rows)} chunks for {q!r} under {root}:\n"
                + "\n".join(lines)
                if top else "no matches"
            ),
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Docs URL crawler — per-seed semantic index of a public documentation site.
#
# Same vector store as doc_index / codebase_search (doc_chunks), but every
# chunk's `path` is prefixed with `url:` so we can filter the index down to
# only URL-sourced material without a separate table. The crawler is
# deliberately minimal: breadth-first, same-origin by default, capped by
# page count. That's enough for "point it at a docs site and let the agent
# ground answers on it" without the surface area of a real spider.
# ---------------------------------------------------------------------------
_DOCS_URL_FETCH_TIMEOUT = 30.0
_DOCS_URL_MAX_BYTES = 2_000_000
_DOCS_URL_PAGE_CAP = 100         # hard upper bound regardless of user request
_DOCS_URL_EMBED_TIMEOUT = 60.0
# Skip link extensions that are clearly not HTML docs — saves us fetching
# every PDF/image the docs site links to. Keep this list conservative; we
# only skip *obvious* binaries so we don't accidentally miss a docs page.
_DOCS_URL_SKIP_SUFFIXES = {
    ".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp", ".ico",
    ".pdf", ".zip", ".tar", ".gz", ".rar", ".7z",
    ".mp3", ".mp4", ".wav", ".mov", ".avi", ".webm",
    ".css", ".js", ".map", ".woff", ".woff2", ".ttf", ".eot",
    ".exe", ".dmg", ".bin",
}
# Track seeds currently being crawled so a second kickoff is a no-op.
_DOCS_URL_INFLIGHT: set[str] = set()


def _extract_links_sync(html: str, base_url: str) -> list[str]:
    """Pull absolute URLs out of every <a href="..."> in `html`.

    Uses stdlib html.parser (no lxml dep) and urljoin to resolve relative
    refs against `base_url`. Drops fragments (`#foo`) and obviously-binary
    suffixes. Duplicates are preserved in-order — the caller is expected
    to dedupe via a `seen` set during BFS.
    """
    from html.parser import HTMLParser
    from urllib.parse import urljoin, urldefrag

    found: list[str] = []

    class _Collector(HTMLParser):
        def handle_starttag(self, tag: str, attrs: list) -> None:
            if tag.lower() != "a":
                return
            for k, v in attrs:
                if k and k.lower() == "href" and v:
                    try:
                        absu = urljoin(base_url, v)
                        absu, _frag = urldefrag(absu)
                    except Exception:
                        continue
                    found.append(absu)
                    break

    try:
        _Collector(convert_charrefs=True).feed(html)
    except Exception:
        # Malformed HTML — return whatever we got before the parser tripped.
        pass
    return found


def _same_origin(a: str, b: str) -> bool:
    """True if `a` and `b` share scheme + host (+ port)."""
    try:
        ua, ub = urlparse(a), urlparse(b)
    except Exception:
        return False
    return (
        ua.scheme == ub.scheme
        and (ua.hostname or "").lower() == (ub.hostname or "").lower()
        and ua.port == ub.port
    )


async def _fetch_page_for_crawl(url: str) -> tuple[str | None, str, str]:
    """Fetch `url` and return (text, final_url, raw_html).

    Returns (None, url, "") on any failure (blocked URL, HTTP error, no
    extractable content). Uses the same SSRF guard as `fetch_url` — crawls
    of private-IP docs sites are refused, same as any other fetch.
    """
    ok, _reason = _is_safe_url(url)
    if not ok:
        return None, url, ""
    hostname = urlparse(url).hostname or ""
    try:
        ipaddress.ip_address(hostname)
    except ValueError:
        resolves_public = await asyncio.to_thread(_resolves_to_public_ip, hostname)
        if not resolves_public:
            return None, url, ""
    headers = {
        "User-Agent": FETCH_USER_AGENT,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.8",
    }
    try:
        async with httpx.AsyncClient(
            timeout=_DOCS_URL_FETCH_TIMEOUT,
            follow_redirects=True,
            headers=headers,
        ) as client:
            async with client.stream("GET", url) as r:
                r.raise_for_status()
                ctype = r.headers.get("content-type", "").lower()
                # Only HTML pages get crawled for outbound links; we'd parse
                # JSON/PDF with different machinery and that's out of scope
                # here.
                chunks: list[bytes] = []
                total = 0
                async for chunk in r.aiter_bytes():
                    total += len(chunk)
                    if total > _DOCS_URL_MAX_BYTES:
                        break
                    chunks.append(chunk)
                raw = b"".join(chunks)
                final_url = str(r.url)
    except Exception:
        return None, url, ""

    if "charset=" in ctype:
        enc = ctype.split("charset=", 1)[1].split(";", 1)[0].strip() or "utf-8"
    else:
        enc = "utf-8"
    try:
        html_body = raw.decode(enc, errors="replace")
    except LookupError:
        html_body = raw.decode("utf-8", errors="replace")

    if not any(t in ctype for t in ("html", "xml")):
        # Non-HTML — embed raw text directly, no link extraction.
        return html_body, final_url, ""

    extracted = await asyncio.to_thread(_trafilatura_extract_sync, html_body, final_url)
    return extracted or None, final_url, html_body


async def _docs_url_crawl_impl(did: str, model: str | None = None) -> dict:
    """BFS-crawl a registered URL seed, chunk + embed every page's prose,
    and write the result into doc_chunks under `path = "url:<page-url>"`.

    State machine (stored on the doc_urls row): pending → crawling →
    ready | error. The `_DOCS_URL_INFLIGHT` set guards against a second
    concurrent kickoff for the same seed; on exit we always drop the id
    from the set, including on error.
    """
    from . import db as _db
    row = _db.get_doc_url(did)
    if not row:
        return {"ok": False, "error": f"doc url {did} not found"}
    seed = (row["url"] or "").strip()
    if not seed:
        _db.update_doc_url(did, status="error", error="empty seed url")
        return {"ok": False, "error": "empty seed url"}

    _DOCS_URL_INFLIGHT.add(did)
    _db.update_doc_url(did, status="crawling", error="")
    # Drop stale chunks from a previous crawl so a re-index starts clean.
    _db.delete_doc_chunks_for_prefix(f"{_db.DOC_URL_CHUNK_PREFIX}{seed.rstrip('/')}")

    m = (model or _DEFAULT_EMBED_MODEL).strip()
    max_pages = max(1, min(int(row.get("max_pages") or 20), _DOCS_URL_PAGE_CAP))
    same_origin_only = bool(row.get("same_origin_only"))

    seen: set[str] = set()
    queue: list[str] = [seed]
    seen.add(seed)
    pages_crawled = 0
    chunks_total = 0
    title_stored = row.get("title")

    try:
        async with httpx.AsyncClient(timeout=_DOCS_URL_EMBED_TIMEOUT) as embed_client:
            while queue and pages_crawled < max_pages:
                url = queue.pop(0)
                text, final_url, html_body = await _fetch_page_for_crawl(url)
                pages_crawled += 1
                if not text or not text.strip():
                    continue

                # Capture a title from the seed's HTML the first time we see
                # one, so the Settings UI has something friendlier than the
                # raw URL. Stop trying after the first hit.
                if not title_stored and html_body:
                    mt = re.search(
                        r"<title[^>]*>(.*?)</title>",
                        html_body,
                        flags=re.IGNORECASE | re.DOTALL,
                    )
                    if mt:
                        t = re.sub(r"\s+", " ", mt.group(1)).strip()[:200]
                        if t:
                            title_stored = t
                            _db.update_doc_url(did, title=t)

                chunks = _chunk_text(
                    text, _DOC_INDEX_CHUNK_CHARS, _DOC_INDEX_CHUNK_OVERLAP,
                )
                chunk_path = f"{_db.DOC_URL_CHUNK_PREFIX}{final_url}"
                for ci, chunk in enumerate(chunks):
                    try:
                        vec = await _ollama_embed(embed_client, chunk, m)
                    except Exception:
                        continue
                    _db.insert_doc_chunk(
                        path=chunk_path,
                        ordinal=ci,
                        text=chunk,
                        vector=vec,
                        model=m,
                    )
                    chunks_total += 1

                # Progress ping so the UI can show a rising count without
                # blocking on the full crawl.
                _db.update_doc_url(
                    did,
                    pages_crawled=pages_crawled,
                    chunk_count=chunks_total,
                )

                # Enqueue outbound links for the next BFS step — only if we
                # actually got HTML back (non-HTML pages have html_body="").
                if html_body and len(queue) + pages_crawled < max_pages:
                    try:
                        links = await asyncio.to_thread(
                            _extract_links_sync, html_body, final_url,
                        )
                    except Exception:
                        links = []
                    for link in links:
                        if link in seen:
                            continue
                        ok, _reason = _is_safe_url(link)
                        if not ok:
                            continue
                        if same_origin_only and not _same_origin(seed, link):
                            continue
                        parsed = urlparse(link)
                        if any(
                            (parsed.path or "").lower().endswith(sfx)
                            for sfx in _DOCS_URL_SKIP_SUFFIXES
                        ):
                            continue
                        seen.add(link)
                        queue.append(link)
                        if len(seen) >= max_pages * 3:
                            # Seen-set cap keeps memory bounded on a sprawling
                            # site — we still only crawl up to `max_pages`.
                            break

        import time as _time
        _db.update_doc_url(
            did,
            status="ready",
            pages_crawled=pages_crawled,
            chunk_count=chunks_total,
            last_indexed_at=_time.time(),
            error="",
        )
        return {
            "ok": True,
            "pages": pages_crawled,
            "chunks": chunks_total,
            "seed": seed,
        }
    except Exception as e:
        _db.update_doc_url(
            did,
            status="error",
            error=f"{type(e).__name__}: {e}",
        )
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}
    finally:
        _DOCS_URL_INFLIGHT.discard(did)


async def docs_search(query: str, top_k: int = 5, url_prefix: str | None = None) -> dict:
    """Semantic search across every indexed documentation URL.

    Same vector math as doc_search / codebase_search, but scoped to rows
    whose path starts with the URL prefix marker (`url:`). Pass
    `url_prefix` to narrow the search to a specific site — e.g.
    "https://docs.python.org/" — otherwise every indexed URL is in scope.
    """
    try:
        from . import db as _db
        q = (query or "").strip()
        if not q:
            return {"ok": False, "output": "", "error": "query required"}
        k = max(1, min(int(top_k or 5), 30))
        m = _DEFAULT_EMBED_MODEL
        async with httpx.AsyncClient(timeout=30.0) as client:
            qvec = await _ollama_embed(client, q, m)
        # Pull all URL-sourced chunks. If the caller passed a prefix we
        # further narrow with a substring filter; otherwise every chunk
        # whose path begins with "url:" is a candidate.
        substr = _db.DOC_URL_CHUNK_PREFIX
        if url_prefix:
            substr = f"{_db.DOC_URL_CHUNK_PREFIX}{url_prefix}"
        rows = _db.all_doc_chunks(path_substr=substr)
        # Post-filter: path_substr is a LIKE match, so we still verify the
        # prefix position to be safe against accidental substring matches.
        rows = [r for r in rows if (r["path"] or "").startswith(substr)]
        if not rows:
            return {
                "ok": True,
                "output": (
                    "no indexed docs yet — add a URL in Settings → Docs "
                    "and wait for crawl to finish."
                ),
            }
        def cos(a, b):
            import math
            s = sum(x * y for x, y in zip(a, b))
            na = math.sqrt(sum(x * x for x in a)) or 1.0
            nb = math.sqrt(sum(x * x for x in b)) or 1.0
            return s / (na * nb)
        scored: list[tuple[float, dict]] = []
        for r in rows:
            try:
                scored.append((cos(qvec, r["vector"]), r))
            except Exception:
                continue
        scored.sort(key=lambda x: x[0], reverse=True)
        top = scored[:k]
        lines: list[str] = []
        for score, r in top:
            snippet = (r["text"] or "").strip().replace("\n", " ")[:240]
            # Strip the "url:" marker for display; the model doesn't need it.
            display = (r["path"] or "").removeprefix(_db.DOC_URL_CHUNK_PREFIX)
            lines.append(f"[{score:.3f}] {display} #{r['ordinal']}: {snippet}")
        return {
            "ok": True,
            "output": (
                f"top {len(top)}/{len(rows)} docs chunks for {q!r}:\n"
                + "\n".join(lines)
                if top else "no matches"
            ),
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


# ---------------------------------------------------------------------------
# Planning tool — `todo_write`
#
# The tool itself is trivial: the agent just echoes the passed list back.
# The value is UX: the agent loop detects todo_write calls and emits a
# `todos_updated` SSE event so the frontend can render a pinned task panel.
# This gives small models a place to externalize multi-step plans instead
# of letting them fall out of context.
# ---------------------------------------------------------------------------
async def todo_write(todos: list[dict]) -> dict:
    if not isinstance(todos, list):
        return {"ok": False, "output": "", "error": "todos must be a list"}
    cleaned: list[dict] = []
    valid_statuses = {"pending", "in_progress", "completed"}
    for i, t in enumerate(todos):
        if not isinstance(t, dict):
            return {"ok": False, "output": "", "error": f"todo[{i}] must be an object"}
        content = str(t.get("content") or "").strip()
        active = str(t.get("activeForm") or t.get("active_form") or content).strip()
        status = str(t.get("status") or "pending").strip().lower()
        if not content:
            return {"ok": False, "output": "", "error": f"todo[{i}].content is required"}
        if status not in valid_statuses:
            status = "pending"
        cleaned.append({"content": content, "activeForm": active, "status": status})
    # Allow at most one in_progress.
    in_prog = [i for i, t in enumerate(cleaned) if t["status"] == "in_progress"]
    if len(in_prog) > 1:
        # Keep only the first; demote the rest to pending.
        for idx in in_prog[1:]:
            cleaned[idx]["status"] = "pending"
    lines = []
    for t in cleaned:
        marker = {"pending": "[ ]", "in_progress": "[~]", "completed": "[x]"}[t["status"]]
        lines.append(f"  {marker} {t['content']}")
    return {"ok": True, "output": "Todo list updated:\n" + "\n".join(lines), "todos": cleaned}


# ---------------------------------------------------------------------------
# Long-term memory tools
#
# `remember` and `forget` give the model a way to stash facts that need to
# outlive the rolling context window: user preferences, project conventions,
# gotchas discovered during the session, paths it keeps having to rediscover.
# The file is injected into the system prompt on every turn, so anything
# written here survives auto-compaction and server restarts.
#
# Stored as simple markdown at data/memory/<conv_id>.md — human-readable,
# easy to edit by hand, and trivially portable across machines.
# ---------------------------------------------------------------------------
def _memory_path(conv_id: str | None) -> Path | None:
    """Resolve the memory file path for one conversation, or None if no conv."""
    if not conv_id:
        return None
    return MEMORY_DIR / f"{conv_id}.md"


def _trim_memory(text: str) -> str:
    """Keep the memory file under MEMORY_MAX_CHARS by dropping oldest lines."""
    if len(text) <= MEMORY_MAX_CHARS:
        return text
    # Work from the tail: keep as many trailing lines as fit in the budget.
    lines = text.splitlines()
    kept: list[str] = []
    total = 0
    for line in reversed(lines):
        if total + len(line) + 1 > MEMORY_MAX_CHARS:
            break
        kept.append(line)
        total += len(line) + 1
    kept.reverse()
    return "# Memory (older entries trimmed)\n\n" + "\n".join(kept)


async def remember(
    conv_id: str | None,
    content: str,
    topic: str | None = None,
    scope: str = "conversation",
) -> dict:
    """Save a fact to long-term memory.

    Three scopes — pick by how widely this fact applies:

      - `conversation` (default) — THIS chat only. Per-chat scratchpad
        at data/memory/<conv_id>.md. Best for in-flight decisions
        ("we decided to use approach X for this refactor", "the user
        wants the dark variant of the icon"). Vanishes if the chat
        is deleted; never bleeds into other chats.
      - `project` — every chat working in the same directory (cwd).
        Two conversations pointed at the same repo automatically
        share this set; no user-side configuration required. Best
        for project-wide rules the AGENTS.md file doesn't already
        cover ("the lint config is in .eslintrc.cjs at root", "API
        tokens for staging are in 1Password entry X"). Stored in
        the `project_memories` SQLite table keyed by the
        conversation's cwd path.
      - `global` — every conversation, every project, forever.
        Stored in `global_memories`. Best for durable user-wide
        facts ("user prefers SCSS over CSS", "user is on Windows +
        Git Bash, not Linux").

    Pick the NARROWEST scope that fits — global memory is injected
    into every system prompt the user ever runs, so noise here
    bloats every conversation.

    `topic` is optional — for conv scope it groups bullets under a
    `## topic` heading; for project / global scope it's stored in
    its own column for the Settings UI to filter by.
    """
    scope = (scope or "conversation").strip().lower()
    if scope not in {"conversation", "project", "global"}:
        return {
            "ok": False,
            "output": "",
            "error": "scope must be 'conversation', 'project', or 'global'",
        }
    if not content or not content.strip():
        return {"ok": False, "output": "", "error": "content is required"}

    # ----- Global path: write to SQLite, no file I/O ----------------------
    if scope == "global":
        try:
            row = db.add_global_memory(content, topic)
            return {
                "ok": True,
                "output": f"remembered globally: {row['content'][:120]}",
            }
        except ValueError as e:
            return {"ok": False, "output": "", "error": str(e)}
        except Exception as e:
            return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}

    # ----- Project path: keyed by the conversation's cwd ------------------
    if scope == "project":
        if not conv_id:
            return {
                "ok": False, "output": "",
                "error": "remember(scope='project') needs a conversation context",
            }
        try:
            conv = db.get_conversation(conv_id)
        except Exception as e:
            return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}
        if not conv:
            return {
                "ok": False, "output": "",
                "error": "conversation not found",
            }
        cwd = (conv.get("cwd") or "").strip()
        if not cwd:
            # Defensive — every conversation should have a cwd, but if
            # one slips through (legacy / mid-migration row), refuse
            # cleanly rather than silently writing under "".
            return {
                "ok": False, "output": "",
                "error": (
                    "this conversation has no working directory set — "
                    "use scope='conversation' to keep the fact local"
                ),
            }
        try:
            row = db.add_project_memory(cwd, content, topic)
            return {
                "ok": True,
                "output": (
                    f"remembered for project at {row['cwd']}: "
                    f"{row['content'][:120]}"
                ),
            }
        except ValueError as e:
            return {"ok": False, "output": "", "error": str(e)}
        except Exception as e:
            return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}

    # ----- Conversation-scoped path (existing file-backed behaviour) -----
    if not conv_id:
        return {"ok": False, "output": "", "error": "remember needs a conversation context"}
    try:
        path = _memory_path(conv_id)
        assert path is not None  # for type narrowing — already checked above
        path.parent.mkdir(parents=True, exist_ok=True)
        existing = path.read_text(encoding="utf-8") if path.is_file() else ""
        bullet = f"- {content.strip()}"
        t = (topic or "").strip()
        if t:
            heading = f"## {t}"
            if heading in existing:
                # Insert the new bullet right below the heading.
                lines = existing.splitlines()
                out: list[str] = []
                inserted = False
                for line in lines:
                    out.append(line)
                    if not inserted and line.strip() == heading:
                        out.append(bullet)
                        inserted = True
                new_text = "\n".join(out)
            else:
                new_text = existing.rstrip() + f"\n\n{heading}\n{bullet}\n"
        else:
            new_text = existing.rstrip() + f"\n{bullet}\n"
        new_text = _trim_memory(new_text)
        # Bytes, not text — same reason as write_file/edit_file: avoid
        # Windows CRLF rewriting and keep memory files byte-stable.
        path.write_bytes(new_text.encode("utf-8"))
        return {
            "ok": True,
            "output": f"remembered: {content.strip()[:120]}",
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


async def forget(
    conv_id: str | None,
    pattern: str,
    scope: str = "conversation",
) -> dict:
    """Remove memory entries containing `pattern` (case-insensitive substring).

    Same three scopes as `remember`: 'conversation' prunes lines from
    the per-conv markdown file; 'project' deletes matching rows tied
    to the conversation's project label; 'global' deletes matching
    rows from the SQLite `global_memories` table.
    """
    scope = (scope or "conversation").strip().lower()
    if scope not in {"conversation", "project", "global"}:
        return {
            "ok": False,
            "output": "",
            "error": "scope must be 'conversation', 'project', or 'global'",
        }
    if not pattern or not pattern.strip():
        return {"ok": False, "output": "", "error": "pattern is required"}

    # ----- Global path: SQL delete by LIKE --------------------------------
    if scope == "global":
        try:
            n = db.delete_global_memories_matching(pattern)
            return {"ok": True, "output": f"forgot {n} global memor{'y' if n == 1 else 'ies'} matching {pattern!r}"}
        except ValueError as e:
            return {"ok": False, "output": "", "error": str(e)}
        except Exception as e:
            return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}

    # ----- Project path: SQL delete scoped to this conv's cwd -------------
    if scope == "project":
        if not conv_id:
            return {
                "ok": False, "output": "",
                "error": "forget(scope='project') needs a conversation context",
            }
        try:
            conv = db.get_conversation(conv_id)
        except Exception as e:
            return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}
        if not conv:
            return {"ok": False, "output": "", "error": "conversation not found"}
        cwd = (conv.get("cwd") or "").strip()
        if not cwd:
            return {
                "ok": False, "output": "",
                "error": "this conversation has no working directory set",
            }
        try:
            n = db.delete_project_memories_matching(cwd, pattern)
            return {
                "ok": True,
                "output": (
                    f"forgot {n} project memor{'y' if n == 1 else 'ies'} "
                    f"for cwd {cwd!r} matching {pattern!r}"
                ),
            }
        except ValueError as e:
            return {"ok": False, "output": "", "error": str(e)}
        except Exception as e:
            return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}

    # ----- Conversation-scoped path (existing) ----------------------------
    if not conv_id:
        return {"ok": False, "output": "", "error": "forget needs a conversation context"}
    try:
        path = _memory_path(conv_id)
        assert path is not None
        if not path.is_file():
            return {"ok": True, "output": "(no memory file to prune)"}
        needle = pattern.strip().lower()
        lines = path.read_text(encoding="utf-8").splitlines()
        kept = [ln for ln in lines if needle not in ln.lower()]
        removed = len(lines) - len(kept)
        # Bytes to keep line endings stable on Windows (see write_file).
        new_text = "\n".join(kept) + ("\n" if kept else "")
        path.write_bytes(new_text.encode("utf-8"))
        return {"ok": True, "output": f"forgot {removed} line(s) matching {pattern!r}"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def load_memory_for_prompt(conv_id: str | None) -> str:
    """Return the memory file wrapped in a system-prompt-friendly section.

    Called by prompts.build_system_prompt. Safe to call when the file doesn't
    exist — returns "" in that case.
    """
    path = _memory_path(conv_id)
    if not path or not path.is_file():
        return ""
    try:
        text = path.read_text(encoding="utf-8").strip()
    except Exception:
        return ""
    if not text:
        return ""
    return (
        "\n\n## Your long-term memory for this conversation\n\n"
        "The notes below were saved earlier via the `remember` tool. Treat "
        "them as authoritative context about the user, this project, and "
        "the work in flight. Use `remember(...)` to add new facts and "
        "`forget(...)` to remove ones that are no longer true.\n\n"
        + text
        + "\n"
    )


def load_project_memory_for_prompt(cwd: str | None) -> str:
    """Return the project memory rows for this cwd, wrapped in a
    system-prompt section.

    Called by `prompts.build_system_prompt` for every conversation —
    `cwd` is mandatory on conversations, so the section is keyed
    automatically with no user-side configuration. Returns "" when
    the cwd has no associated memory rows yet.

    Memories are grouped by `topic` for readability — same shape as
    `load_global_memory_for_prompt` so the prompt visually distinguishes
    project from global only by section heading.
    """
    key = (cwd or "").strip()
    if not key:
        return ""
    try:
        rows = db.list_project_memories(key)
    except Exception:
        # DB hiccup — skip the section quietly. A missing memory block
        # is much better than a system prompt that fails to assemble.
        return ""
    if not rows:
        return ""
    groups: dict[str, list[dict]] = {}
    for r in rows:
        groups.setdefault(r["topic"] or "General", []).append(r)
    body_parts: list[str] = []
    for topic, items in groups.items():
        body_parts.append(f"### {topic}")
        for it in items:
            content = " ".join(it["content"].split())
            body_parts.append(f"- {content}")
        body_parts.append("")
    return (
        f"\n\n## Project memory — applies to every chat in this directory\n\n"
        "The notes below are conventions and facts the user (or a "
        "previous turn) has saved for THIS working directory — every "
        "conversation pointed at the same cwd sees them. Treat as "
        "authoritative for project-wide rules. To add or remove "
        "project-scoped notes, use `remember(scope=\"project\", ...)` and "
        "`forget(scope=\"project\", ...)`.\n\n"
        + "\n".join(body_parts).rstrip()
        + "\n"
    )


def load_global_memory_for_prompt() -> str:
    """Return the global memory rows wrapped in a system-prompt section.

    Called by prompts.build_system_prompt for every conversation (including
    subagents — global memories are durable user preferences, so they apply
    everywhere). Safe to call when the table is empty — returns "".

    Memories are grouped by `topic` for readability. Untopiced entries are
    rendered under a generic "General" heading so the prompt stays well
    structured even when the user hasn't bothered to categorise.
    """
    try:
        rows = db.list_global_memories()
    except Exception:
        # DB not initialized yet (rare — startup race) or some other I/O
        # failure. Fail silently: a missing memory section is much better
        # than a broken system prompt that prevents the agent from running.
        return ""
    if not rows:
        return ""
    # Group by topic, preserving creation order within each group.
    groups: dict[str, list[dict]] = {}
    for r in rows:
        groups.setdefault(r["topic"] or "General", []).append(r)
    body_parts: list[str] = []
    for topic, items in groups.items():
        body_parts.append(f"### {topic}")
        for it in items:
            # Single-line bullets — collapse internal newlines so the prompt
            # markdown stays clean.
            content = " ".join(it["content"].split())
            body_parts.append(f"- {content}")
        body_parts.append("")
    return (
        "\n\n## Global long-term memory (applies to every conversation)\n\n"
        "The notes below are durable user preferences, project conventions, "
        "and facts the user has curated in Settings → Memories. Treat them "
        "as authoritative — they apply to all your work, not just this one "
        "conversation. To add or remove globally-scoped notes, use "
        "`remember(scope=\"global\", ...)` and `forget(scope=\"global\", ...)`.\n\n"
        + "\n".join(body_parts).rstrip()
        + "\n"
    )


# ---------------------------------------------------------------------------
# Lifecycle hooks
#
# Runs user-defined shell commands at well-known points in the agent loop
# (pre-tool, post-tool, user_prompt_submit, turn_done). The command receives
# structured JSON on stdin so the hook can switch on tool name / args /
# result, and its stdout is captured and surfaced back to the agent as a
# system-level note on the next iteration.
#
# Hooks are powerful — they get a shell with the user's full privileges —
# so the UI treats creation like an approve-to-save operation, and we
# hard-cap timeout (1-120s) + never log stdin contents to disk.
# ---------------------------------------------------------------------------
# Single source of truth for valid event names lives in db.HOOK_EVENTS —
# importing it here as a module-level alias keeps `run_hooks(event=...)`
# and the DB's `create_hook(event=...)` validator in lockstep when new
# events are added (e.g. the `tool_error` / `consecutive_failures`
# workflow events). Earlier code shipped a stale local set that quietly
# rejected any event name added after this module was last touched.
from .db import HOOK_EVENTS as _DB_HOOK_EVENTS
HOOK_EVENTS = _DB_HOOK_EVENTS


async def _run_single_hook(hook: dict, payload: dict) -> dict:
    """Execute one hook command. Payload is serialised to JSON on stdin.

    Returns a dict with:
        ok (bool)     — exit code 0
        output (str)  — captured stdout (trimmed) the agent sees
        error (str)   — failure reason when ok=False
    """
    cmd = hook.get("command") or ""
    timeout = int(hook.get("timeout_seconds") or 10)
    stdin_blob = json.dumps(payload, ensure_ascii=False).encode("utf-8")

    # Use asyncio.create_subprocess_shell where available; fall back to a
    # thread-wrapped subprocess.run on Windows selector loops (same pattern
    # as run_bash_blocking above).
    try:
        proc = await asyncio.create_subprocess_shell(
            cmd,
            stdin=asyncio.subprocess.PIPE,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        try:
            stdout_b, stderr_b = await asyncio.wait_for(
                proc.communicate(stdin_blob), timeout=timeout
            )
        except asyncio.TimeoutError:
            # Kill the subprocess and drain its pipes so the Proactor event
            # loop doesn't warn about unclosed transports at shutdown.
            try:
                proc.kill()
                # communicate() after kill() closes pipes; shield with a
                # short wait_for so a truly stuck process can't wedge us.
                await asyncio.wait_for(proc.communicate(), timeout=2)
            except Exception:
                pass
            return {"ok": False, "output": "", "error": f"hook timed out after {timeout}s"}
        out = (stdout_b or b"").decode("utf-8", errors="replace")
        err = (stderr_b or b"").decode("utf-8", errors="replace")
        # Merge stderr into output so hooks that use `echo ... >&2` are still
        # surfaced. Keep output bounded so a chatty hook can't flood context.
        merged = out if not err else (out + ("\n" if out else "") + err)
        return {
            "ok": proc.returncode == 0,
            "output": _clip(merged.strip()),
            "exit_code": proc.returncode,
            "error": None if proc.returncode == 0 else f"exit code {proc.returncode}",
        }
    except NotImplementedError:
        # Selector loop — fall through to blocking subprocess in a thread.
        def _blocking() -> dict:
            try:
                r = subprocess.run(
                    cmd, shell=True, input=stdin_blob, capture_output=True, timeout=timeout
                )
                out = (r.stdout or b"").decode("utf-8", errors="replace")
                err = (r.stderr or b"").decode("utf-8", errors="replace")
                merged = out if not err else (out + ("\n" if out else "") + err)
                return {
                    "ok": r.returncode == 0,
                    "output": _clip(merged.strip()),
                    "exit_code": r.returncode,
                    "error": None if r.returncode == 0 else f"exit code {r.returncode}",
                }
            except subprocess.TimeoutExpired:
                return {"ok": False, "output": "", "error": f"hook timed out after {timeout}s"}
            except Exception as e:
                return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}
        return await asyncio.to_thread(_blocking)
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}


def _hook_matches(hook: dict, tool_name: str | None) -> bool:
    """Return True if this hook should fire for the given tool call.

    `matcher` is an optional case-insensitive substring applied to tool names.
    Empty/None matcher means "fire for every tool". Non-tool events
    (user_prompt_submit, turn_done) pass matcher="".
    """
    m = (hook.get("matcher") or "").strip().lower()
    if not m:
        return True
    if not tool_name:
        return True
    return m in tool_name.lower()


async def run_hooks(
    event: str,
    payload: dict,
    tool_name: str | None = None,
    conv_id: str | None = None,
) -> list[dict]:
    """Fire every enabled hook registered for `event`. Returns a list of
    `{hook_id, command, output, ok, error, capped}` dicts — one per hook
    that *would* have fired (capped ones are reported with `capped=True`
    and no command output, so the caller / UI can show "this hook hit its
    per-conversation cap" instead of silently dropping it).

    Used by agent.py at specific lifecycle points. Hook failures never
    propagate: they're captured into the return value so the agent can
    continue its turn regardless.

    `conv_id` is required for cap enforcement on `max_fires_per_conv`. If
    omitted (e.g. tests, ad-hoc calls outside the agent loop), the cap is
    silently bypassed.
    """
    if event not in HOOK_EVENTS:
        return []
    try:
        from . import db as _db
        hooks = _db.get_hooks_for_event(event)
    except Exception:
        return []
    results: list[dict] = []
    for h in hooks:
        if not _hook_matches(h, tool_name):
            continue
        # Per-conversation fire cap. We check BEFORE running so a maxed
        # hook doesn't pay the subprocess cost. Increment only AFTER a
        # successful start so a failed cap-check (e.g. DB read error)
        # doesn't double-bill.
        cap = h.get("max_fires_per_conv")
        if conv_id and cap is not None:
            try:
                seen = _db.get_hook_fire_count(h["id"], conv_id)
            except Exception:
                seen = 0
            if seen >= cap:
                results.append({
                    "hook_id": h["id"],
                    "command": h["command"],
                    "event": event,
                    "ok": False,
                    "output": "",
                    "error": (
                        f"hook hit its per-conversation cap of {cap} fires; "
                        "delete the hook or raise the cap to allow more."
                    ),
                    "capped": True,
                })
                continue
        res = await _run_single_hook(h, payload)
        if conv_id:
            try:
                _db.incr_hook_fire(h["id"], conv_id)
            except Exception:
                pass
        results.append({
            "hook_id": h["id"],
            "command": h["command"],
            "event": event,
            "ok": bool(res.get("ok")),
            "output": res.get("output", ""),
            "error": res.get("error"),
            "capped": False,
        })
    return results


# ---------------------------------------------------------------------------
# User-defined tools support (execution + validation)
#
# Users register custom Python tools through the Settings → Tools UI. Those
# tools live in SQLite (see db.create_user_tool) and execute in an isolated
# venv (see user_tools_runtime.execute_user_tool). The dispatcher below
# checks the user_tools table before the built-in registry lookup so a
# user-defined name gets routed to the runtime, not "unknown tool".
#
# The model itself CANNOT create, list, or delete user tools — only the user
# can via the settings panel. This is a deliberate safety boundary: giving
# an LLM the ability to mint new tools would let it extend its own
# privileges mid-conversation.
# ---------------------------------------------------------------------------
import ast as _ast  # stdlib; used to parse + validate user-submitted code
from . import user_tools_runtime as _utr


def _validate_user_tool_code(code: str) -> None:
    """Parse the code, reject non-Python, ensure ``def run(args)`` exists.

    Pre-flighting catches the most common "I gave the model a typo" failure
    modes *before* we spend a minute installing deps into the venv. We do
    NOT try to sandbox the AST (no import blocklist, etc.) — subprocess
    isolation + the permission-mode approval gate is the real defence.
    """
    try:
        tree = _ast.parse(code, mode="exec")
    except SyntaxError as e:
        raise ValueError(f"code has a syntax error at line {e.lineno}: {e.msg}") from e
    for node in tree.body:
        if isinstance(node, (_ast.FunctionDef, _ast.AsyncFunctionDef)) and node.name == "run":
            args = node.args
            total = len(args.args) + len(args.posonlyargs) + len(args.kwonlyargs)
            if total < 1 and args.vararg is None and args.kwarg is None:
                raise ValueError("run() must accept at least one argument (args: dict)")
            return
    raise ValueError("code must define `def run(args: dict) -> dict`")


async def create_tool(
    name: str,
    description: str,
    code: str,
    schema: dict | None = None,
    deps: list[str] | None = None,
    category: str = "write",
    timeout_seconds: int = 60,
) -> dict:
    """Register a new user-defined tool and install its dependencies.

    Invoked only from the Settings → Tools HTTP endpoint (``POST /api/user-tools``).
    This function is intentionally NOT exposed as an LLM tool — giving the model
    the ability to mint its own tools at runtime would let it extend its own
    privileges mid-conversation. Tool creation is a user-gated action.

    Execution order: (1) validate name isn't already a built-in or MCP name,
    (2) AST-check the code, (3) pip-install deps into the shared venv, (4)
    persist to SQLite. Installing before persisting means a bad package spec
    doesn't leave a half-registered tool in the palette.
    """
    if _utr.is_disabled():
        return {"ok": False, "output": "", "error": "user tools are disabled via GIGACHAT_DISABLE_USER_TOOLS"}
    clean_name = (name or "").strip().lower()
    if clean_name in TOOL_REGISTRY or clean_name in {"delegate", "delegate_parallel"}:
        return {
            "ok": False,
            "output": "",
            "error": f"name {clean_name!r} collides with a built-in tool; pick a different name",
        }
    if clean_name.startswith("mcp__"):
        return {
            "ok": False,
            "output": "",
            "error": "tool names starting with `mcp__` are reserved for MCP-backed tools",
        }
    try:
        _validate_user_tool_code(code or "")
    except ValueError as e:
        return {"ok": False, "output": "", "error": str(e)}
    try:
        cleaned_deps = _utr.validate_deps(deps or [])
    except ValueError as e:
        return {"ok": False, "output": "", "error": str(e)}
    install_log = ""
    if cleaned_deps:
        install_result = await _utr.install_deps(cleaned_deps)
        install_log = (install_result.get("output") or "")[-4000:]
        if not install_result.get("ok"):
            return {
                "ok": False,
                "output": install_log,
                "error": install_result.get("error") or "pip install failed",
            }
    try:
        row = db.create_user_tool(
            name=clean_name,
            description=description or "",
            code=code or "",
            schema=schema or {},
            deps=cleaned_deps,
            category=category or "write",
            timeout_seconds=int(timeout_seconds or 60),
            enabled=True,
        )
    except ValueError as e:
        return {"ok": False, "output": install_log, "error": str(e)}
    dep_summary = f" with deps {cleaned_deps}" if cleaned_deps else ""
    summary = (
        f"registered tool {clean_name!r}{dep_summary}. "
        "It's immediately callable in this conversation and every future one."
    )
    return {
        "ok": True,
        "output": _clip((install_log + "\n\n" + summary).strip()),
        "tool": {"name": row.get("name"), "id": row.get("id")},
    }


def user_tool_schemas() -> list[dict]:
    """Return Ollama-style function-calling schemas for every enabled user tool.

    Called by agent.py on every turn so the live set of user tools is merged
    into the palette alongside built-ins + MCP tools. Cheap: one SQLite
    SELECT per turn, same cost as ``mcp.tool_schemas_for_agent()``.
    """
    out: list[dict] = []
    try:
        rows = db.list_user_tools(enabled_only=True)
    except Exception:
        return out
    for r in rows:
        schema = r.get("schema") if isinstance(r.get("schema"), dict) else {}
        # Ollama wants a proper ``{"type": "object", "properties": {...}}``
        # block. Accept both "already-wrapped" and "bare-properties" shapes
        # from the model for forgiveness: if the stored schema has no
        # ``type`` key but looks like a properties dict, wrap it.
        if "type" not in schema and "properties" not in schema and schema:
            schema = {"type": "object", "properties": schema}
        schema.setdefault("type", "object")
        schema.setdefault("properties", {})
        # Reason field — every tool surface the model uses requires one; the
        # dispatcher strips it before calling the handler.
        props = dict(schema.get("properties") or {})
        props.setdefault(
            "reason",
            {
                "type": "string",
                "description": "One-sentence justification shown on the approval card.",
            },
        )
        schema["properties"] = props
        out.append({
            "type": "function",
            "function": {
                "name": r["name"],
                "description": r.get("description") or f"User-defined tool {r['name']}",
                "parameters": schema,
            },
        })
    return out


def classify_user_tool(name: str) -> str | None:
    """Return 'read' / 'write' for a user tool, or None if the name is unknown.

    Used by ``classify_tool`` so the permission-mode gate treats user tools
    the same way as built-ins. Falling back to ``None`` (rather than
    defaulting to 'write') lets the caller distinguish "this is a user tool,
    honor its stored category" from "this name matched nothing".
    """
    try:
        row = db.get_user_tool_by_name(name)
    except Exception:
        return None
    if not row:
        return None
    cat = (row.get("category") or "write").strip().lower()
    return cat if cat in {"read", "write"} else "write"


# ---------------------------------------------------------------------------
# Docker / sandboxed containers
#
# Goal: let the agent run ANY language or piece of software (Node, Rust, Go,
# ffmpeg, headless browsers, ML models, ...) inside an isolated container,
# without polluting the host or the shared Python venv used by user tools.
#
# Design choices:
#   * Drive the official ``docker`` CLI via subprocess (no Python SDK
#     dependency — `docker run` is universally available where Docker is
#     installed).
#   * Two runtime modes: foreground (``docker_run``) for short tasks and
#     background (``docker_run_bg``) for long-running services. They mirror
#     the ``bash`` / ``bash_bg`` ergonomics the agent already knows.
#   * Track every container we start in an in-memory registry keyed by the
#     container name we minted. ``docker_logs`` / ``docker_exec`` /
#     ``docker_stop`` will only act on containers in that registry — the
#     model can never accidentally kill a user-owned container.
#   * Sane defaults: ``--rm`` for foreground, ``--security-opt
#     no-new-privileges``. No ``--privileged``, no ``--cap-add``, no
#     ``--device`` are exposed. Memory and CPU are capped (--memory, --cpus).
#     Workspace is mounted READ-ONLY by default at /workspace so a runaway
#     container can't wreck host files. Network is "bridge" (outbound
#     allowed) so package installs work; toggle to "none" for hermetic runs.
#   * Image names are pattern-validated. Combined with argv-list dispatch
#     (no shell=True anywhere), this defeats command-injection attempts via
#     hostile image strings.
# ---------------------------------------------------------------------------

# Registry of containers Gigachat started. Maps the container *name* (the
# `--name` we passed) to a small record. Only entries here are addressable
# via docker_logs / docker_exec / docker_stop — calling those tools with a
# name we don't own is rejected so the agent can't poke at user-owned
# containers running on the same machine.
_DOCKER_CONTAINERS: dict[str, dict] = {}

# Container names we mint. Prefix lets the user identify them with
# ``docker ps`` and clean them up safely if the backend ever crashes mid-run.
_DOCKER_NAME_PREFIX = "gigachat_"

# Image-name validation. Docker accepts a wide alphabet but we restrict to
# the safe subset to defeat any "image name with spaces or `; rm -rf /`"
# tricks. This is purely defence-in-depth — we already pass docker args as
# a list (no shell), so injection is not actually possible at the host
# level either way.
_DOCKER_IMAGE_RE = re.compile(r"^[A-Za-z0-9._/:@-]+$")

# Network modes we'll accept from the model. Anything else is rejected so a
# typo can't accidentally land the container on a privileged network.
_DOCKER_NETWORK_MODES = {"none", "bridge", "host"}


def _docker_available() -> tuple[bool, str]:
    """Return ``(True, "")`` if `docker` is on PATH and the daemon responds.

    We probe with ``docker version`` (5s timeout) so a stopped Docker Desktop
    surfaces as a clear, actionable error to the model instead of a cryptic
    subprocess failure 60s later.
    """
    if shutil.which("docker") is None:
        return False, (
            "docker CLI not found on PATH. Install Docker Desktop "
            "(https://docs.docker.com/get-docker/) and ensure the daemon is running."
        )
    try:
        r = subprocess.run(
            ["docker", "version", "--format", "{{.Server.Version}}"],
            capture_output=True,
            timeout=5,
        )
        if r.returncode != 0:
            stderr = (r.stderr or b"").decode("utf-8", "replace").strip()
            return False, (
                "docker CLI is installed but the daemon isn't responding. "
                "Start Docker Desktop / `dockerd` and try again. "
                f"stderr: {stderr[:200]}"
            )
        return True, ""
    except Exception as e:
        return False, f"docker probe failed: {type(e).__name__}: {e}"


def _validate_docker_image(image: str) -> str | None:
    """Return None if the image name passes our allowlist, else a reason."""
    if not image or not image.strip():
        return "image is required"
    if len(image) > 256:
        return "image name too long (>256 chars)"
    if not _DOCKER_IMAGE_RE.match(image):
        return (
            "image name contains characters outside [A-Za-z0-9._/:@-]. "
            "Use a plain reference like 'python:3.12-slim' or "
            "'ghcr.io/owner/image:tag'."
        )
    return None


def _new_container_name(conv_id: str | None) -> str:
    """Mint a container name that's identifiable as ours and unique."""
    suffix = uuid.uuid4().hex[:10]
    if conv_id:
        # Docker container names accept [a-zA-Z0-9_.-]; sanitise everything else.
        safe = re.sub(r"[^A-Za-z0-9_.-]", "_", str(conv_id))[:24]
        return f"{_DOCKER_NAME_PREFIX}{safe}_{suffix}"
    return f"{_DOCKER_NAME_PREFIX}{suffix}"


def _build_run_argv(
    image: str,
    container_name: str,
    command: str | None,
    workdir: str | None,
    cwd: str,
    mount_workspace: bool,
    mount_mode: str,
    env: dict | None,
    network: str,
    memory: str,
    cpus: str,
    detach: bool,
    auto_remove: bool,
    ports: dict | None,
) -> list[str]:
    """Construct the ``docker run ...`` argv with safe defaults baked in.

    Everything user-influenced (command, env values, workdir, ports) is
    splatted as a separate argv element — never concatenated into a shell
    string. The container's own ``sh -c`` interprets the command, but the
    host shell is never involved, so injection at the host level is
    impossible regardless of what the model sends.
    """
    argv: list[str] = ["docker", "run"]
    if detach:
        argv.append("-d")
    if auto_remove:
        argv.append("--rm")
    argv += ["--name", container_name]
    # Hardening flag — refuses any future setuid/setcap-driven escalation
    # inside the container even if the image has a buggy binary on PATH.
    argv += ["--security-opt", "no-new-privileges"]
    if network in _DOCKER_NETWORK_MODES:
        argv += ["--network", network]
    if memory:
        argv += ["--memory", memory]
    if cpus:
        argv += ["--cpus", cpus]
    if mount_workspace:
        host_path = str(Path(cwd).resolve())
        mode = "rw" if (mount_mode or "ro").lower() == "rw" else "ro"
        argv += ["-v", f"{host_path}:/workspace:{mode}"]
    if workdir:
        argv += ["-w", workdir]
    elif mount_workspace:
        argv += ["-w", "/workspace"]
    for k, v in (env or {}).items():
        if not isinstance(k, str) or not k:
            continue
        argv += ["-e", f"{k}={v}"]
    for host_port, container_port in (ports or {}).items():
        try:
            hp = int(host_port)
            cp = int(container_port)
        except (TypeError, ValueError):
            continue
        if not (1 <= hp <= 65535 and 1 <= cp <= 65535):
            continue
        argv += ["-p", f"{hp}:{cp}"]
    argv.append(image)
    if command:
        # Pass the user's command through the *container's* ``sh -c`` so
        # multi-statement scripts (pipes, ``&&``, redirects) work. The host
        # shell never sees this string — it's argv[3+] inside the container.
        argv += ["sh", "-c", command]
    return argv


async def _docker_pull(image: str, timeout: int = 300) -> dict:
    """Pull an image, returning the same shape as a tool call.

    Used both by ``docker_pull`` (model-callable) and as the optional
    pre-step for ``docker_run`` / ``docker_run_bg`` when ``auto_pull=True``.
    Long timeout because first-time image pulls can be hundreds of MB.
    """
    try:
        proc = await asyncio.create_subprocess_exec(
            "docker", "pull", image,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )
        try:
            stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=timeout)
        except asyncio.TimeoutError:
            proc.kill()
            await proc.wait()
            return {"ok": False, "output": "", "error": f"docker pull timed out after {timeout}s"}
        out = stdout.decode("utf-8", "replace") if stdout else ""
        return {
            "ok": proc.returncode == 0,
            "output": _clip(out),
            "exit_code": proc.returncode,
        }
    except Exception as e:
        return {"ok": False, "output": "", "error": f"docker pull: {type(e).__name__}: {e}"}


async def docker_run(
    image: str,
    command: str | None = None,
    cwd: str = ".",
    workdir: str | None = None,
    mount_workspace: bool = True,
    mount_mode: str = "ro",
    env: dict | None = None,
    network: str = "bridge",
    memory: str = "512m",
    cpus: str = "1.0",
    timeout: int = 120,
    auto_pull: bool = True,
) -> dict:
    """Run a one-shot command in a container and return its captured output.

    Mirrors the ``bash`` ergonomics the agent already knows: synchronous,
    captures stdout+stderr, returns when the container exits or the timeout
    fires. ``--rm`` is set so the container is auto-removed on exit and the
    host stays clean even on crashes.
    """
    ok, msg = _docker_available()
    if not ok:
        return {"ok": False, "output": "", "error": msg}
    bad = _validate_docker_image(image)
    if bad:
        return {"ok": False, "output": "", "error": bad}
    network_norm = (network or "bridge").lower()
    if network_norm not in _DOCKER_NETWORK_MODES:
        return {
            "ok": False, "output": "",
            "error": f"network must be one of {sorted(_DOCKER_NETWORK_MODES)}",
        }
    timeout = max(1, min(int(timeout or 120), 600))

    # Optional pre-pull. If the image is already cached locally this is a
    # near-instant no-op; on first use it can take minutes.
    pull_summary = ""
    if auto_pull:
        pull = await _docker_pull(image)
        if not pull["ok"]:
            return pull
        # Compress the verbose per-layer output to one line so we don't
        # waste context — the model only needs to know that the image is
        # ready.
        for line in (pull.get("output") or "").splitlines():
            if line.startswith("Status:"):
                pull_summary = line.strip()
                break

    container_name = _new_container_name(None)
    argv = _build_run_argv(
        image=image, container_name=container_name, command=command,
        workdir=workdir, cwd=cwd, mount_workspace=mount_workspace,
        mount_mode=mount_mode, env=env, network=network_norm,
        memory=memory, cpus=cpus, detach=False, auto_remove=True,
        ports=None,
    )
    try:
        proc = await asyncio.create_subprocess_exec(
            *argv,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )
        try:
            stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=timeout)
        except asyncio.TimeoutError:
            # The container is still running — kill it via docker so --rm
            # can clean it up. Killing only the host process leaves the
            # container running.
            try:
                kill = await asyncio.create_subprocess_exec(
                    "docker", "kill", container_name,
                    stdout=asyncio.subprocess.DEVNULL,
                    stderr=asyncio.subprocess.DEVNULL,
                )
                await asyncio.wait_for(kill.wait(), timeout=10)
            except Exception:
                pass
            proc.kill()
            await proc.wait()
            return {"ok": False, "output": "", "error": f"docker_run timed out after {timeout}s"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"docker run: {type(e).__name__}: {e}"}

    out = stdout.decode("utf-8", "replace") if stdout else ""
    if pull_summary:
        out = f"[{pull_summary}]\n{out}"
    return {
        "ok": proc.returncode == 0,
        "output": _clip(out),
        "exit_code": proc.returncode,
        "container": container_name,
    }


async def docker_run_bg(
    image: str,
    command: str | None = None,
    cwd: str = ".",
    workdir: str | None = None,
    mount_workspace: bool = True,
    mount_mode: str = "ro",
    env: dict | None = None,
    network: str = "bridge",
    memory: str = "512m",
    cpus: str = "1.0",
    ports: dict | None = None,
    auto_pull: bool = True,
    conv_id: str | None = None,
) -> dict:
    """Launch a long-running container in the background. Returns its name.

    Use for dev servers, watchers, ML inference daemons, anything that
    doesn't terminate quickly. Poll output with ``docker_logs(name)``,
    run extra commands with ``docker_exec(name, ...)``, stop and remove
    with ``docker_stop(name)``.
    """
    ok, msg = _docker_available()
    if not ok:
        return {"ok": False, "output": "", "error": msg}
    bad = _validate_docker_image(image)
    if bad:
        return {"ok": False, "output": "", "error": bad}
    network_norm = (network or "bridge").lower()
    if network_norm not in _DOCKER_NETWORK_MODES:
        return {
            "ok": False, "output": "",
            "error": f"network must be one of {sorted(_DOCKER_NETWORK_MODES)}",
        }

    if auto_pull:
        pull = await _docker_pull(image)
        if not pull["ok"]:
            return pull

    container_name = _new_container_name(conv_id)
    argv = _build_run_argv(
        image=image, container_name=container_name, command=command,
        workdir=workdir, cwd=cwd, mount_workspace=mount_workspace,
        mount_mode=mount_mode, env=env, network=network_norm,
        memory=memory, cpus=cpus, detach=True, auto_remove=False,
        ports=ports,
    )
    try:
        proc = await asyncio.create_subprocess_exec(
            *argv,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )
        stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=60)
        if proc.returncode != 0:
            out = stdout.decode("utf-8", "replace") if stdout else ""
            return {
                "ok": False, "output": _clip(out),
                "error": f"docker run -d failed (exit {proc.returncode})",
            }
    except asyncio.TimeoutError:
        return {"ok": False, "output": "", "error": "docker run -d timed out after 60s"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"docker run -d: {type(e).__name__}: {e}"}

    cid = stdout.decode("utf-8", "replace").strip() if stdout else ""
    _DOCKER_CONTAINERS[container_name] = {
        "name": container_name,
        "id": cid,
        "image": image,
        "command": command or "",
        "started_at": time.time(),
    }
    return {
        "ok": True,
        "output": (
            f"started container {container_name}\n"
            f"image: {image}\n"
            f"poll with docker_logs({container_name!r}); stop with docker_stop({container_name!r})."
        ),
        "container": container_name,
    }


async def docker_logs(name: str, tail: int = 200) -> dict:
    """Return the most recent stdout+stderr from a Gigachat-started container.

    Refuses to read logs from any container we didn't start — the agent
    can't snoop on user-owned containers running on the same Docker daemon.
    """
    if name not in _DOCKER_CONTAINERS:
        return {
            "ok": False, "output": "",
            "error": f"no Gigachat-managed container named {name!r}",
        }
    ok, msg = _docker_available()
    if not ok:
        return {"ok": False, "output": "", "error": msg}
    tail = max(1, min(int(tail or 200), 5000))
    try:
        proc = await asyncio.create_subprocess_exec(
            "docker", "logs", "--tail", str(tail), name,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )
        stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=30)
    except asyncio.TimeoutError:
        return {"ok": False, "output": "", "error": "docker logs timed out"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"docker logs: {type(e).__name__}: {e}"}
    out = stdout.decode("utf-8", "replace") if stdout else ""

    # Probe state too so the model knows whether to keep polling vs. give up.
    status_line = ""
    try:
        state = await asyncio.create_subprocess_exec(
            "docker", "inspect", "-f",
            "{{.State.Status}} (exit {{.State.ExitCode}})", name,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.DEVNULL,
        )
        state_out, _ = await asyncio.wait_for(state.communicate(), timeout=5)
        if state_out:
            status_line = state_out.decode("utf-8", "replace").strip()
    except Exception:
        pass
    header = f"[container {name} status: {status_line or 'unknown'}]\n"
    return {"ok": True, "output": _clip(header + out)}


async def docker_exec(name: str, command: str, timeout: int = 60) -> dict:
    """Run an additional command inside a Gigachat-managed running container.

    Useful for installing a missing package, running a one-off script, or
    poking at the file system of a long-running service. Refuses to exec in
    containers we don't own, and refuses to escalate to ``--privileged``.
    """
    if name not in _DOCKER_CONTAINERS:
        return {
            "ok": False, "output": "",
            "error": f"no Gigachat-managed container named {name!r}",
        }
    if not command or not command.strip():
        return {"ok": False, "output": "", "error": "command is required"}
    ok, msg = _docker_available()
    if not ok:
        return {"ok": False, "output": "", "error": msg}
    timeout = max(1, min(int(timeout or 60), 600))
    try:
        proc = await asyncio.create_subprocess_exec(
            "docker", "exec", name, "sh", "-c", command,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )
        try:
            stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=timeout)
        except asyncio.TimeoutError:
            proc.kill()
            await proc.wait()
            return {"ok": False, "output": "", "error": f"docker exec timed out after {timeout}s"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"docker exec: {type(e).__name__}: {e}"}
    out = stdout.decode("utf-8", "replace") if stdout else ""
    return {
        "ok": proc.returncode == 0,
        "output": _clip(out),
        "exit_code": proc.returncode,
    }


async def docker_stop(name: str, remove: bool = True) -> dict:
    """Stop (and by default remove) a Gigachat-managed background container.

    Sends SIGTERM with a 5s grace period via ``docker stop``; the daemon
    SIGKILLs after that. ``remove=True`` also runs ``docker rm -f`` so the
    container's filesystem layer is reclaimed.
    """
    if name not in _DOCKER_CONTAINERS:
        return {
            "ok": False, "output": "",
            "error": f"no Gigachat-managed container named {name!r}",
        }
    ok, msg = _docker_available()
    if not ok:
        return {"ok": False, "output": "", "error": msg}
    try:
        stop = await asyncio.create_subprocess_exec(
            "docker", "stop", "-t", "5", name,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )
        await asyncio.wait_for(stop.wait(), timeout=20)
    except asyncio.TimeoutError:
        return {"ok": False, "output": "", "error": "docker stop timed out"}
    except Exception as e:
        return {"ok": False, "output": "", "error": f"docker stop: {type(e).__name__}: {e}"}

    if remove:
        try:
            rm = await asyncio.create_subprocess_exec(
                "docker", "rm", "-f", name,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.DEVNULL,
            )
            await asyncio.wait_for(rm.wait(), timeout=15)
        except Exception:
            pass
    _DOCKER_CONTAINERS.pop(name, None)
    return {"ok": True, "output": f"stopped {name}" + (" and removed" if remove else "")}


async def docker_list() -> dict:
    """List Gigachat-managed containers and their current state.

    Reconciles the registry with the daemon: any container that vanished
    (user removed it manually with ``docker rm``) is dropped from the
    in-memory tracker so the list stays accurate.
    """
    if not _DOCKER_CONTAINERS:
        return {"ok": True, "output": "no Gigachat-managed containers running."}
    ok, msg = _docker_available()
    if not ok:
        return {"ok": False, "output": "", "error": msg}
    lines = []
    for name, info in list(_DOCKER_CONTAINERS.items()):
        try:
            proc = await asyncio.create_subprocess_exec(
                "docker", "inspect", "-f",
                "{{.State.Status}}|{{.State.ExitCode}}",
                name,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.DEVNULL,
            )
            stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=5)
            state = stdout.decode("utf-8", "replace").strip() if stdout else ""
            if not state:
                # Vanished from the daemon — drop the stale registry entry.
                _DOCKER_CONTAINERS.pop(name, None)
                continue
        except Exception:
            state = "unknown"
        uptime = time.time() - info.get("started_at", time.time())
        lines.append(
            f"- {name} [{state}] image={info['image']} uptime={uptime:.0f}s"
        )
    if not lines:
        return {"ok": True, "output": "no Gigachat-managed containers running."}
    return {"ok": True, "output": "\n".join(lines)}


async def docker_pull(image: str) -> dict:
    """Proactively download an image so subsequent ``docker_run`` calls are instant.

    Useful when the model knows it'll need a 1+ GB image (e.g. ``node:20``,
    ``rust:slim``, ``tensorflow/tensorflow``) before the first run, so the
    pull progress doesn't eat into the run's timeout budget.
    """
    ok, msg = _docker_available()
    if not ok:
        return {"ok": False, "output": "", "error": msg}
    bad = _validate_docker_image(image)
    if bad:
        return {"ok": False, "output": "", "error": bad}
    return await _docker_pull(image)


# ---------------------------------------------------------------------------
# Lazy tool loading: tool_search + tool_load
# ---------------------------------------------------------------------------
# Big tool palettes (this codebase ships ~70 built-in schemas, ~18 K tokens
# stringified) eat into a small local model's context window every turn.
# Anthropic's Claude UI exposes the same trade-off as a "Load tools when
# needed" / "Tools always loaded" toggle. We always load on demand.
#
# Flow:
#   1. The system prompt embeds a short manifest — name + 1-line summary
#      per tool — so the model can see what's available without paying the
#      full schema cost.
#   2. The model picks one or more tools and calls
#      `tool_load({"names": [...]})`. The agent loop reads the persisted
#      set on the next iteration and includes those schemas in the
#      `tools=[...]` payload to Ollama (or, for stub-template models, in
#      the prompt-space adapter's XML block).
#   3. From that point on, the model can call those tools natively. The
#      set survives a backend restart so a long conversation doesn't have
#      to re-load on every reboot.
#
# `tool_search(query)` is a fuzzy lookup against the manifest — case-
# insensitive substring against the name + summary — so the model can
# narrow down "I need to read a PDF" → `read_doc` without scanning the
# whole list manually.


# Tool bundles for lazy-load expansion. When the model calls
# `tool_load(["X"])` (or dispatch auto-loads X on first use), every entry
# in the bundle keyed by X also gets loaded. The criterion for bundling:
# the tools share state that's only meaningful as a set — e.g. `bash_bg`
# returns a `shell_id` that ONLY `bash_output` and `kill_shell` consume,
# so loading `bash` without those is a footgun (the conversation gets
# stuck the moment a long-running command is needed because the model
# can't easily realize "oh I should also load bash_bg").
#
# We deliberately do NOT bundle merely-related tools (e.g. read_file +
# write_file) because each can be used independently and bundling
# would re-bloat the schemas list, defeating lazy load. Add an entry
# here only when a real "missing companion → conversation wedges"
# bug surfaces.
#
# Each value is the COMPLETE set to load when ANY key in it is touched —
# we look up by member, so calling `tool_load(["bash_output"])` will
# pull the full shell toolkit too. Keeps the load symmetric regardless
# of which member the model happened to discover first.
_TOOL_BUNDLES: list[set[str]] = [
    {"bash", "bash_bg", "bash_output", "kill_shell"},
]


def _expand_with_bundles(names: list[str]) -> list[str]:
    """Expand `names` to include every member of any bundle one of them
    belongs to. Preserves first-seen order so the originally-requested
    names lead, with bundle siblings appended after. Idempotent: passing
    a fully-expanded set in returns it unchanged."""
    out: list[str] = []
    seen: set[str] = set()
    for n in names:
        if not isinstance(n, str) or n in seen:
            continue
        out.append(n)
        seen.add(n)
        for bundle in _TOOL_BUNDLES:
            if n in bundle:
                for sibling in bundle:
                    if sibling not in seen:
                        out.append(sibling)
                        seen.add(sibling)
                break
    return out


def _full_manifest() -> list[dict]:
    """Return [{name, summary, category}] for every loadable tool.

    Sources merged: built-in schemas (`prompts.TOOL_SCHEMAS`), live MCP
    server schemas, and user-defined Python tools. The meta-tools
    themselves (`tool_search`, `tool_load`) are excluded — they're always
    loaded so listing them in the manifest just adds noise.
    """
    from . import mcp as _mcp
    from . import prompts as _prompts
    seen: set[str] = set()
    out: list[dict] = []
    schemas = (
        list(_prompts.TOOL_SCHEMAS)
        + _mcp.tool_schemas_for_agent()
        + user_tool_schemas()
    )
    for s in schemas:
        fn = s.get("function") or {}
        name = fn.get("name")
        if not name or name in seen:
            continue
        if name in ("tool_search", "tool_load"):
            continue
        seen.add(name)
        desc = (fn.get("description") or "").strip()
        # First sentence — up to first '.', '\n', or 140 chars, whichever
        # comes first. Keeps the manifest line dense but readable.
        cut_idx = len(desc)
        for stop in (". ", ".\n", "\n"):
            i = desc.find(stop)
            if i != -1 and i < cut_idx:
                cut_idx = i
        summary = desc[:cut_idx].strip()
        if len(summary) > 140:
            summary = summary[:137] + "..."
        # Required-field hint — without this the model has no way to
        # know that `bash` needs a `command` field, `read_file` needs a
        # `path`, etc. just from the manifest. In adapter mode (the
        # parser accepts any tool name regardless of the loaded set)
        # the model would then call the tool with only `reason` filled
        # and get a stream of "empty command" errors. The hint is
        # cheap (~30 chars per entry, ~2 KB total on 70 tools) and
        # eliminates the "guess the field name" failure mode entirely.
        params = fn.get("parameters") or {}
        required = [r for r in (params.get("required") or []) if r != "reason"]
        out.append({
            "name": name,
            "summary": summary or "(no description)",
            "category": classify_tool(name),
            "required": required,
        })
    out.sort(key=lambda x: x["name"])
    return out


async def tool_search(
    query: str,
    limit: int = 8,
    conv_id: str | None = None,
) -> dict:
    """Fuzzy-search the tool manifest by name + summary.

    Case-insensitive substring match — every word in `query` must appear
    somewhere in the candidate's name or summary. Returns at most `limit`
    matches, sorted with name-hits first, then summary-hits, then by
    name. The returned entries include a `loaded` flag so the model can
    skip a redundant `tool_load` if the schema is already available this
    conversation.
    """
    q = (query or "").strip().lower()
    if not q:
        return {"ok": False, "output": "", "error": "query is required"}
    n = max(1, min(int(limit or 8), 30))
    manifest = _full_manifest()
    loaded: set[str] = set()
    if conv_id:
        try:
            loaded = set(db.get_loaded_tools(conv_id))
        except Exception:
            loaded = set()
    terms = [t for t in q.split() if t]
    scored: list[tuple[int, dict]] = []
    for entry in manifest:
        name_l = entry["name"].lower()
        summary_l = entry["summary"].lower()
        if not all((t in name_l) or (t in summary_l) for t in terms):
            continue
        # Lower score sorts first.
        score = 0 if any(t in name_l for t in terms) else 1
        scored.append((score, entry))
    scored.sort(key=lambda kv: (kv[0], kv[1]["name"]))
    hits = []
    for _score, entry in scored[:n]:
        hits.append({
            **entry,
            "loaded": entry["name"] in loaded,
        })
    if not hits:
        return {
            "ok": True,
            "output": f"No tools match {query!r}. Try fewer or broader terms.",
        }
    lines = [f"Found {len(hits)} matching tool(s):"]
    for h in hits:
        marker = " (loaded)" if h["loaded"] else ""
        lines.append(f"  • {h['name']}{marker} — {h['summary']}")
    lines.append("")
    lines.append(
        'Call tool_load({"names": ["..."]}) on the ones you want before '
        "calling them."
    )
    return {"ok": True, "output": "\n".join(lines)}


async def tool_load(
    names: list[str] | str,
    conv_id: str | None = None,
) -> dict:
    """Add tool schemas to the conversation's loaded set.

    The next agent-loop iteration will include these schemas in the
    `tools=[...]` payload sent to Ollama, so the model can call them
    natively from that point on. Idempotent — re-loading an already-
    loaded tool is reported as such. Unknown names are reported as
    errors so the model can correct itself instead of silently calling
    a non-existent tool.

    Accepts either a list (canonical) or a single string (forgiving for
    smaller models that sometimes pass a bare name).
    """
    if not conv_id:
        return {
            "ok": False,
            "output": "",
            "error": "tool_load requires a conversation context",
        }
    # Normalize input.
    if isinstance(names, str):
        names_list = [names]
    elif isinstance(names, list):
        names_list = [str(n) for n in names if n]
    else:
        return {
            "ok": False,
            "output": "",
            "error": "names must be a list of tool names",
        }
    if not names_list:
        return {
            "ok": False,
            "output": "",
            "error": "names cannot be empty — pass at least one tool name",
        }
    # Expand bundled tools so the model can't end up holding only half of
    # a state-sharing toolkit. Bundle members are appended after the
    # originally-requested names so reporting still leads with what the
    # model asked for. See `_TOOL_BUNDLES` for the criterion.
    names_list = _expand_with_bundles(names_list)
    manifest_names = {e["name"] for e in _full_manifest()}
    already = set(db.get_loaded_tools(conv_id))
    to_load: list[str] = []
    unknown: list[str] = []
    re_load: list[str] = []
    for n in names_list:
        if n not in manifest_names:
            unknown.append(n)
            continue
        if n in already:
            re_load.append(n)
            continue
        to_load.append(n)
    new_set = db.add_loaded_tools(conv_id, to_load) if to_load else list(already)
    lines: list[str] = []
    if to_load:
        lines.append(f"Loaded {len(to_load)} tool(s): {', '.join(to_load)}")
    if re_load:
        lines.append(
            f"Already loaded ({len(re_load)}): {', '.join(re_load)}"
        )
    if unknown:
        lines.append(
            f"Unknown tool name(s) ({len(unknown)}): {', '.join(unknown)}. "
            f"Use tool_search to find the right name."
        )
    lines.append(
        f"Total loaded this conversation: {len(new_set)}."
    )
    if to_load:
        lines.append("Schemas will be available on your next turn.")
    return {
        "ok": not unknown or bool(to_load or re_load),
        "output": "\n".join(lines),
        "loaded": list(new_set),
        "unknown": unknown,
    }


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------
TOOL_REGISTRY = {
    # file / shell
    "bash": run_bash,
    "bash_bg": bash_bg,
    "bash_output": bash_output,
    "kill_shell": kill_shell,
    "python_exec": python_exec,
    "read_file": read_file,
    "write_file": write_file,
    "edit_file": edit_file,
    "list_dir": list_dir,
    "glob": glob_tool,
    "grep": grep_tool,
    # documents
    "read_doc": read_doc,
    # clipboard
    "clipboard_read": clipboard_read,
    "clipboard_write": clipboard_write,
    # computer use
    "screenshot": take_screenshot,
    "list_monitors": list_monitors,
    "computer_click": computer_click,
    "computer_drag": computer_drag,
    "computer_type": computer_type,
    "computer_key": computer_key,
    "computer_scroll": computer_scroll,
    "computer_mouse_move": computer_mouse_move,
    "click_element": click_element,
    "click_element_id": click_element_id,
    "type_into_element": type_into_element,
    "focus_window": focus_window,
    "open_app": open_app,
    "window_action": window_action,
    "window_bounds": window_bounds,
    "inspect_window": inspect_window,
    "ocr_screenshot": ocr_screenshot,
    "ui_wait": ui_wait,
    "computer_batch": computer_batch,
    "screenshot_window": screenshot_window,
    "list_windows": list_windows,
    # browser automation over Chrome DevTools Protocol
    "browser_tabs": browser_tabs,
    "browser_goto": browser_goto,
    "browser_click": browser_click,
    "browser_type": browser_type,
    "browser_text": browser_text,
    "browser_eval": browser_eval,
    # web
    "web_search": web_search,
    "fetch_url": fetch_url,
    "http_request": http_request,
    # scheduling
    "schedule_task": schedule_task,
    "list_scheduled_tasks": list_scheduled_tasks,
    "cancel_scheduled_task": cancel_scheduled_task,
    # self-paced wake-up — resumes THIS conversation after a delay
    "schedule_wakeup": schedule_wakeup,
    # autonomous loop — recurring self-resume until stop_loop is called
    "start_loop": start_loop,
    "stop_loop": stop_loop,
    # side tasks — flag a drive-by issue as a chip without derailing the turn
    "spawn_task": spawn_task,
    # structured inline question — registered as a stub; dispatch() intercepts
    "ask_user_question": None,
    # git worktree isolation — run risky edits in a throwaway checkout
    "create_worktree": create_worktree,
    "list_worktrees": list_worktrees,
    "remove_worktree": remove_worktree,
    # polling / watch-until-condition
    "monitor": monitor,
    # local semantic doc search (Ollama embeddings)
    "doc_index": doc_index,
    "doc_search": doc_search,
    # codebase-aware semantic search — scoped to the active conversation's cwd
    "codebase_search": codebase_search,
    # semantic search across URL-indexed public documentation sites
    "docs_search": docs_search,
    # planning
    "todo_write": todo_write,
    # long-term memory
    "remember": remember,
    "forget": forget,
    # subagent — registered but dispatched specially because it needs
    # reentry into the agent module; see dispatch() below.
    "delegate": None,
    "delegate_parallel": None,
    # docker / sandboxed containers — run any language or piece of software
    # in an isolated container without polluting the host. Foreground +
    # background variants mirror the bash / bash_bg pair the agent already
    # knows. Logs / exec / stop only act on containers Gigachat itself
    # started, tracked in the in-memory _DOCKER_CONTAINERS registry.
    "docker_run": docker_run,
    "docker_run_bg": docker_run_bg,
    "docker_logs": docker_logs,
    "docker_exec": docker_exec,
    "docker_stop": docker_stop,
    "docker_list": docker_list,
    "docker_pull": docker_pull,
    # lazy tool loading meta-tools — always available, never lazy
    "tool_search": tool_search,
    "tool_load": tool_load,
}


# Tools that are always loaded into every conversation's tool palette,
# regardless of the lazy-load gate. Without these the model has no way to
# discover or activate the rest of the toolbelt — they're the bootstrap
# pair that the manifest section of the system prompt instructs the model
# to use first.
ALWAYS_LOADED_TOOLS: tuple[str, ...] = ("tool_search", "tool_load")


# ---------------------------------------------------------------------------
# Tool categorisation for the per-conversation permission_mode
# ---------------------------------------------------------------------------
# Every tool is either `"read"` or `"write"`. The agent loop uses this to:
#
#   * `read_only`     — refuse write tools outright (the model gets a
#                       `permission_denied` error and can try a read instead).
#   * `approve_edits` — run read tools silently, pause write tools for manual
#                       approval. This is the zero-friction default: the user
#                       no longer has to click Approve on harmless `read_file`
#                       calls, but destructive work still prompts.
#   * `allow_all`     — nothing pauses; behaves like the old auto_approve=1.
#
# The classification errs toward "write" whenever a tool has any observable
# external side-effect (files, network mutation, desktop control, shell
# execution, subagent spawn). Pure informational / UI tools (todos, memory)
# are "read" because they don't touch anything outside the agent.
TOOL_CATEGORIES: dict[str, str] = {
    # file / shell
    "bash": "write",
    "bash_bg": "write",
    "bash_output": "read",
    "kill_shell": "write",
    "python_exec": "write",
    "read_file": "read",
    "write_file": "write",
    "edit_file": "write",
    "list_dir": "read",
    "glob": "read",
    "grep": "read",
    # documents
    "read_doc": "read",
    # clipboard
    "clipboard_read": "read",
    "clipboard_write": "write",
    # computer use
    "screenshot": "read",
    "list_monitors": "read",
    "computer_click": "write",
    "computer_drag": "write",
    "computer_type": "write",
    "computer_key": "write",
    "computer_scroll": "write",
    "computer_mouse_move": "write",
    "click_element": "write",
    "click_element_id": "write",
    "type_into_element": "write",
    "focus_window": "write",
    "open_app": "write",
    "window_action": "write",
    "window_bounds": "write",
    "inspect_window": "read",
    "ocr_screenshot": "read",
    "ui_wait": "read",
    "computer_batch": "write",
    "screenshot_window": "read",
    "list_windows": "read",
    # browser automation — `browser_text` / `browser_tabs` just read state,
    # the rest mutate either the tab's URL / DOM or execute JS.
    "browser_tabs": "read",
    "browser_goto": "write",
    "browser_click": "write",
    "browser_type": "write",
    "browser_text": "read",
    "browser_eval": "write",
    # web
    "web_search": "read",
    "fetch_url": "read",
    # http_request is write-class regardless of method — even GET can trigger
    # side effects on a poorly-designed API, and the model shouldn't be able
    # to hit arbitrary endpoints in read-only mode.
    "http_request": "write",
    # scheduling — queuing a future run has real external effect, so write.
    "schedule_task": "write",
    "list_scheduled_tasks": "read",
    "cancel_scheduled_task": "write",
    # schedule_wakeup queues a future re-entry into THIS conversation — still
    # a write because a future turn will fire unattended.
    "schedule_wakeup": "write",
    # start_loop / stop_loop — same rationale as schedule_wakeup: future
    # unattended re-entries. stop_loop is a write because it mutates the
    # scheduled_tasks table even though its net effect is "undo".
    "start_loop": "write",
    "stop_loop": "write",
    # spawn_task only inserts a row; the actual new conversation is opened
    # explicitly by the user clicking the chip. Classify as "read" so plan
    # mode can still flag drive-by issues without being refused.
    "spawn_task": "read",
    # ask_user_question has no external effect — it pauses the turn and
    # returns whatever the user clicked. "read" keeps it usable in plan mode.
    "ask_user_question": "read",
    # worktrees mutate the git working tree (create a branch, check it out
    # in a new directory); "write" across the board.
    "create_worktree": "write",
    "list_worktrees": "read",
    "remove_worktree": "write",
    # polling — documented as read-only, enforced in-tool.
    "monitor": "read",
    # doc search — indexing reads files and stores embeddings in the local
    # DB; no external effect, so treat as a read so research workflows work
    # in `read_only` mode.
    "doc_index": "read",
    "doc_search": "read",
    # codebase_search is read-only: it only consults the pre-built index.
    "codebase_search": "read",
    # docs_search reads URL-indexed docs only; no fetch, no side effects.
    "docs_search": "read",
    # planning / memory are pure agent-internal bookkeeping. Classifying
    # them as reads keeps the model capable of planning + note-taking even
    # in read_only mode without clicks.
    "todo_write": "read",
    "remember": "read",
    "forget": "read",
    # subagents — their effect is the union of their trimmed toolset, which
    # includes shells. Treat as writes so a `read_only` conversation can't
    # bypass the restriction by delegating.
    "delegate": "write",
    "delegate_parallel": "write",
    # docker — every container start executes user-controlled code (the
    # image's CMD or our `sh -c <command>`), which is the textbook
    # definition of a write-class side effect. Pulling an image also writes
    # to the host's image cache. Listing and reading logs from containers
    # we already own are safe reads.
    "docker_run": "write",
    "docker_run_bg": "write",
    "docker_exec": "write",
    "docker_stop": "write",
    "docker_pull": "write",
    "docker_logs": "read",
    "docker_list": "read",
    # lazy tool loading meta-tools — pure bookkeeping, no external effects
    "tool_search": "read",
    "tool_load": "read",
}


def classify_tool(name: str) -> str:
    """Return 'read' or 'write' for a tool name. MCP tools default to 'write'.

    MCP tools are dynamically registered by third-party servers so we can't
    know their side-effects a priori. Err on the side of requiring approval
    (and refusing them in read_only mode) — a user who trusts a particular
    MCP server can bump the conversation to `allow_all` or `approve_edits`
    and approve individual calls as needed.

    User-defined tools honor the category the agent (or the user, via the
    Settings UI) chose at creation time — stored in `user_tools.category` and
    fetched via ``classify_user_tool``.
    """
    if name.startswith("mcp__"):
        return "write"
    if name in TOOL_CATEGORIES:
        return TOOL_CATEGORIES[name]
    user_cat = classify_user_tool(name)
    if user_cat:
        return user_cat
    return "write"


# Common hallucinated / alias names the model sometimes emits instead of the
# real tool name. Mapped to the canonical computer-use / shell tool so we can
# give a useful error message rather than a bare "unknown tool". We do NOT
# silently dispatch to the mapped tool — the args usually don't line up
# (e.g. a CSS selector vs x/y pixel coordinates), so the model still needs
# to retry with the correct shape.
_TOOL_ALIAS_HINTS = {
    "action:click": "computer_click",
    "action:type": "computer_type",
    "action:key": "computer_key",
    "action:scroll": "computer_scroll",
    "click": "computer_click",
    "type": "computer_type",
    "key_press": "computer_key",
    "press": "computer_key",
    "scroll": "computer_scroll",
    "mouse_move": "computer_mouse_move",
    "move_mouse": "computer_mouse_move",
    "screen_shot": "screenshot",
    "take_screenshot": "screenshot",
    "capture": "screenshot",
    "shell": "bash",
    "run": "bash",
    "exec": "bash",
    "cat": "read_file",
    "open_file": "read_file",
    "create_file": "write_file",
    "save_file": "write_file",
    "ls": "list_dir",
    "search": "grep",
    "find": "glob",
    "fetch": "fetch_url",
    "search_web": "web_search",
    # new-tool aliases — keep the model aimed at the canonical names.
    "drag": "computer_drag",
    "mouse_drag": "computer_drag",
    "monitors": "list_monitors",
    "screens": "list_monitors",
    "window_close": "window_action",
    "close_window": "window_action",
    "minimize_window": "window_action",
    "maximize_window": "window_action",
    "move_window": "window_bounds",
    "resize_window": "window_bounds",
    "get_window": "window_bounds",
    "inspect": "inspect_window",
    "a11y_dump": "inspect_window",
    "pdf": "read_doc",
    "read_pdf": "read_doc",
    "read_docx": "read_doc",
    "read_xlsx": "read_doc",
    "ocr": "ocr_screenshot",
    "tabs": "browser_tabs",
    "goto": "browser_goto",
    "navigate": "browser_goto",
    "browser_go": "browser_goto",
    "browser_click_selector": "browser_click",
    "browser_fill": "browser_type",
    "eval_js": "browser_eval",
    "schedule": "schedule_task",
    "cron": "schedule_task",
    "index_docs": "doc_index",
    "search_docs": "doc_search",
    "rag": "doc_search",
    # parallel subagents / polling
    "parallel": "delegate_parallel",
    "delegate_many": "delegate_parallel",
    "fanout": "delegate_parallel",
    "spawn_parallel": "delegate_parallel",
    "watch": "monitor",
    "poll": "monitor",
    "wait_for": "monitor",
    "until": "monitor",
    # ui-wait + batch + element-id aliases
    "wait": "ui_wait",
    "wait_for_window": "ui_wait",
    "wait_for_element": "ui_wait",
    "wait_for_text": "ui_wait",
    "wait_for_image": "ui_wait",
    "wait_for_change": "ui_wait",
    "ui_poll": "ui_wait",
    "batch": "computer_batch",
    "computer_sequence": "computer_batch",
    "do_steps": "computer_batch",
    "click_id": "click_element_id",
    "click_by_id": "click_element_id",
    # type_into_element aliases — common ways the model might phrase it
    # when reaching for a "fill in this field by name" primitive.
    "type_into": "type_into_element",
    "type_in": "type_into_element",
    "fill_element": "type_into_element",
    "fill_field": "type_into_element",
    "fill_input": "type_into_element",
    "set_field": "type_into_element",
    "enter_text": "type_into_element",
    "type_to_element": "type_into_element",
    # screenshot_window + list_windows aliases
    "window_screenshot": "screenshot_window",
    "screenshot_app": "screenshot_window",
    "capture_window": "screenshot_window",
    "windows": "list_windows",
    "list_apps": "list_windows",
    "open_windows": "list_windows",
    "enumerate_windows": "list_windows",
}


def _normalize_arg_keys(args: dict[str, Any]) -> dict[str, Any]:
    """Return a copy of `args` with keys canonicalised for dispatch lookup.

    Weak tool-callers frequently decorate arg names with leftover schema
    markers or casing drift that the `args.get("command", "")`-style lookups
    in `dispatch` never match — so the tool silently sees an empty string
    and returns a terse error the model can't recover from. Seen in the
    wild:

      * `"command*"` — a required-marker that leaked from the schema render
      * `"Command"` / `"COMMAND"` — case drift on smaller models
      * `"command "` / `"command_"` — stray punctuation / whitespace

    This helper strips trailing non-word characters (ASCII
    punctuation + whitespace, keeping letters/digits/underscore) and
    lowercases. If normalisation produces a collision, the first occurrence
    wins — so a correctly-formed `"command"` beats a garbled
    `"command*"` emitted alongside it.

    Safe for every built-in tool because all built-in schemas use clean
    snake_case keys. MCP/user-tool args never reach this code path — they
    dispatch before we call this.
    """
    if not isinstance(args, dict):
        return args  # let the caller surface the shape error as it sees fit
    out: dict[str, Any] = {}
    # Two-pass: canonical (`k == k.lower() == cleaned`) keys claim their slot
    # first, then typo variants fill whatever's still empty. Prevents a
    # mangled `"command*"` from shadowing a correctly-spelled `"command"`
    # when dict insertion order puts the typo first.
    for k, v in args.items():
        if not isinstance(k, str):
            out[k] = v
            continue
        clean = re.sub(r"[^A-Za-z0-9_]+$", "", k).lower()
        if clean and k == clean:
            out[clean] = v
    for k, v in args.items():
        if not isinstance(k, str):
            continue
        clean = re.sub(r"[^A-Za-z0-9_]+$", "", k).lower()
        if clean and clean not in out:
            out[clean] = v
    return out


def _suggest_tool_name(name: str) -> str:
    """Return a short hint for an unknown tool name.

    First looks at a hard-coded alias table for names the model tends to
    hallucinate (e.g. `action:click` for `computer_click`), then falls back
    to difflib's close-match search against the real registry. Returns an
    empty string when no reasonable suggestion exists, which keeps the error
    message clean.
    """
    key = (name or "").strip().lower()
    if key in _TOOL_ALIAS_HINTS:
        return _TOOL_ALIAS_HINTS[key]
    # Strip any "namespace:action" prefix before fuzzy-matching — the model
    # sometimes emits e.g. `browser:click`, `action:type`, `tool:bash`.
    bare = key.split(":")[-1]
    candidates = list(TOOL_REGISTRY.keys())
    close = difflib.get_close_matches(bare, candidates, n=1, cutoff=0.6)
    return close[0] if close else ""


# ---------------------------------------------------------------------------
# Silent tool-name aliasing.
#
# Weak tool-callers (gemma4:e4b, small qwens, llama 3.2 1b/3b) regularly
# invent plausible-sounding names like `google:search`, `run_shell`,
# `take_screenshot` that don't actually exist. The _TOOL_ALIAS_HINTS table
# above only HINTS — the call still fails with a red "Failed" box, and the
# model has to burn a round-trip correcting itself.
#
# For a hand-picked set of aliases where the argument shape is compatible
# (or cheaply coercible), we silently rewrite the call to the canonical
# tool instead. The user never sees the invented name — the tool_call UI
# shows the real one, and execution proceeds normally.
#
# Aliases included here must satisfy TWO conditions:
#   1. The canonical tool accepts the same shape of arguments (after the
#      optional coercer runs), OR the alias is so unambiguous that arg
#      mismatch is unlikely to matter.
#   2. No sensible user would name a custom tool this — redirecting is
#      safe because the path is only taken when the name doesn't resolve
#      to an MCP tool, a user tool, or a built-in.
# ---------------------------------------------------------------------------

def _coerce_web_search_args(args: dict | None) -> dict:
    """Coerce common hallucinated shapes to `web_search`'s canonical args.

    Handles:
      - `queries: ["..."]` (array form — take the first item)
      - `q` / `search` / `search_query` / `search_term` / `term` / `text`
        as alt keys for `query`
      - `limit` / `count` / `n` as alt keys for `max_results`
    """
    if not args:
        return {}
    out = dict(args)
    if "queries" in out and "query" not in out:
        qs = out.pop("queries")
        if isinstance(qs, list) and qs:
            out["query"] = str(qs[0])
        elif isinstance(qs, str):
            out["query"] = qs
    for alt in ("q", "search", "search_query", "search_term", "term", "text"):
        if alt in out and "query" not in out:
            out["query"] = out.pop(alt)
    for alt in ("limit", "count", "n", "num_results"):
        if alt in out and "max_results" not in out:
            try:
                out["max_results"] = int(out.pop(alt))
            except (TypeError, ValueError):
                out.pop(alt, None)
    return out


def _coerce_fetch_url_args(args: dict | None) -> dict:
    """Rewrite `link` / `href` / `u` → `url`. `fetch_url` takes a single URL."""
    if not args:
        return {}
    out = dict(args)
    for alt in ("link", "href", "u", "page_url", "target"):
        if alt in out and "url" not in out:
            out["url"] = out.pop(alt)
    return out


def _coerce_bash_args(args: dict | None) -> dict:
    """Rewrite `cmd` / `shell_command` / `script` → `command`."""
    if not args:
        return {}
    out = dict(args)
    for alt in ("cmd", "shell_command", "script", "code", "line"):
        if alt in out and "command" not in out:
            out["command"] = out.pop(alt)
    return out


def _coerce_read_file_args(args: dict | None) -> dict:
    """Rewrite `file` / `filepath` / `filename` → `path`."""
    if not args:
        return {}
    out = dict(args)
    for alt in ("file", "filepath", "filename", "file_path", "file_name"):
        if alt in out and "path" not in out:
            out["path"] = out.pop(alt)
    return out


def _coerce_write_file_args(args: dict | None) -> dict:
    """Rewrite `file` / `filepath` → `path`, `body` / `text` → `content`."""
    if not args:
        return {}
    out = dict(args)
    for alt in ("file", "filepath", "filename", "file_path", "file_name"):
        if alt in out and "path" not in out:
            out["path"] = out.pop(alt)
    for alt in ("body", "text", "data", "contents"):
        if alt in out and "content" not in out:
            out["content"] = out.pop(alt)
    return out


# Aliases that silently redirect to the canonical name. Keys are lowercased
# pre-lookup. Value is (canonical_name, optional_args_coercer).
_SILENT_TOOL_REDIRECTS: dict[str, tuple[str, Any]] = {
    # Web search — high-frequency hallucination: models imagine a "google"
    # namespace even though our only search tool is DuckDuckGo-backed.
    "google:search": ("web_search", _coerce_web_search_args),
    "google_search": ("web_search", _coerce_web_search_args),
    "google.search": ("web_search", _coerce_web_search_args),
    "search_google": ("web_search", _coerce_web_search_args),
    "search_web": ("web_search", _coerce_web_search_args),
    "search_internet": ("web_search", _coerce_web_search_args),
    "internet_search": ("web_search", _coerce_web_search_args),
    "search:web": ("web_search", _coerce_web_search_args),
    "web:search": ("web_search", _coerce_web_search_args),
    "bing_search": ("web_search", _coerce_web_search_args),
    "bing:search": ("web_search", _coerce_web_search_args),
    "duckduckgo_search": ("web_search", _coerce_web_search_args),
    "duckduckgo:search": ("web_search", _coerce_web_search_args),
    "ddg_search": ("web_search", _coerce_web_search_args),
    "ddg:search": ("web_search", _coerce_web_search_args),
    "websearch": ("web_search", _coerce_web_search_args),
    "query_web": ("web_search", _coerce_web_search_args),
    # URL fetch
    "fetch": ("fetch_url", _coerce_fetch_url_args),
    "http_get": ("fetch_url", _coerce_fetch_url_args),
    "download_url": ("fetch_url", _coerce_fetch_url_args),
    "get_url": ("fetch_url", _coerce_fetch_url_args),
    "load_url": ("fetch_url", _coerce_fetch_url_args),
    "read_url": ("fetch_url", _coerce_fetch_url_args),
    # Shell
    "run_shell": ("bash", _coerce_bash_args),
    "execute_shell": ("bash", _coerce_bash_args),
    "execute_command": ("bash", _coerce_bash_args),
    "run_command": ("bash", _coerce_bash_args),
    "shell_command": ("bash", _coerce_bash_args),
    "shell:exec": ("bash", _coerce_bash_args),
    "sh": ("bash", _coerce_bash_args),
    # Files
    "cat_file": ("read_file", _coerce_read_file_args),
    "open_file": ("read_file", _coerce_read_file_args),
    "file:read": ("read_file", _coerce_read_file_args),
    "read:file": ("read_file", _coerce_read_file_args),
    "create_file": ("write_file", _coerce_write_file_args),
    "save_file": ("write_file", _coerce_write_file_args),
    "file:write": ("write_file", _coerce_write_file_args),
    "write:file": ("write_file", _coerce_write_file_args),
    # Screenshot
    "take_screenshot": ("screenshot", None),
    "screen_shot": ("screenshot", None),
    "capture_screen": ("screenshot", None),
}


def resolve_tool_alias(name: str, args: dict | None) -> tuple[str, dict]:
    """If `name` is a known silent alias, return the canonical (name, args).

    Only rewrites when the raw name doesn't resolve to a real tool — MCP,
    user-defined, built-in, or delegate. That way a user tool called e.g.
    `fetch` still wins over the alias; the redirect is strictly a
    "nothing else matched, try to recover the intent" fallback.

    Returns the input unchanged if no redirect applies.
    """
    raw = (name or "").strip()
    if not raw:
        return name, args or {}
    # Skip redirect if the name already resolves somewhere real. Imported
    # lazily because tools.py is imported from agent.py and we'd otherwise
    # spin a partial-init loop.
    from . import mcp as _mcp
    if _mcp.is_mcp_tool(raw):
        return name, args or {}
    try:
        if db.get_user_tool_by_name(raw):
            return name, args or {}
    except Exception:
        pass
    if raw in TOOL_REGISTRY or raw in {"delegate", "delegate_parallel"}:
        return name, args or {}
    entry = _SILENT_TOOL_REDIRECTS.get(raw.lower())
    if not entry:
        return name, args or {}
    canonical, coercer = entry
    coerced = coercer(args or {}) if coercer else dict(args or {})
    return canonical, coerced


def _default_subagent_model() -> str:
    """Pick the chat model to use when a delegate tool omits `model`.

    Priority mirrors `app._resolve_default_chat_model`:
      1. User-configured default (Settings → Default model).
      2. Auto-tuner's pick for this hardware.
      3. Hardcoded gemma4:e4b fallback.

    Imported lazily to avoid a module-level circular dependency with db and
    ollama_runtime (both of which indirectly import tools).
    """
    try:
        from . import db as _db, ollama_runtime as _ort
    except Exception:
        return "gemma4:e4b"
    chosen = _db.get_setting("default_chat_model") if hasattr(_db, "get_setting") else None
    if isinstance(chosen, str) and chosen.strip():
        return chosen.strip()
    try:
        rec = _ort.get_recommendation().get("chat_model")
    except Exception:
        rec = None
    if isinstance(rec, str) and rec.strip():
        return rec.strip()
    return "gemma4:e4b"


@timed_tool
async def dispatch(
    name: str,
    args: dict[str, Any],
    cwd: str,
    conv_id: str | None = None,
    model: str | None = None,
    tool_call_id: str | None = None,
) -> dict:
    """Route a tool call to the right function with argument coercion.

    We deliberately pull args out explicitly rather than **-splatting so the
    LLM can't pass unexpected kwargs that might surprise pyautogui. The
    optional `conv_id` / `model` params are plumbed through for tools that
    need conversation context (checkpoints, subagent delegation).

    `tool_call_id` — the LLM-assigned id for THIS tool call. Only used by
    `delegate` / `delegate_parallel` to tag subagent progress events so
    the UI can nest them under the originating tool card.

    Wrapped in `@timed_tool` so every call emits one structured log line
    with `{tool, duration_ms, ok, conv_id, arg_size_bytes, error_kind}`
    — see `backend/telemetry.py`.
    """
    # Safety net: if the agent layer didn't already canonicalize, handle
    # invented names here (e.g. `google:search` → `web_search`). No-op when
    # `name` already resolves to an MCP / user tool / built-in.
    name, args = resolve_tool_alias(name, args or {})

    # Auto-load on first call. The lazy-load filter only ships a small
    # `tools=[...]` payload to the model — the rest are listed by name +
    # 1-liner in the system prompt manifest. In adapter mode the model
    # can still call any tool by name (the parser doesn't gate on the
    # advertised set), so the FIRST call may go in with imperfect args
    # because the schema wasn't visible. Auto-adding to the loaded set
    # here means the NEXT turn carries the full schema, which clears
    # every "I forgot a required field" repeat-failure pattern. Skipped
    # for the meta-tools themselves (always loaded) and for unknown
    # names (would just clutter the persisted set with garbage).
    if conv_id and name not in ALWAYS_LOADED_TOOLS:
        try:
            already = set(db.get_loaded_tools(conv_id))
        except Exception:
            already = set()
        if name not in already and name in TOOL_REGISTRY:
            # Expand into the tool's bundle so the model gets the
            # whole state-sharing toolkit at once, not just the
            # specific member it happened to call. Without this, a
            # `bash` call auto-loads only `bash` — but the moment the
            # task needs a long-running command (`bash_bg`) or a
            # poll on it (`bash_output`), the model has to know to
            # `tool_load` again, and small models often don't.
            bundle = _expand_with_bundles([name])
            try:
                db.add_loaded_tools(
                    conv_id,
                    [n for n in bundle if n in TOOL_REGISTRY],
                )
            except Exception:
                pass

    # MCP tools live outside TOOL_REGISTRY — their names are namespaced
    # (`mcp__<server>__<tool>`) so we can route them without a registry
    # lookup. Imported lazily to avoid an import cycle (mcp.py → db.py,
    # and tools.py is one of db's callers indirectly).
    from . import mcp as _mcp
    if _mcp.is_mcp_tool(name):
        return await _mcp.dispatch_tool(name, args or {})

    # User-defined tools — persisted in the user_tools table, executed in a
    # sandboxed venv. Checked before the built-in registry so names are
    # already guaranteed not to collide (create_tool refuses collisions),
    # but this ordering also lets a future admin shadow a built-in on
    # purpose if we ever want that.
    user_tool_row = None
    try:
        user_tool_row = db.get_user_tool_by_name(name)
    except Exception:
        user_tool_row = None
    if user_tool_row and user_tool_row.get("enabled"):
        # Strip the `reason` arg (display-only, like every other tool) and
        # forward the rest untouched. Subprocess timeout comes from the
        # stored per-tool config rather than a model-controlled arg so the
        # model can't bypass the cap by passing a huge value.
        forwarded = {k: v for k, v in (args or {}).items() if k != "reason"}
        return await _utr.execute_user_tool(
            code=user_tool_row.get("code") or "",
            args=forwarded,
            timeout=int(user_tool_row.get("timeout_seconds") or 60),
            cwd=cwd,
        )

    fn = TOOL_REGISTRY.get(name)
    if fn is None and name not in {"delegate", "delegate_parallel"}:
        hint = _suggest_tool_name(name)
        suffix = f" — did you mean `{hint}`?" if hint else ""
        return {
            "ok": False,
            "output": "",
            "error": (
                f"unknown tool: {name!r}{suffix}. Valid tools: "
                + ", ".join(sorted(TOOL_REGISTRY.keys()))
            ),
        }

    # Normalize arg keys against common LLM typos before reading them below.
    # Models routinely emit `"command*"` (asterisks from a required-marker in
    # the schema rendering), `"Command"` (case drift), or `"command "`
    # (trailing whitespace). Every `args.get("command", "")` lookup would
    # silently see None and pass "" to the tool, which then errored with a
    # terse "empty command" that the model couldn't recover from — it just
    # retried with the same broken key. Stripping trailing non-word chars
    # and lowering case here makes dispatch tolerant of that noise. Built-in
    # schemas all use clean snake_case so collisions are safe: the first
    # occurrence of a normalized key wins.
    args = _normalize_arg_keys(args or {})
    try:
        if name == "bash":
            return await fn(
                cwd,
                args.get("command", ""),
                int(args.get("timeout", DEFAULT_TIMEOUT_SEC)),
                conv_id=conv_id,
            )
        if name == "bash_bg":
            return await fn(cwd, args.get("command", ""), conv_id=conv_id)
        if name == "bash_output":
            return await fn(args.get("shell_id", ""))
        if name == "kill_shell":
            return await fn(args.get("shell_id", ""))
        if name == "python_exec":
            return await fn(
                cwd,
                args.get("code", ""),
                int(args.get("timeout", 60)),
            )
        if name == "read_file":
            return await fn(cwd, args.get("path", ""), conv_id=conv_id)
        if name == "write_file":
            # Checkpoint the existing file (if any) before overwriting.
            try:
                resolved = _resolve(cwd, args.get("path", ""), conv_id)
                if resolved.is_file():
                    _checkpoint_file(conv_id, resolved, resolved.read_bytes())
            except Exception:
                pass
            return await fn(
                cwd,
                args.get("path", ""),
                args.get("content", ""),
                conv_id=conv_id,
            )
        if name == "edit_file":
            return await fn(
                cwd,
                args.get("path", ""),
                args.get("old_string", ""),
                args.get("new_string", ""),
                bool(args.get("replace_all", False)),
                conv_id,
            )
        if name == "list_dir":
            # `args.get("path", ".")` returns "" (not ".") when the model
            # emits `{"path": ""}` — a common pattern when it wants "current
            # directory". Normalize to "." so the call resolves to `cwd`
            # instead of the filesystem root.
            return await fn(cwd, args.get("path") or ".", conv_id=conv_id)
        if name == "glob":
            return await fn(
                cwd,
                args.get("pattern", ""),
                args.get("path") or ".",
                conv_id=conv_id,
            )
        if name == "grep":
            return await fn(
                cwd,
                args.get("pattern", ""),
                args.get("path") or ".",
                args.get("glob"),
                bool(args.get("case_insensitive", False)),
                args.get("output_mode", "files_with_matches"),
                int(args.get("head_limit", 100)),
                conv_id=conv_id,
            )
        if name == "clipboard_read":
            return await fn()
        if name == "clipboard_write":
            return await fn(args.get("text", ""))
        if name == "screenshot":
            # `monitor` is optional; coerce to int only when provided so the
            # take_screenshot default (None → primary) is preserved.
            m = args.get("monitor")
            return await fn(
                int(m) if m is not None else None,
                bool(args.get("with_elements", False)),
            )
        if name == "list_monitors":
            return await fn()
        if name == "computer_click":
            return await fn(
                args.get("x", 0),
                args.get("y", 0),
                args.get("button", "left"),
                bool(args.get("double", False)),
            )
        if name == "computer_type":
            return await fn(args.get("text", ""), args.get("interval", 0.02))
        if name == "computer_key":
            return await fn(args.get("keys", ""))
        if name == "computer_scroll":
            return await fn(
                args.get("x", 0),
                args.get("y", 0),
                args.get("direction", "down"),
                args.get("amount", 5),
            )
        if name == "computer_mouse_move":
            return await fn(args.get("x", 0), args.get("y", 0))
        if name == "click_element":
            return await fn(
                args.get("name", ""),
                args.get("match", "contains"),
                args.get("click_type", "left"),
                args.get("timeout", 2.0),
            )
        if name == "click_element_id":
            return await fn(
                args.get("id", ""),
                args.get("click_type", "left"),
            )
        if name == "type_into_element":
            return await fn(
                args.get("name", ""),
                args.get("text", ""),
                args.get("match", "contains"),
                bool(args.get("clear", False)),
                args.get("interval", 0.02),
            )
        if name == "focus_window":
            return await fn(args.get("name", ""))
        if name == "open_app":
            return await fn(args.get("name", ""), args.get("args"))
        if name == "web_search":
            return await fn(
                args.get("query", ""),
                args.get("max_results", 5),
                args.get("region"),
            )
        if name == "fetch_url":
            return await fn(
                args.get("url", ""),
                args.get("max_chars", FETCH_DEFAULT_MAX_CHARS),
            )
        if name == "http_request":
            return await fn(
                args.get("url", ""),
                args.get("method", "GET"),
                args.get("headers"),
                args.get("body"),
                args.get("query"),
                args.get("timeout", HTTP_REQUEST_TIMEOUT_SEC),
                bool(args.get("allow_private", False)),
                args.get("max_output_chars", HTTP_REQUEST_DEFAULT_OUTPUT_CHARS),
            )
        if name == "todo_write":
            return await fn(args.get("todos", []))
        if name == "remember":
            return await fn(conv_id, args.get("content", ""), args.get("topic"))
        if name == "forget":
            return await fn(conv_id, args.get("pattern", ""))
        if name == "tool_search":
            return await fn(
                args.get("query", ""),
                int(args.get("limit", 8) or 8),
                conv_id=conv_id,
            )
        if name == "tool_load":
            return await fn(
                args.get("names") or args.get("name") or [],
                conv_id=conv_id,
            )
        if name == "computer_drag":
            return await fn(
                args.get("x1", 0),
                args.get("y1", 0),
                args.get("x2", 0),
                args.get("y2", 0),
                args.get("duration", 0.4),
                args.get("button", "left"),
            )
        if name == "window_action":
            return await fn(args.get("name", ""), args.get("action", ""))
        if name == "window_bounds":
            return await fn(
                args.get("name", ""),
                args.get("x"),
                args.get("y"),
                args.get("width"),
                args.get("height"),
            )
        if name == "inspect_window":
            return await fn(
                args.get("name"),
                int(args.get("max_depth", 12)),
                int(args.get("max_nodes", 500)),
                bool(args.get("overlay", True)),
            )
        if name == "screenshot_window":
            return await fn(
                args.get("name", ""),
                bool(args.get("with_elements", False)),
            )
        if name == "list_windows":
            return await fn(int(args.get("max_count", 40)))
        if name == "ui_wait":
            return await fn(
                args.get("kind", ""),
                args.get("target", ""),
                int(args.get("timeout_seconds", 15)),
                float(args.get("interval_seconds", 1.0)),
                bool(args.get("require_enabled", False)),
            )
        if name == "computer_batch":
            return await fn(
                args.get("steps") or [],
                bool(args.get("screenshot", True)),
            )
        if name == "read_doc":
            return await fn(
                args.get("path", ""),
                args.get("pages"),
                args.get("sheets"),
            )
        if name == "ocr_screenshot":
            return await fn(args.get("image_path"), args.get("match"))
        if name == "browser_tabs":
            return await fn(int(args.get("port", 9222)))
        if name == "browser_goto":
            return await fn(
                args.get("url", ""),
                args.get("tab_index"),
                int(args.get("port", 9222)),
            )
        if name == "browser_click":
            return await fn(
                args.get("selector", ""),
                args.get("tab_index"),
                int(args.get("port", 9222)),
            )
        if name == "browser_type":
            return await fn(
                args.get("selector", ""),
                args.get("text", ""),
                bool(args.get("press_enter", False)),
                args.get("tab_index"),
                int(args.get("port", 9222)),
            )
        if name == "browser_text":
            return await fn(
                args.get("selector", "body"),
                args.get("tab_index"),
                int(args.get("port", 9222)),
                int(args.get("max_chars", 15000)),
            )
        if name == "browser_eval":
            return await fn(
                args.get("expression", ""),
                args.get("tab_index"),
                int(args.get("port", 9222)),
            )
        if name == "schedule_task":
            return await fn(
                args.get("name", ""),
                args.get("prompt", ""),
                args.get("run_at"),
                args.get("every_minutes"),
                args.get("cwd") or cwd,
            )
        if name == "list_scheduled_tasks":
            return await fn()
        if name == "cancel_scheduled_task":
            return await fn(args.get("id", ""))
        if name == "schedule_wakeup":
            return await fn(
                args.get("delay_seconds", 60),
                args.get("note", ""),
                conversation_id=conv_id,
                cwd=args.get("cwd") or cwd,
            )
        if name == "start_loop":
            return await fn(
                args.get("goal", ""),
                interval_seconds=args.get("interval_seconds", 300),
                conversation_id=conv_id,
                cwd=args.get("cwd") or cwd,
            )
        if name == "stop_loop":
            return await fn(conversation_id=conv_id)
        if name == "spawn_task":
            return await fn(
                args.get("title", ""),
                args.get("prompt", ""),
                args.get("tldr", ""),
                conversation_id=conv_id,
            )
        if name == "ask_user_question":
            # The tool function itself is a stub — the agent loop intercepts
            # this name BEFORE dispatch() and routes it through the
            # await_user_answer SSE path. If we got here it means the agent
            # layer didn't intercept (e.g. a subagent tried to call it).
            return {
                "ok": False, "output": "",
                "error": (
                    "ask_user_question can only be used from the main "
                    "conversation loop — subagents cannot prompt the user. "
                    "Return control to the parent agent first."
                ),
            }
        if name == "create_worktree":
            return await fn(
                args.get("branch", ""),
                args.get("base_ref", "HEAD"),
                conversation_id=conv_id,
                cwd=args.get("cwd") or cwd,
            )
        if name == "list_worktrees":
            return await fn(conversation_id=conv_id)
        if name == "remove_worktree":
            return await fn(args.get("id", ""), conversation_id=conv_id)
        if name == "doc_index":
            return await fn(
                args.get("path", ""),
                args.get("extensions"),
                args.get("model"),
            )
        if name == "doc_search":
            return await fn(
                args.get("query", ""),
                int(args.get("top_k", 5)),
                args.get("path_glob"),
                args.get("model"),
            )
        if name == "codebase_search":
            return await fn(
                args.get("query", ""),
                top_k=int(args.get("top_k", 8)),
                conversation_id=conv_id,
                cwd=args.get("cwd") or cwd,
            )
        if name == "delegate":
            # Lazy import breaks a circular dependency (agent imports tools).
            from .agent import run_subagent, SUBAGENT_TYPES
            subtype = str(args.get("type") or "general").strip().lower()
            if subtype not in SUBAGENT_TYPES:
                return {
                    "ok": False, "output": "",
                    "error": (
                        f"unknown subagent type {subtype!r}. Valid: "
                        + ", ".join(SUBAGENT_TYPES.keys())
                    ),
                }
            return await run_subagent(
                task=args.get("task", ""),
                cwd=cwd,
                model=model or _default_subagent_model(),
                max_iterations=int(args.get("max_iterations", 10)),
                subagent_type=subtype,
                # Subagent publishes progress events to the parent turn's
                # bus so the UI can show nested activity under this call.
                parent_conv_id=conv_id,
                parent_tool_call_id=tool_call_id,
            )
        if name == "delegate_parallel":
            # Fan-out delegate. Same circular-import workaround as `delegate`.
            from .agent import run_subagents_parallel, SUBAGENT_TYPES
            subtype = str(args.get("type") or "general").strip().lower()
            if subtype not in SUBAGENT_TYPES:
                return {
                    "ok": False, "output": "",
                    "error": (
                        f"unknown subagent type {subtype!r}. Valid: "
                        + ", ".join(SUBAGENT_TYPES.keys())
                    ),
                }
            return await run_subagents_parallel(
                tasks=args.get("tasks") or [],
                cwd=cwd,
                model=model or _default_subagent_model(),
                max_iterations=int(args.get("max_iterations", 10)),
                subagent_type=subtype,
                parent_conv_id=conv_id,
                parent_tool_call_id=tool_call_id,
            )
        if name == "monitor":
            return await fn(
                args.get("target", ""),
                args.get("condition", ""),
                int(args.get("interval_seconds", 5)),
                int(args.get("timeout_seconds", 120)),
                cwd,
            )
        if name == "docker_run":
            return await fn(
                args.get("image", ""),
                args.get("command"),
                cwd,
                args.get("workdir"),
                bool(args.get("mount_workspace", True)),
                args.get("mount_mode", "ro"),
                args.get("env") if isinstance(args.get("env"), dict) else None,
                args.get("network", "bridge"),
                args.get("memory", "512m"),
                args.get("cpus", "1.0"),
                int(args.get("timeout", 120)),
                bool(args.get("auto_pull", True)),
            )
        if name == "docker_run_bg":
            return await fn(
                args.get("image", ""),
                args.get("command"),
                cwd,
                args.get("workdir"),
                bool(args.get("mount_workspace", True)),
                args.get("mount_mode", "ro"),
                args.get("env") if isinstance(args.get("env"), dict) else None,
                args.get("network", "bridge"),
                args.get("memory", "512m"),
                args.get("cpus", "1.0"),
                args.get("ports") if isinstance(args.get("ports"), dict) else None,
                bool(args.get("auto_pull", True)),
                conv_id,
            )
        if name == "docker_logs":
            return await fn(args.get("name", ""), int(args.get("tail", 200)))
        if name == "docker_exec":
            return await fn(
                args.get("name", ""),
                args.get("command", ""),
                int(args.get("timeout", 60)),
            )
        if name == "docker_stop":
            return await fn(args.get("name", ""), bool(args.get("remove", True)))
        if name == "docker_list":
            return await fn()
        if name == "docker_pull":
            return await fn(args.get("image", ""))
    except Exception as e:
        return {"ok": False, "output": "", "error": f"{type(e).__name__}: {e}"}
    # Tool was found in TOOL_REGISTRY but no `if name == ...` branch above
    # routed it. This should never happen in a well-maintained dispatcher —
    # it almost always means somebody added a tool to TOOL_REGISTRY without
    # wiring a branch here. Surface that loudly instead of hiding it.
    return {
        "ok": False,
        "output": "",
        "error": (
            f"internal error: tool {name!r} is registered but has no dispatch "
            f"branch in tools.dispatch(). Add one."
        ),
    }


def describe_tool_call(name: str, args: dict[str, Any]) -> str:
    """Short one-line label shown in the UI on every tool-call card.

    Kept intentionally terse — the full argument object is rendered beneath
    the label by the frontend when the card is expanded.
    """
    # MCP tool call — show the namespaced name plus a hint of the first
    # positional-ish arg so users can tell calls apart in the tool panel.
    if name.startswith("mcp__"):
        hint = ""
        for key in ("query", "q", "name", "path", "url", "text"):
            v = args.get(key)
            if isinstance(v, str) and v:
                hint = f": {v[:80]}" + ("..." if len(v) > 80 else "")
                break
        return f"{name}{hint}"
    if name == "bash":
        cmd = args.get("command", "")
        return f"bash: {cmd[:120]}" + ("..." if len(cmd) > 120 else "")
    if name == "bash_bg":
        cmd = args.get("command", "")
        return f"bash_bg: {cmd[:100]}" + ("..." if len(cmd) > 100 else "")
    if name == "bash_output":
        return f"bash_output: {args.get('shell_id', '')}"
    if name == "kill_shell":
        return f"kill_shell: {args.get('shell_id', '')}"
    if name == "read_file":
        return f"read_file: {args.get('path', '')}"
    if name == "write_file":
        return f"write_file: {args.get('path', '')}"
    if name == "edit_file":
        return f"edit_file: {args.get('path', '')}"
    if name == "list_dir":
        return f"list_dir: {args.get('path', '.')}"
    if name == "glob":
        return f"glob: {args.get('pattern', '')}"
    if name == "grep":
        p = str(args.get("pattern", ""))
        return f"grep: {p[:80]}" + ("..." if len(p) > 80 else "")
    if name == "clipboard_read":
        return "clipboard_read"
    if name == "clipboard_write":
        t = str(args.get("text", ""))
        return f"clipboard_write: {t[:60]}" + ("..." if len(t) > 60 else "")
    if name == "screenshot":
        return "screenshot"
    if name == "computer_click":
        btn = args.get("button", "left")
        dbl = " (double)" if args.get("double") else ""
        return f"computer_click: {btn}{dbl} at ({args.get('x', '?')}, {args.get('y', '?')})"
    if name == "computer_type":
        t = str(args.get("text", ""))
        preview = t if len(t) <= 60 else (t[:60] + "…")
        return f"computer_type: {preview!r}"
    if name == "computer_key":
        return f"computer_key: {args.get('keys', '')}"
    if name == "computer_scroll":
        return (
            f"computer_scroll: {args.get('direction', 'down')} x{args.get('amount', 5)} "
            f"at ({args.get('x', '?')}, {args.get('y', '?')})"
        )
    if name == "computer_mouse_move":
        return f"computer_mouse_move: ({args.get('x', '?')}, {args.get('y', '?')})"
    if name == "web_search":
        q = str(args.get("query", ""))
        return f"web_search: {q[:100]}" + ("..." if len(q) > 100 else "")
    if name == "fetch_url":
        return f"fetch_url: {args.get('url', '')}"
    if name == "http_request":
        return f"http_request: {args.get('method', 'GET')} {args.get('url', '')}"
    if name == "todo_write":
        n = len(args.get("todos") or [])
        return f"todo_write: {n} item{'s' if n != 1 else ''}"
    if name == "remember":
        topic = args.get("topic")
        content = str(args.get("content", ""))
        preview = content[:80] + ("..." if len(content) > 80 else "")
        return f"remember{f' [{topic}]' if topic else ''}: {preview}"
    if name == "forget":
        return f"forget: {args.get('pattern', '')}"
    if name == "delegate":
        t = str(args.get("task", ""))
        return f"delegate: {t[:100]}" + ("..." if len(t) > 100 else "")
    if name == "delegate_parallel":
        tasks = args.get("tasks") or []
        n = len(tasks) if isinstance(tasks, list) else 0
        return f"delegate_parallel: {n} task{'s' if n != 1 else ''}"
    if name == "monitor":
        t = str(args.get("target", ""))
        c = str(args.get("condition", ""))
        return f"monitor: {t[:60]} until {c[:40]}"
    # New tools — keep labels terse so the UI card doesn't wrap awkwardly.
    if name == "list_monitors":
        return "list_monitors"
    if name == "computer_drag":
        btn = args.get("button", "left")
        return (
            f"computer_drag: {btn} from ({args.get('x1', '?')}, {args.get('y1', '?')}) "
            f"to ({args.get('x2', '?')}, {args.get('y2', '?')})"
        )
    if name == "window_action":
        return f"window_action: {args.get('action', '')} {args.get('name', '')!r}"
    if name == "window_bounds":
        if all(args.get(k) is None for k in ("x", "y", "width", "height")):
            return f"window_bounds: read {args.get('name', '')!r}"
        return (
            f"window_bounds: set {args.get('name', '')!r} → "
            f"pos=({args.get('x')},{args.get('y')}) size=({args.get('width')}x{args.get('height')})"
        )
    if name == "inspect_window":
        target = args.get("name") or "(foreground)"
        suffix = "" if args.get("overlay", True) else " no-overlay"
        return f"inspect_window: {target!r}{suffix}"
    if name == "screenshot_window":
        return f"screenshot_window: {args.get('name', '')!r}"
    if name == "list_windows":
        cap = args.get("max_count", 40)
        return f"list_windows (cap {cap})"
    if name == "click_element_id":
        ct = args.get("click_type", "left")
        return f"click_element_id: {ct} {args.get('id', '?')!r}"
    if name == "type_into_element":
        # Show the target element + a short text preview; matches how the
        # other typing tools render in the UI card.
        target = str(args.get("name", "?"))
        t = str(args.get("text", ""))
        preview = t if len(t) <= 40 else (t[:40] + "…")
        clear_flag = " (cleared)" if args.get("clear") else ""
        return f"type_into_element: {target!r}{clear_flag} ← {preview!r}"
    if name == "ui_wait":
        kind = args.get("kind", "")
        target = args.get("target", "")
        if kind == "pixel_change":
            return f"ui_wait: {kind} (timeout {args.get('timeout_seconds', 15)}s)"
        return f"ui_wait: {kind} ~ {str(target)[:60]!r}"
    if name == "computer_batch":
        steps = args.get("steps") or []
        n = len(steps) if isinstance(steps, list) else 0
        first = ""
        if n and isinstance(steps[0], dict):
            first = f" first={steps[0].get('action', '?')}"
        return f"computer_batch: {n} step{'s' if n != 1 else ''}{first}"
    if name == "read_doc":
        return f"read_doc: {args.get('path', '')}"
    if name == "ocr_screenshot":
        m = args.get("match")
        return f"ocr_screenshot" + (f": match={m!r}" if m else "")
    if name == "browser_tabs":
        return f"browser_tabs (port {args.get('port', 9222)})"
    if name == "browser_goto":
        return f"browser_goto: {args.get('url', '')}"
    if name == "browser_click":
        return f"browser_click: {args.get('selector', '')}"
    if name == "browser_type":
        sel = args.get("selector", "")
        t = str(args.get("text", ""))
        preview = t if len(t) <= 40 else (t[:40] + "…")
        enter = " +Enter" if args.get("press_enter") else ""
        return f"browser_type: {sel} ← {preview!r}{enter}"
    if name == "browser_text":
        return f"browser_text: {args.get('selector', 'body')}"
    if name == "browser_eval":
        e = str(args.get("expression", ""))
        return f"browser_eval: {e[:80]}" + ("..." if len(e) > 80 else "")
    if name == "schedule_task":
        when = args.get("run_at") or (f"every {args.get('every_minutes')}min")
        return f"schedule_task: {args.get('name', '')!r} ({when})"
    if name == "list_scheduled_tasks":
        return "list_scheduled_tasks"
    if name == "cancel_scheduled_task":
        return f"cancel_scheduled_task: {args.get('id', '')}"
    if name == "schedule_wakeup":
        return f"schedule_wakeup: in {args.get('delay_seconds', 0)}s"
    if name == "start_loop":
        return f"start_loop: every {args.get('interval_seconds', 300)}s"
    if name == "stop_loop":
        return "stop_loop"
    if name == "spawn_task":
        return f"spawn_task: {str(args.get('title', ''))[:60]}"
    if name == "ask_user_question":
        q = str(args.get("question", ""))
        return f"ask_user_question: {q[:60]}" + ("..." if len(q) > 60 else "")
    if name == "create_worktree":
        return f"create_worktree: branch={args.get('branch', '')!r} off {args.get('base_ref', 'HEAD')}"
    if name == "list_worktrees":
        return "list_worktrees"
    if name == "remove_worktree":
        return f"remove_worktree: {args.get('id', '')}"
    if name == "doc_index":
        return f"doc_index: {args.get('path', '')}"
    if name == "doc_search":
        q = str(args.get("query", ""))
        return f"doc_search: {q[:80]}" + ("..." if len(q) > 80 else "")
    if name == "codebase_search":
        q = str(args.get("query", ""))
        return f"codebase_search: {q[:80]}" + ("..." if len(q) > 80 else "")
    if name == "docker_run":
        cmd = str(args.get("command") or "(image default CMD)")
        return f"docker_run: {args.get('image', '?')} → {cmd[:80]}" + ("..." if len(cmd) > 80 else "")
    if name == "docker_run_bg":
        cmd = str(args.get("command") or "(image default CMD)")
        return f"docker_run_bg: {args.get('image', '?')} → {cmd[:60]}" + ("..." if len(cmd) > 60 else "")
    if name == "docker_logs":
        return f"docker_logs: {args.get('name', '?')} (tail {args.get('tail', 200)})"
    if name == "docker_exec":
        cmd = str(args.get("command", ""))
        return f"docker_exec: {args.get('name', '?')} → {cmd[:80]}" + ("..." if len(cmd) > 80 else "")
    if name == "docker_stop":
        return f"docker_stop: {args.get('name', '?')}"
    if name == "docker_list":
        return "docker_list"
    if name == "docker_pull":
        return f"docker_pull: {args.get('image', '?')}"
    # User-defined tool invocation — render the tool name + a short arg hint
    # so the card is distinguishable from a "mysterious unknown name" error.
    try:
        if db.get_user_tool_by_name(name):
            preview = ""
            for k, v in (args or {}).items():
                if k == "reason":
                    continue
                if isinstance(v, str) and v:
                    preview = f": {k}={v[:60]}" + ("..." if len(v) > 60 else "")
                    break
            return f"{name}{preview}"
    except Exception:
        pass
    return f"{name}: {args}"


# ---------------------------------------------------------------------------
# Checkpoint restore
#
# Called by the /api/conversations/{id}/restore/{message_id} endpoint.
# Reads the sidecar .path file, copies the .bin back to its original
# location. Best-effort: missing checkpoints are a no-op.
# ---------------------------------------------------------------------------
def restore_checkpoint(conv_id: str, stamp: str) -> dict:
    """Restore every file snapshotted under CHECKPOINT_DIR/<conv_id>/<stamp>."""
    folder = CHECKPOINT_DIR / conv_id / stamp
    if not folder.is_dir():
        return {"ok": False, "error": "checkpoint not found"}
    restored: list[str] = []
    for bin_path in folder.glob("*.bin"):
        sidecar = bin_path.with_suffix(".path")
        if not sidecar.is_file():
            continue
        try:
            target = Path(sidecar.read_text(encoding="utf-8").strip())
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes(bin_path.read_bytes())
            restored.append(str(target))
        except Exception:
            continue
    return {"ok": True, "restored": restored}
